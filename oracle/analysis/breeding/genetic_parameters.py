import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
import matplotlib.pyplot as plt
import seaborn as sns
import base64

class GeneticParameterAnalyzer:
    def __init__(self, df, genotype_col, rep_col, traits):
        self.df = df
        self.genotype_col = genotype_col
        self.rep_col = rep_col
        self.traits = traits
        self.results = {}

    def validate(self):
        self.df[self.genotype_col] = self.df[self.genotype_col].astype(str)
        self.df[self.rep_col] = self.df[self.rep_col].astype(str)
        for trait in self.traits:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        
        if self.df[self.traits].isnull().any().any():
            self.df = self.df.dropna(subset=self.traits)
            
        self.g = len(self.df[self.genotype_col].unique())
        self.r = len(self.df[self.rep_col].unique())
        
        if len(self.df) != self.g * self.r:
            raise ValueError("Data must be a balanced RCBD (every genotype in every replication).")

    def run_analysis(self):
        for trait in self.traits:
            self.results[trait] = self._analyze_single_trait(trait)

    def _analyze_single_trait(self, trait):
        y = self.df[trait].values
        g_factor = self.df[self.genotype_col].values
        r_factor = self.df[self.rep_col].values
        
        g_num = self.g
        r_num = self.r
        
        total_mean = np.mean(y)
        total_sum = np.sum(y)
        cf = (total_sum ** 2) / (g_num * r_num)
        
        ss_total = np.sum(y ** 2) - cf
        
        rep_sums = self.df.groupby(self.rep_col)[trait].sum()
        ss_rep = np.sum(rep_sums ** 2) / g_num - cf
        
        geno_sums = self.df.groupby(self.genotype_col)[trait].sum()
        ss_geno = np.sum(geno_sums ** 2) / r_num - cf
        
        ss_error = ss_total - ss_rep - ss_geno
        
        df_rep = r_num - 1
        df_geno = g_num - 1
        df_error = (g_num - 1) * (r_num - 1)
        df_total = (g_num * r_num) - 1
        
        ms_rep = ss_rep / df_rep
        ms_geno = ss_geno / df_geno
        ms_error = ss_error / df_error
        
        f_geno = ms_geno / ms_error
        p_geno = 1 - stats.f.cdf(f_geno, df_geno, df_error)
        
        f_rep = ms_rep / ms_error
        p_rep = 1 - stats.f.cdf(f_rep, df_rep, df_error)
        
        # Variance Components
        sigma2_e = ms_error
        sigma2_g = (ms_geno - ms_error) / r_num
        if sigma2_g < 0: sigma2_g = 0
        
        sigma2_p = sigma2_g + sigma2_e
        
        # Genetic Parameters
        h2 = sigma2_g / sigma2_p if sigma2_p > 0 else 0
        gcv = (np.sqrt(sigma2_g) / total_mean) * 100 if total_mean != 0 else 0
        pcv = (np.sqrt(sigma2_p) / total_mean) * 100 if total_mean != 0 else 0
        ecv = (np.sqrt(sigma2_e) / total_mean) * 100 if total_mean != 0 else 0
        
        k = 2.06 # 5% selection intensity
        ga = k * np.sqrt(sigma2_p) * h2
        ga_percent = (ga / total_mean) * 100 if total_mean != 0 else 0
        
        # LSD and Grouping
        sem = np.sqrt(ms_error / r_num)
        sed = np.sqrt((2 * ms_error) / r_num)
        t_05 = stats.t.ppf(0.975, df_error)
        lsd_05 = t_05 * sed
        
        # Genotype Means and Grouping
        means_df = self.df.groupby(self.genotype_col)[trait].agg(['mean', 'std']).reset_index()
        means_df = means_df.sort_values('mean', ascending=False)
        means_df['sem'] = sem
        
        # Basic Grouping Algorithm
        means = means_df['mean'].values
        genotypes = means_df[self.genotype_col].values
        groups = [""] * len(means)
        current_letter = 'a'
        
        for i in range(len(means)):
            if groups[i] == "":
                groups[i] = chr(ord('a') + len(set(filter(None, groups))))
                for j in range(i + 1, len(means)):
                    if abs(means[i] - means[j]) <= lsd_05:
                        groups[j] += groups[i]
        
        means_df['group'] = groups
        
        # Generate Plot
        plt.figure(figsize=(10, 6))
        sns.boxplot(x=self.genotype_col, y=trait, data=self.df)
        plt.xticks(rotation=45)
        plt.title(f"Genotypic Distribution - {trait}")
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        plt.close()
        buf.seek(0)
        img_b64 = base64.b64encode(buf.read()).decode('utf-8')
        
        return {
            "anova": {
                "SS": [ss_rep, ss_geno, ss_error, ss_total],
                "DF": [df_rep, df_geno, df_error, df_total],
                "MS": [ms_rep, ms_geno, ms_error, None],
                "F": [f_rep, f_geno, None, None],
                "P": [p_rep, p_geno, None, None]
            },
            "parameters": {
                "Mean": total_mean,
                "Min": np.min(y),
                "Max": np.max(y),
                "sigma2_e": sigma2_e,
                "sigma2_g": sigma2_g,
                "sigma2_p": sigma2_p,
                "h2": h2,
                "GCV": gcv,
                "PCV": pcv,
                "ECV": ecv,
                "GA": ga,
                "GAM": ga_percent
            },
            "means": means_df.to_dict(orient='records'),
            "lsd_05": lsd_05,
            "plot": img_b64
        }

    def get_summary_table(self):
        summary = []
        for trait, res in self.results.items():
            p = res['parameters']
            summary.append({
                "Trait": trait,
                "Mean": p['Mean'],
                "Max": p['Max'],
                "Min": p['Min'],
                "GCV": p['GCV'],
                "PCV": p['PCV'],
                "ECV": p['ECV'],
                "h2": p['h2'],
                "GA": p['GA'],
                "GAM": p['GAM']
            })
        return pd.DataFrame(summary)

    def create_report(self):
        doc = Document()
        doc.add_heading('Genetic Parameter Estimation Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        for trait, res in self.results.items():
            doc.add_heading(f"Trait: {trait}", level=1)
            
            # ANOVA
            doc.add_heading("ANOVA Table", level=2)
            anova = res['anova']
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, h in enumerate(['Source', 'DF', 'SS', 'MS', 'F', 'P-value']): hdr[i].text = h
            
            sources = ['Replication', 'Genotype', 'Error', 'Total']
            for i in range(4):
                row = table.add_row().cells
                row[0].text = sources[i]
                row[1].text = str(anova['DF'][i])
                row[2].text = f"{anova['SS'][i]:.4f}"
                row[3].text = f"{anova['MS'][i]:.4f}" if anova['MS'][i] else "-"
                row[4].text = f"{anova['F'][i]:.4f}" if anova['F'][i] else "-"
                row[5].text = f"{anova['P'][i]:.4f}" if anova['P'][i] else "-"

            # Parameters
            doc.add_heading("Genetic Parameters", level=2)
            p = res['parameters']
            p_table = doc.add_table(rows=0, cols=2)
            p_table.style = 'Table Grid'
            param_data = [
                ("Mean", f"{p['Mean']:.4f}"),
                ("Environmental Variance (sigma2e)", f"{p['sigma2_e']:.4f}"),
                ("Genotypic Variance (sigma2g)", f"{p['sigma2_g']:.4f}"),
                ("Phenotypic Variance (sigma2p)", f"{p['sigma2_p']:.4f}"),
                ("Broad Sense Heritability (h2)", f"{p['h2']:.4f}"),
                ("GCV (%)", f"{p['GCV']:.4f}"),
                ("PCV (%)", f"{p['PCV']:.4f}"),
                ("ECV (%)", f"{p['ECV']:.4f}"),
                ("Genetic Advance (GA)", f"{p['GA']:.4f}"),
                ("GA as % of Mean (GAM)", f"{p['GAM']:.4f}%")
            ]
            for name, val in param_data:
                row = p_table.add_row().cells
                row[0].text = name
                row[1].text = val
                
            # Mean Comparison
            doc.add_heading("Mean Comparison (LSD)", level=2)
            m_table = doc.add_table(rows=1, cols=5)
            m_table.style = 'Table Grid'
            for i, h in enumerate(['Genotype', 'Mean', 'Std Dev', 'Std Error', 'Group']): m_table.rows[0].cells[i].text = h
            for m in res['means']:
                row = m_table.add_row().cells
                row[0].text = str(m[self.genotype_col])
                row[1].text = f"{m['mean']:.4f}"
                row[2].text = f"{m['std']:.4f}"
                row[3].text = f"{m['sem']:.4f}"
                row[4].text = m['group']

        doc.add_heading("Summary All Traits", level=1)
        summary_df = self.get_summary_table()
        s_table = doc.add_table(rows=1, cols=len(summary_df.columns))
        s_table.style = 'Table Grid'
        for i, col in enumerate(summary_df.columns): s_table.rows[0].cells[i].text = col
        for _, r in summary_df.iterrows():
            row = s_table.add_row().cells
            for i, val in enumerate(r):
                row[i].text = f"{val:.4f}" if isinstance(val, float) else str(val)

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            self.get_summary_table().to_excel(writer, sheet_name='Summary_All_Traits', index=False)
            for trait, res in self.results.items():
                pd.DataFrame(res['means']).to_excel(writer, sheet_name=f'Means_{trait[:25]}', index=False)
        output.seek(0)
        return output
