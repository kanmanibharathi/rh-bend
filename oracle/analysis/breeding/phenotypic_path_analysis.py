import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import seaborn as sns

class PhenotypicPathAnalyzer:
    def __init__(self, df, genotype_col, rep_col, dependent_var, independent_vars):
        self.df = df
        self.genotype_col = genotype_col
        self.rep_col = rep_col
        self.dependent_var = dependent_var
        self.independent_vars = independent_vars
        self.all_traits = independent_vars + [dependent_var]
        
        self.n = 0 # Number of genotypes
        self.r = 0 # Number of replications
        
        # Results
        self.correlation_matrix = None # All traits
        self.direct_effects = []
        self.indirect_effects_matrix = None
        self.residual_effect = 0
        self.explained_variation = 0
        self.unexplained_variation = 0
        self.variances = {} # trait -> sigma2p
        
    def validate(self):
        self.df[self.genotype_col] = self.df[self.genotype_col].astype(str)
        self.df[self.rep_col] = self.df[self.rep_col].astype(str)
        
        for trait in self.all_traits:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        
        if self.df[self.all_traits].isnull().any().any():
            raise ValueError("Data contains missing or non-numeric values in traits.")
            
        self.genotypes = sorted(self.df[self.genotype_col].unique())
        self.n = len(self.genotypes)
        self.reps = sorted(self.df[self.rep_col].unique())
        self.r = len(self.reps)
        
        # Check balance
        counts = self.df.groupby([self.genotype_col, self.rep_col]).size()
        if len(counts) != self.n * self.r:
            raise ValueError("Unbalanced RCBD detected. All genotypes must exist in all replications.")

    def _compute_phenotypic_correlation(self):
        """Computes phenotypic correlation matrix using genotype mean squares."""
        traits = self.all_traits
        num_traits = len(traits)
        n, r = self.n, self.r
        
        # 1. Phenotypic Variance for each trait (sigma_p2 = MS_genotype)
        trait_stats = {}
        for trait in traits:
            y = self.df[trait].values
            G_total = y.sum()
            CF = (G_total**2) / (n * r)
            
            genotype_sums = self.df.groupby(self.genotype_col)[trait].sum()
            SS_Geno = (genotype_sums**2).sum() / r - CF
            MS_G = SS_Geno / (n - 1)
            
            # Phenotypic variance is the Genotype Mean Square
            var_p = MS_G
            if var_p <= 0: var_p = 1e-9 # Prevent math errors
            
            trait_stats[trait] = {"var_p": var_p, "MS_G": MS_G}
            self.variances[trait] = {"sigma2_p": var_p}

        # 2. Phenotypic Correlation Matrix
        rp_matrix = np.eye(num_traits)
        for i in range(num_traits):
            for j in range(i + 1, num_traits):
                t1, t2 = traits[i], traits[j]
                
                # Cross-product ANOVA for covariance (Cov_p = MS_genotype(XY))
                X = self.df[t1].values
                Y = self.df[t2].values
                
                # Sum of traits trait
                Z = X + Y
                G_total_Z = Z.sum()
                CF_Z = (G_total_Z**2) / (n * r)
                geno_sum_Z = self.df.groupby(self.genotype_col).apply(lambda d: (d[t1] + d[t2]).sum())
                SS_Geno_Z = (geno_sum_Z**2).sum() / r - CF_Z
                MS_G_Z = SS_Geno_Z / (n - 1)
                
                # MS_G(XY) = (MS_G(X+Y) - MS_G(X) - MS_G(Y)) / 2
                cov_p = (MS_G_Z - trait_stats[t1]["MS_G"] - trait_stats[t2]["MS_G"]) / 2
                
                rp = cov_p / np.sqrt(trait_stats[t1]["var_p"] * trait_stats[t2]["var_p"])
                rp = np.clip(rp, -1.0, 1.0)
                
                rp_matrix[i, j] = rp
                rp_matrix[j, i] = rp
                
        self.correlation_matrix = pd.DataFrame(rp_matrix, index=traits, columns=traits)

    def run_analysis(self):
        self._compute_phenotypic_correlation()
        
        indep = self.independent_vars
        dep = self.dependent_var
        
        # 1. R matrix (correlations among independents)
        R = self.correlation_matrix.loc[indep, indep].values
        
        # 2. r_y vector (correlations of independents with dependent)
        r_y = self.correlation_matrix.loc[indep, dep].values
        
        # 3. Direct Effects
        try:
            inv_R = np.linalg.pinv(R)
            P_direct = np.dot(inv_R, r_y)
        except Exception as e:
            raise ValueError(f"Matrix inversion failed: {str(e)}. Check for high multi-collinearity.")
            
        self.direct_effects = P_direct
        
        # 4. Indirect Effects Matrix
        num_indep = len(indep)
        indirect_matrix = np.zeros((num_indep, num_indep))
        for i in range(num_indep):
            for j in range(num_indep):
                if i == j:
                    indirect_matrix[i, j] = P_direct[i]
                else:
                    indirect_matrix[i, j] = R[i, j] * P_direct[j]
                    
        self.indirect_effects_matrix = pd.DataFrame(indirect_matrix, index=indep, columns=indep)
        
        # 5. Residual Effect
        sum_explained = np.sum(r_y * P_direct)
        self.residual_effect = np.sqrt(max(0, 1 - sum_explained))
        self.explained_variation = (1 - self.residual_effect**2) * 100
        self.unexplained_variation = (self.residual_effect**2) * 100

    def get_path_table(self):
        indep = self.independent_vars
        dep = self.dependent_var
        data = []
        for i, trait_i in enumerate(indep):
            row = {"Trait": trait_i}
            row["Direct Effect"] = self.direct_effects[i]
            total_indirect = 0
            for j, trait_j in enumerate(indep):
                if i == j:
                    row[f"via {trait_j}"] = f"({self.direct_effects[i]:.4f})"
                else:
                    val = self.indirect_effects_matrix.loc[trait_i, trait_j]
                    row[f"via {trait_j}"] = val
                    total_indirect += val
            row["Total Indirect"] = total_indirect
            row[f"r_p with {dep}"] = self.correlation_matrix.loc[trait_i, dep]
            data.append(row)
        return pd.DataFrame(data)

    def generate_path_diagram(self):
        """Generates a fan-style Path Diagram for Phenotypic Effects."""
        num_traits = len(self.independent_vars)
        plt.figure(figsize=(14, 10))
        ax = plt.gca()
        ax.set_xlim(-2, 10)
        ax.set_ylim(-1, num_traits)
        plt.axis('off')

        dep_x, dep_y = 0, (num_traits - 1) / 2
        trait_x = 5
        trait_y_coords = np.linspace(num_traits - 1, 0, num_traits)
        
        # Residual
        res_x, res_y = -1.5, dep_y - 1.5
        plt.plot([res_x, res_x, dep_x], [res_y - 0.5, res_y, dep_y], color='blue', lw=1.5)
        plt.text(res_x - 0.2, res_y, f"Residual effect\nSQRT(1-{1-self.residual_effect**2:.4f})", 
                 rotation=90, va='center', ha='right', fontsize=9)
        plt.scatter([res_x], [res_y - 0.5], marker='s', s=100, color='white', edgecolors='blue', zorder=5)

        for i, trait in enumerate(self.independent_vars):
            ty = trait_y_coords[i]
            direct_eff = self.direct_effects[i]
            plt.plot([trait_x, dep_x], [ty, dep_y], color='blue', alpha=0.6, lw=1)
            mid_x, mid_y = (trait_x + dep_x) / 2 - 1.5, (ty + dep_y) / 2
            angle = np.degrees(np.arctan2(ty - dep_y, trait_x - dep_x))
            plt.text(mid_x, mid_y, f"{direct_eff:.3f}  {trait}", rotation=angle, va='bottom', ha='center', fontsize=9)
            plt.scatter([trait_x], [ty], s=150, color='white', edgecolors='blue', zorder=5)
            
        plt.scatter([dep_x], [dep_y], s=200, color='white', edgecolors='blue', zorder=6)
        plt.text(dep_x - 0.5, dep_y, f" {self.dependent_var}", ha='right', va='center', fontweight='bold')

        for i in range(num_traits):
            for j in range(i + 1, num_traits):
                y1, y2 = trait_y_coords[i], trait_y_coords[j]
                rp = self.correlation_matrix.loc[self.independent_vars[i], self.independent_vars[j]]
                center_y, dist = (y1 + y2) / 2, abs(y1 - y2)
                arc_x = trait_x + 0.5 + (dist * 0.4)
                coeffs = np.polyfit([y1, center_y, y2], [trait_x, arc_x, trait_x], 2)
                y_fine = np.linspace(y1, y2, 50)
                plt.plot(np.polyval(coeffs, y_fine), y_fine, color='blue', alpha=0.3, lw=0.8)
                plt.text(arc_x + 0.1, center_y, f"{rp:.3f}", va='center', ha='left', fontsize=8)

        plt.title(f"Phenotypical Path Diagram for {self.dependent_var}", loc='left', fontsize=14, fontweight='bold')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf

    def create_report(self):
        doc = Document()
        title = doc.add_heading("Phenotypic Path Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading("1. Analysis Parameters", level=1)
        doc.add_paragraph(f"Dependent Variable: {self.dependent_var}")
        doc.add_paragraph(f"Independent Variables: {', '.join(self.independent_vars)}")
        doc.add_paragraph(f"Method: RCBD Phenotypic Covariance Partitioning")
        
        doc.add_heading("2. Phenotypic Path Coefficients", level=1)
        path_df = self.get_path_table()
        table = doc.add_table(rows=1, cols=len(path_df.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(path_df.columns): table.cell(0, i).text = col
        for _, row in path_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = f"{val:.4f}" if isinstance(val, (float, np.float64)) else str(val)

        doc.add_heading("3. Residual Analysis", level=1)
        doc.add_paragraph(f"Residual Effect: {self.residual_effect:.4f}")
        doc.add_paragraph(f"Explained Variation: {self.explained_variation:.2f}%")
        
        doc.add_heading("4. Path Influence Diagram", level=1)
        buf = self.generate_path_diagram()
        doc.add_picture(buf, width=Inches(6))
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            self.get_path_table().to_excel(writer, sheet_name='Path_Summary', index=False)
            self.correlation_matrix.to_excel(writer, sheet_name='Correlation_Matrix')
            pd.DataFrame(self.direct_effects, index=self.independent_vars, columns=['Direct_Effect']).to_excel(writer, sheet_name='Direct_Effects')
            self.indirect_effects_matrix.to_excel(writer, sheet_name='Indirect_Effects')
            pd.DataFrame({"Parameter": ["Residual Effect", "Explained %"], "Value": [self.residual_effect, self.explained_variation]}).to_excel(writer, sheet_name='Residual_Effect', index=False)
        output.seek(0)
        return output
