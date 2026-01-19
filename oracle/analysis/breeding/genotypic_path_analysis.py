import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import seaborn as sns

class GenotypicPathAnalyzer:
    def __init__(self, df, genotype_col, rep_col, dependent_var, independent_vars):
        self.df = df
        self.genotype_col = genotype_col
        self.rep_col = rep_col
        self.dependent_var = dependent_var
        self.independent_vars = independent_vars
        self.all_traits = independent_vars + [dependent_var]
        
        self.n = 0 # Number of genotypes
        self.r = 0 # Number of replications
        
        # Results
        self.correlation_matrix = None # All traits
        self.direct_effects = []
        self.indirect_effects_matrix = None
        self.residual_effect = 0
        self.explained_variation = 0
        self.unexplained_variation = 0
        self.variances = {} # trait -> sigma2g, sigma2e
        
    def validate(self):
        self.df[self.genotype_col] = self.df[self.genotype_col].astype(str)
        self.df[self.rep_col] = self.df[self.rep_col].astype(str)
        
        for trait in self.all_traits:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        
        if self.df[self.all_traits].isnull().any().any():
            raise ValueError("Data contains missing or non-numeric values in traits.")
            
        self.genotypes = sorted(self.df[self.genotype_col].unique())
        self.n = len(self.genotypes)
        self.reps = sorted(self.df[self.rep_col].unique())
        self.r = len(self.reps)
        
        # Check balance
        counts = self.df.groupby([self.genotype_col, self.rep_col]).size()
        if len(counts) != self.n * self.r:
            raise ValueError("Unbalanced RCBD detected. All genotypes must exist in all replications.")

    def _compute_genotypic_correlation(self):
        """Computes genotypic correlation matrix using partitioning of covariance."""
        traits = self.all_traits
        num_traits = len(traits)
        n, r = self.n, self.r
        
        # 1. Variance components for each trait
        trait_stats = {}
        for trait in traits:
            y = self.df[trait].values
            G_total = y.sum()
            CF = (G_total**2) / (n * r)
            
            SS_Total = (y**2).sum() - CF
            rep_sums = self.df.groupby(self.rep_col)[trait].sum()
            SS_Rep = (rep_sums**2).sum() / n - CF
            genotype_sums = self.df.groupby(self.genotype_col)[trait].sum()
            SS_Geno = (genotype_sums**2).sum() / r - CF
            SS_Error = SS_Total - SS_Rep - SS_Geno
            
            MS_G = SS_Geno / (n - 1)
            MS_E = SS_Error / ((n - 1) * (r - 1))
            
            var_g = (MS_G - MS_E) / r
            if var_g < 0: var_g = 1e-9 # Prevent sqrt(negative)
            
            trait_stats[trait] = {"var_g": var_g, "MS_G": MS_G, "MS_E": MS_E}
            self.variances[trait] = {"sigma2_g": var_g, "sigma2_e": MS_E}

        # 2. Genotypic Correlation Matrix
        rg_matrix = np.eye(num_traits)
        for i in range(num_traits):
            for j in range(i + 1, num_traits):
                t1, t2 = traits[i], traits[j]
                
                # Cross-product ANOVA for covariance
                X = self.df[t1].values
                Y = self.df[t2].values
                CP_Total = (X * Y).sum() - (X.sum() * Y.sum()) / (n * r)
                
                rep_sum_X = self.df.groupby(self.rep_col)[t1].sum()
                rep_sum_Y = self.df.groupby(self.rep_col)[t2].sum()
                CP_Rep = (rep_sum_X * rep_sum_Y).sum() / n - (X.sum() * Y.sum()) / (n * r)
                
                geno_sum_X = self.df.groupby(self.genotype_col)[t1].sum()
                geno_sum_Y = self.df.groupby(self.genotype_col)[t2].sum()
                CP_Geno = (geno_sum_X * geno_sum_Y).sum() / r - (X.sum() * Y.sum()) / (n * r)
                
                CP_Error = CP_Total - CP_Rep - CP_Geno
                
                MCP_G = CP_Geno / (n - 1)
                MCP_E = CP_Error / ((n - 1) * (r - 1))
                
                cov_g = (MCP_G - MCP_E) / r
                
                rg = cov_g / np.sqrt(trait_stats[t1]["var_g"] * trait_stats[t2]["var_g"])
                rg = np.clip(rg, -1.0, 1.0)
                
                rg_matrix[i, j] = rg
                rg_matrix[j, i] = rg
                
        self.correlation_matrix = pd.DataFrame(rg_matrix, index=traits, columns=traits)

    def run_analysis(self):
        self._compute_genotypic_correlation()
        
        # Path Analysis Logic
        indep = self.independent_vars
        dep = self.dependent_var
        
        # 1. R matrix (correlations among independents)
        R = self.correlation_matrix.loc[indep, indep].values
        
        # 2. r_y vector (correlations of independents with dependent)
        r_y = self.correlation_matrix.loc[indep, dep].values
        
        # 3. Direct Effects (Path Coefficients)
        # Solve R * P = r_y  => P = R^-1 * r_y
        try:
            # Use pseudo-inverse for stability if R is near singular
            inv_R = np.linalg.pinv(R)
            P_direct = np.dot(inv_R, r_y)
        except Exception as e:
            raise ValueError(f"Matrix inversion failed: {str(e)}. Check for multi-collinearity among traits.")
            
        self.direct_effects = P_direct
        
        # 4. Total and Specific Indirect Effects
        # Indirect(Xi via Xj) = r_ij * P_jy
        num_indep = len(indep)
        indirect_matrix = np.zeros((num_indep, num_indep))
        
        for i in range(num_indep):
            for j in range(num_indep):
                if i == j:
                    indirect_matrix[i, j] = P_direct[i] # Store direct on diagonal for reporting convenience
                else:
                    indirect_matrix[i, j] = R[i, j] * P_direct[j]
                    
        self.indirect_effects_matrix = pd.DataFrame(indirect_matrix, index=indep, columns=indep)
        
        # 5. Residual Effect
        # Residual = sqrt(1 - Σ(r_iy * P_iy))
        sum_explained = np.sum(r_y * P_direct)
        self.residual_effect = np.sqrt(max(0, 1 - sum_explained))
        self.explained_variation = (1 - self.residual_effect**2) * 100
        self.unexplained_variation = (self.residual_effect**2) * 100

    def get_path_table(self):
        """Constructs the standard path analysis output table."""
        indep = self.independent_vars
        dep = self.dependent_var
        data = []
        
        for i, trait_i in enumerate(indep):
            row = {"Trait": trait_i}
            row["Direct Effect"] = self.direct_effects[i]
            
            total_indirect = 0
            for j, trait_j in enumerate(indep):
                if i == j:
                    row[f"via {trait_j}"] = f"({self.direct_effects[i]:.4f})" # Direct in parentheses
                else:
                    val = self.indirect_effects_matrix.loc[trait_i, trait_j]
                    row[f"via {trait_j}"] = val
                    total_indirect += val
                    
            row["Total Indirect"] = total_indirect
            row[f"r_g with {dep}"] = self.correlation_matrix.loc[trait_i, dep]
            data.append(row)
            
        return pd.DataFrame(data)

    def create_report(self):
        doc = Document()
        title = doc.add_heading("Genotypic Path Analysis Report (RCBD)", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading("1. Analysis Parameters", level=1)
        doc.add_paragraph(f"Dependent Variable (Effect): {self.dependent_var}")
        doc.add_paragraph(f"Independent Variables (Causes): {', '.join(self.independent_vars)}")
        doc.add_paragraph(f"Sample Size: {self.n} Genotypes, {self.r} Replications")
        
        doc.add_heading("2. Genotypic Path Coefficients", level=1)
        path_df = self.get_path_table()
        
        # Add table to doc
        table = doc.add_table(rows=1, cols=len(path_df.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(path_df.columns):
            table.cell(0, i).text = col
            
        for _, row in path_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, (float, np.float64)):
                    row_cells[i].text = f"{val:.4f}"
                else:
                    row_cells[i].text = str(val)

        doc.add_heading("3. Residual Analysis", level=1)
        doc.add_paragraph(f"Residual Effect (R): {self.residual_effect:.4f}")
        doc.add_paragraph(f"Explained Variation: {self.explained_variation:.2f}%")
        doc.add_paragraph(f"Unexplained Variation: {self.unexplained_variation:.2f}%")
        
        doc.add_heading("4. Path Influence Diagram", level=1)
        buf = self.generate_path_diagram()
        doc.add_picture(buf, width=Inches(6))
        
        # Interpretation
        doc.add_heading("5. Statistical Summary", level=1)
        top_trait_idx = np.argmax(np.abs(self.direct_effects))
        top_trait = self.independent_vars[top_trait_idx]
        direct_val = self.direct_effects[top_trait_idx]
        
        summary = f"Among the traits studied, '{top_trait}' exhibited the strongest direct effect ({direct_val:.4f}) on '{self.dependent_var}'. "
        if self.residual_effect < 0.3:
            summary += "The low residual effect suggests that the chosen independent variables account for most of the variation in the dependent variable."
        else:
            summary += "The residual effect is relatively high, indicating other traits not included in the model might significantly influence the dependent variable."
            
        doc.add_paragraph(summary)
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Sheet 1: Path Table
            self.get_path_table().to_excel(writer, sheet_name='Path Coefficients', index=False)
            
            # Sheet 2: Genotypic Correlations
            self.correlation_matrix.to_excel(writer, sheet_name='Genotypic Correlations')
            
            # Sheet 3: Variances & Residual
            res_df = pd.DataFrame({
                "Parameter": ["Residual Effect", "Explained %", "Unexplained %"],
                "Value": [self.residual_effect, self.explained_variation, self.unexplained_variation]
            })
            res_df.to_excel(writer, sheet_name='Model Summary', index=False)
            
        output.seek(0)
        return output

    def generate_path_diagram(self):
        """Generates a fan-style Path Diagram matching the user's specified layout."""
        num_traits = len(self.independent_vars)
        plt.figure(figsize=(14, 10))
        ax = plt.gca()
        ax.set_xlim(-2, 10)
        ax.set_ylim(-1, num_traits)
        plt.axis('off')

        # 1. Coordinate setup
        # Dependent variable (Effect) on the left
        dep_x, dep_y = 0, (num_traits - 1) / 2
        
        # Independent variables in a vertical line
        trait_x = 5
        trait_y_coords = np.linspace(num_traits - 1, 0, num_traits)
        
        # 2. Draw residual effect
        res_x, res_y = -1.5, dep_y - 1.5
        plt.plot([res_x, res_x, dep_x], [res_y - 0.5, res_y, dep_y], color='blue', lw=1.5)
        plt.text(res_x - 0.2, res_y, f"Residual effect\nSQRT(1-{1-self.residual_effect**2:.4f})", 
                 rotation=90, va='center', ha='right', fontsize=9)
        plt.scatter([res_x], [res_y - 0.5], marker='s', s=100, color='white', edgecolors='blue', zorder=5)

        # 3. Draw Trait Nodes and Direct Paths
        for i, trait in enumerate(self.independent_vars):
            ty = trait_y_coords[i]
            direct_eff = self.direct_effects[i]
            
            # Draw line to dependent var
            plt.plot([trait_x, dep_x], [ty, dep_y], color='blue', alpha=0.6, lw=1)
            
            # Text for Direct Effect and Trait Name
            # Position text along the line
            mid_x = (trait_x + dep_x) / 2 - 1.5
            mid_y = (ty + dep_y) / 2
            angle = np.degrees(np.arctan2(ty - dep_y, trait_x - dep_x))
            
            label = f"{direct_eff:.3f}  {trait}"
            plt.text(mid_x, mid_y, label, rotation=angle, va='bottom', ha='center', fontsize=9)
            
            # Nodes
            plt.scatter([trait_x], [ty], s=150, color='white', edgecolors='blue', zorder=5)
            
        # Dependent variable node
        plt.scatter([dep_x], [dep_y], s=200, color='white', edgecolors='blue', zorder=6)
        plt.text(dep_x - 0.5, dep_y, f" {self.dependent_var}", ha='right', va='center', fontweight='bold')

        # 4. Draw Correlation Arcs (Inter-correlations)
        # We only draw arcs for independent traits on the right side
        for i in range(num_traits):
            for j in range(i + 1, num_traits):
                y1, y2 = trait_y_coords[i], trait_y_coords[j]
                rg = self.correlation_matrix.loc[self.independent_vars[i], self.independent_vars[j]]
                
                # Draw arc on the right
                # Using a bezier-like curve path
                center_y = (y1 + y2) / 2
                dist = abs(y1 - y2)
                arc_x = trait_x + 0.5 + (dist * 0.4) # Arc outward based on distance
                
                # Plot arc
                # For simplicity in matplotlib without complex Path, we can use a small segment of sine/parabola
                curve_pts_x = [trait_x, arc_x, trait_x]
                curve_pts_y = [y1, center_y, y2]
                
                # Polynomial interpolation (parabola)
                coeffs = np.polyfit(curve_pts_y, curve_pts_x, 2)
                y_fine = np.linspace(y1, y2, 50)
                x_fine = np.polyval(coeffs, y_fine)
                
                plt.plot(x_fine, y_fine, color='blue', alpha=0.3, lw=0.8)
                
                # Annotation for correlation
                plt.text(arc_x + 0.1, center_y, f"{rg:.3f}", va='center', ha='left', fontsize=8)

        plt.title(f"Genotypical Path Diagram for {self.dependent_var}", loc='left', fontsize=14, fontweight='bold')
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
