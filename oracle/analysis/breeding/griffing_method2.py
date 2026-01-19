import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import seaborn as sns

class GriffingMethod2Analyzer:
    def __init__(self, df, female_col, male_col, rep_col, trait_cols):
        self.df = df
        self.female_col = female_col
        self.male_col = male_col
        self.rep_col = rep_col
        self.trait_cols = trait_cols
        
        self.p = 0 # Number of parents
        self.r = 0 # Number of replications
        self.parents = []
        
        # Results storage per trait
        self.results = {}

    def validate(self):
        # 1. Cleaning
        self.df[self.female_col] = self.df[self.female_col].astype(str).str.strip()
        self.df[self.male_col] = self.df[self.male_col].astype(str).str.strip()
        
        for trait in self.trait_cols:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        
        if self.df[self.trait_cols].isnull().any().any():
            raise ValueError("Traits contain missing or non-numeric values.")

        # 2. Identify parents
        all_p = sorted(list(set(self.df[self.female_col].unique()) | set(self.df[self.male_col].unique())))
        self.parents = all_p
        self.p = len(all_p)
        
        # 3. Handle structure (Half Diallel: i <= j)
        # We ensure that (P1, P2) is same as (P2, P1) for half diallel logic
        # But Method II strictly requires parents (i=j) and F1s (i<j). 
        # No reciprocals (j>i) should be provided, or if provided, they must be treated as duplicates?
        # Usually user provides one way. Let's force-sort the parent pair.
        self.df['__p1'] = self.df[[self.female_col, self.male_col]].min(axis=1)
        self.df['__p2'] = self.df[[self.female_col, self.male_col]].max(axis=1)
        
        # Check for missing crosses
        entry_counts = self.df.groupby(['__p1', '__p2']).size()
        required = []
        for i in range(self.p):
            for j in range(i, self.p):
                required.append((self.parents[i], self.parents[j]))
        
        missing = [f"{e[0]} x {e[1]}" for e in required if e not in entry_counts.index]
        if missing:
            raise ValueError(f"Missing half-diallel entries (including parents): {', '.join(missing[:5])}...")
            
        # Check reciprocals (if user provided strictly according to Female/Male columns)
        # In a curated half-diallel, j < i shouldn't exist as separate entries if they are meant to be reciprocals.
        # But we've normalized them to __p1 and __p2. So duplicate (P1, P2) in original would mean more reps.
        
        reps = entry_counts.unique()
        if len(reps) > 1:
            raise ValueError(f"Unbalanced replications detected: {reps}. Ensure all crosses have same number of observations.")
        self.r = int(reps[0])

    def analyze_trait(self, trait):
        p = self.p
        r = self.r
        parents = self.parents
        
        # Means for each entry
        means_df = self.df.groupby(['__p1', '__p2'])[trait].mean().unstack()
        # Fill symmetric part for calculation convenience (Method II logic uses Y_i. which includes all crosses parent i is in)
        # We create a full symmetric matrix for calculation
        full_means = np.zeros((p, p))
        for i in range(p):
            for j in range(i, p):
                val = means_df.loc[parents[i], parents[j]]
                full_means[i, j] = val
                full_means[j, i] = val
        
        # Grand Mean and Row Totals of Means
        Y_dot_dot_means = 0
        for i in range(p):
            for j in range(i, p):
                Y_dot_dot_means += full_means[i, j]
        # Y_.. in formulas is usually the sum of entry means for Diallel ANOVA
        # Prompt says: Grand Mean Y_..
        grand_mean_diallel = Y_dot_dot_means / (p * (p + 1) / 2)
        
        # Prompt estimation logic:
        # g_i = (1 / (p + 2)) * [ (p + 1)Ȳ_i. - 2Ȳ_.. ]
        # where Ȳ_i. is mean of all crosses involving parent i
        Y_i_bar = []
        for i in range(p):
            # Sum of entry means involving parent i
            # Method II: parent i is in P combinations (one self, P-1 hybrids)
            s = full_means[i, :].sum()
            Y_i_bar.append(s / p)
            
        # Grand mean of all entries Ȳ_..
        sum_all_means = 0
        for i in range(p):
            for j in range(i, p):
                sum_all_means += full_means[i, j]
        total_grand_mean = sum_all_means / (p * (p + 1) / 2)

        # --- ANOVA (Genotypes, Replications, Error) ---
        y_all = self.df[trait].values
        G_total_sum = y_all.sum()
        N = len(y_all)
        CF = (G_total_sum**2) / N
        SS_Total = (y_all**2).sum() - CF
        
        rep_sums = self.df.groupby(self.rep_col)[trait].sum()
        SS_Rep = (rep_sums**2).sum() / (p * (p + 1) / 2) - CF
        
        entry_sums = self.df.groupby(['__p1', '__p2'])[trait].sum()
        SS_Geno = (entry_sums**2).sum() / r - CF
        
        SS_Error = SS_Total - SS_Rep - SS_Geno
        
        df_rep = r - 1
        df_geno = (p * (p + 1) / 2) - 1
        df_error = df_rep * df_geno
        
        MS_Geno = SS_Geno / df_geno
        MS_Error = SS_Error / df_error
        
        F_Geno = MS_Geno / MS_Error if MS_Error > 0 else 0
        P_Geno = 1 - stats.f.cdf(F_Geno, df_geno, df_error)

        # --- Combining Ability ANOVA ---
        # GCA SS = [1/(p+2)] * [ Σ(X_i + Y_ii)^2 - (4/p) X^2 ]  <-- Standard Griffing II
        # Let's use the sums X_i = sum of entries containing parent i
        X_i = full_means.sum(axis=1) # Note: this is sum of means
        X = sum_all_means # Sum of all means
        
        # Griffing Method II Sum of Squares (using entry means)
        # SS GCA = r/(p+2) * [ Σ(X_i + Y_ii)^2 - (4/p) X^2 ]
        # where Y_ii is the mean of parent i
        term1_gca = 0
        for i in range(p):
            term1_gca += (X_i[i] + full_means[i, i])**2
        
        SS_GCA = (r / (p + 2)) * (term1_gca - (4 / p) * (X**2))
        
        SS_SCA = SS_Geno - SS_GCA
        
        df_gca = p - 1
        df_sca = p * (p - 1) // 2
        
        MS_GCA = SS_GCA / df_gca
        MS_SCA = SS_SCA / df_sca
        
        F_GCA = MS_GCA / MS_Error
        F_SCA = MS_SCA / MS_Error
        P_GCA = 1 - stats.f.cdf(F_GCA, df_gca, df_error)
        P_SCA = 1 - stats.f.cdf(F_SCA, df_sca, df_error)

        # --- Effects ---
        # g_i = (1 / (p + 2)) * [ (X_i + Y_ii) - (2/p) X ]
        gca_effects = []
        for i in range(p):
            gi = (1 / (p + 2)) * ( (X_i[i] + full_means[i, i]) - (2 / p) * X )
            gca_effects.append(gi)
            
        sca_effects = np.zeros((p, p))
        for i in range(p):
            for j in range(i, p):
                # s_ij = Y_ij - [1/(p+2)] (X_i + Y_ii + X_j + Y_jj) + [2 / ((p+1)(p+2))] X
                s_ij = full_means[i, j] - (1 / (p + 2)) * (X_i[i] + full_means[i, i] + X_i[j] + full_means[j, j]) + (2 / ((p + 1) * (p + 2))) * X
                sca_effects[i, j] = s_ij
        
        # Standard Errors
        se_gi = np.sqrt(((p - 1) * MS_Error) / (p * (p + 2) * r))
        se_sij_cross = np.sqrt(((p**2 + p + 2) * MS_Error) / ((p + 1) * (p + 2) * r)) # Specific for i != j? 
        # Actually Griffing II SE(s_ij):
        # i=j: se = sqrt( p(p-1) MS_E / ( (p+1)(p+2)r ) )? 
        # Let's use simpler ones if not specified.
        se_sij = np.sqrt(((p**2 - 1) * MS_Error) / (2 * p * (p + 2) * r)) # General proxy or specific for i!=j
        
        # --- Genetic Parameters ---
        # σ²_GCA = (MS_GCA − MS_Error) / (p + 2)
        var_gca = max(0, (MS_GCA - MS_Error) / (p + 2))
        var_sca = max(0, (MS_SCA - MS_Error))
        var_a = 2 * var_gca
        var_d = var_sca
        var_p = var_a + var_d + MS_Error
        
        h2_broad = (var_a + var_d) / var_p if var_p > 0 else 0
        h2_narrow = var_a / var_p if var_p > 0 else 0
        predictability = (2 * var_gca) / (2 * var_gca + var_sca) if (2 * var_gca + var_sca) > 0 else 0

        # --- Heterosis ---
        mph = {}
        hb = {}
        for i in range(p):
            for j in range(i + 1, p):
                Pi = full_means[i, i]
                Pj = full_means[j, j]
                F1 = full_means[i, j]
                MP = (Pi + Pj) / 2
                BP = max(Pi, Pj)
                
                mph_val = ((F1 - MP) / MP) * 100 if MP != 0 else 0
                hb_val = ((F1 - BP) / BP) * 100 if BP != 0 else 0
                
                # t-tests
                sed_mph = np.sqrt(1.5 * MS_Error / r)
                sed_hb = np.sqrt(2 * MS_Error / r)
                
                t_mph = (F1 - MP) / sed_mph if sed_mph > 0 else 0
                t_hb = (F1 - BP) / sed_hb if sed_hb > 0 else 0
                
                mph[f"{parents[i]} x {parents[j]}"] = {"val": mph_val, "t": t_mph, "p": 2 * (1 - stats.t.cdf(abs(t_mph), df_error))}
                hb[f"{parents[i]} x {parents[j]}"] = {"val": hb_val, "t": t_hb, "p": 2 * (1 - stats.t.cdf(abs(t_hb), df_error))}

        res = {
            "anova_geno": {
                "Replication": {"df": df_rep, "SS": SS_Rep, "MS": SS_Rep / df_rep, "F": (SS_Rep / df_rep) / MS_Error, "P": 1 - stats.f.cdf((SS_Rep / df_rep) / MS_Error, df_rep, df_error)},
                "Genotypes": {"df": df_geno, "SS": SS_Geno, "MS": MS_Geno, "F": F_Geno, "P": P_Geno},
                "Error": {"df": df_error, "SS": SS_Error, "MS": MS_Error, "F": None, "P": None},
                "Total": {"df": N - 1, "SS": SS_Total, "MS": None, "F": None, "P": None}
            },
            "anova_comb": {
                "GCA": {"df": df_gca, "SS": SS_GCA, "MS": MS_GCA, "F": F_GCA, "P": P_GCA},
                "SCA": {"df": df_sca, "SS": SS_SCA, "MS": MS_SCA, "F": F_SCA, "P": P_SCA},
                "Error": {"df": df_error, "SS": SS_Error, "MS": MS_Error, "F": None, "P": None}
            },
            "gca_effects": [{"parent": parents[i], "effect": gca_effects[i], "t": gca_effects[i] / se_gi, "p": 2 * (1 - stats.t.cdf(abs(gca_effects[i] / se_gi), df_error)), "se": se_gi} for i in range(p)],
            "sca_matrix": sca_effects.tolist(),
            "variances": {
                "h2_broad": h2_broad, "h2_narrow": h2_narrow, "predictability": predictability,
                "sigma2_gca": var_gca, "sigma2_sca": var_sca,
                "sigma2_a": var_a, "sigma2_d": var_d
            },
            "heterosis": {"mph": mph, "hb": hb},
            "parents": parents
        }
        self.results[trait] = res
        return res

    def get_sig(self, p):
        if p is None: return ""
        if p <= 0.01: return "**"
        if p <= 0.05: return "*"
        return "ns"

    def create_report(self):
        doc = Document()
        doc.add_heading("Griffing's Method II Diallel Analysis Report", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading("1. Analysis Overview", level=1)
        doc.add_paragraph("Method: Griffing (1956) Method II (Half Diallel: Parents + F1s, no reciprocals)")
        doc.add_paragraph("Model: Fixed Effects Model")
        doc.add_paragraph(f"Number of Parents: {self.p} | Replications: {self.r}")

        for trait in self.trait_cols:
            res = self.results[trait]
            doc.add_page_break()
            doc.add_heading(f"Trait: {trait}", level=1)
            
            # ANOVA
            doc.add_heading("A. ANOVA Table", level=2)
            table = doc.add_table(rows=1, cols=7); table.style = 'Table Grid'
            for i, text in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'p-value', 'Sig']):
                table.rows[0].cells[i].text = text
            for src, d in res['anova_geno'].items():
                r = table.add_row().cells
                r[0].text = src; r[1].text = str(d['df']); r[2].text = f"{d['SS']:.4f}"
                r[3].text = f"{d['MS']:.4f}" if d['MS'] else ""; r[4].text = f"{d['F']:.4f}" if d['F'] else ""
                r[5].text = f"{d['P']:.4f}" if d['P'] is not None else ""; r[6].text = self.get_sig(d['P'])

            doc.add_heading("B. Combining Ability ANOVA", level=2)
            table2 = doc.add_table(rows=1, cols=7); table2.style = 'Table Grid'
            for i, text in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'p-value', 'Sig']):
                table2.rows[0].cells[i].text = text
            for src in ["GCA", "SCA", "Error"]:
                d = res['anova_comb'][src]
                r = table2.add_row().cells
                r[0].text = src; r[1].text = str(d['df']); r[2].text = f"{d['SS']:.4f}"
                r[3].text = f"{d['MS']:.4f}" if d['MS'] else ""; r[4].text = f"{d['F']:.4f}" if d['F'] else ""
                r[5].text = f"{d['P']:.4f}" if d['P'] is not None else ""; r[6].text = self.get_sig(d['P'])

            doc.add_heading("C. Genetic Parameters", level=2)
            v = res['variances']
            table_v = doc.add_table(rows=6, cols=2); table_v.style = 'Table Grid'
            v_items = [("V_A", v['sigma2_a']), ("V_D", v['sigma2_d']), ("H² (Broad)", v['h2_broad']), ("h² (Narrow)", v['h2_narrow']), ("Predictability Ratio", v['predictability'])]
            for i, (lab, val) in enumerate(v_items):
                table_v.cell(i, 0).text = lab; table_v.cell(i, 1).text = f"{val:.4f}"

            doc.add_heading("D. Interpretation", level=2)
            p_gca = res['anova_comb']['GCA']['P']
            p_sca = res['anova_comb']['SCA']['P']
            interp = f"The GCA effects were {self.get_sig(p_gca).replace('ns', 'non-significant')} and SCA effects were {self.get_sig(p_sca).replace('ns', 'non-significant')}. "
            if v['predictability'] > 0.5: interp += "The high predictability ratio indicates that additive gene action is predominant."
            else: interp += "The low predictability ratio indicates that non-additive gene action is more influential."
            doc.add_paragraph(interp)

        f = io.BytesIO()
        doc.save(f); f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for trait in self.trait_cols:
                res = self.results[trait]
                pd.DataFrame.from_dict(res['anova_comb'], orient='index').to_excel(writer, sheet_name=f"{trait[:10]}_ANOVA")
                pd.DataFrame(res['gca_effects']).to_excel(writer, sheet_name=f"{trait[:10]}_GCA")
                pd.DataFrame(res['sca_matrix'], index=self.parents, columns=self.parents).to_excel(writer, sheet_name=f"{trait[:10]}_SCA")
                pd.DataFrame.from_dict(res['variances'], orient='index').to_excel(writer, sheet_name=f"{trait[:10]}_Genetic")
        output.seek(0)
        return output
