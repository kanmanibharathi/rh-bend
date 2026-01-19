import pandas as pd
import numpy as np
from scipy import stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime

class LineTesterAnalyzer:
    def __init__(self, df, line_col, tester_col, rep_col, trait_col, alpha=0.05):
        self.df = df
        self.line_col = line_col
        self.tester_col = tester_col
        self.rep_col = rep_col
        self.trait_col = trait_col
        self.alpha = float(alpha)
        
        # Results storage
        self.anova_table = []
        self.gca_lines = []
        self.gca_testers = []
        self.sca_effects = []
        self.genetic_variances = {}
        self.summary_stats = {}
        
    def validate(self):
        # Numeric check for trait
        self.df[self.trait_col] = pd.to_numeric(self.df[self.trait_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.line_col, self.tester_col, self.rep_col, self.trait_col])
        
        l_count = self.df[self.line_col].nunique()
        t_count = self.df[self.tester_col].nunique()
        r_count = self.df[self.rep_col].nunique()
        
        if l_count < 2: raise ValueError("At least 2 lines required.")
        if t_count < 2: raise ValueError("At least 2 testers required.")
        if r_count < 2: raise ValueError("At least 2 replications required.")
        
        # Check for balanced data
        expected_size = l_count * t_count * r_count
        if len(self.df) != expected_size:
            # Check for duplicates or missing combinations
            counts = self.df.groupby([self.line_col, self.tester_col, self.rep_col]).size()
            if any(counts > 1):
                raise ValueError("Duplicate entries found for some Line x Tester x Rep combinations.")
            else:
                raise ValueError(f"Incomplete data. Expected {expected_size} observations, but found {len(self.df)}.")

    def run_analysis(self):
        l = self.df[self.line_col].nunique()
        t = self.df[self.tester_col].nunique()
        r = self.df[self.rep_col].nunique()
        N = len(self.df)
        
        grand_total = self.df[self.trait_col].sum()
        grand_mean = grand_total / N
        CF = (grand_total ** 2) / N
        
        # Sum of Squares
        SS_total = (self.df[self.trait_col] ** 2).sum() - CF
        
        # Rep SS
        rep_sums = self.df.groupby(self.rep_col)[self.trait_col].sum()
        SS_rep = (rep_sums ** 2).sum() / (l * t) - CF
        
        # Hybrid (L x T) SS
        hybrid_sums = self.df.groupby([self.line_col, self.tester_col])[self.trait_col].sum()
        SS_hybrids = (hybrid_sums ** 2).sum() / r - CF
        
        # Individual effects (Lines, Testers, LxT)
        line_sums = self.df.groupby(self.line_col)[self.trait_col].sum()
        SS_lines = (line_sums ** 2).sum() / (t * r) - CF
        
        tester_sums = self.df.groupby(self.tester_col)[self.trait_col].sum()
        SS_testers = (tester_sums ** 2).sum() / (l * r) - CF
        
        SS_lxt = SS_hybrids - SS_lines - SS_testers
        
        # Error SS
        SS_error = SS_total - SS_rep - SS_hybrids
        
        # DFs
        df_rep = r - 1
        df_lines = l - 1
        df_testers = t - 1
        df_lxt = (l - 1) * (t - 1)
        df_error = (l * t - 1) * (r - 1)
        df_total = N - 1
        
        # MS
        MS_rep = SS_rep / df_rep
        MS_lines = SS_lines / df_lines
        MS_testers = SS_testers / df_testers
        MS_lxt = SS_lxt / df_lxt
        MS_error = SS_error / df_error
        
        # F-stats (using Error MS)
        F_lines = MS_lines / MS_error
        F_testers = MS_testers / MS_error
        F_lxt = MS_lxt / MS_error
        F_rep = MS_rep / MS_error
        
        p_lines = 1 - stats.f.cdf(F_lines, df_lines, df_error)
        p_testers = 1 - stats.f.cdf(F_testers, df_testers, df_error)
        p_lxt = 1 - stats.f.cdf(F_lxt, df_lxt, df_error)
        p_rep = 1 - stats.f.cdf(F_rep, df_rep, df_error)
        
        # Build ANOVA
        self.anova_table = [
            {"Source": "Replications", "DF": df_rep, "SS": SS_rep, "MS": MS_rep, "F": F_rep, "p": p_rep},
            {"Source": "Lines", "DF": df_lines, "SS": SS_lines, "MS": MS_lines, "F": F_lines, "p": p_lines},
            {"Source": "Testers", "DF": df_testers, "SS": SS_testers, "MS": MS_testers, "F": F_testers, "p": p_testers},
            {"Source": "Lines x Testers", "DF": df_lxt, "SS": SS_lxt, "MS": MS_lxt, "F": F_lxt, "p": p_lxt},
            {"Source": "Error", "DF": df_error, "SS": SS_error, "MS": MS_error, "F": None, "p": None},
            {"Source": "Total", "DF": df_total, "SS": SS_total, "MS": None, "F": None, "p": None}
        ]
        
        # Variance Components
        sigma2_e = MS_error
        sigma2_sca = max(0, (MS_lxt - MS_error) / r)
        sigma2_gca_lines = max(0, (MS_lines - MS_lxt) / (r * t))
        sigma2_gca_testers = max(0, (MS_testers - MS_lxt) / (r * l))
        
        sigma2_gca = (sigma2_gca_lines + sigma2_gca_testers) / 2
        sigma2_A = 2 * sigma2_gca
        sigma2_D = sigma2_sca
        
        dom_ratio = 0
        if sigma2_A > 0:
            dom_ratio = np.sqrt(sigma2_D / sigma2_A)
            
        self.genetic_variances = {
            "sigma2_gca_lines": sigma2_gca_lines,
            "sigma2_gca_testers": sigma2_gca_testers,
            "sigma2_sca": sigma2_sca,
            "sigma2_A": sigma2_A,
            "sigma2_D": sigma2_D,
            "Degree_of_Dominance": dom_ratio,
            "GeneAction": "Additive" if sigma2_gca > sigma2_sca else "Non-Additive"
        }
        
        # Effects calculation
        # SE
        se_line = np.sqrt(sigma2_e / (r * t))
        se_tester = np.sqrt(sigma2_e / (r * l))
        se_sca = np.sqrt(sigma2_e / r)
        t_crit = stats.t.ppf(1 - self.alpha/2, df_error)
        
        # GCA Lines
        for line, val in line_sums.items():
            eff = (val / (t * r)) - grand_mean
            t_val = eff / se_line
            self.gca_lines.append({
                "Line": str(line),
                "Effect": eff,
                "SE": se_line,
                "t_value": t_val,
                "Sig": "*" if abs(t_val) > t_crit else "ns"
            })
            
        # GCA Testers
        for tester, val in tester_sums.items():
            eff = (val / (l * r)) - grand_mean
            t_val = eff / se_tester
            self.gca_testers.append({
                "Tester": str(tester),
                "Effect": eff,
                "SE": se_tester,
                "t_value": t_val,
                "Sig": "*" if abs(t_val) > t_crit else "ns"
            })
            
        # SCA
        hybrid_means = self.df.groupby([self.line_col, self.tester_col])[self.trait_col].mean()
        line_means = self.df.groupby(self.line_col)[self.trait_col].mean()
        tester_means = self.df.groupby(self.tester_col)[self.trait_col].mean()
        
        for (line, tester), h_mean in hybrid_means.items():
            eff = h_mean - line_means[line] - tester_means[tester] + grand_mean
            t_val = eff / se_sca
            self.sca_effects.append({
                "Hybrid": f"{line} x {tester}",
                "Effect": eff,
                "SE": se_sca,
                "t_value": t_val,
                "Sig": "*" if abs(t_val) > t_crit else "ns"
            })
            
        self.summary_stats = {
            "Mean": grand_mean,
            "CV": (np.sqrt(sigma2_e) / grand_mean) * 100 if grand_mean != 0 else 0
        }

    def get_interpretation(self):
        v = self.genetic_variances
        gene_action = "non-additive (dominance/epistasis)" if v['sigma2_sca'] > ((v['sigma2_gca_lines'] + v['sigma2_gca_testers'])/2) else "additive"
        
        text = f"The analysis of variance revealed genetic components of variation. The SCA variance ({v['sigma2_sca']:.4f}) "
        text += f"as compared to GCA variance indicates that {gene_action} gene action is predominant for this trait. "
        
        if v['Degree_of_Dominance'] > 1.1:
            text += "The degree of dominance suggests over-dominance. "
        elif v['Degree_of_Dominance'] > 0.9:
            text += "The degree of dominance suggests complete dominance. "
        else:
            text += "The degree of dominance suggests partial dominance. "
            
        return text

    def create_report(self):
        doc = Document()
        doc.add_heading("Line x Tester Analysis Report", 0)
        
        doc.add_heading("1. Analysis of Variance (ANOVA)", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text = "Source", "DF", "SS", "MS", "F", "P-value"
        for row in self.anova_table:
            r = table.add_row().cells
            r[0].text = str(row['Source'])
            r[1].text = str(row['DF'])
            r[2].text = f"{row['SS']:.4f}"
            r[3].text = f"{row['MS']:.4f}" if row['MS'] else "-"
            r[4].text = f"{row['F']:.4f}" if row['F'] else "-"
            r[5].text = f"{row['p']:.4g}" if row['p'] else "-"

        doc.add_heading("2. Genetic Variance Components", level=1)
        v_table = doc.add_table(rows=1, cols=2)
        v_table.style = 'Table Grid'
        v = self.genetic_variances
        data = [
            ("σ² GCA (Lines)", f"{v['sigma2_gca_lines']:.4f}"),
            ("σ² GCA (Testers)", f"{v['sigma2_gca_testers']:.4f}"),
            ("σ² SCA", f"{v['sigma2_sca']:.4f}"),
            ("Additive Variance (σ²A)", f"{v['sigma2_A']:.4f}"),
            ("Dominance Variance (σ²D)", f"{v['sigma2_D']:.4f}"),
            ("Degree of Dominance", f"{v['Degree_of_Dominance']:.4f}"),
            ("Predominant Gene Action", v['GeneAction'])
        ]
        for lab, val in data:
            row = v_table.add_row().cells
            row[0].text, row[1].text = lab, val

        doc.add_heading("3. GCA Effects (Lines)", level=1)
        l_table = doc.add_table(rows=1, cols=5)
        l_table.style = 'Table Grid'
        l_table.rows[0].cells[0].text = "Line"
        l_table.rows[0].cells[1].text = "Effect"
        l_table.rows[0].cells[2].text = "SE"
        l_table.rows[0].cells[3].text = "t-value"
        l_table.rows[0].cells[4].text = "Sig"
        for line in self.gca_lines:
            r = l_table.add_row().cells
            r[0].text, r[1].text, r[2].text, r[3].text, r[4].text = str(line['Line']), f"{line['Effect']:.4f}", f"{line['SE']:.4f}", f"{line['t_value']:.4f}", line['Sig']

        doc.add_heading("4. Interpretation", level=1)
        doc.add_paragraph(self.get_interpretation())
        
        doc.add_heading("5. Scientific References", level=1)
        refs = ["Kempthorne, O. (1957). An introduction to genetical statistics. John Wiley & Sons.", 
                "Singh, R. K., & Chaudhary, B. D. (1979). Biometrical methods in quantitative genetic analysis."]
        for ref in refs: doc.add_paragraph(ref)

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
        
    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            pd.DataFrame(self.anova_table).to_excel(writer, sheet_name='ANOVA', index=False)
            pd.DataFrame(self.gca_lines).to_excel(writer, sheet_name='GCA Lines', index=False)
            pd.DataFrame(self.gca_testers).to_excel(writer, sheet_name='GCA Testers', index=False)
            pd.DataFrame(self.sca_effects).to_excel(writer, sheet_name='SCA Effects', index=False)
            pd.DataFrame([self.genetic_variances]).to_excel(writer, sheet_name='Genetic Variances', index=False)
        output.seek(0)
        return output
