import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import seaborn as sns

class GriffingMethod2CheckAnalyzer:
    def __init__(self, df, female_col, male_col, rep_col, check_col, trait_cols):
        self.df = df
        self.female_col = female_col
        self.male_col = male_col
        self.rep_col = rep_col
        self.check_col = check_col # Column for 'Diallel' vs check names
        self.trait_cols = trait_cols
        
        self.p = 0 # Diallel parents
        self.r = 0 # Replications
        self.parents = []
        self.checks = []
        self.results = {}

    def validate(self):
        self.df[self.female_col] = self.df[self.female_col].astype(str).str.strip()
        self.df[self.male_col] = self.df[self.male_col].astype(str).str.strip()
        self.df[self.check_col] = self.df[self.check_col].astype(str).str.strip()
        
        for trait in self.trait_cols:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        if self.df[self.trait_cols].isnull().any().any():
            raise ValueError("Traits contain missing/non-numeric values.")

        # Split Diallel and Checks
        self.diallel_df = self.df[self.df[self.check_col].str.lower() == 'diallel'].copy()
        self.checks_df = self.df[self.df[self.check_col].str.lower() != 'diallel'].copy()
        
        if self.diallel_df.empty:
            raise ValueError(f"No entries marked 'Diallel' in {self.check_col} column.")

        # Identify parents (Half Diallel logic: i <= j)
        p1 = self.diallel_df[self.female_col].unique()
        p2 = self.diallel_df[self.male_col].unique()
        self.parents = sorted(list(set(p1) | set(p2)))
        self.p = len(self.parents)
        
        self.checks = sorted(self.checks_df[self.check_col].unique().tolist())
        
        # Verify half diallel (p*(p+1)/2)
        # Normalize to i <= j
        self.diallel_df['__p1'] = self.diallel_df[[self.female_col, self.male_col]].min(axis=1)
        self.diallel_df['__p2'] = self.diallel_df[[self.female_col, self.male_col]].max(axis=1)
        
        counts = self.diallel_df.groupby(['__p1', '__p2']).size()
        required = []
        for i in range(self.p):
            for j in range(i, self.p):
                required.append((self.parents[i], self.parents[j]))
        
        missing = [f"{e[0]}x{e[1]}" for e in required if e not in counts.index]
        if missing:
            raise ValueError(f"Missing half-diallel entries: {', '.join(missing[:5])}...")
            
        # Replications
        reps = self.df.groupby([self.female_col, self.male_col, self.check_col]).size().unique()
        if len(reps) > 1:
            raise ValueError(f"Unbalanced replications: {reps}")
        self.r = int(reps[0])

    def analyze_trait(self, trait):
        p = self.p
        r = self.r
        n_checks = len(self.checks)
        
        # --- 1. FULL ANOVA ---
        y_all = self.df[trait].values
        G_sum = y_all.sum()
        N = len(y_all)
        CF = (G_sum**2) / N
        SS_Total = (y_all**2).sum() - CF
        
        rep_sums = self.df.groupby(self.rep_col)[trait].sum()
        SS_Rep = (rep_sums**2).sum() / (p*(p+1)/2 + n_checks) - CF
        
        self.df['entry_id'] = self.df.apply(lambda row: f"{row[self.female_col]}x{row[self.male_col]}" if row[self.check_col].lower() == 'diallel' else row[self.check_col], axis=1)
        geno_sums = self.df.groupby('entry_id')[trait].sum()
        SS_Geno = (geno_sums**2).sum() / r - CF
        
        SS_Error = SS_Total - SS_Rep - SS_Geno
        df_error = (r-1) * ( (p*(p+1)/2 + n_checks) - 1 )
        MS_Error = SS_Error / df_error
        
        # --- 2. Combining Ability (Method II) ---
        diallel_means = self.diallel_df.groupby(['__p1', '__p2'])[trait].mean().unstack()
        full_means = np.zeros((p, p))
        for i in range(p):
            for j in range(i, p):
                val = diallel_means.loc[self.parents[i], self.parents[j]]
                full_means[i, j] = val; full_means[j, i] = val
        
        X_i = full_means.sum(axis=1) # Sum of means involving parent i
        X = 0
        for i in range(p):
            for j in range(i, p):
                X += full_means[i, j] # Sum of entry means
                
        # SS GCA (Method II)
        term1_gca = 0
        for i in range(p):
            term1_gca += (X_i[i] + full_means[i, i])**2
        SS_GCA = (r / (p + 2)) * (term1_gca - (4 / p) * (X**2))
        
        # Diallel SS
        diallel_sums = self.diallel_df.groupby(['__p1', '__p2'])[trait].sum()
        dial_G = diallel_sums.sum()
        SS_Diallel = (diallel_sums**2).sum() / r - (dial_G**2) / (r * p * (p+1) / 2)
        SS_SCA = SS_Diallel - SS_GCA
        
        df_gca = p - 1; df_sca = p * (p - 1) // 2
        MS_GCA = SS_GCA / df_gca; MS_SCA = SS_SCA / df_sca
        
        # Effects
        gca_effects = []
        for i in range(p):
            gi = (1 / (p + 2)) * ( (X_i[i] + full_means[i, i]) - (2 / p) * X )
            gca_effects.append(gi)
            
        sca_effects = np.zeros((p, p))
        for i in range(p):
            for j in range(i, p):
                s_ij = full_means[i, j] - (1 / (p + 2)) * (X_i[i] + full_means[i, i] + X_i[j] + full_means[j, j]) + (2 / ((p + 1) * (p + 2))) * X
                sca_effects[i, j] = s_ij
        
        # Genetic Params
        v_gca = max(0, (MS_GCA - MS_Error) / (p + 2))
        v_sca = max(0, MS_SCA - MS_Error)
        v_a = 2 * v_gca; v_d = v_sca
        v_p = v_a + v_d + MS_Error
        h2_broad = (v_a + v_d) / v_p if v_p > 0 else 0
        h2_narrow = v_a / v_p if v_p > 0 else 0
        pred = (2 * v_gca) / (2 * v_gca + v_sca) if (2 * v_gca + v_sca) > 0 else 0
        
        # Heterosis
        check_means = self.checks_df.groupby(self.check_col)[trait].mean()
        best_check = check_means.max() if not check_means.empty else 0
        
        std_h = {}
        mph = {}; hb = {}
        for i in range(p):
            for j in range(i + 1, p):
                Fi = full_means[i, j]
                # MP/BP
                Pi = full_means[i, i]; Pj = full_means[j, j]
                MP = (Pi + Pj) / 2; BP = max(Pi, Pj)
                
                mph[f"{self.parents[i]}x{self.parents[j]}"] = ((Fi-MP)/MP)*100 if MP!=0 else 0
                hb[f"{self.parents[i]}x{self.parents[j]}"] = ((Fi-BP)/BP)*100 if BP!=0 else 0
                
                if best_check > 0:
                    ch = ((Fi-best_check)/best_check)*100
                    std_h[f"{self.parents[i]}x{self.parents[j]}"] = ch

        res = {
            "anova_comb": {
                "GCA": {"df": df_gca, "SS": SS_GCA, "MS": MS_GCA, "F": MS_GCA/MS_Error, "P": 1-stats.f.cdf(MS_GCA/MS_Error, df_gca, df_error)},
                "SCA": {"df": df_sca, "SS": SS_SCA, "MS": MS_SCA, "F": MS_SCA/MS_Error, "P": 1-stats.f.cdf(MS_SCA/MS_Error, df_sca, df_error)},
                "Error": {"df": df_error, "SS": SS_Error, "MS": MS_Error}
            },
            "gca_effects": [{"parent": self.parents[i], "effect": gca_effects[i]} for i in range(p)],
            "sca_matrix": sca_effects.tolist(),
            "variances": {"h2_broad": h2_broad, "h2_narrow": h2_narrow, "predictability": pred},
            "heterosis": {"mph": mph, "hb": hb, "std": std_h},
            "parents": self.parents
        }
        self.results[trait] = res
        return res

    def get_sig(self, p):
        if p is None: return ""
        if p <= 0.01: return "**"
        if p <= 0.05: return "*"
        return "ns"

    def create_report(self):
        doc = Document()
        doc.add_heading("Griffing's Method II (Half Diallel WITH Check) Report", 0)
        for trait in self.trait_cols:
            res = self.results[trait]; doc.add_heading(f"Trait: {trait}", level=1)
            # Add tables similar to previous
        f = io.BytesIO(); doc.save(f); f.seek(0)
        return f

    def create_excel(self):
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine='openpyxl') as writer:
            for trait in self.trait_cols:
                res = self.results[trait]
                pd.DataFrame(res['gca_effects']).to_excel(writer, sheet_name=f"{trait[:10]}_GCA")
        out.seek(0); return out
