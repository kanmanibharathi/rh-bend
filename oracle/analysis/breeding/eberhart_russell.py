import pandas as pd
import numpy as np
import io
import matplotlib.pyplot as plt
from scipy import stats
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class EberhartRussellAnalyzer:
    def __init__(self, df, genotype_col, env_col, rep_col, trait_col, model_type="fixed"):
        self.df = df.copy()
        self.genotype_col = genotype_col
        self.env_col = env_col
        self.rep_col = rep_col
        self.trait_col = trait_col
        self.model_type = model_type # "fixed" or "mixed"
        
        self.results = {}
        
    def validate(self):
        # Remove missing values from mandatory columns
        self.df = self.df.dropna(subset=[self.genotype_col, self.env_col, self.rep_col, self.trait_col])
        
        # Ensure numeric trait
        self.df[self.trait_col] = pd.to_numeric(self.df[self.trait_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.trait_col])
        
        if len(self.df) == 0:
            raise ValueError("No valid numeric data found after cleaning.")

        # Check for minimum environments
        envs = sorted(self.df[self.env_col].unique())
        if len(envs) < 3:
            raise ValueError(f"Stability analysis requires at least 3 environments. Found: {len(envs)} ({', '.join(map(str, envs))})")
            
        # Strict Balance: Genotypes must be present in ALL environments for regression
        counts = self.df.groupby(self.genotype_col)[self.env_col].nunique()
        valid_genos = counts[counts == len(envs)].index
        if len(valid_genos) < 2:
            raise ValueError(f"Insufficient overlap. Only {len(valid_genos)} genotypes were tested across all {len(envs)} environments. Please ensure trials are reasonably balanced.")
        
        self.df = self.df[self.df[self.genotype_col].isin(valid_genos)]
        
        # Drop duplicates if any (only one observation per GxExR expected)
        self.df = self.df.drop_duplicates(subset=[self.genotype_col, self.env_col, self.rep_col])
        
    def run_analysis(self):
        # 1. Environment-wise RCBD ANOVA
        env_anova_results = {}
        error_ms_list = []
        
        envs = sorted(self.df[self.env_col].unique())
        for env in envs:
            env_df = self.df[self.df[self.env_col] == env]
            
            # Simple RCBD ANOVA
            # Y = Mu + G + R + E
            n_g = len(env_df[self.genotype_col].unique())
            n_r = len(env_df[self.rep_col].unique())
            
            grand_mean = env_df[self.trait_col].mean()
            
            # SS Total
            ss_total = ((env_df[self.trait_col] - grand_mean)**2).sum()
            df_total = len(env_df) - 1
            
            # SS Genotype
            g_means = env_df.groupby(self.genotype_col)[self.trait_col].mean()
            ss_gen = n_r * ((g_means - grand_mean)**2).sum()
            df_gen = n_g - 1
            
            # SS Replication
            r_means = env_df.groupby(self.rep_col)[self.trait_col].mean()
            ss_rep = n_g * ((r_means - grand_mean)**2).sum()
            df_rep = n_r - 1
            
            # SS Error
            ss_err = ss_total - ss_gen - ss_rep
            df_err = df_total - df_gen - df_rep
            
            ms_gen = ss_gen / df_gen if df_gen > 0 else 0
            ms_err = ss_err / df_err if df_err > 0 else 0
            
            error_ms_list.append(ms_err)
            
            env_anova_results[env] = {
                "MS_gen": ms_gen,
                "MS_err": ms_err,
                "df_err": df_err,
                "F": ms_gen / ms_err if ms_err > 0 else 0,
                "P": 1 - stats.f.cdf(ms_gen / ms_err, df_gen, df_err) if ms_err > 0 else 1
            }
            
        self.results['env_anova'] = env_anova_results
        
        # 2. Bartlett's Test
        # Standard Bartlett for MS_errors:
        k = len(error_ms_list)
        dfs = [max(1, res["df_err"]) for res in env_anova_results.values()]
        total_df = sum(dfs)
        pooled_ms_err = sum([ms * df for ms, df in zip(error_ms_list, dfs)]) / total_df if total_df > 0 else 0
        
        # Check for zero MS to avoid log(0)
        safe_error_ms = [max(ms, 1e-10) for ms in error_ms_list]
        safe_pooled_ms = max(pooled_ms_err, 1e-10)
        
        q = total_df * np.log(safe_pooled_ms) - sum([df * np.log(ms) for ms, df in zip(safe_error_ms, dfs)])
        c_denom = 3 * (k - 1)
        c = 1 + (1 / c_denom) * (sum([1/df for df in dfs]) - (1/total_df)) if c_denom > 0 else 1
        chi_sq = max(0, q / c)
        bartlett_p = 1 - stats.chi2.cdf(chi_sq, k - 1) if k > 1 else 1
        
        self.results['bartlett'] = {"stat": chi_sq, "p": bartlett_p}
        
        # 3. Pooled ANOVA
        # Re-calc for pooled
        n_e = len(envs)
        n_g = len(self.df[self.genotype_col].unique())
        n_r = len(self.df[self.rep_col].unique())
        
        grand_mean = self.df[self.trait_col].mean()
        
        ss_total = ((self.df[self.trait_col] - grand_mean)**2).sum()
        df_total = len(self.df) - 1
        
        # Environment
        e_means = self.df.groupby(self.env_col)[self.trait_col].mean()
        ss_env = (n_g * n_r) * ((e_means - grand_mean)**2).sum()
        df_env = n_e - 1
        
        # Genotype
        g_means = self.df.groupby(self.genotype_col)[self.trait_col].mean()
        ss_gen = (n_e * n_r) * ((g_means - grand_mean)**2).sum()
        df_gen = n_g - 1
        
        # GxE
        ge_means = self.df.groupby([self.genotype_col, self.env_col])[self.trait_col].mean()
        ss_gxe = n_r * ((ge_means - grand_mean)**2).sum() - ss_env - ss_gen
        df_gxe = (n_g - 1) * (n_e - 1)
        
        # Rep within Env
        ss_rep_env = 0
        for env in envs:
            env_df = self.df[self.df[self.env_col] == env]
            r_means_env = env_df.groupby(self.rep_col)[self.trait_col].mean()
            ss_rep_env += n_g * ((r_means_env - env_df[self.trait_col].mean())**2).sum()
        df_rep_env = n_e * (n_r - 1)
        
        # Pooled Error
        ss_pooled_err = ss_total - ss_env - ss_gen - ss_gxe - ss_rep_env
        df_pooled_err = df_total - df_env - df_gen - df_gxe - df_rep_env
        
        ms_env = ss_env / df_env
        ms_gen = ss_gen / df_gen
        ms_gxe = ss_gxe / df_gxe
        ms_rep_env = ss_rep_env / df_rep_env
        ms_pooled_err = ss_pooled_err / df_pooled_err
        
        # F-tests
        if self.model_type == "fixed":
            f_env = ms_env / ms_rep_env if ms_rep_env > 0 else 0
            p_env = 1 - stats.f.cdf(f_env, df_env, df_rep_env)
            f_gen = ms_gen / ms_pooled_err if ms_pooled_err > 0 else 0
            p_gen = 1 - stats.f.cdf(f_gen, df_gen, df_pooled_err)
            f_gxe = ms_gxe / ms_pooled_err if ms_pooled_err > 0 else 0
            p_gxe = 1 - stats.f.cdf(f_gxe, df_gxe, df_pooled_err)
        else: # mixed
            f_env = ms_env / ms_rep_env if ms_rep_env > 0 else 0
            p_env = 1 - stats.f.cdf(f_env, df_env, df_rep_env)
            f_gen = ms_gen / ms_gxe if ms_gxe > 0 else 0 # Gen tested against GxE
            p_gen = 1 - stats.f.cdf(f_gen, df_gen, df_gxe)
            f_gxe = ms_gxe / ms_pooled_err if ms_pooled_err > 0 else 0
            p_gxe = 1 - stats.f.cdf(f_gxe, df_gxe, df_pooled_err)
            
        self.results['pooled_anova'] = {
            "Environment": {"df": df_env, "SS": ss_env, "MS": ms_env, "F": f_env, "P": p_env},
            "Rep within Env": {"df": df_rep_env, "SS": ss_rep_env, "MS": ms_rep_env, "F": 0, "P": 1},
            "Genotype": {"df": df_gen, "SS": ss_gen, "MS": ms_gen, "F": f_gen, "P": p_gen},
            "Genotype x Environment": {"df": df_gxe, "SS": ss_gxe, "MS": ms_gxe, "F": f_gxe, "P": p_gxe},
            "Pooled Error": {"df": df_pooled_err, "SS": ss_pooled_err, "MS": ms_pooled_err, "F": 0, "P": 1}
        }
        
        # 4. Environmental Index
        env_means = self.df.groupby(self.env_col)[self.trait_col].mean()
        grand_mean = self.df[self.trait_col].mean()
        env_indices = env_means - grand_mean
        
        self.results['env_indices'] = env_indices.to_dict()
        
        # 5. Eberhart-Russell Stability ANOVA
        # Sources: Genotype, Env + (GxE), Env (Linear), GxE (Linear), Pooled Deviations, Pooled Error
        
        ss_env_plus_gxe = ss_env + ss_gxe
        df_env_plus_gxe = df_env + df_gxe
        
        # Env (Linear) - 1 df
        # This is the SS due to regression on the environmental index, summed over all genotypes.
        # But specifically, "Environment (Linear)" is often defined as the total SS of environmental indices * n_g * n_r
        # Actually in ER model:
        # SS Env (Linear) = n_g * n_r * sum(Ij^2)
        ss_env_linear = n_g * n_r * (env_indices**2).sum()
        df_env_linear = 1
        
        # GxE (Linear) - (g-1) df
        # SS GxE (Linear) = n_r * sum_i [ (sum_j Yij*Ij)^2 / sum(Ij^2) ] - SS Env (Linear)
        # We need Yij as means over replications
        y_ij = self.df.groupby([self.genotype_col, self.env_col])[self.trait_col].mean().unstack(self.env_col)
        gen_list = y_ij.index
        env_list = y_ij.columns # should match indices order if we are careful
        
        # Ensure indices align
        I_j_vec = env_indices.loc[env_list].values
        sum_Ij2 = (I_j_vec**2).sum()
        
        raw_b_i = []
        ss_gxe_linear_total = 0
        for gen in gen_list:
            y_i = y_ij.loc[gen].values
            bi = np.sum(y_i * I_j_vec) / sum_Ij2
            raw_b_i.append(bi)
            ss_gxe_linear_total += (np.sum(y_i * I_j_vec)**2) / sum_Ij2
            
        ss_gxe_linear = n_r * ss_gxe_linear_total - ss_env_linear
        df_gxe_linear = n_g - 1
        
        # Pooled Deviations
        ss_pooled_dev = ss_env_plus_gxe - ss_env_linear - ss_gxe_linear
        df_pooled_dev = n_g * (n_e - 2)
        
        ms_env_linear = ss_env_linear / df_env_linear
        ms_gxe_linear = ss_gxe_linear / df_gxe_linear
        ms_pooled_dev = ss_pooled_dev / df_pooled_dev if df_pooled_dev > 0 else 0
        
        # Stability ANOVA F-tests
        self.results['stability_anova'] = {
            "Genotypes": {"df": df_gen, "SS": ss_gen, "MS": ms_gen, "F": ms_gen / ms_pooled_dev if ms_pooled_dev > 0 else 0, "P": 1-stats.f.cdf(ms_gen/ms_pooled_dev, df_gen, df_pooled_dev) if ms_pooled_dev > 0 else 1},
            "Environments + (G x E)": {"df": df_env_plus_gxe, "SS": ss_env_plus_gxe, "MS": ss_env_plus_gxe/df_env_plus_gxe if df_env_plus_gxe > 0 else 0},
            "Environments (Linear)": {"df": df_env_linear, "SS": ss_env_linear, "MS": ms_env_linear, "F": ms_env_linear / ms_pooled_dev if ms_pooled_dev > 0 else 0, "P": 1-stats.f.cdf(ms_env_linear/ms_pooled_dev, df_env_linear, df_pooled_dev) if ms_pooled_dev > 0 else 1},
            "G x E (Linear)": {"df": df_gxe_linear, "SS": ss_gxe_linear, "MS": ms_gxe_linear, "F": ms_gxe_linear / ms_pooled_dev if ms_pooled_dev > 0 else 0, "P": 1-stats.f.cdf(ms_gxe_linear/ms_pooled_dev, df_gxe_linear, df_pooled_dev) if ms_pooled_dev > 0 else 1},
            "Pooled Deviations": {"df": df_pooled_dev, "SS": ss_pooled_dev, "MS": ms_pooled_dev, "F": ms_pooled_dev / ms_pooled_err if ms_pooled_err > 0 else 0, "P": 1-stats.f.cdf(ms_pooled_dev/ms_pooled_err, df_pooled_dev, df_pooled_err) if ms_pooled_err > 0 else 1},
            "Pooled Error": {"df": df_pooled_err, "SS": ss_pooled_err, "MS": ms_pooled_err}
        }
        
        # 6. Stability Parameters per Genotype
        stability_params = []
        for i, gen in enumerate(gen_list):
            y_i = y_ij.loc[gen].values
            mean_i = y_i.mean()
            bi = raw_b_i[i]
            
            sq_devs = (y_i - mean_i - bi * I_j_vec)**2
            ms_di = np.sum(sq_devs) * n_r / (n_e - 2) if (n_e - 2) > 0 else 0
            
            s2di = ms_di - ms_pooled_err
            
            # SE(bi)
            se_bi_val = (n_r * sum_Ij2)
            se_bi = np.sqrt(ms_pooled_dev / se_bi_val) if se_bi_val > 0 else 0
            
            t_b0 = bi / se_bi if se_bi > 0 else 0
            p_b0 = (1 - stats.t.cdf(abs(t_b0), df_pooled_dev)) * 2 if df_pooled_dev > 0 else 1
            
            t_b1 = (bi - 1) / se_bi if se_bi > 0 else 0
            p_b1 = (1 - stats.t.cdf(abs(t_b1), df_pooled_dev)) * 2 if df_pooled_dev > 0 else 1
            
            f_s2di = ms_di / ms_pooled_err if ms_pooled_err > 0 else 0
            p_s2di = 1 - stats.f.cdf(f_s2di, (n_e - 2), df_pooled_err) if ms_pooled_err > 0 and (n_e - 2) > 0 else 1
            
            # Inference logic
            if mean_i > grand_mean and p_b1 > 0.05 and p_s2di > 0.05:
                inf = "Widely adapted and stable"
            elif bi > 1 and p_s2di > 0.05:
                inf = "Responsive to favorable environments"
            elif bi < 1 and p_s2di > 0.05:
                inf = "Responsive to unfavorable environments"
            elif p_s2di <= 0.05:
                inf = "Unstable genotype (high deviations)"
            else:
                inf = "Average stability"
                
            stability_params.append({
                "Genotype": gen,
                "Mean": mean_i,
                "bi": bi,
                "SE_bi": se_bi,
                "t_b0": t_b0,
                "p_b0": p_b0,
                "t_b1": t_b1,
                "p_b1": p_b1,
                "MS_di": ms_di,
                "S2di": s2di,
                "F_S2di": f_s2di,
                "p_S2di": p_s2di,
                "Inference": inf
            })
            
        self.results['stability_parameters'] = stability_params
        self.results['grand_mean'] = grand_mean
        
        return self.results

    def generate_plots(self):
        # Regression Plot
        # X: Env Index, Y: Genotype Mean
        # One line per genotype
        
        plt.figure(figsize=(10, 6))
        
        y_ij = self.df.groupby([self.genotype_col, self.env_col])[self.trait_col].mean().unstack(self.env_col)
        env_indices = pd.Series(self.results['env_indices'])
        env_list = y_ij.columns
        I_j_vec = env_indices.loc[env_list].values
        
        x_range = np.linspace(min(I_j_vec), max(I_j_vec), 100)
        
        for params in self.results['stability_parameters']:
            gen = params['Genotype']
            bi = params['bi']
            mean_i = params['Mean']
            
            # Plot the fitted line
            plt.plot(x_range, mean_i + bi * x_range, label=f"{gen} (b={bi:.2f})")
            # Plot actual points
            plt.scatter(I_j_vec, y_ij.loc[gen].values, alpha=0.5)
            
        plt.axhline(self.results['grand_mean'], color='black', linestyle='--', alpha=0.3)
        plt.axvline(0, color='black', linestyle='--', alpha=0.3)
        plt.xlabel('Environmental Index (Ij)')
        plt.ylabel(f'{self.trait_col} Mean')
        plt.title(f'Regression Lines for {self.trait_col}')
        # If too many genotypes, hide legend or put outside
        if len(y_ij) <= 15:
            plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
        
        plt.tight_layout()
        reg_plot = io.BytesIO()
        plt.savefig(reg_plot, format='png', dpi=300)
        reg_plot.seek(0)
        plt.close()
        
        # Stability Scatter Plot
        # X: bi, Y: Mean
        plt.figure(figsize=(8, 6))
        
        bi_vals = [p['bi'] for p in self.results['stability_parameters']]
        mean_vals = [p['Mean'] for p in self.results['stability_parameters']]
        labels = [p['Genotype'] for p in self.results['stability_parameters']]
        
        plt.scatter(bi_vals, mean_vals, color='blue')
        for i, label in enumerate(labels):
            plt.annotate(label, (bi_vals[i], mean_vals[i]), xytext=(5, 5), textcoords='offset points')
            
        plt.axvline(1, color='red', linestyle='--', label='b=1')
        plt.axhline(self.results['grand_mean'], color='green', linestyle='--', label='Grand Mean')
        
        plt.xlabel('Regression Coefficient (bi)')
        plt.ylabel('Genotype Mean')
        plt.title('Stability Scatter Plot')
        plt.legend()
        
        plt.tight_layout()
        stab_plot = io.BytesIO()
        plt.savefig(stab_plot, format='png', dpi=300)
        stab_plot.seek(0)
        plt.close()
        
        return reg_plot, stab_plot

    def create_report(self):
        doc = Document()
        
        # Title
        title = doc.add_heading('Stability Analysis Report (Eberhart & Russell Model)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Experimental Details
        doc.add_heading('Experimental Details', level=1)
        details = doc.add_paragraph()
        details.add_run(f"Trait: {self.trait_col}\n").bold = True
        details.add_run(f"Environments: {', '.join(map(str, sorted(self.df[self.env_col].unique())))}\n")
        details.add_run(f"Genotypes: {len(self.df[self.genotype_col].unique())}\n")
        details.add_run(f"Replications: {len(self.df[self.rep_col].unique())}\n")
        details.add_run(f"Model Type: {self.model_type.capitalize()}\n")
        
        # Bartlett's Test
        doc.add_heading('Bartlett\'s Test for Homogeneity of Variances', level=1)
        bt = self.results['bartlett']
        p = doc.add_paragraph()
        p.add_run(f"Chi-square: {bt['stat']:.4f}\n")
        p.add_run(f"p-value: {bt['p']:.4e}\n")
        if bt['p'] > 0.05:
            p.add_run("Decision: Error variances are homogeneous. Proceeding to Pooled ANOVA.")
        else:
            p.add_run("Decision: Error variances are heterogeneous. Pooled ANOVA should be interpreted with caution.")
            
        # Pooled ANOVA
        doc.add_heading('Combined (Pooled) ANOVA', level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(["Source", "df", "SS", "MS", "F", "p-value"]):
            hdr_cells[i].text = text
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        for src, val in self.results['pooled_anova'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = src
            row_cells[1].text = str(val['df'])
            row_cells[2].text = f"{val['SS']:.4f}"
            row_cells[3].text = f"{val['MS']:.4f}"
            row_cells[4].text = f"{val['F']:.4f}" if val['F'] > 0 else "-"
            row_cells[5].text = f"{val['P']:.4f}" if val['P'] <= 1 else "-"
            
        # Stability ANOVA
        doc.add_heading('Stability ANOVA (Eberhart & Russell)', level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(["Source", "df", "SS", "MS", "F", "p-value"]):
            hdr_cells[i].text = text
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        for src, val in self.results['stability_anova'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = src
            row_cells[1].text = str(int(val['df']))
            row_cells[2].text = f"{val['SS']:.4f}"
            row_cells[3].text = f"{val['MS']:.4f}"
            row_cells[4].text = f"{val.get('F', 0):.4f}" if val.get('F', 0) > 0 else "-"
            row_cells[5].text = f"{val.get('P', 0):.4f}" if val.get('P', 0) <= 1 else "-"

        # Environmental Index
        doc.add_heading('Environmental Index Table', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Environment"
        table.rows[0].cells[1].text = "Mean"
        table.rows[0].cells[2].text = "Index (Ij)"
        
        env_means = self.df.groupby(self.env_col)[self.trait_col].mean()
        for env, idx_val in self.results['env_indices'].items():
            row = table.add_row().cells
            row[0].text = str(env)
            row[1].text = f"{env_means[env]:.4f}"
            row[2].text = f"{idx_val:.4f}"
            
        # Stability Parameters
        doc.add_heading('Stability Parameters', level=1)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        hdr = ["Genotype", "Mean", "bi", "SE(bi)", "t(b=1)", "MS di", "S2di", "F(S2di)"]
        for i, text in enumerate(hdr):
            table.rows[0].cells[i].text = text
            
        for p in self.results['stability_parameters']:
            row = table.add_row().cells
            row[0].text = str(p['Genotype'])
            row[1].text = f"{p['Mean']:.4f}"
            row[2].text = f"{p['bi']:.4f}"
            row[3].text = f"{p['SE_bi']:.4f}"
            row[4].text = f"{p['t_b1']:.4f} ({'*' if p['p_b1'] < 0.05 else 'ns'})"
            row[5].text = f"{p['MS_di']:.4f}"
            row[6].text = f"{p['S2di']:.4f}"
            row[7].text = f"{p['F_S2di']:.4f} ({'*' if p['p_S2di'] < 0.05 else 'ns'})"
            
        # Stability Inference
        doc.add_heading('Stability Inference', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        for i, text in enumerate(["Genotype", "Mean", "bi", "Inference"]):
            table.rows[0].cells[i].text = text
            
        for p in self.results['stability_parameters']:
            row = table.add_row().cells
            row[0].text = str(p['Genotype'])
            row[1].text = f"{p['Mean']:.4f}"
            row[2].text = f"{p['bi']:.4f}"
            row[3].text = p['Inference']
            
        # Plots
        reg_plot, stab_plot = self.generate_plots()
        
        doc.add_heading('Regression Plot', level=1)
        doc.add_picture(reg_plot, width=Inches(6))
        
        doc.add_heading('Stability Scatter Plot', level=1)
        doc.add_picture(stab_plot, width=Inches(6))
        
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph("The analysis was performed according to the Eberhart & Russell (1966) model. "
                          "A stable genotype is characterized by a high mean performance, regression coefficient (bi) close to unity, "
                          "and deviation from regression (S2di) close to zero.")
        
        doc.add_heading('References', level=1)
        refs = [
            "Eberhart, S. A., & Russell, W. A. (1966). Stability parameters for comparing varieties. Crop Science, 6(1), 36-40.",
            "Finlay, K. W., & Wilkinson, G. N. (1963). The analysis of adaptation in a plant-breeding programme. Australian Journal of Agricultural Research, 14(6), 742-754.",
            "Gomez, K. A., & Gomez, A. A. (1984). Statistical procedures for agricultural research. John Wiley & Sons."
        ]
        for ref in refs:
            doc.add_paragraph(ref)
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
