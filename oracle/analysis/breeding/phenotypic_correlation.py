import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
import matplotlib.pyplot as plt
import seaborn as sns

class PhenotypicCorrelationAnalyzer:
    def __init__(self, df, genotype_col, rep_col, trait_cols):
        self.df = df
        self.genotype_col = genotype_col
        self.rep_col = rep_col
        self.trait_cols = trait_cols
        
        self.n = 0 # Number of genotypes
        self.r = 0 # Number of replications
        self.N = 0 # Total observations
        self.genotypes = []
        
        # Results storage
        self.variances = {} # trait -> {sigma2_p}
        self.correlation_matrix = None
        self.p_values = None
        self.t_values = None
        self.covariance_matrix = None
        
    def validate(self):
        # Ensure column types
        self.df[self.genotype_col] = self.df[self.genotype_col].astype(str)
        self.df[self.rep_col] = self.df[self.rep_col].astype(str)
        
        for trait in self.trait_cols:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        
        # Check for missing values
        if self.df[self.trait_cols].isnull().any().any():
            raise ValueError("Some trait values are missing or non-numeric.")
            
        # Check RCBD structure
        self.genotypes = sorted(self.df[self.genotype_col].unique())
        self.n = len(self.genotypes)
        reps = sorted(self.df[self.rep_col].unique())
        self.r = len(reps)
        self.N = self.n * self.r
        
        # Check balance
        counts = self.df.groupby([self.genotype_col, self.rep_col]).size()
        if len(counts) != self.N:
            raise ValueError("Experimental design is not a balanced RCBD. All genotypes must appear in all replications.")
        
        if (counts != 1).any():
            raise ValueError("Duplicate entries detected for some Genotype-Replication combinations.")

    def run_analysis(self):
        traits = self.trait_cols
        num_traits = len(traits)
        N = self.N
        
        # 1. Phenotypic Variances
        for trait in traits:
            # Phenotypic variance is the total variance of the trait across all plots
            y = self.df[trait].values
            var_p = np.var(y, ddof=1) # Sample variance (n-1)
            
            self.variances[trait] = {
                "sigma2_p": var_p,
                "mean": np.mean(y)
            }

        # 2. Covariance and Correlation Matrices
        rp_matrix = np.zeros((num_traits, num_traits))
        cov_p_matrix = np.zeros((num_traits, num_traits))
        p_matrix = np.zeros((num_traits, num_traits))
        t_matrix = np.zeros((num_traits, num_traits))
        
        for i in range(num_traits):
            for j in range(num_traits):
                if i == j:
                    rp_matrix[i, j] = 1.0
                    cov_p_matrix[i, j] = self.variances[traits[i]]["sigma2_p"]
                    p_matrix[i, j] = 0.0
                    t_matrix[i, j] = np.inf
                    continue
                
                traitX = traits[i]
                traitY = traits[j]
                
                # Phenotypic Covariance (Pearson based on all observations)
                X = self.df[traitX].values
                Y = self.df[traitY].values
                
                # Cov(X,Y) = Σ[(X_ij - X̄)(Y_ij - Ȳ)] / (N - 1)
                cov_p = np.cov(X, Y)[0, 1]
                
                # Phenotypic Correlation
                var_pX = self.variances[traitX]["sigma2_p"]
                var_pY = self.variances[traitY]["sigma2_p"]
                
                if var_pX > 0 and var_pY > 0:
                    rp = cov_p / np.sqrt(var_pX * var_pY)
                    rp = np.clip(rp, -1.0, 1.0)
                else:
                    rp = 0.0
                
                # Significance Testing (t-test)
                # df = N - 2 where N = total observatons (g * r)
                deg_f = N - 2
                if abs(rp) < 1.0:
                    t_val = rp * np.sqrt(deg_f) / np.sqrt(1 - rp**2)
                else:
                    t_val = np.inf if rp > 0 else -np.inf
                
                p_val = 2 * (1 - stats.t.cdf(abs(t_val), deg_f))
                
                rp_matrix[i, j] = rp
                cov_p_matrix[i, j] = cov_p
                t_matrix[i, j] = t_val
                p_matrix[i, j] = p_val
                
        self.correlation_matrix = pd.DataFrame(rp_matrix, index=traits, columns=traits)
        self.covariance_matrix = pd.DataFrame(cov_p_matrix, index=traits, columns=traits)
        self.t_values = pd.DataFrame(t_matrix, index=traits, columns=traits)
        self.p_values = pd.DataFrame(p_matrix, index=traits, columns=traits)

    def get_interpretation(self):
        interpretations = []
        traits = self.trait_cols
        
        for i in range(len(traits)):
            for j in range(i + 1, len(traits)):
                t1 = traits[i]
                t2 = traits[j]
                rp = self.correlation_matrix.loc[t1, t2]
                p = self.p_values.loc[t1, t2]
                
                strength = ""
                abs_r = abs(rp)
                if abs_r > 0.8: strength = "a very strong"
                elif abs_r > 0.6: strength = "a strong"
                elif abs_r > 0.4: strength = "a moderate"
                else: strength = "a weak"
                
                direction = "positive" if rp > 0 else "negative"
                
                if p <= 0.01:
                    text = f"The phenotypic expression of '{t1}' exhibited {strength} {direction} and highly significant association with '{t2}' (rp = {rp:.3f}**)."
                elif p <= 0.05:
                    text = f"The trait '{t1}' showed {strength} {direction} and significant phenotypic correlation with '{t2}' (rp = {rp:.3f}*)."
                else:
                    text = f"The phenotypic relationship between '{t1}' and '{t2}' was non-significant (rp = {rp:.3f}, p > 0.05)."
                
                interpretations.append(text)
        return interpretations

    def create_report(self):
        doc = Document()
        
        # Heading
        title = doc.add_heading("Phenotypic Correlation Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 1. Experimental Details
        doc.add_heading("1. Analysis Overview", level=1)
        doc.add_paragraph(f"Experimental Design: Randomized Complete Block Design (RCBD)")
        doc.add_paragraph(f"Total Observations (N): {self.N} ({self.n} Genotypes × {self.r} Replications)")
        doc.add_paragraph("Phenotypic correlation provides insights into the overall observable association between traits, combining both genetic and environmental influences.")
        
        # 2. Phenotypic Variances
        doc.add_heading("2. Phenotypic Variances", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Trait'
        hdr_cells[1].text = 'Mean'
        hdr_cells[2].text = 'Phenotypic Var (σ²p)'
        
        for trait in self.trait_cols:
            v = self.variances[trait]
            row_cells = table.add_row().cells
            row_cells[0].text = trait
            row_cells[1].text = f"{v['mean']:.4f}"
            row_cells[2].text = f"{v['sigma2_p']:.4f}"
            
        # 3. Phenotypic Correlation Matrix
        doc.add_heading("3. Phenotypic Correlation Matrix", level=1)
        num_traits = len(self.trait_cols)
        table_rp = doc.add_table(rows=num_traits + 1, cols=num_traits + 1)
        table_rp.style = 'Table Grid'
        
        # Headers
        for i, trait in enumerate(self.trait_cols):
            table_rp.cell(0, i + 1).text = trait
            table_rp.cell(i + 1, 0).text = trait
            
        # Values
        for i in range(num_traits):
            for j in range(num_traits):
                rp = self.correlation_matrix.iloc[i, j]
                p = self.p_values.iloc[i, j]
                sig = "**" if p <= 0.01 else ("*" if p <= 0.05 else "") if i != j else ""
                table_rp.cell(i + 1, j + 1).text = f"{rp:.3f}{sig}"
                
        doc.add_paragraph("Significance levels: * p ≤ 0.05, ** p ≤ 0.01")
        
        # 4. Heatmap
        doc.add_heading("4. Correlation Heatmap", level=1)
        heatmap_buf = self.generate_heatmap()
        doc.add_picture(heatmap_buf, width=Inches(6))
        
        # 5. Interpretation
        doc.add_heading("5. Automated Interpretation", level=1)
        interps = self.get_interpretation()
        for p_text in interps:
            doc.add_paragraph(p_text, style='List Bullet')
            
        # Footer
        doc.add_page_break()
        doc.add_paragraph("Statistical Note:")
        doc.add_paragraph(f"Degrees of freedom (df) for t-test assessment = (N - 2) = {self.N - 2}.")
        doc.add_paragraph("Report generated by Research Hub Statistical Analysis Engine.")
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Sheet 1: Raw Data
            self.df.to_excel(writer, sheet_name='Raw Data', index=False)
            
            # Sheet 2: Phenotypic Variances
            var_df = pd.DataFrame.from_dict(self.variances, orient='index')
            var_df.to_excel(writer, sheet_name='Phenotypic Variances')
            
            # Sheet 3: Covariance Matrix
            self.covariance_matrix.to_excel(writer, sheet_name='Phenotypic Covariance')
            
            # Sheet 4: Correlation Matrix
            self.correlation_matrix.to_excel(writer, sheet_name='Phenotypic Correlation')
            
            # Sheet 5: Significance
            sig_data = []
            for i in range(num_traits := len(self.trait_cols)):
                for j in range(i+1, num_traits):
                    t1 = self.trait_cols[i]
                    t2 = self.trait_cols[j]
                    sig_data.append({
                        "Trait 1": t1,
                        "Trait 2": t2,
                        "Correlation (rp)": self.correlation_matrix.loc[t1, t2],
                        "t-value": self.t_values.loc[t1, t2],
                        "p-value": self.p_values.loc[t1, t2],
                        "Significance": "Highly Significant (**)" if self.p_values.loc[t1, t2] <= 0.01 else ("Significant (*)" if self.p_values.loc[t1, t2] <= 0.05 else "")
                    })
            sig_df = pd.DataFrame(sig_data)
            sig_df.to_excel(writer, sheet_name='Significance Testing', index=False)
            
        output.seek(0)
        return output

    def generate_heatmap(self):
        plt.figure(figsize=(10, 8))
        sns.heatmap(self.correlation_matrix, annot=True, fmt=".3f", cmap="coolwarm", center=0, vmin=-1, vmax=1)
        plt.title("Phenotypic Correlation Heatmap")
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
