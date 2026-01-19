import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.cluster.hierarchy import dendrogram, linkage
from sklearn.decomposition import PCA
import base64

class MahalanobisD2Analyzer:
    def __init__(self, df, genotype_col, rep_col, trait_cols):
        self.df = df
        self.genotype_col = genotype_col
        self.rep_col = rep_col
        self.trait_cols = trait_cols
        
        # Results
        self.standardized_data = None
        self.d2_matrix = None
        self.clusters = {} # Cluster ID -> List of Genotypes
        self.intra_distances = {}
        self.inter_distances = None
        self.cluster_means = None
        self.trait_contributions = None
        self.genotypes = []
        
    def validate(self):
        self.df[self.genotype_col] = self.df[self.genotype_col].astype(str)
        for col in self.trait_cols:
            self.df[col] = pd.to_numeric(self.df[col], errors='coerce')
        
        if self.df[self.trait_cols].isnull().any().any():
            # Mean imputation as requested
            self.df[self.trait_cols] = self.df[self.trait_cols].fillna(self.df[self.trait_cols].mean())
            
        self.genotypes = self.df[self.genotype_col].tolist()
        if len(self.genotypes) < 3:
            raise ValueError("At least 3 genotypes are required for cluster analysis.")

    def run_analysis(self):
        # 1. Group by Genotype and take means (in case of reps, though Tocher usually expects means)
        # If the input has reps, we should probably average them first. 
        # But Tocher can also be done on raw data if it's an unreplicated set.
        # We'll assume one row per genotype for D2 calculation, or average if duplicates.
        data_means = self.df.groupby(self.genotype_col)[self.trait_cols].mean()
        self.genotypes = data_means.index.tolist()
        X = data_means.values
        n_geno = len(self.genotypes)
        
        # 2. Data Standardization
        means = np.mean(X, axis=0)
        stds = np.std(X, axis=0, ddof=1)
        stds[stds == 0] = 1e-9 # Prevent division by zero
        Z = (X - means) / stds
        self.standardized_data = pd.DataFrame(Z, index=self.genotypes, columns=self.trait_cols)
        
        # 3. Variance-Covariance Matrix (S) and its Inverse
        # Since we used Z, this is effectively the correlation matrix of X
        S = np.cov(Z, rowvar=False)
        # Add small value to diagonal for stability if near singular
        S += np.eye(S.shape[0]) * 1e-6
        try:
            S_inv = np.linalg.inv(S)
        except np.linalg.LinAlgError:
            S_inv = np.linalg.pinv(S) # Use pseudo-inverse if singular
            
        # 4. Compute Mahalanobis D2 Distance Matrix
        d2_mat = np.zeros((n_geno, n_geno))
        for i in range(n_geno):
            for j in range(i + 1, n_geno):
                diff = Z[i] - Z[j]
                d2 = diff.T @ S_inv @ diff
                d2_mat[i, j] = d2
                d2_mat[j, i] = d2
        
        self.d2_matrix = pd.DataFrame(d2_mat, index=self.genotypes, columns=self.genotypes)
        
        # 5. Tocher's Algorithm for Clustering
        self._apply_tocher_clustering(d2_mat)
        
        # 6. Statistical Summaries
        self._compute_cluster_stats(data_means)
        self._compute_trait_contributions(Z)

    def _apply_tocher_clustering(self, d2_mat):
        n = len(self.genotypes)
        unassigned = list(range(n))
        cluster_list = []
        
        # Overall mean D2 for threshold
        overall_mean_d2 = np.mean(d2_mat[np.triu_indices(n, k=1)])
        
        while unassigned:
            if not cluster_list or not unassigned:
                # Find pair with minimum distance among unassigned
                min_d2 = float('inf')
                pair = None
                for i in range(len(unassigned)):
                    for j in range(i + 1, len(unassigned)):
                        idx1, idx2 = unassigned[i], unassigned[j]
                        if d2_mat[idx1, idx2] < min_d2:
                            min_d2 = d2_mat[idx1, idx2]
                            pair = (idx1, idx2)
                
                if pair:
                    current_cluster = [pair[0], pair[1]]
                    unassigned.remove(pair[0])
                    unassigned.remove(pair[1])
                else:
                    # Only one genotype left
                    current_cluster = [unassigned.pop(0)]
                cluster_list.append(current_cluster)
            
            # Try adding more to the latest cluster
            changed = True
            while changed and unassigned:
                changed = False
                best_candidate = None
                min_avg_d2 = float('inf')
                
                current_cluster = cluster_list[-1]
                for cand in unassigned:
                    avg_d2 = np.mean([d2_mat[cand, member] for member in current_cluster])
                    if avg_d2 < min_avg_d2:
                        min_avg_d2 = avg_d2
                        best_candidate = cand
                
                # Logic: If avg distance to cluster <= overall mean, add to cluster
                if best_candidate is not None and min_avg_d2 <= overall_mean_d2:
                    current_cluster.append(best_candidate)
                    unassigned.remove(best_candidate)
                    changed = True
                else:
                    # Start a new cluster in the next iteration of the outer while loop
                    break

        # Convert indices back to genotype names
        self.clusters = {f"Cluster {i+1}": [self.genotypes[idx] for idx in c] for i, c in enumerate(cluster_list)}

    def _compute_cluster_stats(self, raw_data_means):
        # Cluster Means
        means_list = []
        for name, members in self.clusters.items():
            c_mean = raw_data_means.loc[members].mean()
            c_mean.name = name
            means_list.append(c_mean)
        self.cluster_means = pd.DataFrame(means_list)
        
        # Intra & Inter distances
        cluster_names = list(self.clusters.keys())
        n_clusters = len(cluster_names)
        inter_mat = np.zeros((n_clusters, n_clusters))
        
        for i, name_i in enumerate(cluster_names):
            members_i = self.clusters[name_i]
            # Intra
            if len(members_i) > 1:
                sub_mat = self.d2_matrix.loc[members_i, members_i].values
                self.intra_distances[name_i] = np.mean(sub_mat[np.triu_indices(len(members_i), k=1)])
            else:
                self.intra_distances[name_i] = 0.0
                
            # Inter
            for j, name_j in enumerate(cluster_names):
                if i == j: continue
                members_j = self.clusters[name_j]
                inter_val = self.d2_matrix.loc[members_i, members_j].values.mean()
                inter_mat[i, j] = inter_val
                
        self.inter_distances = pd.DataFrame(inter_mat, index=cluster_names, columns=cluster_names)

    def _compute_trait_contributions(self, Z):
        # Singh & Chaudhary method
        n = Z.shape[0]
        n_traits = Z.shape[1]
        counts = np.zeros(n_traits)
        
        for i in range(n):
            for j in range(i + 1, n):
                diffs = np.abs(Z[i] - Z[j])
                ranked_trait = np.argmax(diffs)
                counts[ranked_trait] += 1
                
        total_pairs = n * (n - 1) / 2
        percentages = (counts / total_pairs) * 100
        self.trait_contributions = pd.DataFrame({
            "Trait": self.trait_cols,
            "Count": counts,
            "Contribution %": percentages
        }).sort_values("Contribution %", ascending=False)

    def generate_dendrogram(self):
        plt.figure(figsize=(10, 7))
        # Use simple linkage for standard breeding dendrogram
        Z_link = linkage(self.d2_matrix.values[np.triu_indices(len(self.genotypes), k=1)], method='average')
        dendrogram(Z_link, labels=self.genotypes, leaf_rotation=90)
        plt.title("Hierarchical Clustering Dendrogram (D2 Distance)")
        plt.ylabel("D2 Distance")
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
        buf.seek(0)
        plt.close()
        return buf

    def generate_cluster_plot(self):
        # 2D PCA for visualization
        pca = PCA(n_components=2)
        X_pca = pca.fit_transform(self.standardized_data)
        
        plt.figure(figsize=(10, 8))
        # Use modern colormap access
        cmap = plt.get_cmap('tab10')
        n_clusters = len(self.clusters)
        
        for i, (name, members) in enumerate(self.clusters.items()):
            indices = [self.genotypes.index(m) for m in members]
            # Cycle through colors if many clusters
            color = cmap(i % 10)
            plt.scatter(X_pca[indices, 0], X_pca[indices, 1], label=name, s=100, color=color)
            for idx in indices:
                plt.annotate(self.genotypes[idx], (X_pca[idx, 0], X_pca[idx, 1]), 
                             xytext=(5, 5), textcoords='offset points', fontsize=8, opacity=0.7)
                
        plt.title("Cluster Visualization (PCA Projection of D2 Statistics)")
        plt.xlabel(f"PC1 ({pca.explained_variance_ratio_[0]*100:.1f}%)")
        plt.ylabel(f"PC2 ({pca.explained_variance_ratio_[1]*100:.1f}%)")
        plt.legend()
        plt.grid(alpha=0.2)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
        buf.seek(0)
        plt.close()
        return buf

    def create_report(self):
        doc = Document()
        title = doc.add_heading("Mahalanobis D2 and Tocher Cluster Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading("1. Cluster Groupings", level=1)
        for name, members in self.clusters.items():
            p = doc.add_paragraph()
            p.add_run(f"{name}: ").bold = True
            p.add_run(", ".join(members))
            
        doc.add_heading("2. Intra and Inter-Cluster Distances", level=1)
        table = doc.add_table(rows=1, cols=len(self.clusters) + 1)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Cluster"
        for i, name in enumerate(self.clusters.keys()):
            hdr[i+1].text = name
            
        for name_i in self.clusters.keys():
            row = table.add_row().cells
            row[0].text = name_i
            for j, name_j in enumerate(self.clusters.keys()):
                if name_i == name_j:
                    row[j+1].text = f"{self.intra_distances[name_i]:.2f}"
                else:
                    row[j+1].text = f"{self.inter_distances.loc[name_i, name_j]:.2f}"

        doc.add_heading("3. Trait Contribution to Divergence", level=1)
        t_table = doc.add_table(rows=1, cols=3)
        t_table.style = 'Table Grid'
        t_table.rows[0].cells[0].text = "Trait"
        t_table.rows[0].cells[1].text = "Rank 1 Count"
        t_table.rows[0].cells[2].text = "Contribution %"
        
        for _, r in self.trait_contributions.iterrows():
            row = t_table.add_row().cells
            row[0].text = str(r['Trait'])
            row[1].text = str(int(r['Count']))
            row[2].text = f"{r['Contribution %']:.2f}%"

        doc.add_heading("4. Visualizations", level=1)
        doc.add_heading("4.1 Dendrogram", level=2)
        doc.add_picture(self.generate_dendrogram(), width=Inches(6))
        
        doc.add_heading("4.2 Cluster Plot", level=2)
        doc.add_picture(self.generate_cluster_plot(), width=Inches(6))
        
        doc.add_heading("5. Interpretation", level=1)
        # Find max inter distance
        max_inter = self.inter_distances.stack().idxmax()
        max_val = self.inter_distances.loc[max_inter]
        top_trait = self.trait_contributions.iloc[0]['Trait']
        
        interpretation = (f"Maximum inter-cluster distance ({max_val:.2f}) was observed between {max_inter[0]} and {max_inter[1]}, "
                          f"indicating wide genetic divergence. The trait '{top_trait}' contributed the maximum ({self.trait_contributions.iloc[0]['Contribution %']:.2f}%) "
                          "to total divergence. Genotypes from these divergent clusters may be prioritized for hybridization programs.")
        doc.add_paragraph(interpretation)
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            self.d2_matrix.to_excel(writer, sheet_name='D2_Matrix')
            
            # Cluster Groups Sheet
            group_data = []
            for name, members in self.clusters.items():
                group_data.append({"Cluster": name, "Genotypes": ", ".join(members), "Size": len(members)})
            pd.DataFrame(group_data).to_excel(writer, sheet_name='Cluster_Groups', index=False)
            
            self.cluster_means.to_excel(writer, sheet_name='Cluster_Means')
            self.inter_distances.to_excel(writer, sheet_name='Intra_Inter_Distances')
            self.trait_contributions.to_excel(writer, sheet_name='Trait_Contributions', index=False)
            
        output.seek(0)
        return output
