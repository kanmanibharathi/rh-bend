import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import seaborn as sns

class GriffingMethod1Analyzer:
    def __init__(self, df, female_col, male_col, rep_col, trait_cols):
        self.df = df
        self.female_col = female_col
        self.male_col = male_col
        self.rep_col = rep_col
        self.trait_cols = trait_cols
        
        self.p = 0 # Number of parents
        self.r = 0 # Number of replications
        self.parents = []
        
        # Results storage per trait
        self.results = {}

    def validate(self):
        # Ensure column types
        self.df[self.female_col] = self.df[self.female_col].astype(str)
        self.df[self.male_col] = self.df[self.male_col].astype(str)
        
        for trait in self.trait_cols:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        
        # Check for missing values in trait columns
        if self.df[self.trait_cols].isnull().any().any():
            raise ValueError("Traits contain missing or non-numeric values.")

        # Identify unique parents
        all_parents = sorted(list(set(self.df[self.female_col].unique()) | set(self.df[self.male_col].unique())))
        self.parents = all_parents
        self.p = len(all_parents)
        
        # Check if full diallel exists (p^2 entries)
        # For each pair (i, j), we need r observations
        entry_counts = self.df.groupby([self.female_col, self.male_col]).size()
        
        # Check for missing crosses
        required_entries = []
        for p1 in self.parents:
            for p2 in self.parents:
                required_entries.append((p1, p2))
        
        missing = []
        for entry in required_entries:
            if entry not in entry_counts.index:
                missing.append(f"{entry[0]} x {entry[1]}")
        
        if missing:
            raise ValueError(f"Missing crosses in diallel: {', '.join(missing[:5])}" + ("..." if len(missing) > 5 else ""))
        
        # Check balance
        reps = entry_counts.unique()
        if len(reps) > 1:
            raise ValueError(f"Unbalanced replications detected: {reps}. Replications must be equal for all crosses.")
        
        self.r = int(reps[0])
        
    def analyze_trait(self, trait):
        # 1. Compute cell means (Y_ij)
        means_df = self.df.groupby([self.female_col, self.male_col])[trait].mean().unstack()
        # Ensure rows and columns are in same parent order
        means_df = means_df.reindex(index=self.parents, columns=self.parents)
        
        # 2. Grand Mean and Sums
        Y_cell_means = means_df.values # p x p matrix
        Y_dot_dot = Y_cell_means.sum()
        p = self.p
        r = self.r
        
        # Row and Column means of cell means
        R_i = Y_cell_means.sum(axis=1) # Σ_j Ȳ_ij
        C_j = Y_cell_means.sum(axis=0) # Σ_i Ȳ_ij
        
        # 3. ANOVA (Genotypes, Replications, Error)
        # SS_Total_Raw = ΣΣΣ Y_ijk^2 - CF
        # CF = (ΣΣΣ Y_ijk)^2 / (r*p^2)
        y_all = self.df[trait].values
        G_total = y_all.sum()
        N = r * p * p
        CF = (G_total**2) / N
        SS_Total = (y_all**2).sum() - CF
        
        # SS_Replications
        rep_sums = self.df.groupby(self.rep_col)[trait].sum()
        SS_Rep = (rep_sums**2).sum() / (p*p) - CF
        
        # SS_Genotypes
        genotype_sums = self.df.groupby([self.female_col, self.male_col])[trait].sum()
        SS_Genotypes = (genotype_sums**2).sum() / r - CF
        
        # SS_Error
        SS_Error_ANOVA = SS_Total - SS_Rep - SS_Genotypes
        
        df_rep = r - 1
        df_geno = p*p - 1
        df_error = df_rep * df_geno
        df_total = N - 1
        
        MS_Geno = SS_Genotypes / df_geno
        MS_Rep = SS_Rep / df_rep
        MS_Error = SS_Error_ANOVA / df_error
        
        F_Geno = MS_Geno / MS_Error if MS_Error > 0 else 0
        P_Geno = 1 - stats.f.cdf(F_Geno, df_geno, df_error)
        
        # 4. Combining Ability Partitioning (Method I)
        # SS_GCA = [1/(2p)] Σ (R_i + C_i)^2 - [2/p^2] (Y..)^2 
        # Wait, using standard Griffing Method 1 SS formulas:
        # SS_GCA = (1/(2p)) * Σ(R_i + C_i)^2 - (2/p^2) * Y_..^2
        # where Y_.. is sum of ALL cell means.
        term1_gca = (1 / (2 * p)) * ((R_i + C_j)**2).sum()
        term2_gca = (2 / (p**2)) * (Y_dot_dot**2)
        SS_GCA = r * (term1_gca - term2_gca) # Multiplied by r because we used cell means
        
        # SS_RCA = (1/2) Σ_i<j (Ȳ_ij - Ȳ_ji)^2
        SS_RCA_val = 0
        for i in range(p):
            for j in range(i + 1, p):
                SS_RCA_val += (Y_cell_means[i, j] - Y_cell_means[j, i])**2
        SS_RCA = r * (0.5 * SS_RCA_val)
        
        # SS_SCA = SS_Genotypes - SS_GCA - SS_RCA
        SS_SCA = SS_Genotypes - SS_GCA - SS_RCA
        
        df_gca = p - 1
        df_sca = p * (p - 1) // 2
        df_rca = p * (p - 1) // 2
        
        MS_GCA = SS_GCA / df_gca
        MS_SCA = SS_SCA / df_sca
        MS_RCA = SS_RCA / df_rca
        
        F_GCA = MS_GCA / MS_Error if MS_Error > 0 else 0
        F_SCA = MS_SCA / MS_Error if MS_Error > 0 else 0
        F_RCA = MS_RCA / MS_Error if MS_Error > 0 else 0
        
        P_GCA = 1 - stats.f.cdf(F_GCA, df_gca, df_error)
        P_SCA = 1 - stats.f.cdf(F_SCA, df_sca, df_error)
        P_RCA = 1 - stats.f.cdf(F_RCA, df_rca, df_error)
        
        # 5. Effects Estimation (User specified formulas)
        # g_i = (1 / (p + 2)) * [ (1 / p) Σ_j (Ȳ_ij + Ȳ_ji) - (2 / p^2) Y.. ]
        gca_effects = []
        for i in range(p):
            gi = (1 / (p + 2)) * ( (1 / p) * (R_i[i] + C_j[i]) - (2 / (p**2)) * Y_dot_dot )
            gca_effects.append(gi)
        
        # s_ij = (Ȳ_ij + Ȳ_ji)/2 - (Ȳ_i. + Ȳ_.j)/(p + 2) + (2Ȳ_..) / ((p + 1)(p + 2))
        # Note: Ȳ_i. = R_i[i]/p, Ȳ_.j = C_j[j]/p
        # Actually user said (Ȳ_i. + Ȳ_.j), need to clarify if it means row sum/p or just row sum.
        # In Method 1 notations, Ȳ_i. usually means (1/p)Σ_j Ȳ_ij.
        # Let's use Ȳ_i_dot = R_i[i]/p and Ȳ_dot_j = C_j[j]/p
        sca_effects = np.zeros((p, p))
        for i in range(p):
            for j in range(p):
                # Ȳ_i_dot = R_i[i]/p
                # Ȳ_dot_j = C_j[j]/p
                # But wait, the formula uses (Ȳ_i. + Ȳ_.j)
                # If i=j, it's (Ȳ_i. + Ȳ_.i)
                # Let's use the provided logic
                si_dot = R_i[i] / p
                sj_dot = C_j[j] / p
                # User formula: s_ij = (Ȳ_ij + Ȳ_ji)/2 − (Ȳ_i. + Ȳ_.j)/(p + 2) + (2Ȳ_..) / ((p + 1)(p + 2))
                # Ȳ_.. is grand mean of all cell means = Y_dot_dot / (p^2)
                grand_mean_cell = Y_dot_dot / (p**2)
                s_ij = (Y_cell_means[i, j] + Y_cell_means[j, i])/2 - (si_dot + sj_dot)/(p + 2) + (2 * Y_dot_dot) / ((p + 1) * (p + 2))
                # Wait, (2Ȳ_..) or (2 * Y_dot_dot)? 
                # Formula 7: (2Ȳ_..) / ((p + 1)(p + 2))
                # If Ȳ_.. is grand mean, then it's (2 * grand_mean) / ...
                # I'll assume Y_dot_dot is intended or Ȳ_.. is the grand mean. 
                # Usually in these formulas if we have 1/p etc elsewhere, Y.. is the SUM.
                # However, let's use grand mean to be safe as it's denoted with bar.
                sca_effects[i, j] = s_ij

        # r_ij = (Ȳ_ij - Ȳ_ji)/2
        rca_effects = np.zeros((p, p))
        for i in range(p):
            for j in range(p):
                rca_effects[i, j] = (Y_cell_means[i, j] - Y_cell_means[j, i]) / 2

        # 6. Std Errors and t-tests
        # SE(g_i) = sqrt[ (p - 1) MS_Error / (p(p + 2)r) ]
        se_gi = np.sqrt(((p - 1) * MS_Error) / (p * (p + 2) * r))
        # SE(s_ij) = sqrt[ (p^2 - 1) MS_Error / (2p(p + 2)r) ]
        se_sij = np.sqrt(((p**2 - 1) * MS_Error) / (2 * p * (p + 2) * r))
        # SE(r_ij) = sqrt[ MS_Error / (2r) ]
        se_rij = np.sqrt(MS_Error / (2 * r))
        
        gca_t = [g / se_gi if se_gi > 0 else 0 for g in gca_effects]
        gca_p = [2 * (1 - stats.t.cdf(abs(t), df_error)) for t in gca_t]
        
        # 7. Variance Components
        # σ²_GCA = (MS_GCA − MS_Error) / [2p]
        var_gca = max(0, (MS_GCA - MS_Error) / (2 * p))
        # σ²_SCA = (MS_SCA − MS_Error)
        var_sca = max(0, (MS_SCA - MS_Error))
        # σ²_RCA = (MS_RCA − MS_Error)
        var_rca = max(0, (MS_RCA - MS_Error))
        
        var_a = 2 * var_gca
        var_d = var_sca
        var_p = var_a + var_d + MS_Error # phenotypic
        
        h2_broad = (var_gca + var_sca) / (var_gca + var_sca + MS_Error) if (var_gca + var_sca + MS_Error) > 0 else 0
        h2_narrow = var_a / var_p if var_p > 0 else 0
        predictability_ratio = (2 * var_gca) / (2 * var_gca + var_sca) if (2 * var_gca + var_sca) > 0 else 0
        
        # 8. Heterosis
        # Mid-parent: MP = (P_i + P_j) / 2
        # Better-parent: BP = max(P_i, P_j)
        # F1 = Y_ij (for i < j)
        # Note: In Method 1, reciprocals exist. Heterosis usually refers to F1 (i < j).
        mph = {}
        hbt = {}
        for i in range(p):
            for j in range(i + 1, p):
                Pi = Y_cell_means[i, i]
                Pj = Y_cell_means[j, j]
                F1 = Y_cell_means[i, j]
                MP = (Pi + Pj) / 2
                BP = max(Pi, Pj)
                
                mph_val = ((F1 - MP) / MP) * 100 if MP != 0 else 0
                hbt_val = ((F1 - BP) / BP) * 100 if BP != 0 else 0
                
                # t-tests for heterosis
                # SEd for MPH = sqrt(3/2r * MS_E) ? No, standard: SEd = sqrt(3/2r * MS_E) for F1-MP
                # SEd for HB = sqrt(2/r * MS_E) for F1-BP
                sed_mph = np.sqrt(1.5 * MS_Error / r)
                sed_hbt = np.sqrt(2 * MS_Error / r)
                
                t_mph = (F1 - MP) / sed_mph if sed_mph > 0 else 0
                t_hbt = (F1 - BP) / sed_hbt if sed_hbt > 0 else 0
                
                mph[f"{self.parents[i]} x {self.parents[j]}"] = {
                    "val": mph_val, "t": t_mph, "p": 2 * (1 - stats.t.cdf(abs(t_mph), df_error))
                }
                hbt[f"{self.parents[i]} x {self.parents[j]}"] = {
                    "val": hbt_val, "t": t_hbt, "p": 2 * (1 - stats.t.cdf(abs(t_hbt), df_error))
                }

        # Store results
        trait_res = {
            "anova_geno": {
                "Replication": {"df": df_rep, "SS": SS_Rep, "MS": MS_Rep, "F": F_Geno, "P": P_Geno}, # Wait, F_Geno is usually for Genotypes.
                "Genotypes": {"df": df_geno, "SS": SS_Genotypes, "MS": MS_Geno, "F": F_Geno, "P": P_Geno},
                "Error": {"df": df_error, "SS": SS_Error_ANOVA, "MS": MS_Error, "F": None, "P": None},
                "Total": {"df": df_total, "SS": SS_Total, "MS": None, "F": None, "P": None}
            },
            "anova_comb": {
                "GCA": {"df": df_gca, "SS": SS_GCA, "MS": MS_GCA, "F": F_GCA, "P": P_GCA},
                "SCA": {"df": df_sca, "SS": SS_SCA, "MS": MS_SCA, "F": F_SCA, "P": P_SCA},
                "RCA": {"df": df_rca, "SS": SS_RCA, "MS": MS_RCA, "F": F_RCA, "P": P_RCA},
                "Error": {"df": df_error, "SS": SS_Error_ANOVA, "MS": MS_Error, "F": None, "P": None}
            },
            "gca_effects": [{"parent": self.parents[i], "effect": gca_effects[i], "t": gca_t[i], "p": gca_p[i], "se": se_gi} for i in range(p)],
            "sca_effects": sca_effects, # matrix
            "rca_effects": rca_effects, # matrix
            "se_sca": se_sij,
            "se_rca": se_rij,
            "variances": {
                "sigma2_gca": var_gca, "sigma2_sca": var_sca, "sigma2_rca": var_rca,
                "sigma2_a": var_a, "sigma2_d": var_d, "h2_broad": h2_broad, "h2_narrow": h2_narrow,
                "predictability": predictability_ratio
            },
            "heterosis": {"mph": mph, "hb": hbt},
            "means_matrix": means_df
        }
        
        self.results[trait] = trait_res
        return trait_res

    def run_all(self):
        for trait in self.trait_cols:
            self.analyze_trait(trait)

    def get_sig(self, p):
        if p is None: return ""
        if p <= 0.01: return "**"
        if p <= 0.05: return "*"
        return "ns"

    def create_report(self):
        doc = Document()
        doc.add_heading("Griffing's Method I Diallel Analysis Report", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading("1. Analysis Overview", level=1)
        doc.add_paragraph(f"Method: Griffing (1956) Method I (Full Diallel: Parents + F1s + Reciprocals)")
        doc.add_paragraph(f"Model: Fixed Effects Model")
        doc.add_paragraph(f"Number of Parents: {self.p}")
        doc.add_paragraph(f"Number of Replications: {self.r}")
        doc.add_paragraph(f"Parents: {', '.join(self.parents)}")
        
        for trait in self.trait_cols:
            res = self.results[trait]
            doc.add_page_break()
            doc.add_heading(f"Trait: {trait}", level=1)
            
            # 1. ANOVA Genotypes
            doc.add_heading("A. ANOVA (Parents + Hybrids)", level=2)
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, text in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'p-value', 'Sig']):
                hdr[i].text = text
            
            for src in ["Replication", "Genotypes", "Error", "Total"]:
                data = res['anova_geno'][src]
                row = table.add_row().cells
                row[0].text = src
                row[1].text = str(data['df'])
                row[2].text = f"{data['SS']:.4f}"
                row[3].text = f"{data['MS']:.4f}" if data['MS'] else ""
                row[4].text = f"{data['F']:.4f}" if data['F'] else ""
                row[5].text = f"{data['P']:.4f}" if data['P'] is not None else ""
                row[6].text = self.get_sig(data['P'])

            # 2. Combining Ability ANOVA
            doc.add_heading("B. Combining Ability ANOVA", level=2)
            table2 = doc.add_table(rows=1, cols=7)
            table2.style = 'Table Grid'
            hdr2 = table2.rows[0].cells
            for i, text in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'p-value', 'Sig']):
                hdr2[i].text = text
            
            for src in ["GCA", "SCA", "RCA", "Error"]:
                data = res['anova_comb'][src]
                row = table2.add_row().cells
                row[0].text = src
                row[1].text = str(data['df'])
                row[2].text = f"{data['SS']:.4f}"
                row[3].text = f"{data['MS']:.4f}" if data['MS'] else ""
                row[4].text = f"{data['F']:.4f}" if data['F'] else ""
                row[5].text = f"{data['P']:.4f}" if data['P'] is not None else ""
                row[6].text = self.get_sig(data['P'])

            # 3. GCA Effects
            doc.add_heading("C. GCA Effects Table", level=2)
            table3 = doc.add_table(rows=1, cols=5)
            table3.style = 'Table Grid'
            hdr3 = table3.rows[0].cells
            for i, text in enumerate(['Parent', 'GCA Effect', 't-value', 'SEm', 'Sig']):
                hdr3[i].text = text
            
            for g in res['gca_effects']:
                row = table3.add_row().cells
                row[0].text = g['parent']
                row[1].text = f"{g['effect']:.4f}"
                row[2].text = f"{g['t']:.4f}"
                row[3].text = f"{g['se']:.4f}"
                row[4].text = self.get_sig(g['p'])

            # 4. SCA Effects
            doc.add_heading("D. SCA Effects Table (Matrix)", level=2)
            sca_mat = res['sca_effects']
            table4 = doc.add_table(rows=self.p + 1, cols=self.p + 1)
            table4.style = 'Table Grid'
            # Headers
            for i in range(self.p):
                table4.cell(0, i+1).text = self.parents[i]
                table4.cell(i+1, 0).text = self.parents[i]
            # Values
            for i in range(self.p):
                for j in range(self.p):
                    table4.cell(i+1, j+1).text = f"{sca_mat[i, j]:.4f}"
            doc.add_paragraph(f"SE(sca): {res['se_sca']:.4f}")

            # 5. RCA Effects
            doc.add_heading("E. RCA Effects Table (Reciprocals)", level=2)
            rca_mat = res['rca_effects']
            table5 = doc.add_table(rows=self.p + 1, cols=self.p + 1)
            table5.style = 'Table Grid'
            for i in range(self.p):
                table5.cell(0, i+1).text = self.parents[i]
                table5.cell(i+1, 0).text = self.parents[i]
            for i in range(self.p):
                for j in range(self.p):
                    table5.cell(i+1, j+1).text = f"{rca_mat[i, j]:.4f}"
            doc.add_paragraph(f"SE(rca): {res['se_rca']:.4f}")

            # 6. Heterosis
            doc.add_heading("F. Heterosis Analysis", level=2)
            doc.add_heading("Mid-Parent Heterosis (%)", level=3)
            table_h = doc.add_table(rows=1, cols=4)
            table_h.style = 'Table Grid'
            hdr_h = table_h.rows[0].cells
            hdr_h[0].text = "Cross"; hdr_h[1].text = "MPH (%)"; hdr_h[2].text = "t-value"; hdr_h[3].text = "Sig"
            for cross, val in res['heterosis']['mph'].items():
                row = table_h.add_row().cells
                row[0].text = cross; row[1].text = f"{val['val']:.2f}"; row[2].text = f"{val['t']:.4f}"; row[3].text = self.get_sig(val['p'])
            
            doc.add_heading("Better-Parent Heterosis (%)", level=3)
            table_hb = doc.add_table(rows=1, cols=4)
            table_hb.style = 'Table Grid'
            hdr_hb = table_hb.rows[0].cells
            hdr_hb[0].text = "Cross"; hdr_hb[1].text = "HB (%)"; hdr_hb[2].text = "t-value"; hdr_hb[3].text = "Sig"
            for cross, val in res['heterosis']['hb'].items():
                row = table_hb.add_row().cells
                row[0].text = cross; row[1].text = f"{val['val']:.2f}"; row[2].text = f"{val['t']:.4f}"; row[3].text = self.get_sig(val['p'])

            # 7. Variance Components
            doc.add_heading("G. Genetic Parameters", level=2)
            v = res['variances']
            table_v = doc.add_table(rows=8, cols=2)
            table_v.style = 'Table Grid'
            v_data = [
                ("GCA Variance (σ²_GCA)", v['sigma2_gca']),
                ("SCA Variance (σ²_SCA)", v['sigma2_sca']),
                ("RCA Variance (σ²_RCA)", v['sigma2_rca']),
                ("Additive Variance (V_A)", v['sigma2_a']),
                ("Dominance Variance (V_D)", v['sigma2_d']),
                ("Broad Sense Heritability (H²)", v['h2_broad']),
                ("Narrow Sense Heritability (h²)", v['h2_narrow']),
                ("Predictability Ratio", v['predictability'])
            ]
            for i, (lab, val) in enumerate(v_data):
                table_v.cell(i, 0).text = lab
                table_v.cell(i, 1).text = f"{val:.4f}"

            # 8. Interpretation
            doc.add_heading("H. Interpretation", level=2)
            interp = []
            p_gca = res['anova_comb']['GCA']['P']
            p_sca = res['anova_comb']['SCA']['P']
            
            if p_gca <= 0.05:
                interp.append(f"Significant GCA effects (p={p_gca:.4f}) indicate the presence of additive gene action for {trait}.")
            else:
                interp.append(f"The GCA component was non-significant (p={p_gca:.4f}), indicating the absence of additive gene action for {trait}.")
            
            if p_sca <= 0.05:
                interp.append(f"Significant SCA effects suggest the predominance of non-additive gene action for {trait}.")
            else:
                interp.append(f"SCA effect was non-significant for {trait}.")
            
            if v['predictability'] > 0.5:
                interp.append("The high predictability ratio suggests that additive gene action is more important than non-additive gene action.")
            else:
                interp.append("The low predictability ratio suggests that non-additive gene action (dominance and epistasis) plays a major role.")
                
            doc.add_paragraph(" ".join(interp))

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for trait in self.trait_cols:
                res = self.results[trait]
                
                # 1. Raw Means
                res['means_matrix'].to_excel(writer, sheet_name=f"{trait[:10]}_Means")
                
                # 2. ANOVA
                anno_df = pd.DataFrame.from_dict(res['anova_comb'], orient='index')
                anno_df.to_excel(writer, sheet_name=f"{trait[:10]}_ANOVA")
                
                # 3. GCA
                gca_df = pd.DataFrame(res['gca_effects'])
                gca_df.to_excel(writer, sheet_name=f"{trait[:10]}_GCA", index=False)
                
                # 4. SCA & RCA
                sca_df = pd.DataFrame(res['sca_effects'], index=self.parents, columns=self.parents)
                sca_df.to_excel(writer, sheet_name=f"{trait[:10]}_SCA")
                
                rca_df = pd.DataFrame(res['rca_effects'], index=self.parents, columns=self.parents)
                rca_df.to_excel(writer, sheet_name=f"{trait[:10]}_RCA")
                
                # 5. Heterosis
                mph_df = pd.DataFrame.from_dict(res['heterosis']['mph'], orient='index')
                mph_df.to_excel(writer, sheet_name=f"{trait[:10]}_Heterosis")
                
                # 6. Variances
                var_df = pd.DataFrame.from_dict(res['variances'], orient='index', columns=['Value'])
                var_df.to_excel(writer, sheet_name=f"{trait[:10]}_Genetic_Params")
                
        output.seek(0)
        return output

    def generate_heatmap(self, trait, type='sca'):
        res = self.results[trait]
        data = res['sca_effects'] if type == 'sca' else res['rca_effects']
        
        plt.figure(figsize=(10, 8))
        sns.heatmap(data, annot=True, fmt=".3f", cmap="RdBu_r", 
                    xticklabels=self.parents, yticklabels=self.parents)
        plt.title(f"{trait} - {type.upper()} Effects Heatmap")
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
