import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import seaborn as sns

class GenotypicCorrelationAnalyzer:
    def __init__(self, df, genotype_col, rep_col, trait_cols):
        self.df = df
        self.genotype_col = genotype_col
        self.rep_col = rep_col
        self.trait_cols = trait_cols
        
        self.n = 0 # Number of genotypes
        self.r = 0 # Number of replications
        self.genotypes = []
        
        # Results storage
        self.variances = {} # trait -> {sigma2_g, sigma2_e}
        self.correlation_matrix = None
        self.p_values = None
        self.t_values = None
        self.covariance_matrix = None
        
    def validate(self):
        # Ensure column types
        self.df[self.genotype_col] = self.df[self.genotype_col].astype(str)
        self.df[self.rep_col] = self.df[self.rep_col].astype(str)
        
        for trait in self.trait_cols:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        
        # Check for missing values
        if self.df[self.trait_cols].isnull().any().any():
            raise ValueError("Some trait values are missing or non-numeric.")
            
        # Check RCBD structure
        self.genotypes = sorted(self.df[self.genotype_col].unique())
        self.n = len(self.genotypes)
        reps = sorted(self.df[self.rep_col].unique())
        self.r = len(reps)
        
        # Check if all genotypes appear in all replications
        counts = self.df.groupby([self.genotype_col, self.rep_col]).size()
        if len(counts) != self.n * self.r:
            raise ValueError("Experimental design is not a balanced RCBD. All genotypes must appear in all replications.")
        
        if (counts != 1).any():
            raise ValueError("Duplicate entries detected for some Genotype-Replication combinations.")

    def run_analysis(self):
        traits = self.trait_cols
        num_traits = len(traits)
        n = self.n
        r = self.r
        
        # 1. Component ANOVA for each trait (Variance components)
        for trait in traits:
            # Grand Mean and Sums
            y = self.df[trait].values
            G_total = y.sum()
            N = n * r
            CF = (G_total**2) / N
            
            # SS Total
            SS_Total = (y**2).sum() - CF
            
            # SS Replications
            rep_sums = self.df.groupby(self.rep_col)[trait].sum()
            SS_Rep = (rep_sums**2).sum() / n - CF
            
            # SS Genotypes
            genotype_sums = self.df.groupby(self.genotype_col)[trait].sum()
            SS_Geno = (genotype_sums**2).sum() / r - CF
            
            # SS Error
            SS_Error = SS_Total - SS_Rep - SS_Geno
            
            df_geno = n - 1
            df_error = (n - 1) * (r - 1)
            
            MS_Geno = SS_Geno / df_geno
            MS_Error = SS_Error / df_error
            
            # Genotypic Variance
            var_g = (MS_Geno - MS_Error) / r
            if var_g < 0: var_g = 0 # Handle negative estimates
            
            # Environmental Variance
            var_e = MS_Error
            
            self.variances[trait] = {
                "sigma2_g": var_g,
                "sigma2_e": var_e,
                "MS_G": MS_Geno,
                "MS_E": MS_Error
            }

        # 2. Covariance and Correlation Matrices
        rg_matrix = np.zeros((num_traits, num_traits))
        cov_g_matrix = np.zeros((num_traits, num_traits))
        p_matrix = np.zeros((num_traits, num_traits))
        t_matrix = np.zeros((num_traits, num_traits))
        
        for i in range(num_traits):
            for j in range(num_traits):
                if i == j:
                    rg_matrix[i, j] = 1.0
                    cov_g_matrix[i, j] = self.variances[traits[i]]["sigma2_g"]
                    p_matrix[i, j] = 0.0
                    t_matrix[i, j] = np.inf
                    continue
                
                traitX = traits[i]
                traitY = traits[j]
                
                # Cross Product calculations
                X = self.df[traitX].values
                Y = self.df[traitY].values
                
                G_total_X = X.sum()
                G_total_Y = Y.sum()
                N = n * r
                CF_XY = (G_total_X * G_total_Y) / N
                
                # CP Total
                CP_Total = (X * Y).sum() - CF_XY
                
                # CP Replications
                rep_sums_X = self.df.groupby(self.rep_col)[traitX].sum()
                rep_sums_Y = self.df.groupby(self.rep_col)[traitY].sum()
                CP_Rep = (rep_sums_X * rep_sums_Y).sum() / n - CF_XY
                
                # CP Genotypes
                genotype_sums_X = self.df.groupby(self.genotype_col)[traitX].sum()
                genotype_sums_Y = self.df.groupby(self.genotype_col)[traitY].sum()
                CP_Geno = (genotype_sums_X * genotype_sums_Y).sum() / r - CF_XY
                
                # CP Error
                CP_Error = CP_Total - CP_Rep - CP_Geno
                
                df_geno = n - 1
                df_error = (n - 1) * (r - 1)
                
                MCP_Geno = CP_Geno / df_geno
                MCP_Error = CP_Error / df_error
                
                # Genotypic Covariance
                cov_g = (MCP_Geno - MCP_Error) / r
                
                # Genotypic Correlation
                var_gX = self.variances[traitX]["sigma2_g"]
                var_gY = self.variances[traitY]["sigma2_g"]
                
                if var_gX > 0 and var_gY > 0:
                    rg = cov_g / np.sqrt(var_gX * var_gY)
                    # Clip to [-1, 1] due to numerical precision or near-zero variances
                    rg = np.clip(rg, -1.0, 1.0)
                else:
                    rg = 0.0
                
                # Significance Testing (t-test)
                # t = rg * sqrt(n - 2) / sqrt(1 - rg^2)
                # n is the number of GENOTYPES
                deg_f = n - 2
                if abs(rg) < 1.0:
                    t_val = rg * np.sqrt(deg_f) / np.sqrt(1 - rg**2)
                else:
                    t_val = np.inf if rg > 0 else -np.inf
                
                p_val = 2 * (1 - stats.t.cdf(abs(t_val), deg_f))
                
                rg_matrix[i, j] = rg
                cov_g_matrix[i, j] = cov_g
                t_matrix[i, j] = t_val
                p_matrix[i, j] = p_val
                
        self.correlation_matrix = pd.DataFrame(rg_matrix, index=traits, columns=traits)
        self.covariance_matrix = pd.DataFrame(cov_g_matrix, index=traits, columns=traits)
        self.t_values = pd.DataFrame(t_matrix, index=traits, columns=traits)
        self.p_values = pd.DataFrame(p_matrix, index=traits, columns=traits)

    def get_interpretation(self):
        interpretations = []
        traits = self.trait_cols
        
        for i in range(len(traits)):
            for j in range(i + 1, len(traits)):
                t1 = traits[i]
                t2 = traits[j]
                rg = self.correlation_matrix.loc[t1, t2]
                p = self.p_values.loc[t1, t2]
                
                strength = "low"
                if abs(rg) > 0.7: strength = "strong"
                elif abs(rg) > 0.4: strength = "moderate"
                
                sign = "positive" if rg > 0 else "negative"
                
                if p <= 0.01:
                    text = f"The trait {t1} exhibited a {strength} {sign} and highly significant genotypic association with {t2} (rg = {rg:.3f}**)."
                elif p <= 0.05:
                    text = f"The trait {t1} showed a {strength} {sign} and significant genotypic correlation with {t2} (rg = {rg:.3f}*)."
                else:
                    text = f"The association between {t1} and {t2} was non-significant (rg = {rg:.3f}, p > 0.05)."
                
                interpretations.append(text)
        return interpretations

    def create_report(self):
        doc = Document()
        
        # Heading
        title = doc.add_heading("Genotypic Correlation Analysis Report (RCBD)", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 1. Experimental Details
        doc.add_heading("1. Experimental Details", level=1)
        doc.add_paragraph(f"Experimental Design: Randomized Complete Block Design (RCBD)")
        doc.add_paragraph(f"Number of Genotypes: {self.n}")
        doc.add_paragraph(f"Number of Replications: {self.r}")
        doc.add_paragraph(f"Methodology: Genotypic correlations were estimated by partitioning covariance components from cross-product ANOVA.")
        
        # 2. ANOVA Summary and Genetic Variances
        doc.add_heading("2. Genetic Parameters", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Trait'
        hdr_cells[1].text = 'MS_Geno'
        hdr_cells[2].text = 'MS_Error'
        hdr_cells[3].text = 'Genotypic Var (σ²g)'
        
        for trait in self.trait_cols:
            v = self.variances[trait]
            row_cells = table.add_row().cells
            row_cells[0].text = trait
            row_cells[1].text = f"{v['MS_G']:.4f}"
            row_cells[2].text = f"{v['MS_E']:.4f}"
            row_cells[3].text = f"{v['sigma2_g']:.4f}"
            
        # 3. Genotypic Correlation Matrix
        doc.add_heading("3. Genotypic Correlation Matrix", level=1)
        num_traits = len(self.trait_cols)
        table_rg = doc.add_table(rows=num_traits + 1, cols=num_traits + 1)
        table_rg.style = 'Table Grid'
        
        # Headers
        for i, trait in enumerate(self.trait_cols):
            table_rg.cell(0, i + 1).text = trait
            table_rg.cell(i + 1, 0).text = trait
            
        # Values
        for i in range(num_traits):
            for j in range(num_traits):
                rg = self.correlation_matrix.iloc[i, j]
                p = self.p_values.iloc[i, j]
                sig = "**" if p <= 0.01 else ("*" if p <= 0.05 else "") if i != j else ""
                table_rg.cell(i + 1, j + 1).text = f"{rg:.3f}{sig}"
                
        doc.add_paragraph("Note: * Significant at 5%; ** Significant at 1%")
        
        # 4. Heatmap
        doc.add_heading("4. Heatmap Visualization", level=1)
        heatmap_buf = self.generate_heatmap()
        doc.add_picture(heatmap_buf, width=Inches(6))
        
        # 5. Interpretation
        doc.add_heading("5. Statistical Interpretation", level=1)
        interps = self.get_interpretation()
        for p_text in interps:
            doc.add_paragraph(p_text, style='List Bullet')
            
        # Footer
        doc.add_page_break()
        doc.add_paragraph("Summary of Degrees of Freedom:")
        doc.add_paragraph(f"Degrees of freedom for t-test (n - 2) = {self.n - 2}")
        doc.add_paragraph(f"Analysis performed by Research Hub Data Analysis Engine.")
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Sheet 1: Raw Data
            self.df.to_excel(writer, sheet_name='Raw Data', index=False)
            
            # Sheet 2: Genetic Variances
            var_df = pd.DataFrame.from_dict(self.variances, orient='index')
            var_df.to_excel(writer, sheet_name='Genetic Variances')
            
            # Sheet 3: Covariance Matrix
            self.covariance_matrix.to_excel(writer, sheet_name='Genotypic Covariance')
            
            # Sheet 4: Correlation Matrix
            self.correlation_matrix.to_excel(writer, sheet_name='Genotypic Correlation')
            
            # Sheet 5: Significance
            sig_data = []
            for i in range(len(self.trait_cols)):
                for j in range(i+1, len(self.trait_cols)):
                    t1 = self.trait_cols[i]
                    t2 = self.trait_cols[j]
                    sig_data.append({
                        "Trait 1": t1,
                        "Trait 2": t2,
                        "Correlation": self.correlation_matrix.loc[t1, t2],
                        "t-value": self.t_values.loc[t1, t2],
                        "p-value": self.p_values.loc[t1, t2],
                        "Significance": "Highly Significant (**)" if self.p_values.loc[t1, t2] <= 0.01 else ("Significant (*)" if self.p_values.loc[t1, t2] <= 0.05 else "")
                    })
            sig_df = pd.DataFrame(sig_data)
            sig_df.to_excel(writer, sheet_name='Significance Testing', index=False)
            
        output.seek(0)
        return output

    def generate_heatmap(self):
        plt.figure(figsize=(10, 8))
        sns.heatmap(self.correlation_matrix, annot=True, fmt=".3f", cmap="coolwarm", center=0, vmin=-1, vmax=1)
        plt.title("Genotypic Correlation Heatmap")
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
