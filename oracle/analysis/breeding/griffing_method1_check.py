import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import seaborn as sns

class GriffingMethod1CheckAnalyzer:
    def __init__(self, df, female_col, male_col, rep_col, check_col, trait_cols):
        self.df = df
        self.female_col = female_col
        self.male_col = male_col
        self.rep_col = rep_col
        self.check_col = check_col # Column that marks "Check" or "Diallel"
        self.trait_cols = trait_cols
        
        self.p = 0 # Number of diallel parents
        self.r = 0 # Number of replications
        self.parents = []
        self.checks = []
        
        # Results storage per trait
        self.results = {}

    def validate(self):
        # 1. Basic Cleaning
        self.df[self.female_col] = self.df[self.female_col].astype(str)
        self.df[self.male_col] = self.df[self.male_col].astype(str)
        self.df[self.check_col] = self.df[self.check_col].astype(str).str.strip()
        
        for trait in self.trait_cols:
            self.df[trait] = pd.to_numeric(self.df[trait], errors='coerce')
        
        if self.df[self.trait_cols].isnull().any().any():
            raise ValueError("Traits contain missing or non-numeric values.")

        # 2. Separate Diallel and Checks
        # Definition: entries where check_col is not 'Diallel' (or whatever marker we use)
        # Let's assume user provides a column where entries are marked 'Check' or 'Diallel'
        # Or we can infer: Diallel entries have both Female and Male parent specified.
        # Checks might have only one name or be marked explicitly.
        
        # We will follow the logic: 
        # Diallel entries: check_col value is 'Diallel'
        # Check entries: check_col value is anything else (the name of the check)
        
        self.diallel_df = self.df[self.df[self.check_col].str.lower() == 'diallel'].copy()
        self.checks_df = self.df[self.df[self.check_col].str.lower() != 'diallel'].copy()
        
        if self.diallel_df.empty:
            found_values = self.df[self.check_col].unique()[:5]
            raise ValueError(f"No entries with marker 'Diallel' found in column '{self.check_col}'. "
                             f"Found values in this column: {found_values}. "
                             "Please ensure your diallel crosses are marked exactly as 'Diallel'.")
        
        # Identify parents in diallel
        self.parents = sorted(list(set(self.diallel_df[self.female_col].unique()) | set(self.diallel_df[self.male_col].unique())))
        self.p = len(self.parents)
        
        # Identify check names
        self.checks = sorted(self.checks_df[self.check_col].unique().tolist())
        
        # 3. Validate Diallel Completeness (p^2 crosses)
        entry_counts = self.diallel_df.groupby([self.female_col, self.male_col]).size()
        required = []
        for p1 in self.parents:
            for p2 in self.parents:
                required.append((p1, p2))
        
        missing = [f"{e[0]} x {e[1]}" for e in required if e not in entry_counts.index]
        if missing:
            raise ValueError(f"Missing diallel crosses: {', '.join(missing[:5])}...")
            
        # 4. Replication balance
        reps = self.df.groupby([self.female_col, self.male_col, self.check_col]).size().unique()
        if len(reps) > 1:
            raise ValueError(f"Unbalanced replications: {reps}. All entries (diallel and checks) must have equal replications.")
        self.r = int(reps[0])

    def analyze_trait(self, trait):
        p = self.p
        r = self.r
        n_checks = len(self.checks)
        
        # --- 1. FULL ANOVA (Genotypes, Replications, Error) ---
        # Genotypes = p^2 diallel entries + n_checks
        y_all = self.df[trait].values
        G_sum = y_all.sum()
        N = len(y_all)
        CF = (G_sum**2) / N
        SS_Total = (y_all**2).sum() - CF
        
        rep_sums = self.df.groupby(self.rep_col)[trait].sum()
        SS_Rep = (rep_sums**2).sum() / (p*p + n_checks) - CF
        
        # Genotype Sums (combining female+male+check_col to uniquely identify entries)
        # For diallel, identity is (P1, P2). For checks, identity is check name.
        self.df['entry_id'] = self.df.apply(lambda row: f"{row[self.female_col]}x{row[self.male_col]}" if row[self.check_col].lower() == 'diallel' else row[self.check_col], axis=1)
        geno_sums = self.df.groupby('entry_id')[trait].sum()
        SS_Geno = (geno_sums**2).sum() / r - CF
        
        SS_Error = SS_Total - SS_Rep - SS_Geno
        
        df_rep = r - 1
        df_geno = (p*p + n_checks) - 1
        df_error = df_rep * df_geno
        df_total = N - 1
        
        MS_Geno = SS_Geno / df_geno
        MS_Error = SS_Error / df_error
        F_Geno = MS_Geno / MS_Error if MS_Error > 0 else 0
        P_Geno = 1 - stats.f.cdf(F_Geno, df_geno, df_error)
        
        # --- 2. PARTITIONING GENOTYPES ---
        # Diallel entries sum of squares
        diallel_sums = self.diallel_df.groupby([self.female_col, self.male_col])[trait].sum()
        diallel_G = diallel_sums.sum()
        SS_Diallel = (diallel_sums**2).sum() / r - (diallel_G**2) / (r * p * p)
        
        # Checks sum of squares
        check_sums = self.checks_df.groupby(self.check_col)[trait].sum()
        check_G = check_sums.sum()
        SS_Checks = ((check_sums**2).sum() / r - (check_G**2) / (r * n_checks)) if n_checks > 0 else 0
        
        # Diallel vs Checks
        # CF_total = (G_sum**2)/N
        # We can also get it as SS_Geno - SS_Diallel - SS_Checks if carefully calculated
        # Better: SS_D_vs_C = [ (G_diallel^2 / N_diallel) + (G_checks^2 / N_checks) ] - CF_geno_total
        N_d = r * p * p
        N_c = r * n_checks
        SS_D_vs_C = SS_Geno - SS_Diallel - SS_Checks
        
        # --- 3. COMBINING ABILITY ON DIALLEL ENTRIES ---
        # Compute cell means for diallel
        means_df = self.diallel_df.groupby([self.female_col, self.male_col])[trait].mean().unstack()
        means_df = means_df.reindex(index=self.parents, columns=self.parents)
        Y_cell_means = means_df.values
        Y_dot_dot = Y_cell_means.sum()
        R_i = Y_cell_means.sum(axis=1) # Σ_j Ȳ_ij
        C_i = Y_cell_means.sum(axis=0) # Σ_i Ȳ_ji
        
        # SS GCA = r * [ (1/2p) Σ(R_i + C_i)^2 - (2/p^2) Y..^2 ]
        term1_gca = (1 / (2 * p)) * ((R_i + C_i)**2).sum()
        term2_gca = (2 / (p**2)) * (Y_dot_dot**2)
        SS_GCA = r * (term1_gca - term2_gca)
        
        # SS RCA = r * [ (1/2) Σ_i<j (Ȳ_ij - Ȳ_ji)^2 ]
        ss_rca_raw = 0
        for i in range(p):
            for j in range(i + 1, p):
                ss_rca_raw += (Y_cell_means[i, j] - Y_cell_means[j, i])**2
        SS_RCA = r * 0.5 * ss_rca_raw
        
        SS_SCA = SS_Diallel - SS_GCA - SS_RCA
        
        df_gca = p - 1
        df_sca = p * (p - 1) // 2
        df_rca = p * (p - 1) // 2
        
        MS_GCA = SS_GCA / df_gca
        MS_SCA = SS_SCA / df_sca
        MS_RCA = SS_RCA / df_rca
        
        F_GCA = MS_GCA / MS_Error
        F_SCA = MS_SCA / MS_Error
        F_RCA = MS_RCA / MS_Error
        
        P_GCA = 1 - stats.f.cdf(F_GCA, df_gca, df_error)
        P_SCA = 1 - stats.f.cdf(F_SCA, df_sca, df_error)
        P_RCA = 1 - stats.f.cdf(F_RCA, df_rca, df_error)
        
        # --- 4. EFFECTS ESTIMATION ---
        # g_i = (1 / (p + 2)) * [ (1 / p) Σ_j (Ȳ_ij + Ȳ_ji) - (2 / p^2) Y.. ]
        gca_effects = []
        for i in range(p):
            gi = (1 / (p + 2)) * ( (1/p)*(R_i[i] + C_i[i]) - (2/(p**2))*Y_dot_dot )
            gca_effects.append(gi)
            
        sca_effects = np.zeros((p, p))
        for i in range(p):
            for j in range(p):
                si_dot = R_i[i] / p
                sj_dot = C_i[j] / p
                s_ij = (Y_cell_means[i, j] + Y_cell_means[j, i])/2 - (si_dot + sj_dot)/(p + 2) + (2 * Y_dot_dot) / ((p + 1) * (p + 2))
                sca_effects[i, j] = s_ij
                
        rca_effects = np.zeros((p, p))
        for i in range(p):
            for j in range(p):
                rca_effects[i, j] = (Y_cell_means[i, j] - Y_cell_means[j, i]) / 2

        # SEs
        se_gi = np.sqrt(((p - 1) * MS_Error) / (p * (p + 2) * r))
        se_sij = np.sqrt(((p**2 - 1) * MS_Error) / (2 * p * (p + 2) * r))
        se_rij = np.sqrt(MS_Error / (2 * r))
        
        # --- 5. HETEROSIS (Standard/Economic Heterosis over best check) ---
        check_means = self.checks_df.groupby(self.check_col)[trait].mean()
        best_check_mean = check_means.max() if not check_means.empty else None
        
        std_heterosis = {}
        if best_check_mean:
            # For each diallel cross F1 (i < j)
            for i in range(p):
                for j in range(i + 1, p):
                    f1_mean = Y_cell_means[i, j]
                    h_val = ((f1_mean - best_check_mean) / best_check_mean) * 100
                    # SE for std heterosis = sqrt(2/r * MS_E)
                    sed = np.sqrt(2 * MS_Error / r)
                    t_val = (f1_mean - best_check_mean) / sed if sed > 0 else 0
                    p_val = 2 * (1 - stats.t.cdf(abs(t_val), df_error))
                    std_heterosis[f"{self.parents[i]} x {self.parents[j]}"] = {
                        "val": h_val, "t": t_val, "p": p_val
                    }

        # --- 6. GENETIC PARAMETERS ---
        var_gca = max(0, (MS_GCA - MS_Error) / (2 * p))
        var_sca = max(0, (MS_SCA - MS_Error))
        var_rca = max(0, (MS_RCA - MS_Error))
        var_a = 2 * var_gca
        var_d = var_sca
        var_p = var_a + var_d + MS_Error
        
        h2_broad = (var_gca + var_sca) / (var_gca + var_sca + MS_Error) if (var_gca + var_sca + MS_Error) > 0 else 0
        h2_narrow = var_a / var_p if var_p > 0 else 0
        predictability = (2 * var_gca) / (2 * var_gca + var_sca) if (2 * var_gca + var_sca) > 0 else 0

        # --- 7. COMPLETE RESULT PACK ---
        res = {
            "anova_geno": {
                "Replication": {"df": df_rep, "SS": SS_Rep, "MS": SS_Rep / df_rep, "F": (SS_Rep / df_rep)/MS_Error, "P": 1 - stats.f.cdf((SS_Rep / df_rep)/MS_Error, df_rep, df_error)},
                "Genotypes": {"df": df_geno, "SS": SS_Geno, "MS": MS_Geno, "F": F_Geno, "P": P_Geno},
                "    Diallel": {"df": p*p-1, "SS": SS_Diallel, "MS": SS_Diallel/(p*p-1), "F": (SS_Diallel/(p*p-1))/MS_Error, "P": 1 - stats.f.cdf((SS_Diallel/(p*p-1))/MS_Error, p*p-1, df_error)},
                "    Checks": {"df": max(0, n_checks-1), "SS": SS_Checks, "MS": SS_Checks/max(1, n_checks-1) if n_checks > 1 else 0, "F": (SS_Checks/max(1, n_checks-1))/MS_Error if n_checks > 1 else 0, "P": 1-stats.f.cdf((SS_Checks/max(1, n_checks-1))/MS_Error, max(1, n_checks-1), df_error) if n_checks>1 else 1},
                "    D vs C": {"df": 1 if n_checks > 0 else 0, "SS": SS_D_vs_C, "MS": SS_D_vs_C, "F": SS_D_vs_C/MS_Error if n_checks > 0 else 0, "P": 1-stats.f.cdf(SS_D_vs_C/MS_Error, 1, df_error) if n_checks>0 else 1},
                "Error": {"df": df_error, "SS": SS_Error, "MS": MS_Error, "F": None, "P": None},
                "Total": {"df": df_total, "SS": SS_Total, "MS": None, "F": None, "P": None}
            },
            "anova_comb": {
                "GCA": {"df": df_gca, "SS": SS_GCA, "MS": MS_GCA, "F": F_GCA, "P": P_GCA},
                "SCA": {"df": df_sca, "SS": SS_SCA, "MS": MS_SCA, "F": F_SCA, "P": P_SCA},
                "RCA": {"df": df_rca, "SS": SS_RCA, "MS": MS_RCA, "F": F_RCA, "P": P_RCA},
                "Error": {"df": df_error, "SS": SS_Error, "MS": MS_Error, "F": None, "P": None}
            },
            "gca_effects": [{"parent": self.parents[i], "effect": gca_effects[i], "t": gca_effects[i]/se_gi, "p": 2*(1-stats.t.cdf(abs(gca_effects[i]/se_gi), df_error)), "se": se_gi} for i in range(p)],
            "sca_matrix": sca_effects.tolist(),
            "rca_matrix": rca_effects.tolist(),
            "se_sca": se_sij,
            "se_rca": se_rij,
            "variances": {
                "h2_broad": h2_broad, "h2_narrow": h2_narrow, "predictability": predictability,
                "sigma2_gca": var_gca, "sigma2_sca": var_sca, "sigma2_rca": var_rca
            },
            "std_heterosis": std_heterosis,
            "check_means": check_means.to_dict(),
            "parents": self.parents
        }
        self.results[trait] = res
        return res

    def get_sig(self, p):
        if p is None: return ""
        if p <= 0.01: return "**"
        if p <= 0.05: return "*"
        return "ns"

    def create_report(self):
        doc = Document()
        doc.add_heading("Griffing's Method I (Full Diallel WITH Checks) Report", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        for trait in self.trait_cols:
            res = self.results[trait]
            doc.add_page_break()
            doc.add_heading(f"Trait Analysis: {trait}", level=1)
            
            doc.add_heading("I. Consolidated ANOVA", level=2)
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            for i, text in enumerate(['Source', 'DF', 'SS', 'MS', 'F-Value', 'P-Value', 'Sig']):
                table.rows[0].cells[i].text = text
            
            for src, d in res['anova_geno'].items():
                row = table.add_row().cells
                row[0].text = src
                row[1].text = str(d['df'])
                row[2].text = f"{d['SS']:.4f}"
                row[3].text = f"{d['MS']:.4f}" if d['MS'] else ""
                row[4].text = f"{d['F']:.4f}" if d['F'] else ""
                row[5].text = f"{d['P']:.4f}" if d['P'] is not None else ""
                row[6].text = self.get_sig(d['P'])

            doc.add_heading("II. Combining Ability ANOVA", level=2)
            table2 = doc.add_table(rows=1, cols=7)
            table2.style = 'Table Grid'
            for i, text in enumerate(['Source', 'DF', 'SS', 'MS', 'F-Value', 'P-Value', 'Sig']):
                table2.rows[0].cells[i].text = text
            for src in ["GCA", "SCA", "RCA", "Error"]:
                d = res['anova_comb'][src]
                row = table2.add_row().cells
                row[0].text = src; row[1].text = str(d['df']); row[2].text = f"{d['SS']:.4f}"
                row[3].text = f"{d['MS']:.4f}" if d['MS'] else ""; row[4].text = f"{d['F']:.4f}" if d['F'] else ""
                row[5].text = f"{d['P']:.4f}" if d['P'] is not None else ""; row[6].text = self.get_sig(d['P'])

            # Effects ... (summarizing for space)
            doc.add_heading("III. GCA Effects", level=2)
            for g in res['gca_effects']:
                doc.add_paragraph(f"Parent {g['parent']}: {g['effect']:.4f} ({self.get_sig(g['p'])})")

            doc.add_heading("IV. Standard Heterosis over Best Check", level=2)
            if not res['std_heterosis']:
                doc.add_paragraph("No checks available for heterosis calculation.")
            else:
                table_h = doc.add_table(rows=1, cols=4)
                table_h.style = 'Table Grid'
                table_h.rows[0].cells[0].text = "Cross"; table_h.rows[0].cells[1].text = "H (%)"; table_h.rows[0].cells[2].text = "t"; table_h.rows[0].cells[3].text = "Sig"
                for cross, h in res['std_heterosis'].items():
                    r = table_h.add_row().cells
                    r[0].text = cross; r[1].text = f"{h['val']:.2f}"; r[2].text = f"{h['t']:.2f}"; r[3].text = self.get_sig(h['p'])

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for trait in self.trait_cols:
                res = self.results[trait]
                pd.DataFrame.from_dict(res['anova_geno'], orient='index').to_excel(writer, sheet_name=f"{trait[:10]}_ANOVA")
                pd.DataFrame(res['gca_effects']).to_excel(writer, sheet_name=f"{trait[:10]}_GCA")
                pd.DataFrame(res['sca_matrix'], index=self.parents, columns=self.parents).to_excel(writer, sheet_name=f"{trait[:10]}_SCA")
                pd.DataFrame.from_dict(res['variances'], orient='index').to_excel(writer, sheet_name=f"{trait[:10]}_Genetic")
        output.seek(0)
        return output
