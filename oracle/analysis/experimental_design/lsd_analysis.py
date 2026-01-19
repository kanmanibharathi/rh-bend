import pandas as pd
import numpy as np
import scipy.stats as stats
import scikit_posthocs as sp
import networkx as nx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from .duncan_util import get_duncan_q

class LSDAnalyzer:
    def __init__(self, df, row_col, col_col, treat_col, resp_col):
        self.df = df
        self.row_col = row_col
        self.col_col = col_col
        self.treat_col = treat_col
        self.resp_col = resp_col
        self.t = 0
        self.anova_table = {}
        self.p_val_treat = 1.0
        self.df_E = 0
        self.MS_E = 0
        self.SE_m = 0
        self.SE_d = 0
        self.CV = 0
        self.means = None
        self.grouping = None
        self.post_hoc_method = "None"
        self.alpha = 0.05
        
    def validate(self):
        # Convert columns to string/cat for factors
        self.df[self.row_col] = self.df[self.row_col].astype(str)
        self.df[self.col_col] = self.df[self.col_col].astype(str)
        self.df[self.treat_col] = self.df[self.treat_col].astype(str)
        
        # Ensure response is numeric
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        if self.df[self.resp_col].isnull().any():
            raise ValueError(f"Response column '{self.resp_col}' contains non-numeric or missing values.")

        rows = self.df[self.row_col].unique()
        cols = self.df[self.col_col].unique()
        treats = self.df[self.treat_col].unique()
        
        t_r = len(rows)
        t_c = len(cols)
        t_t = len(treats)
        
        if not (t_r == t_c == t_t):
             raise ValueError(f"Design Unbalanced: Rows={t_r}, Cols={t_c}, Treatments={t_t}. Must be equal (t).")
        
        self.t = t_r
        if len(self.df) != self.t ** 2:
             raise ValueError(f"Data count {len(self.df)} does not match t^2 ({self.t**2}).")

        # Check for Latin Square property (One treat per row/col)
        # Check uniqueness in Rows
        for r in rows:
            sub = self.df[self.df[self.row_col] == r]
            if len(sub[self.treat_col].unique()) != self.t:
                raise ValueError(f"Row {r} does not contain exactly one of each treatment.")

        return True

    def run_anova(self):
        y = self.df[self.resp_col]
        G = y.sum()
        CF = (G**2) / (self.t**2)
        
        # Total SS
        SS_T = (y**2).sum() - CF
        
        # Row SS
        rows_sum = self.df.groupby(self.row_col)[self.resp_col].sum()
        SS_R = (rows_sum**2).sum() / self.t - CF
        
        # Col SS
        cols_sum = self.df.groupby(self.col_col)[self.resp_col].sum()
        SS_C = (cols_sum**2).sum() / self.t - CF
        
        # Treat SS
        treats_sum = self.df.groupby(self.treat_col)[self.resp_col].sum()
        SS_Tr = (treats_sum**2).sum() / self.t - CF
        
        # Error SS
        SS_E = SS_T - (SS_R + SS_C + SS_Tr)
        
        # DF
        df_r = self.t - 1
        df_c = self.t - 1
        df_t = self.t - 1
        df_e = (self.t - 1) * (self.t - 2)
        df_tot = (self.t**2) - 1
        
        # MS
        MS_R = SS_R / df_r
        MS_C = SS_C / df_c
        MS_Tr = SS_Tr / df_t
        MS_E = SS_E / df_e
        
        # F
        F_R = MS_R / MS_E
        F_C = MS_C / MS_E
        F_Tr = MS_Tr / MS_E
        
        # P
        P_R = 1 - stats.f.cdf(F_R, df_r, df_e)
        P_C = 1 - stats.f.cdf(F_C, df_c, df_e)
        P_Tr = 1 - stats.f.cdf(F_Tr, df_t, df_e)
        
        self.anova_table = {
            "Rows": {"df": df_r, "SS": SS_R, "MS": MS_R, "F": F_R, "P": P_R},
            "Columns": {"df": df_c, "SS": SS_C, "MS": MS_C, "F": F_C, "P": P_C},
            "Treatments": {"df": df_t, "SS": SS_Tr, "MS": MS_Tr, "F": F_Tr, "P": P_Tr},
            "Error": {"df": df_e, "SS": SS_E, "MS": MS_E, "F": None, "P": None},
            "Total": {"df": df_tot, "SS": SS_T, "MS": None, "F": None, "P": None}
        }
        
        self.MS_E = MS_E
        self.df_E = df_e
        self.p_val_treat = P_Tr
        
        mean_y = y.mean()
        self.SE_m = np.sqrt(MS_E / self.t)
        self.SE_d = np.sqrt(2 * MS_E / self.t)
        self.CV = (np.sqrt(MS_E) / mean_y) * 100
        
        return self.anova_table

    def run_post_hoc(self, method='lsd', alpha=0.05, order='desc'):
        self.post_hoc_method = method
        self.alpha = alpha
        
        # Sort based on order
        is_ascending = True if order == 'asc' else False
        means = self.df.groupby(self.treat_col)[self.resp_col].mean().sort_values(ascending=is_ascending)
        
        self.means = means
        treat_labels = means.index.tolist()
        mean_values = means.values
        n = len(treat_labels)
        
        # Calculate CD (LSD value) for the selected alpha
        t_crit = stats.t.ppf(1 - alpha/2, self.df_E)
        self.CD = t_crit * self.SE_d
        
        # Initialize grouping (all 'a' if not sig)
        if self.p_val_treat > alpha:
            self.grouping = {t: 'a' for t in treat_labels}
            return self.grouping

        # Determine Significance Matrix
        significance_matrix = set()
        
        # LSD Logic
        if method.lower() == 'lsd':
            t_crit = stats.t.ppf(1 - alpha/2, self.df_E)
            lsd_val = t_crit * self.SE_d
            self.CD = lsd_val # Set CD for LSD
            for i in range(n):
                for j in range(i+1, n):
                    if abs(mean_values[i] - mean_values[j]) >= lsd_val:
                        significance_matrix.add((i, j))
        
        # Tukey Logic
        elif method.lower() == 'tukey':
            q_crit = stats.studentized_range.ppf(1-alpha, self.t, self.df_E)
            hsd_val = q_crit * self.SE_m 
            self.CD = hsd_val # Set CD for Tukey
            for i in range(n):
                for j in range(i+1, n):
                    if abs(mean_values[i] - mean_values[j]) >= hsd_val:
                        significance_matrix.add((i, j))
                        
        # Duncan Logic
        elif method.lower() == 'duncan':
             # Manual Duncan using q tables
             self.CD = get_duncan_q(2, self.df_E, alpha) * self.SE_m
             for i in range(n):
                 for j in range(i+1, n):
                     p = j - i + 1
                     # SE used for Duncan is typically SE_mean
                     q_val = get_duncan_q(p, self.df_E, alpha)
                     D_p = q_val * self.SE_m
                     
                     if abs(mean_values[i] - mean_values[j]) >= D_p:
                         significance_matrix.add((i, j))

        # Assign Letters using Clique Cover on Non-Significant Graph
        G_ns = nx.Graph()
        G_ns.add_nodes_from(range(n))
        for i in range(n):
            for j in range(i+1, n):
                if (i, j) not in significance_matrix:
                    G_ns.add_edge(i, j)
        
        # Find maximal cliques
        cliques = list(nx.find_cliques(G_ns))
        # Sort cliques to ensure 'a' goes to the highest mean's group
        # Sort key: (max_mean_in_clique, size_of_clique)
        cliques.sort(key=lambda c: (-max(c), -len(c))) # Heuristic sort

        letters_vocab = "abcdefghijklmnopqrstuvwxyz"
        grouping_letters = {i: "" for i in range(n)}
        
        # Optimization: We need to pick a set of cliques that cover all nodes/edges correctly
        # Actually, standard CLD: For each clique, assign a letter.
        # But we merge cliques if possible? No, standard is one letter per maximal clique.
        # But we only need to show letters for cliques strictly needed.
        # This is complex. Stick to assigning a letter to every maximal clique.
        
        for idx, clq in enumerate(cliques):
            if idx < len(letters_vocab):
                let = letters_vocab[idx]
                for node in clq:
                    grouping_letters[node] += let
                    
        # Map back to labels
        self.grouping = {treat_labels[i]: grouping_letters[i] for i in range(n)}
        
        # Sort letters internally (e.g. "ba" -> "ab")
        for t in self.grouping:
            self.grouping[t] = "".join(sorted(self.grouping[t]))
            
        return self.grouping

    def interpret(self):
        # Generate text
        p = self.anova_table["Treatments"]["P"]
        sig_text = "significant" if p <= self.alpha else "not significant"
        top_treat = self.means.index[0]
        
        lines = []
        lines.append(f"Analysis of Variance revealed that the effect of Treatments was {sig_text} (p = {p:.4f}).")
        if p <= self.alpha:
            lines.append(f"Post-hoc testing ({self.post_hoc_method.title()}) indicated statistical differences among means.")
            lines.append(f"Treatment '{top_treat}' recorded the highest mean value ({self.means.iloc[0]:.2f}).")
            # Find statistical parity
            # Get top group letter
            top_let = self.grouping[top_treat]
            parity = [t for t, l in self.grouping.items() if any(char in l for char in top_let) and t != top_treat]
            if parity:
                lines.append(f"It was statistically at par with: {', '.join(parity)}.")
            else:
                lines.append("It was significantly superior to all other treatments.")
        else:
            lines.append("No significant differences were observed between treatments.")
            
        return " ".join(lines)

    def create_report(self):
        doc = Document()
        
        # Title
        title = doc.add_heading('Latin Square Design Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Method
        doc.add_heading('1. Experimental Design & Model', level=1)
        p = doc.add_paragraph()
        p.add_run("The data was analyzed using a Latin Square Design (LSD) with the model: ").bold = False
        # Add equation manually or text
        p.add_run("Y_ijk = Mean + Row_i + Col_j + Treat_k + Error_ijk")
        
        # ANOVA Table
        doc.add_heading('2. ANOVA Summary', level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdrs = ['Source', 'DF', 'SS', 'MS', 'F-value', 'P-value']
        for i, h in enumerate(hdrs):
            hdr_cells[i].text = h
            
        sources = ['Rows', 'Columns', 'Treatments', 'Error', 'Total']
        for src in sources:
            row_cells = table.add_row().cells
            dat = self.anova_table[src]
            row_cells[0].text = src
            row_cells[1].text = str(dat['df'])
            row_cells[2].text = f"{dat['SS']:.4f}"
            if dat['MS']: row_cells[3].text = f"{dat['MS']:.4f}"
            if dat['F']: row_cells[4].text = f"{dat['F']:.4f}"
            if dat['P']: 
                pv = dat['P']
                sig = "*" if pv <= 0.05 else "ns"
                if pv <= 0.01: sig = "**"
                row_cells[5].text = f"{pv:.4f} {sig}"
        
        # Means Table
        doc.add_heading('3. Mean Comparison', level=1)
        table2 = doc.add_table(rows=1, cols=4)
        table2.style = 'Table Grid'
        hdr2 = table2.rows[0].cells
        hdr2[0].text = "Treatment"
        hdr2[1].text = f"Mean ({self.resp_col})"
        hdr2[2].text = "Std Deviation" # Note: We didn't calc SD per treatment, but we can
        hdr2[3].text = "Group"
        
        # Get SD per treatment
        sds = self.df.groupby(self.treat_col)[self.resp_col].std()
        
        for treat in self.means.index:
            r = table2.add_row().cells
            r[0].text = str(treat)
            r[1].text = f"{self.means[treat]:.4f}"
            r[2].text = f"{sds.get(treat, 0):.4f}"
            r[3].text = self.grouping.get(treat, "-")
            
        # Stats
        doc.add_paragraph(f"\nStandard Error of Mean (SEm): {self.SE_m:.4f}")
        doc.add_paragraph(f"Standard Error of Diff (SEd): {self.SE_d:.4f}")
        doc.add_paragraph(f"Coefficient of Variation (CV%): {self.CV:.2f}%")
        
        # Interpretation
        doc.add_heading('4. Interpretation', level=1)
        doc.add_paragraph(self.interpret())
        
        # Footer
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = "Analysis performed by Antigravity Analysis Module | Sci-Accurate Stats"
        
        # Save to buffer
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
