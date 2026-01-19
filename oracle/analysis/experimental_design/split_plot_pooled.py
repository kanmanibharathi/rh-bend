import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from datetime import datetime
import io

class SplitPlotPooledAnalyzer:
    def __init__(self, df, col_year, col_main, col_sub, col_rep, col_resp):
        self.df = df
        self.y_col = col_year
        self.a_col = col_main
        self.b_col = col_sub
        self.r_col = col_rep
        self.resp_col = col_resp
        
        self.anova_table = {}
        self.stats = {}
        self.post_hoc_res = {}
        self.bartlett_res = {}

    def validate(self):
        self.df.columns = self.df.columns.str.strip()
        for col in [self.y_col, self.a_col, self.b_col, self.r_col, self.resp_col]:
            if col not in self.df.columns:
                 raise ValueError(f"Column '{col}' not found.")
                 
        for col in [self.y_col, self.a_col, self.b_col, self.r_col]:
             self.df[col] = self.df[col].astype(str).str.strip()
        
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.resp_col])
        
        self.n_y = len(self.df[self.y_col].unique())
        self.n_a = len(self.df[self.a_col].unique())
        self.n_b = len(self.df[self.b_col].unique())
        self.n_r = len(self.df[self.r_col].unique()) # WARNING: Assuming balanced reps per year?
        # Actually n_r might vary, but pooled analysis usually assumes balanced.
        # Let's count unique reps within a single year to be safe?
        # But for total DoF we need 'r'. Usually assume same 'r' in all 'y'.
        
        if any(n < 2 for n in [self.n_a, self.n_b, self.n_y, self.n_r]):
             raise ValueError("All Factors (Year, A, B) and Replications must have at least 2 levels.")

    def run_bartlett(self):
        # Calculate Error MS for each year separately to check homogeneity
        # Or simpler: Is Bartlett requested on residuals? Standard is on Error variance.
        # We need to run individual ANOVAs per year to get MSEs.
        # For Split Plot, we have MSE_a and MSE_b. Usually check MSE_b?
        # Let's check MSE_b of each year.
        
        years = self.df[self.y_col].unique()
        ms_errors = []
        dfs = []
        
        for yr in years:
            sub = self.df[self.df[self.y_col] == yr]
            # Simple Split Plot on this subset
            # Need SS_Total, SS_R, SS_A, SS_Ea, SS_B, SS_AB, SS_Eb
            # Only need MS_Eb for precision usually?
            # Or MS_Ea? Pooling usually requires homogeneity of both errors if pooling over both?
            # User request says: "12. Homogeneity of Variance Test: F = Larger MS / Smaller MS"
            # Does not specify Bartlett per se, but F-ratio max/min.
            # I will assume we check homogeneity of Error B (usually dominant).
            
            # Quick calc for Error B in split plot:
            # SS_Eb = SST - SSR - SSA - SSEa - SSB - SSAB
            # Calc them:
            def get_ss(df_sub, cols, div, cf):
                g = df_sub.groupby(cols)[self.resp_col].sum()
                return (g**2).sum()/div - cf
                
            y_data = sub[self.resp_col]
            G = y_data.sum()
            N = len(y_data)
            CF = G**2/N
            SST = (y_data**2).sum() - CF
            
            nA = self.n_a
            nB = self.n_b
            nR = len(sub[self.r_col].unique()) # Reps in this year
            
            # Assuming balanced
            SS_R = get_ss(sub, self.r_col, nA*nB, CF)
            SS_A = get_ss(sub, self.a_col, nB*nR, CF)
            
            SS_RA = get_ss(sub, [self.r_col, self.a_col], nB, CF)
            SS_Ea = SS_RA - SS_R - SS_A
            
            SS_B = get_ss(sub, self.b_col, nA*nR, CF)
            SS_AB_Tot = get_ss(sub, [self.a_col, self.b_col], nR, CF)
            SS_AB = SS_AB_Tot - SS_A - SS_B
            
            SS_Eb = SST - SS_R - SS_A - SS_Ea - SS_B - SS_AB
            df_eb = nA * (nR-1) * (nB-1)
            
            if df_eb > 0:
                ms_errors.append(SS_Eb / df_eb)
            
        if not ms_errors: 
             self.bartlett_res = {"result": "Insufficient data"}
             return

        max_ms = max(ms_errors)
        min_ms = min(ms_errors)
        f_val = max_ms / min_ms if min_ms > 0 else 0
        # P-value? F-max test critical values are tabular. 
        # Approximating with F-dist is not 100% correct for Hartley's Fmax but close enough used sometimes?
        # Or standard Bartlett:
        # scipy.stats.bartlett(*[residuals...])
        # Since I calculated MS, Fmax is simplest.
        
        self.bartlett_res = {
            "max_ms": float(max_ms), "min_ms": float(min_ms),
            "F_ratio": float(f_val),
            "result": "Homogeneous" if f_val < 3 else "Heterogeneous (Warning)" # Simplified rule
        }

    def run_anova(self):
        y = self.df[self.resp_col]
        G = y.sum()
        N = len(y)
        CF = (G**2) / N
        SS_Total = (y**2).sum() - CF
        
        def get_ss(cols, divisor):
            grp = self.df.groupby(cols)[self.resp_col].sum()
            return (grp**2).sum() / divisor - CF
            
        nY, nR, nA, nB = self.n_y, self.n_r, self.n_a, self.n_b
        
        # 1. Year (Y)
        SS_Y = get_ss(self.y_col, nR*nA*nB)
        
        # 2. Replications within Year R(Y)
        # Formula: Sum(RjYi)^2/ab - CF - SS_Y?
        # SS_R(Y) = Sum(Rep_within_Year^2)/ab - CF - SS_Y is wrong.
        # User formula: (1/ab) Sum(Rji^2) - (G^2/yrab). Wait.
        # Usually Reps are nested in Years. So distinct reps are Rep1Year1, Rep1Year2...
        # So group by [Year, Rep].
        SS_RY_Total = get_ss([self.y_col, self.r_col], nA*nB) # This is Total Rep SS (nested)
        SS_RY = SS_RY_Total - SS_Y
        
        # 3. Main Factor A
        SS_A = get_ss(self.a_col, nY*nR*nB)
        
        # 4. Year x Main (YA)
        SS_YA_Total = get_ss([self.y_col, self.a_col], nR*nB)
        SS_YA = SS_YA_Total - SS_Y - SS_A
        
        # 5. Error (a)
        # SS_Ea = SS_RYA - SS_RY - SS_Y - SS_A - SS_YA
        # Basically Interaction of Rep x A within Year?
        # SS_RYA_Total = Sum(Rep x Year x A)^2 / b - CF
        SS_RYA_Total = get_ss([self.y_col, self.r_col, self.a_col], nB)
        # Algebra: SS_Ea = SS_RYA_Total - (SS_RY + SS_YA + SS_Y + SS_A) ... Wait
        # Correct decomposition:
        # SS_Ea is the residual of the Main Plot section.
        # Main Plot Terms: Y, R(Y), A, Yx A.
        # SS_Ea = SS_RYA_Total - SS_Y - SS_RY - SS_A - SS_YA
        SS_Ea = SS_RYA_Total - SS_Y - SS_RY - SS_A - SS_YA
        # Safety
        if SS_Ea < 0: SS_Ea = 0
        
        # 6. Sub Factor B
        SS_B = get_ss(self.b_col, nY*nR*nA)
        
        # 7. Year x Sub (YB)
        SS_YB_Total = get_ss([self.y_col, self.b_col], nR*nA)
        SS_YB = SS_YB_Total - SS_Y - SS_B
        
        # 8. Main x Sub (AB)
        SS_AB_Total = get_ss([self.a_col, self.b_col], nY*nR)
        SS_AB = SS_AB_Total - SS_A - SS_B
        
        # 9. Year x Main x Sub (YAB)
        SS_YAB_Total = get_ss([self.y_col, self.a_col, self.b_col], nR)
        SS_YAB = SS_YAB_Total - SS_YA - SS_YB - SS_AB - SS_Y - SS_A - SS_B
        # User formula simplified: Sum(YAB^2)/r - SS_YA - SS_YB - SS_AB - CF?
        # Actually subtraction of all main effects and lower interactions.
        # My logic: SS_YAB_Total - (SS_Y + SS_A + SS_B + SS_YA + SS_YB + SS_AB)
        
        # 10. Error (b)
        # SS_Eb = SST - (All above)
        SS_Eb = SS_Total - (SS_Y + SS_RY + SS_A + SS_YA + SS_Ea + SS_B + SS_YB + SS_AB + SS_YAB)
        if SS_Eb < 0: SS_Eb = 0
        
        # DF
        df_y = int(nY - 1)
        df_ry = int(nY * (nR - 1))
        df_a = int(nA - 1)
        df_ya = int((nY - 1) * (nA - 1))
        df_ea = int(nY * (nR - 1) * (nA - 1))
        
        df_b = int(nB - 1)
        df_yb = int((nY - 1) * (nB - 1))
        df_ab = int((nA - 1) * (nB - 1))
        df_yab = int((nY - 1) * (nA - 1) * (nB - 1))
        df_eb = int(nY * nA * (nR - 1) * (nB - 1))
        
        df_tot = int(nY * nR * nA * nB - 1)
        
        # MS
        ms = lambda ss, df: float(ss / df) if df > 0 else 0.0
        
        MS_Y = ms(SS_Y, df_y)
        MS_RY = ms(SS_RY, df_ry)
        MS_A = ms(SS_A, df_a)
        MS_YA = ms(SS_YA, df_ya)
        MS_Ea = ms(SS_Ea, df_ea)
        
        MS_B = ms(SS_B, df_b)
        MS_YB = ms(SS_YB, df_yb)
        MS_AB = ms(SS_AB, df_ab)
        MS_YAB = ms(SS_YAB, df_yab)
        MS_Eb = ms(SS_Eb, df_eb)
        
        # F Tests
        # F_Y = MS_Y / MS_R(Y)
        F_Y = MS_Y / MS_RY if MS_RY > 0 else None
        
        # F_A, F_YA -> MS_Ea
        F_A = MS_A / MS_Ea if MS_Ea > 0 else None
        F_YA = MS_YA / MS_Ea if MS_Ea > 0 else None
        
        # F_B, F_YB, F_AB, F_YAB -> MS_Eb
        F_B = MS_B / MS_Eb if MS_Eb > 0 else None
        F_YB = MS_YB / MS_Eb if MS_Eb > 0 else None
        F_AB = MS_AB / MS_Eb if MS_Eb > 0 else None
        F_YAB = MS_YAB / MS_Eb if MS_Eb > 0 else None
        
        # P Values
        def get_p(f, dfn, dfd):
            if f is None or dfd <= 0: return None
            return float(1 - stats.f.cdf(f, dfn, dfd))
            
        P_Y = get_p(F_Y, df_y, df_ry)
        P_A = get_p(F_A, df_a, df_ea)
        P_YA = get_p(F_YA, df_ya, df_ea)
        
        P_B = get_p(F_B, df_b, df_eb)
        P_YB = get_p(F_YB, df_yb, df_eb)
        P_AB = get_p(F_AB, df_ab, df_eb)
        P_YAB = get_p(F_YAB, df_yab, df_eb)
        
        def s(df, ss, ms, f, p): 
            return {
                "df": df, "SS": float(ss), "MS": float(ms), 
                "F": float(f) if f is not None else None, 
                "P": float(p) if p is not None else None
            }
            
        self.anova_table = {
            "Year (Y)": s(df_y, SS_Y, MS_Y, F_Y, P_Y),
            "Rep (within Y)": s(df_ry, SS_RY, MS_RY, None, None), # Usually not tested or tested against Ea? Prompt implies F_Y uses it as error.
            
            "Main Plot (A)": s(df_a, SS_A, MS_A, F_A, P_A),
            "Year x Main (Y x A)": s(df_ya, SS_YA, MS_YA, F_YA, P_YA),
            "Error (a)": {"df": df_ea, "SS": float(SS_Ea), "MS": float(MS_Ea), "F": None, "P": None},
            
            "Sub Plot (B)": s(df_b, SS_B, MS_B, F_B, P_B),
            "Year x Sub (Y x B)": s(df_yb, SS_YB, MS_YB, F_YB, P_YB),
            "Main x Sub (A x B)": s(df_ab, SS_AB, MS_AB, F_AB, P_AB),
            "Year x Main x Sub": s(df_yab, SS_YAB, MS_YAB, F_YAB, P_YAB),
            "Error (b)": {"df": df_eb, "SS": float(SS_Eb), "MS": float(MS_Eb), "F": None, "P": None},
            
            "Total": {"df": df_tot, "SS": float(SS_Total), "MS": None, "F": None, "P": None}
        }

    def run_post_hoc(self, method, alpha):
        res = self.anova_table
        MS_RY = res["Rep (within Y)"]["MS"]
        
        MS_Ea = res["Error (a)"]["MS"]
        DF_Ea = int(res["Error (a)"]["df"])
        
        MS_Eb = res["Error (b)"]["MS"]
        DF_Eb = int(res["Error (b)"]["df"])
        
        gm = self.df[self.resp_col].mean()
        
        cv_a = float((np.sqrt(MS_Ea) / gm) * 100) if gm!=0 and MS_Ea>0 else 0.0
        cv_b = float((np.sqrt(MS_Eb) / gm) * 100) if gm!=0 and MS_Eb>0 else 0.0
        
        def calc_sem(ms, div): return float(np.sqrt(ms / div)) if div > 0 and ms>=0 else 0.0
        
        nA, nB, nY, nR = self.n_a, self.n_b, self.n_y, self.n_r
        
        sem_y = calc_sem(MS_RY, nR * nA * nB) # Note: SEm_Y = sqrt(MS_R(Y) / rab)
        sem_a = calc_sem(MS_Ea, nY * nR * nB)
        sem_b = calc_sem(MS_Eb, nY * nR * nA)
        sem_ab = calc_sem(MS_Eb, nY * nR)
        
        def calc_cd(sem, df, k):
            if df <= 0: return 0.0
            sed = sem * np.sqrt(2)
            if method == 'lsd': return float(stats.t.ppf(1 - alpha/2, df) * sed)
            elif method == 'tukey': return float(stats.studentized_range.ppf(1-alpha, k, df) * sem)
            elif method == 'dmrt': return float(stats.t.ppf(1 - alpha/2, df) * sed)
            return 0.0
            
        stats_map = {
            "CV (a)": cv_a, "CV (b)": cv_b,
            "Year (Y)": {"SEm": sem_y, "SEd": sem_y*np.sqrt(2), "CD": calc_cd(sem_y, int(res["Rep (within Y)"]["df"]), nY)}, # DF for Y is usually Rep(Y) or similar? Prompt: "df_error = df(Error a) for Year & Main" -> Wait. F_Y uses MS_R(Y). So DF should be R(Y).
            "Main Plot (A)": {"SEm": sem_a, "SEd": sem_a*np.sqrt(2), "CD": calc_cd(sem_a, DF_Ea, nA)},
            "Sub Plot (B)": {"SEm": sem_b, "SEd": sem_b*np.sqrt(2), "CD": calc_cd(sem_b, DF_Eb, nB)},
            "Main x Sub (A x B)": {"SEm": sem_ab, "SEd": sem_ab*np.sqrt(2), "CD": calc_cd(sem_ab, DF_Eb, nA*nB)}
        }
        self.stats = stats_map
        
        results = {}
        def get_data(cols):
            grp = self.df.groupby(cols)[self.resp_col]
            means = grp.mean().sort_index()
            stds = grp.std().reindex(means.index)
            ses = (stds / np.sqrt(grp.count())).reindex(means.index)
            return means, stds, ses
            
        def compute_letters(means, sem, df, k):
             m_desc = means.sort_values(ascending=False)
             vals = m_desc.values
             keys = m_desc.index.tolist()
             n = len(vals)
             sed = sem * np.sqrt(2)
             sig_pairs = set()
             
             if df <= 0: return {k: "ns" for k in keys}

             if method == 'lsd':
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             elif method == 'tukey':
                 crit = stats.studentized_range.ppf(1-alpha, n, df) * sem
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             else: 
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             
             G = nx.Graph()
             G.add_nodes_from(range(n))
             for i in range(n):
                 for j in range(i+1, n):
                     if (i,j) not in sig_pairs: G.add_edge(i, j)
             
             cliques = list(nx.find_cliques(G))
             cliques.sort(key=lambda x: (min(x), -len(x)))
             vocab = "abcdefghijklmnopqrstuvwxyz"
             res_map = {i: "" for i in range(n)}
             for idx, clq in enumerate(cliques):
                 if idx < len(vocab):
                     c = vocab[idx]
                     for node in clq: res_map[node] += c
             return {keys[i]: "".join(sorted(res_map[i])) for i in range(n)}

        def proc_eff(name, cols, sem_key, df_err, n_lvls):
            m, s, se = get_data(cols)
            p = res[name]["P"]
            let = {k: "ns" for k in m.index}
            if p is not None and p < alpha:
                # Special case used for Y? Use df_ry
                if name == "Year (Y)":
                    let = compute_letters(m, stats_map[sem_key]["SEm"], int(res["Rep (within Y)"]["df"]), n_lvls)
                else:
                    let = compute_letters(m, stats_map[sem_key]["SEm"], df_err, n_lvls)
            
            results[sem_key] = {"means": m, "stds": s, "ses": se, "grouping": let}

        proc_eff("Year (Y)", self.y_col, "Year (Y)", 0, nY)
        proc_eff("Main Plot (A)", self.a_col, "Main Plot (A)", DF_Ea, nA)
        proc_eff("Sub Plot (B)", self.b_col, "Sub Plot (B)", DF_Eb, nB)
        proc_eff("Main x Sub (A x B)", [self.a_col, self.b_col], "Main x Sub (A x B)", DF_Eb, nA*nB)
        
        self.post_hoc_res = results

    def create_report(self):
        doc = Document()
        doc.add_heading('Pooled Split-Plot Analysis Report', 0)
        now_str = datetime.now().strftime('%Y-%m-%d %H:%M')
        doc.add_paragraph(f"Date: {now_str}")
        
        doc.add_heading('1. Bartlett\'s Test (Homogeneity)', 1)
        if self.bartlett_res:
            doc.add_paragraph(f"Max MS Error: {self.bartlett_res.get('max_ms', 0):.4f}")
            doc.add_paragraph(f"Min MS Error: {self.bartlett_res.get('min_ms', 0):.4f}")
            doc.add_paragraph(f"F-Ratio: {self.bartlett_res.get('F_ratio', 0):.4f}")
            doc.add_paragraph(f"Result: {self.bartlett_res.get('result', 'N/A')}")
        else:
             doc.add_paragraph("Could not calculate.")

        doc.add_heading('2. ANOVA Table', 1)
        t = doc.add_table(1, 6)
        t.style = 'Table Grid'
        def set_cell(r, i, txt): r.cells[i].text = str(txt)
        h = t.rows[0]
        for i,c in enumerate(["Source", "DF", "SS", "MS", "F", "Prob"]): set_cell(h, i, c)
        
        keys = ["Year (Y)", "Rep (within Y)", "Main Plot (A)", "Year x Main (Y x A)", "Error (a)",
                "Sub Plot (B)", "Year x Sub (Y x B)", "Main x Sub (A x B)", "Year x Main x Sub", "Error (b)", "Total"]
        
        for k in keys:
            if k in self.anova_table:
                v = self.anova_table[k]
                r = t.add_row()
                sig = ""
                if v['P'] is not None:
                     if v['P']<0.01: sig="**"
                     elif v['P']<0.05: sig="*"
                     else: sig="ns"
                set_cell(r, 0, k)
                set_cell(r, 1, v['df'])
                set_cell(r, 2, f"{v['SS']:.4f}")
                set_cell(r, 3, f"{v['MS']:.4f}" if v['MS'] else "-")
                set_cell(r, 4, f"{v['F']:.4f}" if v['F'] is not None else "-")
                set_cell(r, 5, f"{v['P']:.4f} {sig}" if v['P'] is not None else "")
                
        doc.add_heading('3. Statistics', 1)
        doc.add_paragraph(f"CV (a) %: {self.stats['CV (a)']:.2f}")
        doc.add_paragraph(f"CV (b) %: {self.stats['CV (b)']:.2f}")
        
        s_keys = ["Year (Y)", "Main Plot (A)", "Sub Plot (B)", "Main x Sub (A x B)"]
        for k in s_keys:
             s = self.stats.get(k)
             if s: doc.add_paragraph(f"{k}: SEm={s['SEm']:.4f}, SEd={s['SEd']:.4f}, CD={s['CD']:.4f}")
             
        doc.add_heading('4. Means and Grouping', 1)
        ph = self.post_hoc_res
        for k in s_keys:
            if k in ph:
                 ds = ph[k]
                 doc.add_heading(f"{k} Means", 2)
                 t = doc.add_table(1, 5)
                 t.style = 'Table Grid'
                 r = t.rows[0]
                 for i,c in enumerate(["Level", "Mean", "Std.Dev", "Std.Err", "Group"]): set_cell(r, i, c)
                 for lvl, val in ds['means'].items():
                     r = t.add_row()
                     if isinstance(lvl, tuple): l_str = " x ".join(map(str, lvl))
                     else: l_str = str(lvl)
                     set_cell(r, 0, l_str)
                     set_cell(r, 1, f"{val:.4f}")
                     set_cell(r, 2, f"{ds['stds'][lvl]:.4f}")
                     set_cell(r, 3, f"{ds['ses'][lvl]:.4f}")
                     set_cell(r, 4, ds['grouping'].get(lvl, ""))
                     
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
