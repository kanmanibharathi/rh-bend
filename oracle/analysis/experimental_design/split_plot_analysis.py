import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
import io
from datetime import datetime
from .duncan_util import get_duncan_q

class SplitPlotAnalyzer:
    def __init__(self, df, main_col, sub_col, rep_col, resp_col):
        self.df = df
        self.a_col = main_col
        self.b_col = sub_col
        self.r_col = rep_col
        self.resp_col = resp_col
        
        self.anova_table = {}
        self.stats = {}
        self.post_hoc_res = {}
        
    def validate(self):
        # Convert Cols
        for col in [self.a_col, self.b_col, self.r_col]:
             self.df[col] = self.df[col].astype(str).str.strip()
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.resp_col])
        
        self.a_levels = sorted(self.df[self.a_col].unique())
        self.b_levels = sorted(self.df[self.b_col].unique())
        self.r_levels = sorted(self.df[self.r_col].unique())
        
        self.n_a = len(self.a_levels)
        self.n_b = len(self.b_levels)
        self.n_r = len(self.r_levels)
        
        if any(n < 2 for n in [self.n_a, self.n_b, self.n_r]):
             raise ValueError("Main Plot, Sub Plot and Replications must have at least 2 levels.")

    def run_anova(self):
        y = self.df[self.resp_col]
        G = y.sum()
        N = len(y)
        CF = (G**2) / N
        SS_Total = (y**2).sum() - CF
        
        def get_ss(cols, divisor):
            grp = self.df.groupby(cols)[self.resp_col].sum()
            return (grp**2).sum() / divisor - CF
            
        nA, nB, nR = self.n_a, self.n_b, self.n_r
        
        # 1. Replication
        SS_R = get_ss(self.r_col, nA*nB)
        
        # 2. Main Plot (A)
        SS_A = get_ss(self.a_col, nB*nR)
        
        # 3. Error (a) - Main Plot Error
        # Formula: Sum(RiAj)^2/b - SS_R - SS_A - CF
        # wait, my helper subtacts CF. So SS_RA_Total = get_ss([rep, a], b).
        # SS_Err_A = SS_RA_Total - SS_R - SS_A
        SS_RA_Total = get_ss([self.r_col, self.a_col], nB)
        SS_Err_A = SS_RA_Total - SS_R - SS_A
        
        # 4. Sub Plot (B)
        SS_B = get_ss(self.b_col, nA*nR)
        
        # 5. Interaction (AxB)
        SS_AB_Total = get_ss([self.a_col, self.b_col], nR)
        SS_AxB = SS_AB_Total - SS_A - SS_B
        
        # 6. Error (b) - Sub Plot Error
        # Formula: SS_Total - SS_R - SS_A - SS_Err_A - SS_B - SS_AxB
        SS_Err_B = SS_Total - SS_R - SS_A - SS_Err_A - SS_B - SS_AxB
        
        # DF
        df_r = nR - 1
        df_a = nA - 1
        df_err_a = (nR - 1) * (nA - 1)
        
        df_b = nB - 1
        df_axb = (nA - 1) * (nB - 1)
        df_err_b = nA * (nR - 1) * (nB - 1)
        
        df_tot = nA * nB * nR - 1
        
        # MS
        MS_R = SS_R / df_r
        MS_A = SS_A / df_a
        MS_Err_A = SS_Err_A / df_err_a
        
        MS_B = SS_B / df_b
        MS_AxB = SS_AxB / df_axb
        MS_Err_B = SS_Err_B / df_err_b
        
        # F stats
        # Rep, A use Err_A
        F_R = MS_R / MS_Err_A
        F_A = MS_A / MS_Err_A
        
        # B, AxB use Err_B
        F_B = MS_B / MS_Err_B
        F_AxB = MS_AxB / MS_Err_B
        
        # P values
        P_R = 1 - stats.f.cdf(F_R, df_r, df_err_a)
        P_A = 1 - stats.f.cdf(F_A, df_a, df_err_a)
        
        P_B = 1 - stats.f.cdf(F_B, df_b, df_err_b)
        P_AxB = 1 - stats.f.cdf(F_AxB, df_axb, df_err_b)
        
        def ser(df, ss, ms, f, p, sig_err=None):
            return {"df": df, "SS": ss, "MS": ms, "F": f, "P": p}
            
        self.anova_table = {
            "Replication": ser(df_r, SS_R, MS_R, F_R, P_R),
            "Main Plot (A)": ser(df_a, SS_A, MS_A, F_A, P_A),
            "Error (a)": {"df": df_err_a, "SS": SS_Err_A, "MS": MS_Err_A, "F": None, "P": None},
            "Sub Plot (B)": ser(df_b, SS_B, MS_B, F_B, P_B),
            "Interaction A x B": ser(df_axb, SS_AxB, MS_AxB, F_AxB, P_AxB),
            "Error (b)": {"df": df_err_b, "SS": SS_Err_B, "MS": MS_Err_B, "F": None, "P": None},
            "Total": {"df": df_tot, "SS": SS_Total, "MS": None, "F": None, "P": None}
        }

    def run_post_hoc(self, method, alpha):
        res = self.anova_table
        MS_Ea = res["Error (a)"]["MS"]
        DF_Ea = int(res["Error (a)"]["df"])
        
        MS_Eb = res["Error (b)"]["MS"]
        DF_Eb = int(res["Error (b)"]["df"])
        
        gm = self.df[self.resp_col].mean()
        
        # CVs
        cv_a = (np.sqrt(MS_Ea) / gm) * 100
        cv_b = (np.sqrt(MS_Eb) / gm) * 100
        
        # SEm Func
        def calc_sem(ms, div): return np.sqrt(ms / div)
        
        # Main A -> Err A
        sem_a = calc_sem(MS_Ea, self.n_r * self.n_b)
        sed_a = sem_a * np.sqrt(2)
        
        # Sub B -> Err B
        sem_b = calc_sem(MS_Eb, self.n_r * self.n_a)
        sed_b = sem_b * np.sqrt(2)
        
        # AxB -> Err B (Assumed valid for simple split plot interaction comparison)
        # Note: Interaction CD is complex if comparing A levels at same B vs B at same A.
        # User spec: SEm_AxB = sqrt(MS_Error(b) / r)
        sem_axb = calc_sem(MS_Eb, self.n_r)
        sed_axb = sem_axb * np.sqrt(2)
        
        # CD Func
        def calc_cd(sem, sed, df, k):
            if method == 'lsd': return stats.t.ppf(1 - alpha/2, df) * sed
            elif method == 'tukey': return stats.studentized_range.ppf(1-alpha, k, df) * sem
            elif method == 'dmrt': return stats.t.ppf(1 - alpha/2, df) * sed # Simplfied
            return 0
            
        cd_a = calc_cd(sem_a, sed_a, DF_Ea, self.n_a)
        cd_b = calc_cd(sem_b, sed_b, DF_Eb, self.n_b)
        cd_axb = calc_cd(sem_axb, sed_axb, DF_Eb, self.n_a * self.n_b) # using DF_Eb for interaction
        
        self.stats = {
            "CV (a)": cv_a,
            "CV (b)": cv_b,
            "Main Plot (A)": {"SEm": sem_a, "SEd": sed_a, "CD": cd_a},
            "Sub Plot (B)": {"SEm": sem_b, "SEd": sed_b, "CD": cd_b},
            "Interaction A x B": {"SEm": sem_axb, "SEd": sed_axb, "CD": cd_axb}
        }
        
        # Grouping
        # Helper: Get Table Data
        def get_data(cols, sort_idx=True):
             grp = self.df.groupby(cols)[self.resp_col]
             means = grp.mean()
             if sort_idx: means = means.sort_index()
             stds = grp.std().reindex(means.index)
             ses = (stds / np.sqrt(grp.count())).reindex(means.index)
             return means, stds, ses
             
        # Helper: Compute Letters
        def compute_letters(means, sem, sed, df, k):
             m_desc = means.sort_values(ascending=False)
             vals = m_desc.values
             keys = m_desc.index.tolist()
             n = len(vals)
             sig_pairs = set()
             
             if method == 'lsd':
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             elif method == 'tukey':
                 crit = stats.studentized_range.ppf(1-alpha, n, df) * sem
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             elif method == 'dmrt':
                 for i in range(n):
                     for j in range(i+1, n):
                         p = j - i + 1
                         rng = get_duncan_q(p, df, alpha) * sem
                         if abs(vals[i]-vals[j]) >= rng: sig_pairs.add((i,j))
                         
             G = nx.Graph()
             G.add_nodes_from(range(n))
             for i in range(n):
                 for j in range(i+1, n):
                     if (i,j) not in sig_pairs: G.add_edge(i, j)
             
             cliques = list(nx.find_cliques(G))
             cliques.sort(key=lambda x: (min(x), -len(x)))
             vocab = "abcdefghijklmnopqrstuvwxyz"
             res_map = {i: "" for i in range(n)}
             for idx, clq in enumerate(cliques):
                 if idx < len(vocab):
                     c = vocab[idx]
                     for node in clq: res_map[node] += c
             
             return {keys[i]: "".join(sorted(res_map[i])) for i in range(n)}
             
        results = {}
        
        # User Preference: "Always include" for previous tools. 
        # But here user stated "MANDATORY ... exact hierarchy".
        # However, typically users want to see the table even if grouping is "ns".
        # So I will generate the tables, but suppress LETTERS if not significant.
        # This matches the "Two Factor" updated logic.
        
        # Main A
        ma, sa, sea = get_data(self.a_col)
        let_a = {k: "ns" for k in ma.index}
        if res["Main Plot (A)"]["P"] < alpha:
             let_a = compute_letters(ma, sem_a, sed_a, DF_Ea, self.n_a)
        results["Main Plot (A)"] = {"means": ma, "stds": sa, "ses": sea, "grouping": let_a}
        
        # Sub B
        mb, sb, seb = get_data(self.b_col)
        let_b = {k: "ns" for k in mb.index}
        if res["Sub Plot (B)"]["P"] < alpha:
             let_b = compute_letters(mb, sem_b, sed_b, DF_Eb, self.n_b)
        results["Sub Plot (B)"] = {"means": mb, "stds": sb, "ses": seb, "grouping": let_b}
        
        # Interaction AxB
        mab, sab, seab = get_data([self.a_col, self.b_col])
        let_ab = {k: "ns" for k in mab.index}
        if res["Interaction A x B"]["P"] < alpha:
             let_ab = compute_letters(mab, sem_axb, sed_axb, DF_Eb, self.n_a*self.n_b)
        results["Interaction A x B"] = {"means": mab, "stds": sab, "ses": seab, "grouping": let_ab}
        
        self.post_hoc_res = results

    def create_report(self):
        doc = Document()
        doc.add_heading('Simple Split Plot Analysis Report', 0)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # 1. ANOVA
        doc.add_heading('1. ANOVA Table', 1)
        t = doc.add_table(1, 6)
        t.style = 'Table Grid'
        def set_cell(r, i, txt): r.cells[i].text = str(txt)
        
        h = t.rows[0]
        for i, c in enumerate(["Source", "DF", "SS", "MS", "F", "Prob"]): set_cell(h, i, c)
        
        keys = ["Replication", "Main Plot (A)", "Error (a)", "Sub Plot (B)", "Interaction A x B", "Error (b)", "Total"]
        for k in keys:
             if k in self.anova_table:
                 v = self.anova_table[k]
                 r = t.add_row()
                 sig = ""
                 if v['P'] is not None:
                     if v['P']<0.01: sig="**"
                     elif v['P']<0.05: sig="*"
                     else: sig="ns"
                 set_cell(r, 0, k)
                 set_cell(r, 1, v['df'])
                 set_cell(r, 2, f"{v['SS']:.4f}")
                 set_cell(r, 3, f"{v['MS']:.4f}" if v['MS'] else "-")
                 set_cell(r, 4, f"{v['F']:.4f}" if v['F'] else "-")
                 set_cell(r, 5, f"{v['P']:.4f} {sig}" if v['P'] is not None else "")

        # 2. Stats
        doc.add_heading('2. Statistics', 1)
        doc.add_paragraph(f"CV (a) %: {self.stats['CV (a)']:.2f}")
        doc.add_paragraph(f"CV (b) %: {self.stats['CV (b)']:.2f}")
        
        for k in ["Main Plot (A)", "Sub Plot (B)", "Interaction A x B"]:
            s = self.stats[k]
            doc.add_paragraph(f"{k}: SEm={s['SEm']:.4f}, SEd={s['SEd']:.4f}, CD={s['CD']:.4f}")

        # 3. Means
        doc.add_heading('3. Means and Grouping', 1)
        ph = self.post_hoc_res
        
        def add_table(name, ds):
            doc.add_heading(name, 2)
            t = doc.add_table(1, 5)
            t.style = 'Table Grid'
            r = t.rows[0]
            for i, c in enumerate(["Level", "Mean", "Std.Dev", "Std.Err", "Group"]): set_cell(r, i, c)
            
            for lvl, val in ds['means'].items():
                r = t.add_row()
                if isinstance(lvl, tuple): l_str = " x ".join(map(str, lvl))
                else: l_str = str(lvl)
                set_cell(r, 0, l_str)
                set_cell(r, 1, f"{val:.4f}")
                set_cell(r, 2, f"{ds['stds'][lvl]:.4f}")
                set_cell(r, 3, f"{ds['ses'][lvl]:.4f}")
                set_cell(r, 4, ds['grouping'].get(lvl, ""))
        
        for k in ["Main Plot (A)", "Sub Plot (B)", "Interaction A x B"]:
            add_table(k, ph[k])
            
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
