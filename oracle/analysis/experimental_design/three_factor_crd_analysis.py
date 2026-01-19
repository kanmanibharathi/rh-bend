import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from .duncan_util import get_duncan_q

class ThreeFactorCRDAnalyzer:
    def __init__(self, df, a_col, b_col, c_col, resp_col, rep_col=None):
        self.df = df
        self.a_col = a_col
        self.b_col = b_col
        self.c_col = c_col
        self.resp_col = resp_col
        self.rep_col = rep_col
        
        self.a = 0
        self.b = 0
        self.c = 0
        self.n = 0
        
        self.anova_table = {}
        self.results = {}
        
        self.MS_E = 0
        self.df_E = 0
        self.grand_mean = 0
        self.r_bar = 0
        self.alpha = 0.05

    def validate(self):
        # Type conversion
        self.df[self.a_col] = self.df[self.a_col].astype(str).str.strip()
        self.df[self.b_col] = self.df[self.b_col].astype(str).str.strip()
        self.df[self.c_col] = self.df[self.c_col].astype(str).str.strip()
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        
        if self.rep_col and self.rep_col in self.df.columns:
            self.df[self.rep_col] = self.df[self.rep_col].astype(str).str.strip()
        
        # Drop NaNs
        initial_len = len(self.df)
        self.df = self.df.dropna(subset=[self.resp_col])
        if len(self.df) < initial_len:
             print(f"Warning: Dropped {initial_len - len(self.df)} rows due to missing values.")

        if self.df.empty:
             raise ValueError("No valid data remaining after removing missing values.")
             
        self.n = len(self.df)
        self.a = self.df[self.a_col].nunique()
        self.b = self.df[self.b_col].nunique()
        self.c = self.df[self.c_col].nunique()
        
        # Check DF
        if self.n <= self.a * self.b * self.c:
            raise ValueError(f"Insufficient degrees of freedom. Observations ({self.n}) must exceed treatments ({self.a * self.b * self.c}).")

    def run_anova(self):
        """
        Perform three-factor ANOVA using direct formulas from specification.
        
        Formulas:
        SS_A = b*c*r * Σ(Ȳ_i... − Ȳ....)²
        SS_B = a*c*r * Σ(Ȳ_.j.. − Ȳ....)²
        SS_C = a*b*r * Σ(Ȳ_..k. − Ȳ....)²
        SS_AB = c*r * ΣΣ(Ȳ_ij.. − Ȳ_i... − Ȳ_.j.. + Ȳ....)²
        SS_AC = b*r * ΣΣ(Ȳ_i.k. − Ȳ_i... − Ȳ_..k. + Ȳ....)²
        SS_BC = a*r * ΣΣ(Ȳ_.jk. − Ȳ_.j.. − Ȳ_..k. + Ȳ....)²
        SS_ABC = r * ΣΣΣ(Ȳ_ijk. − Ȳ_ij.. − Ȳ_i.k. − Ȳ_.jk. + Ȳ_i... + Ȳ_.j.. + Ȳ_..k. − Ȳ....)²
        """
        y = self.df[self.resp_col]
        self.grand_mean = y.mean()
        
        # Total SS
        SS_Total = np.sum((y - self.grand_mean) ** 2)
        
        # Calculate means at different levels
        # Ȳ_i... (mean for each level of A)
        mean_A = self.df.groupby(self.a_col)[self.resp_col].mean()
        # Ȳ_.j.. (mean for each level of B)
        mean_B = self.df.groupby(self.b_col)[self.resp_col].mean()
        # Ȳ_..k. (mean for each level of C)
        mean_C = self.df.groupby(self.c_col)[self.resp_col].mean()
        
        # Ȳ_ij.. (mean for each AB combination)
        mean_AB = self.df.groupby([self.a_col, self.b_col])[self.resp_col].mean()
        # Ȳ_i.k. (mean for each AC combination)
        mean_AC = self.df.groupby([self.a_col, self.c_col])[self.resp_col].mean()
        # Ȳ_.jk. (mean for each BC combination)
        mean_BC = self.df.groupby([self.b_col, self.c_col])[self.resp_col].mean()
        
        # Ȳ_ijk. (mean for each ABC combination)
        mean_ABC = self.df.groupby([self.a_col, self.b_col, self.c_col])[self.resp_col].mean()
        
        # Calculate r_bar (average replications per treatment combination)
        self.r_bar = self.n / (self.a * self.b * self.c)
        
        # Main Effects SS
        # SS_A = b*c*r * Σ(Ȳ_i... − Ȳ....)²
        SS_A = self.b * self.c * self.r_bar * np.sum((mean_A - self.grand_mean) ** 2)
        
        # SS_B = a*c*r * Σ(Ȳ_.j.. − Ȳ....)²
        SS_B = self.a * self.c * self.r_bar * np.sum((mean_B - self.grand_mean) ** 2)
        
        # SS_C = a*b*r * Σ(Ȳ_..k. − Ȳ....)²
        SS_C = self.a * self.b * self.r_bar * np.sum((mean_C - self.grand_mean) ** 2)
        
        # Two-way Interaction SS
        # SS_AB = c*r * ΣΣ(Ȳ_ij.. − Ȳ_i... − Ȳ_.j.. + Ȳ....)²
        SS_AB_components = []
        for (a_val, b_val), ab_mean in mean_AB.items():
            component = ab_mean - mean_A[a_val] - mean_B[b_val] + self.grand_mean
            SS_AB_components.append(component ** 2)
        SS_AB = self.c * self.r_bar * np.sum(SS_AB_components)
        
        # SS_AC = b*r * ΣΣ(Ȳ_i.k. − Ȳ_i... − Ȳ_..k. + Ȳ....)²
        SS_AC_components = []
        for (a_val, c_val), ac_mean in mean_AC.items():
            component = ac_mean - mean_A[a_val] - mean_C[c_val] + self.grand_mean
            SS_AC_components.append(component ** 2)
        SS_AC = self.b * self.r_bar * np.sum(SS_AC_components)
        
        # SS_BC = a*r * ΣΣ(Ȳ_.jk. − Ȳ_.j.. − Ȳ_..k. + Ȳ....)²
        SS_BC_components = []
        for (b_val, c_val), bc_mean in mean_BC.items():
            component = bc_mean - mean_B[b_val] - mean_C[c_val] + self.grand_mean
            SS_BC_components.append(component ** 2)
        SS_BC = self.a * self.r_bar * np.sum(SS_BC_components)
        
        # Three-way Interaction SS
        # SS_ABC = r * ΣΣΣ(Ȳ_ijk. − Ȳ_ij.. − Ȳ_i.k. − Ȳ_.jk. + Ȳ_i... + Ȳ_.j.. + Ȳ_..k. − Ȳ....)²
        SS_ABC_components = []
        for (a_val, b_val, c_val), abc_mean in mean_ABC.items():
            component = (abc_mean 
                        - mean_AB[(a_val, b_val)] 
                        - mean_AC[(a_val, c_val)] 
                        - mean_BC[(b_val, c_val)]
                        + mean_A[a_val] 
                        + mean_B[b_val] 
                        + mean_C[c_val]
                        - self.grand_mean)
            SS_ABC_components.append(component ** 2)
        SS_ABC = self.r_bar * np.sum(SS_ABC_components)
        
        # Error SS
        SS_Error = SS_Total - (SS_A + SS_B + SS_C + SS_AB + SS_AC + SS_BC + SS_ABC)
        
        # DFs
        df_A = self.a - 1
        df_B = self.b - 1
        df_C = self.c - 1
        
        df_AB = df_A * df_B
        df_AC = df_A * df_C
        df_BC = df_B * df_C
        
        df_ABC = df_A * df_B * df_C
        
        df_Error = (self.n - 1) - (df_A + df_B + df_C + df_AB + df_AC + df_BC + df_ABC)
        df_Total = self.n - 1
        
        self.MS_E = SS_Error / df_Error if df_Error > 0 else 0
        self.df_E = df_Error
        
        anova = {}
        
        for name, ss, df in [
            ("Factor A", SS_A, df_A),
            ("Factor B", SS_B, df_B),
            ("Factor C", SS_C, df_C),
            ("Interaction AxB", SS_AB, df_AB),
            ("Interaction AxC", SS_AC, df_AC),
            ("Interaction BxC", SS_BC, df_BC),
            ("Interaction AxBxC", SS_ABC, df_ABC),
            ("Error", SS_Error, df_Error),
        ]:
            if name == "Error":
                anova[name] = {"df": df, "SS": ss, "MS": ss/df if df>0 else 0, "F": None, "P": None}
                continue
                
            ms = ss / df if df > 0 else 0
            f = ms / self.MS_E if self.MS_E > 0 else 0
            p = 1 - stats.f.cdf(f, df, df_Error)
            anova[name] = {"df": df, "SS": ss, "MS": ms, "F": f, "P": p}
            
        anova["Total"] = {"df": df_Total, "SS": SS_Total, "MS": None, "F": None, "P": None}
        self.anova_table = anova
        return anova

    def run_post_hoc(self, method='lsd', alpha=0.05, order='desc'):
        self.alpha = alpha
        results = {}
        is_asc = (order == 'asc')
        
        # Helper for analysis
        def analyze_effect(effect_name, group_cols, se_divisor, n_means, SE_denom_is_r_bar=False):
            # Means
            if isinstance(group_cols, list) and len(group_cols) > 1:
                # Create phantom column
                col_name = " : ".join(group_cols)
                self.df[col_name] = self.df.apply(lambda x: " : ".join([str(x[c]) for c in group_cols]), axis=1)
                grp_field = col_name
            else:
                grp_field = group_cols[0]
                
            means = self.df.groupby(grp_field)[self.resp_col].mean().sort_values(ascending=is_asc)
            means_display = means.sort_index()
            
            # SE Calculation based on effect type
            # SE(mean) = sqrt(MS_E / n) where n depends on the effect
            # se_divisor is passed from the calling code
            SEm = np.sqrt(self.MS_E / se_divisor)
            SEd = np.sqrt(2 * self.MS_E / se_divisor)
            CV = (np.sqrt(self.MS_E) / self.grand_mean) * 100
            
            # Create SE dictionary - all levels have the same pooled SE
            # This is the scientifically correct approach for factorial ANOVA
            ses_pooled = {level: SEm for level in means_display.index}
            
            # Post-hoc Grouping
            grouping =  {k: "-" for k in means.index}
            
            # CD
            CD = None
            if method != 'dunnett':
                 CD = self._get_cd(method, alpha, self.df_E, SEm, n_means)
            
            if self.anova_table[effect_name]["P"] <= alpha or method == 'dunnett': 
                 if method == 'dunnett' and self.control_col:
                     # self.control_col now supposedly contains "ValA : ValB : ValC"
                     # We need to map this to the current effect's grouping.
                     
                     ctrl_level = None
                     
                     # 1. Parse the global control string
                     # It's constructed as "ValA : ValB : ValC"
                     # We assume ' : ' separator
                     parts = []
                     if self.control_col:
                         parts = [p.strip() for p in self.control_col.split(' : ')]
                     
                     if len(parts) >= 3:
                         val_a, val_b, val_c = parts[0], parts[1], parts[2]
                         
                         # 2. Determine which factors are in this effect
                         # We can check if `group_cols` contains a_col, b_col, etc.
                         # group_cols is passed to analyze_effect
                         
                         relevant_parts = []
                         
                         # Check strict column identity
                         # Note: This relies on group_cols elements being exactly self.a_col, etc.
                         # which they are in the calling code.
                         
                         for col in group_cols:
                             if col == self.a_col:
                                 relevant_parts.append(val_a)
                             elif col == self.b_col:
                                 relevant_parts.append(val_b)
                             elif col == self.c_col:
                                 relevant_parts.append(val_c)
                         
                         if relevant_parts:
                             if len(relevant_parts) == 1:
                                 ctrl_level = relevant_parts[0]
                             else:
                                 ctrl_level = " : ".join(relevant_parts)
                     
                     # Fallback if parsing failed or didn't match (e.g. string format mismatch)
                     if ctrl_level is None or ctrl_level not in means.index:
                         # Try simple heuristic or first level
                         if self.control_col in means.index:
                              ctrl_level = self.control_col
                         else:
                              # Just take the first one or try to find one with 'control'
                              candidates = [lvl for lvl in means.index.tolist() if 'control' in str(lvl).lower()]
                              ctrl_level = candidates[0] if candidates else means.index[0]

                     # Dunnett's Logic with Notation Support
                     # Calculate Dunnett's CD (Bonferroni T)
                     k = n_means - 1
                     t_crit = stats.t.ppf(1 - alpha/(2*k), self.df_E)
                     d_crit = t_crit * SEm * np.sqrt(2)
                     CD = d_crit
                     
                     # Grouping
                     grouping = {}
                     try:
                         ctrl_mean = means_display.loc[ctrl_level]
                         
                         # Determine notation style
                         notation = getattr(self, 'notation', 'symbol')
                         if not notation: notation = 'symbol'

                         for lvl, val in means_display.items():
                             # Ensure comparison is robust against type differences
                             is_control = str(lvl).strip() == str(ctrl_level).strip()
                             if is_control:
                                 grouping[lvl] = "a" if notation == 'alphabet' else "Control"
                             else:
                                 diff = abs(val - ctrl_mean)
                                 is_sig = diff > d_crit
                                 
                                 if notation == 'alphabet':
                                     # Logic: Control is 'a'. Sig diff is 'b'. Non-sig is 'a'.
                                     # This is not perfect CLD but simplest robust representation for Control vs Treatment
                                     grouping[lvl] = "b" if is_sig else "a"
                                 else:
                                     # Symbol logic
                                     grouping[lvl] = "*" if is_sig else "ns"
                                     
                     except KeyError:
                         grouping = {k: "?" for k in means.index}
                             
                 else:
                     # Standard grouping
                     grouping = self._compute_grouping(means, method, alpha, SEm, self.df_E)
            
            return {
                "means": means_display, "ses": ses_pooled,
                "grouping": grouping, "SE": SEm, "SEd": SEd, "CV": CV, "CD": CD
            }
        


        # Main Effects
        # SE(mean_A) = sqrt(MS_E / (b*c*r))
        results["Factor A"] = analyze_effect("Factor A", [self.a_col], self.b * self.c * self.r_bar, self.a)
        # SE(mean_B) = sqrt(MS_E / (a*c*r))
        results["Factor B"] = analyze_effect("Factor B", [self.b_col], self.a * self.c * self.r_bar, self.b)
        # SE(mean_C) = sqrt(MS_E / (a*b*r))
        results["Factor C"] = analyze_effect("Factor C", [self.c_col], self.a * self.b * self.r_bar, self.c)
        
        # Interactions
        # SE(mean_AB) = sqrt(MS_E / (c*r))
        results["Interaction AxB"] = analyze_effect("Interaction AxB", [self.a_col, self.b_col], self.c * self.r_bar, self.a * self.b)
        # SE(mean_AC) = sqrt(MS_E / (b*r))
        results["Interaction AxC"] = analyze_effect("Interaction AxC", [self.a_col, self.c_col], self.b * self.r_bar, self.a * self.c)
        # SE(mean_BC) = sqrt(MS_E / (a*r))
        results["Interaction BxC"] = analyze_effect("Interaction BxC", [self.b_col, self.c_col], self.a * self.r_bar, self.b * self.c)
        # SE(mean_ABC) = sqrt(MS_E / r)
        results["Interaction AxBxC"] = analyze_effect("Interaction AxBxC", [self.a_col, self.b_col, self.c_col], self.r_bar, self.a * self.b * self.c)
        
        self.results = results
        return results

    def _get_cd(self, method, alpha, df, SE, n_means=2):
        if method == 'lsd':
            t = stats.t.ppf(1 - alpha/2, df)
            return t * (np.sqrt(2) * SE)
        elif method == 'tukey':
            q = stats.studentized_range.ppf(1-alpha, n_means, df)
            return q * SE
        elif method == 'duncan':
            q = get_duncan_q(2, df, alpha)
            return q * SE
        return None

    def _compute_grouping(self, means, method, alpha, SE, df):
        vals = means.values
        labels = means.index.tolist()
        n = len(vals)
        sig_set = set()
        
        if method == 'lsd':
            limit = stats.t.ppf(1 - alpha/2, df) * np.sqrt(2) * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= limit:
                        sig_set.add((i, j))
        elif method == 'tukey':
            q = stats.studentized_range.ppf(1-alpha, n, df)
            limit = q * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= limit:
                        sig_set.add((i, j))
        elif method == 'duncan':
            for i in range(n):
                for j in range(i+1, n):
                    p = j - i + 1
                    q = get_duncan_q(p, df, alpha)
                    if abs(vals[i] - vals[j]) >= (q * SE):
                        sig_set.add((i, j))
                        
        G = nx.Graph()
        G.add_nodes_from(range(n))
        for i in range(n):
            for j in range(i+1, n):
                if (i, j) not in sig_set:
                    G.add_edge(i, j)
                    
        cliques = list(nx.find_cliques(G))
        cliques.sort(key=lambda c: (-max(c), -len(c)))
        vocab = "abcdefghijklmnopqrstuvwxyz"
        letters = {i: "" for i in range(n)}
        
        for idx, clq in enumerate(cliques):
            if idx < len(vocab):
                char = vocab[idx]
                for node in clq:
                    letters[node] += char
                    
        return {labels[i]: "".join(sorted(letters[i])) for i in range(n)}

    def create_report(self):
        doc = Document()
        doc.add_heading('Three-Factor CRD Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now()}")
        self.append_to_report(doc)
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def append_to_report(self, doc):
        doc.add_page_break()
        doc.add_heading(f"Analysis for: {self.resp_col}", level=1)
        
        # ANOVA
        doc.add_heading('ANOVA Summary', 1)
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = 'Table Grid'
        
        hdr = tbl.rows[0].cells
        for i, h in enumerate(['Source', 'DF', 'SS', 'MS', 'F-val', 'Result']):
            hdr[i].text = h
            
        order = ["Factor A", "Factor B", "Factor C", 
                 "Interaction AxB", "Interaction AxC", "Interaction BxC", 
                 "Interaction AxBxC", "Error", "Total"]
                 
        for k in order:
            if k not in self.anova_table: continue
            row = tbl.add_row().cells
            d = self.anova_table[k]
            row[0].text = k
            row[1].text = str(d['df'])
            row[2].text = f"{d['SS']:.4f}"
            row[3].text = f"{d['MS']:.4f}" if d['MS'] else ""
            row[4].text = f"{d['F']:.4f}" if d['F'] else ""
            if d['P'] is not None:
                sig = "**" if d['P'] <= 0.01 else ("*" if d['P'] <= 0.05 else "ns")
                row[5].text = f"{d['P']:.4f} {sig}"
        
        # Helper for Mean Tables
        def add_res_table(title, key):
            if key not in self.results: return
            doc.add_heading(title, 2)
            res = self.results[key]
            
            doc.add_paragraph(f"SE(m): {res['SE']:.4f} | SE(d): {res['SEd']:.4f} | CV: {res['CV']:.2f}% | CD: {res['CD'] if res['CD'] else 'ns'}")
            
            t = doc.add_table(rows=1, cols=4)
            t.style = 'Table Grid'
            h = t.rows[0].cells
            h[0].text = "Level"
            h[1].text = "Mean"
            h[2].text = "Std err"
            h[3].text = "Group"
            
            for lvl, val in res['means'].items():
                r = t.add_row().cells
                r[0].text = str(lvl)
                r[1].text = f"{val:.4f}"
                r[2].text = f"{res['ses'][lvl]:.4f}"
                r[3].text = res['grouping'][lvl]

        add_res_table("Factor A Means", "Factor A")
        add_res_table("Factor B Means", "Factor B")
        add_res_table("Factor C Means", "Factor C")
        add_res_table("AxB Interaction", "Interaction AxB")
        add_res_table("AxC Interaction", "Interaction AxC")
        add_res_table("BxC Interaction", "Interaction BxC")
        add_res_table("AxBxC Interaction", "Interaction AxBxC")
