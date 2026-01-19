import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from .duncan_util import get_duncan_q

class PooledCRDAnalyzer:
    def __init__(self, df, treat_col, year_col, resp_col):
        self.df = df
        self.treat_col = treat_col
        self.year_col = year_col
        self.resp_col = resp_col
        
        # Results storage
        self.bartlett_res = {}
        self.anova_table = {}
        self.post_hoc_res = {}
        
        # Metadata
        self.n_treats = 0
        self.n_years = 0
        self.n_reps_harm = 0

    def validate(self):
        # Data Cleanup
        self.df[self.treat_col] = self.df[self.treat_col].astype(str).str.strip()
        self.df[self.year_col] = self.df[self.year_col].astype(str).str.strip()
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        
        # Drop Nulls
        if self.df[self.resp_col].isnull().any():
            self.df = self.df.dropna(subset=[self.resp_col])
            
        self.treats = sorted(self.df[self.treat_col].unique())
        self.years = sorted(self.df[self.year_col].unique())
        self.n_treats = len(self.treats)
        self.n_years = len(self.years)

        if self.n_treats < 2:
            raise ValueError("Treatment must have at least 2 levels.")
        if self.n_years < 2:
            raise ValueError("Year/Location must have at least 2 levels.")
            
        # Check Balance (Optional, but good for calc)
        counts = self.df.groupby(self.treat_col)[self.resp_col].count()
        if len(counts) > 0:
            recip = sum(1.0/c for c in counts)
            self.n_reps_harm = len(counts) / recip
        else:
            self.n_reps_harm = 1

    def run_bartlett_test(self):
        mse_list = []
        df_list = []
        
        years_group = self.df.groupby(self.year_col)
        for name, group in years_group:
            y = group[self.resp_col]
            n = len(y)
            t = group[self.treat_col].nunique()
            gf = y.sum()
            cf_i = (gf**2)/n
            ss_tot = (y**2).sum() - cf_i
            grp_t = group.groupby(self.treat_col)[self.resp_col]
            ss_tr = sum((g.sum()**2 / len(g)) for _, g in grp_t) - cf_i
            ss_err = ss_tot - ss_tr
            df_err = n - t
            
            if df_err > 0 and ss_err > 1e-9:
                ms_err = ss_err / df_err
                mse_list.append(ms_err)
                df_list.append(df_err)
        
        k = len(mse_list)
        if k < 2:
             self.bartlett_res = {"valid": True, "stat": 0.0, "p": 1.0, "df": 0, "msg": "Insufficient groups."}
             return

        sum_df = sum(df_list)
        numerator = sum(d * s for d, s in zip(df_list, mse_list))
        sp_sq = numerator / sum_df
        
        term1 = sum_df * np.log(sp_sq)
        term2 = sum(d * np.log(s) for d, s in zip(df_list, mse_list))
        chi_sq = term1 - term2
        
        inv_sum = sum(1/d for d in df_list)
        C = 1 + (1 / (3 * (k - 1))) * (inv_sum - (1 / sum_df))
        
        corr_chi_sq = chi_sq / C
        p_val = 1 - stats.chi2.cdf(corr_chi_sq, k - 1)
        
        is_homo = bool(p_val >= 0.05)
        msg = "Variances are homogeneous." if is_homo else "Variances are heterogeneous."
        
        self.bartlett_res = {
            "valid": is_homo,
            "stat": float(corr_chi_sq),
            "p": float(p_val),
            "df": int(k - 1),
            "msg": msg
        }

    def run_pooled_anova(self):
        y_vals = self.df[self.resp_col]
        N = len(y_vals)
        G = y_vals.sum()
        CF = (G**2) / N
        
        SS_Total = (y_vals**2).sum() - CF
        
        t_groups = self.df.groupby(self.treat_col)[self.resp_col]
        SS_T = sum((g.sum()**2 / len(g)) for _, g in t_groups) - CF
        
        y_groups = self.df.groupby(self.year_col)[self.resp_col]
        SS_Y = sum((g.sum()**2 / len(g)) for _, g in y_groups) - CF
        
        ty_groups = self.df.groupby([self.treat_col, self.year_col])[self.resp_col]
        SS_Cells = sum((g.sum()**2 / len(g)) for _, g in ty_groups) - CF
        
        SS_TY = SS_Cells - SS_T - SS_Y
        SS_Error = SS_Total - SS_Cells 
        
        t = self.n_treats
        y = self.n_years
        n_cells = len(ty_groups) 
        
        df_Total = N - 1
        df_T = t - 1
        df_Y = y - 1
        df_TY = (t - 1) * (y - 1)
        df_Error = int(N - n_cells)
        
        MS_T = SS_T / df_T
        MS_Y = SS_Y / df_Y
        MS_TY = SS_TY / df_TY
        MS_E = SS_Error / df_Error
        
        F_T = MS_T / MS_E
        F_Y = MS_Y / MS_E
        F_TY = MS_TY / MS_E
        
        P_T = 1 - stats.f.cdf(F_T, df_T, df_Error)
        P_Y = 1 - stats.f.cdf(F_Y, df_Y, df_Error)
        P_TY = 1 - stats.f.cdf(F_TY, df_TY, df_Error)
        
        self.anova_table = {
            "Treatment": {"df": int(df_T), "SS": float(SS_T), "MS": float(MS_T), "F": float(F_T), "P": float(P_T)},
            "Year": {"df": int(df_Y), "SS": float(SS_Y), "MS": float(MS_Y), "F": float(F_Y), "P": float(P_Y)},
            "Interaction (T x Y)": {"df": int(df_TY), "SS": float(SS_TY), "MS": float(MS_TY), "F": float(F_TY), "P": float(P_TY)},
            "Error": {"df": int(df_Error), "SS": float(SS_Error), "MS": float(MS_E), "F": None, "P": None},
            "Total": {"df": int(df_Total), "SS": float(SS_Total), "MS": None, "F": None, "P": None}
        }

    def run_post_hoc(self, method, alpha, order='desc'):
        res = self.anova_table
        # Common Interaction Check
        p_inte = res["Interaction (T x Y)"]["P"]
        ms_error = res["Error"]["MS"]
        df_err = res["Error"]["df"]
        inte_ns = p_inte >= 0.05
        
        # 1. Treatment Analysis
        p_treat = res["Treatment"]["P"]
        treat_sig = p_treat < alpha
        run_treat = bool(treat_sig and inte_ns)
        
        reason_treat = ""
        if not treat_sig: reason_treat = f"Treatment effect not significant (p={p_treat:.4f})."
        elif not inte_ns: reason_treat = f"Interaction effect significant (p={p_inte:.4f})."

        t_means = self._get_means_stats(self.treat_col, order)
        
        # Pooled SEm for Treatment (sqrt(MSE / (N/t)))
        n_obs_t = len(self.df) / self.n_treats 
        sem_t_pooled = np.sqrt(ms_error / n_obs_t)
        
        t_grouping = {}
        t_cd = 0
        t_sed = sem_t_pooled * np.sqrt(2)
        
        if run_treat:
            t_grouping = self._compute_grouping(t_means['means'], method, alpha, df_err, ms_error, n_obs_t)
            t_cd = self._compute_cd(method, alpha, df_err, ms_error, n_obs_t, self.n_treats)
        else:
            t_grouping = {k: "" for k in t_means['means'].index}

        # 2. Year Analysis
        p_year = res["Year"]["P"]
        year_sig = p_year < alpha
        run_year = bool(year_sig and inte_ns)
        
        reason_year = ""
        if not year_sig: reason_year = f"Year effect not significant (p={p_year:.4f})."
        elif not inte_ns: reason_year = f"Interaction effect significant (p={p_inte:.4f})."

        y_means = self._get_means_stats(self.year_col, order)
        
        # Pooled SEm for Year
        n_obs_y = len(self.df) / self.n_years
        sem_y_pooled = np.sqrt(ms_error / n_obs_y)
        
        y_grouping = {}
        y_cd = 0
        y_sed = sem_y_pooled * np.sqrt(2)
        
        if run_year:
            y_grouping = self._compute_grouping(y_means['means'], method, alpha, df_err, ms_error, n_obs_y)
            y_cd = self._compute_cd(method, alpha, df_err, ms_error, n_obs_y, self.n_years)
        else:
            y_grouping = {k: "" for k in y_means['means'].index}
            
        CV = (np.sqrt(ms_error) / abs(self.df[self.resp_col].mean())) * 100
        
        self.post_hoc_res = {
            "Treatment": {
                "means": t_means['means'],
                "sds": t_means['sds'],
                "ses": t_means['ses'],
                "grouping": t_grouping,
                "sem_pooled": float(sem_t_pooled),
                "sed": float(t_sed),
                "cd": float(t_cd),
                "test_performed": run_treat,
                "reason": reason_treat
            },
            "Year": {
                "means": y_means['means'],
                "sds": y_means['sds'],
                "ses": y_means['ses'],
                "grouping": y_grouping,
                "sem_pooled": float(sem_y_pooled),
                "sed": float(y_sed),
                "cd": float(y_cd),
                "test_performed": run_year,
                "reason": reason_year
            },
            "CV": float(CV),
            "method": method.upper()
        }

    def _get_means_stats(self, col, order):
        grp = self.df.groupby(col)[self.resp_col]
        means = grp.mean()
        sds = grp.std().fillna(0) # SD of raw data
        counts = grp.count()
        ses = sds / np.sqrt(counts) # Individual SE
        
        if order == 'desc': 
            # Original request was to order by value, but user now wants alphabetical order of labels
            # We will ignore 'order' parameter for the LISTING, but keep it for logic if needed (which currently isn't)
            # Sorting by index (Treatment Name)
            means = means.sort_index(ascending=True) 
        else: 
            means = means.sort_index(ascending=True)
            
        return {"means": means, "sds": sds, "ses": ses}

    def _compute_cd(self, method, alpha, df, MSE, n_reps, k):
        SE = np.sqrt(MSE / n_reps)
        if method == 'lsd':
            crit = stats.t.ppf(1 - alpha/2, df)
            return crit * np.sqrt(2) * SE
        elif method == 'tukey':
             crit = stats.studentized_range.ppf(1-alpha, k, df)
             return crit * SE
        return 0

    def _compute_grouping(self, means, method, alpha, df, MSE, n_reps):
        vals = means.values
        labels = means.index.tolist()
        n = len(vals)
        sig = set()
        
        SE = np.sqrt(MSE / n_reps)
        
        if method == 'lsd':
            crit = stats.t.ppf(1 - alpha/2, df)
            LSD = crit * np.sqrt(2) * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= LSD: sig.add((i, j))
        elif method == 'tukey':
            crit = stats.studentized_range.ppf(1-alpha, n, df)
            HSD = crit * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= HSD: sig.add((i, j))
        elif method == 'dmrt':
            for i in range(n):
                for j in range(i+1, n):
                    p = j - i + 1
                    q_val = get_duncan_q(p, df, alpha)
                    Rp = q_val * SE
                    if abs(vals[i] - vals[j]) >= Rp: sig.add((i, j))
                    
        G = nx.Graph()
        G.add_nodes_from(range(n))
        for i in range(n):
            for j in range(i+1, n):
                if (i, j) not in sig: G.add_edge(i, j)
        
        cliques = list(nx.find_cliques(G))
        # Sort cliques so that 'a' is assigned to the clique containing the first element (index 0)
        # We sort by the smallest index in the clique.
        cliques.sort(key=lambda c: (min(c), -len(c)))
        vocab = "abcdefghijklmnopqrstuvwxyz"
        res = {i: "" for i in range(n)}
        
        for idx, clq in enumerate(cliques):
            if idx < len(vocab):
                char = vocab[idx]
                for node in clq: res[node] += char
                
        return {labels[i]: "".join(sorted(res[i])) for i in range(n)}

    def create_report(self):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        doc.add_heading('One Factor Pooled CRD Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # 1. Bartlett
        doc.add_heading('1. Homogeneity of Variance (Bartlett’s Test)', level=1)
        if self.bartlett_res:
            b = self.bartlett_res
            t = doc.add_table(2, 4)
            t.style = 'Table Grid'
            r = t.rows[0].cells
            r[0].text = "Statistic"
            r[1].text = "DF"
            r[2].text = "P-Value"
            r[3].text = "Decision"
            r2 = t.rows[1].cells
            r2[0].text = f"{b['stat']:.4f}"
            r2[1].text = str(b['df'])
            r2[2].text = f"{b['p']:.4f}"
            r2[3].text = b['msg']
            doc.add_paragraph(f"Conclusion: {b['msg']}")
            
        # 2. ANOVA
        doc.add_heading('2. Pooled Analysis of Variance', level=1)
        if self.anova_table:
            t = doc.add_table(1, 6)
            t.style = 'Table Grid'
            h = t.rows[0].cells
            for i, c in enumerate(["Source", "DF", "SS", "MS", "F", "Prob"]): h[i].text = c
            for src in ["Treatment", "Year", "Interaction (T x Y)", "Error", "Total"]:
                if src in self.anova_table:
                    row = t.add_row().cells
                    d = self.anova_table[src]
                    row[0].text = src
                    row[1].text = str(d['df'])
                    row[2].text = f"{d['SS']:.4f}"
                    row[3].text = f"{d['MS']:.4f}" if d['MS'] else ""
                    row[4].text = f"{d['F']:.4f}" if d['F'] else ""
                    if d['P'] is not None:
                        sig = "**" if d['P'] < 0.01 else "*" if d['P'] < 0.05 else "ns"
                        row[5].text = f"{d['P']:.4f} {sig}"
        
        # 3. Means (Treatment)
        if self.post_hoc_res:
            ph = self.post_hoc_res
            tr = ph["Treatment"]
            doc.add_heading('3a. Treatment Means', level=1)
            
            if not tr['test_performed']:
                doc.add_paragraph(f"Note: {tr['reason']}")
            
            doc.add_paragraph(f"Pooled SEm: {tr['sem_pooled']:.4f} | SEd: {tr['sed']:.4f} | CD ({alpha}): {tr['cd']:.4f} | CV%: {ph['CV']:.2f}")
            
            t = doc.add_table(1, 5)
            t.style = 'Table Grid'
            h = t.rows[0].cells
            h[0].text = "Treatment"
            h[1].text = "Mean"
            h[2].text = "St.Dev"
            h[3].text = "St.Err"
            h[4].text = "Group"
            
            for lvl, val in tr['means'].items():
                r = t.add_row().cells
                r[0].text = str(lvl)
                r[1].text = f"{val:.4f}"
                r[2].text = f"{tr['sds'].get(lvl,0):.4f}"
                r[3].text = f"{tr['ses'].get(lvl,0):.4f}"
                r[4].text = tr['grouping'].get(lvl, "")

            # 3b. Year Means
            yr = ph["Year"]
            doc.add_heading('3b. Year / Location Means', level=1)
            if not yr['test_performed']:
                 doc.add_paragraph(f"Note: {yr['reason']}")
            doc.add_paragraph(f"Pooled SEm: {yr['sem_pooled']:.4f} | SEd: {yr['sed']:.4f} | CD ({alpha}): {yr['cd']:.4f}")
            
            t2 = doc.add_table(1, 5)
            t2.style = 'Table Grid'
            h = t2.rows[0].cells
            h[0].text = "Year / Loc"
            h[1].text = "Mean"
            h[2].text = "St.Dev"
            h[3].text = "St.Err"
            h[4].text = "Group"
            
            for lvl, val in yr['means'].items():
                r = t2.add_row().cells
                r[0].text = str(lvl)
                r[1].text = f"{val:.4f}"
                r[2].text = f"{yr['sds'].get(lvl,0):.4f}"
                r[3].text = f"{yr['ses'].get(lvl,0):.4f}"
                r[4].text = yr['grouping'].get(lvl, "")

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
