import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
import io
from datetime import datetime
from .duncan_util import get_duncan_q

class SplitSplitPlotAnalyzer:
    def __init__(self, df, col_main, col_sub, col_subsub, col_rep, col_resp):
        self.df = df
        self.a_col = col_main
        self.b_col = col_sub
        self.c_col = col_subsub
        self.r_col = col_rep
        self.resp_col = col_resp
        
        self.anova_table = {}
        self.stats = {}
        self.post_hoc_res = {}

    def validate(self):
        for col in [self.a_col, self.b_col, self.c_col, self.r_col]:
             self.df[col] = self.df[col].astype(str).str.strip()
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.resp_col])
        
        self.n_a = len(self.df[self.a_col].unique())
        self.n_b = len(self.df[self.b_col].unique())
        self.n_c = len(self.df[self.c_col].unique())
        self.n_r = len(self.df[self.r_col].unique())
        
        if any(n < 2 for n in [self.n_a, self.n_b, self.n_c, self.n_r]):
             raise ValueError("All Factors and Replications must have at least 2 levels.")

    def run_anova(self):
        y = self.df[self.resp_col]
        G = y.sum()
        N = len(y)
        CF = (G**2) / N
        SS_Total = (y**2).sum() - CF
        
        def get_ss(cols, divisor):
            grp = self.df.groupby(cols)[self.resp_col].sum()
            return (grp**2).sum() / divisor - CF
            
        nA, nB, nC, nR = self.n_a, self.n_b, self.n_c, self.n_r
        
        # 1. Replication
        SS_R = get_ss(self.r_col, nA*nB*nC)
        
        # 2. Main Plot (A)
        SS_A = get_ss(self.a_col, nB*nC*nR)
        
        # 3. Error (a)
        # SS_Ea = SS_RA - SS_R - SS_A
        SS_RA_Total = get_ss([self.r_col, self.a_col], nB*nC)
        SS_Err_A = SS_RA_Total - SS_R - SS_A
        
        # 4. Sub Plot (B)
        SS_B = get_ss(self.b_col, nA*nC*nR)
        
        # 5. Interaction AxB
        SS_AB_Total = get_ss([self.a_col, self.b_col], nC*nR)
        SS_AxB = SS_AB_Total - SS_A - SS_B
        
        # 6. Error (b)
        # SS_Eb = SS_RAB - SS_RA - SS_B - SS_AxB
        # where SS_RAB is Rep x Main x Sub interaction total
        SS_RAB_Total = get_ss([self.r_col, self.a_col, self.b_col], nC)
        # SS_RA is already calculated as SS_RA_Total, wait... 
        # SS formulas in user request: SS_Ea = SS_RA - SSR - SSA. 
        # SS_Eb = SS_RAB - SS_RA - SSB - SSAB.
        # Here SS_RA refers to the total SS of Rep x A combination (called SS_RA_Total in my code).
        SS_Err_B = SS_RAB_Total - SS_RA_Total - SS_B - SS_AxB
        
        # 7. Sub-Sub Plot (C)
        SS_C = get_ss(self.c_col, nA*nB*nR)
        
        # 8. Interaction AxC
        SS_AC_Total = get_ss([self.a_col, self.c_col], nB*nR)
        SS_AxC = SS_AC_Total - SS_A - SS_C
        
        # 9. Interaction BxC
        SS_BC_Total = get_ss([self.b_col, self.c_col], nA*nR)
        SS_BxC = SS_BC_Total - SS_B - SS_C
        
        # 10. Interaction AxBxC
        SS_ABC_Total = get_ss([self.a_col, self.b_col, self.c_col], nR)
        SS_AxBxC = SS_ABC_Total - SS_A - SS_B - SS_C - SS_AxB - SS_AxC - SS_BxC
        
        # 11. Error (c)
        # SS_Ec = SST - (Everything else)
        # Explicitly: SST - (SSR + SSA + SS_Ea + SSB + SSAB + SS_Eb + SSC + SSAC + SSBC + SSABC)
        # Note: (SSR + SSA + SS_Ea) = SS_RA_Total
        # (SSB + SSAB + SS_Eb) = SS_RAB_Total - SS_RA_Total
        # So (SSR + ... + SS_Eb) = SS_RAB_Total
        # So SS_Ec = SS_Total - SS_RAB_Total - SS_C - SS_AC - SS_BC - SS_ABC
        SS_Err_C = SS_Total - SS_RAB_Total - SS_C - SS_AxC - SS_BxC - SS_AxBxC
        
        # DF
        df_r = nR - 1
        df_a = nA - 1
        df_err_a = (nR - 1) * (nA - 1)
        
        df_b = nB - 1
        df_axb = (nA - 1) * (nB - 1)
        df_err_b = nA * (nR - 1) * (nB - 1)
        
        df_c = nC - 1
        df_axc = (nA - 1) * (nC - 1)
        df_bxc = (nB - 1) * (nC - 1)
        df_axbxc = (nA - 1) * (nB - 1) * (nC - 1)
        df_err_c = nA * nB * (nR - 1) * (nC - 1)
        
        df_tot = nA * nB * nC * nR - 1
        
        # MS
        MS_R = SS_R / df_r
        MS_A = SS_A / df_a
        MS_Err_A = SS_Err_A / df_err_a
        
        MS_B = SS_B / df_b
        MS_AxB = SS_AxB / df_axb
        MS_Err_B = SS_Err_B / df_err_b
        
        MS_C = SS_C / df_c
        MS_AxC = SS_AxC / df_axc
        MS_BxC = SS_BxC / df_bxc
        MS_AxBxC = SS_AxBxC / df_axbxc
        MS_Err_C = SS_Err_C / df_err_c
        
        # F Tests
        F_A = MS_A / MS_Err_A
        
        F_B = MS_B / MS_Err_B
        F_AxB = MS_AxB / MS_Err_B
        
        F_C = MS_C / MS_Err_C
        F_AxC = MS_AxC / MS_Err_C
        F_BxC = MS_BxC / MS_Err_C
        F_AxBxC = MS_AxBxC / MS_Err_C
        
        # P Values
        def get_p(f, dfn, dfd): return 1 - stats.f.cdf(f, dfn, dfd)
        
        P_A = get_p(F_A, df_a, df_err_a)
        
        P_B = get_p(F_B, df_b, df_err_b)
        P_AxB = get_p(F_AxB, df_axb, df_err_b)
        
        P_C = get_p(F_C, df_c, df_err_c)
        P_AxC = get_p(F_AxC, df_axc, df_err_c)
        P_BxC = get_p(F_BxC, df_bxc, df_err_c)
        P_AxBxC = get_p(F_AxBxC, df_axbxc, df_err_c)
        
        def s(df, ss, ms, f, p): return {"df": df, "SS": ss, "MS": ms, "F": f, "P": p}
        
        self.anova_table = {
            "Replication": {"df": df_r, "SS": SS_R, "MS": MS_R, "F": None, "P": None},
            "Main Plot (A)": s(df_a, SS_A, MS_A, F_A, P_A),
            "Error (a)": {"df": df_err_a, "SS": SS_Err_A, "MS": MS_Err_A, "F": None, "P": None},
            
            "Sub Plot (B)": s(df_b, SS_B, MS_B, F_B, P_B),
            "Interaction A x B": s(df_axb, SS_AxB, MS_AxB, F_AxB, P_AxB),
            "Error (b)": {"df": df_err_b, "SS": SS_Err_B, "MS": MS_Err_B, "F": None, "P": None},
            
            "Sub-Sub Plot (C)": s(df_c, SS_C, MS_C, F_C, P_C),
            "Interaction A x C": s(df_axc, SS_AxC, MS_AxC, F_AxC, P_AxC),
            "Interaction B x C": s(df_bxc, SS_BxC, MS_BxC, F_BxC, P_BxC),
            "Interaction A x B x C": s(df_axbxc, SS_AxBxC, MS_AxBxC, F_AxBxC, P_AxBxC),
            "Error (c)": {"df": df_err_c, "SS": SS_Err_C, "MS": MS_Err_C, "F": None, "P": None},
            
            "Total": {"df": df_tot, "SS": SS_Total, "MS": None, "F": None, "P": None}
        }

    def run_post_hoc(self, method, alpha):
        res = self.anova_table
        MS_Ea = res["Error (a)"]["MS"]
        DF_Ea = int(res["Error (a)"]["df"])
        
        MS_Eb = res["Error (b)"]["MS"]
        DF_Eb = int(res["Error (b)"]["df"])
        
        MS_Ec = res["Error (c)"]["MS"]
        DF_Ec = int(res["Error (c)"]["df"])
        
        gm = self.df[self.resp_col].mean()
        
        cv_a = (np.sqrt(MS_Ea) / gm) * 100
        cv_b = (np.sqrt(MS_Eb) / gm) * 100
        cv_c = (np.sqrt(MS_Ec) / gm) * 100
        
        # SEm Helper
        def calc_sem(ms, div): return np.sqrt(ms / div)
        
        # --- Error A Group ---
        # Main A
        sem_a = calc_sem(MS_Ea, self.n_r * self.n_b * self.n_c)
        
        # --- Error B Group ---
        # Sub B
        sem_b = calc_sem(MS_Eb, self.n_r * self.n_a * self.n_c)
        # Interaction AxB
        sem_axb = calc_sem(MS_Eb, self.n_r * self.n_c)
        
        # --- Error C Group ---
        # Sub-Sub C
        sem_c = calc_sem(MS_Ec, self.n_r * self.n_a * self.n_b)
        # Interactions with C
        sem_axc = calc_sem(MS_Ec, self.n_r * self.n_b)
        sem_bxc = calc_sem(MS_Ec, self.n_r * self.n_a)
        sem_axbxc = calc_sem(MS_Ec, self.n_r)
        
        def calc_cd(sem, df, k):
            sed = sem * np.sqrt(2)
            if method == 'lsd': return stats.t.ppf(1 - alpha/2, df) * sed
            elif method == 'tukey': return stats.studentized_range.ppf(1-alpha, k, df) * sem
            elif method == 'dmrt': return stats.t.ppf(1 - alpha/2, df) * sed
            return 0
        
        stats_map = {
            "CV (a)": cv_a, "CV (b)": cv_b, "CV (c)": cv_c,
            "Main Plot (A)": {"SEm": sem_a, "SEd": sem_a*np.sqrt(2), "CD": calc_cd(sem_a, DF_Ea, self.n_a)},
            
            "Sub Plot (B)": {"SEm": sem_b, "SEd": sem_b*np.sqrt(2), "CD": calc_cd(sem_b, DF_Eb, self.n_b)},
            "Interaction A x B": {"SEm": sem_axb, "SEd": sem_axb*np.sqrt(2), "CD": calc_cd(sem_axb, DF_Eb, self.n_a*self.n_b)},
            
            "Sub-Sub Plot (C)": {"SEm": sem_c, "SEd": sem_c*np.sqrt(2), "CD": calc_cd(sem_c, DF_Ec, self.n_c)},
            "Interaction A x C": {"SEm": sem_axc, "SEd": sem_axc*np.sqrt(2), "CD": calc_cd(sem_axc, DF_Eb, self.n_a*self.n_c)},
            "Interaction B x C": {"SEm": sem_bxc, "SEd": sem_bxc*np.sqrt(2), "CD": calc_cd(sem_bxc, DF_Eb, self.n_b*self.n_c)},
            "Interaction A x B x C": {"SEm": sem_axbxc, "SEd": sem_axbxc*np.sqrt(2), "CD": calc_cd(sem_axbxc, DF_Eb, self.n_a*self.n_b*self.n_c)},
        }
        self.stats = stats_map

        # Results
        results = {}
        def get_data(cols):
            grp = self.df.groupby(cols)[self.resp_col]
            means = grp.mean().sort_index()
            stds = grp.std().reindex(means.index)
            ses = (stds / np.sqrt(grp.count())).reindex(means.index)
            return means, stds, ses

        def compute_letters(means, sem, df, k):
             m_desc = means.sort_values(ascending=False)
             vals = m_desc.values
             keys = m_desc.index.tolist()
             n = len(vals)
             sed = sem * np.sqrt(2)
             sig_pairs = set()
             
             if method == 'lsd':
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             elif method == 'tukey':
                 crit = stats.studentized_range.ppf(1-alpha, n, df) * sem
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             else: # DMRT
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             
             G = nx.Graph()
             G.add_nodes_from(range(n))
             for i in range(n):
                 for j in range(i+1, n):
                     if (i,j) not in sig_pairs: G.add_edge(i, j)
             
             cliques = list(nx.find_cliques(G))
             cliques.sort(key=lambda x: (min(x), -len(x)))
             vocab = "abcdefghijklmnopqrstuvwxyz"
             res_map = {i: "" for i in range(n)}
             for idx, clq in enumerate(cliques):
                 if idx < len(vocab):
                     c = vocab[idx]
                     for node in clq: res_map[node] += c
             return {keys[i]: "".join(sorted(res_map[i])) for i in range(n)}

        def proc_eff(name, cols, sem_key, df_err, n_lvls):
            m, s, se = get_data(cols)
            p = res[name]["P"]
            let = {k: "ns" for k in m.index}
            if p < alpha:
                 let = compute_letters(m, stats_map[sem_key]["SEm"], df_err, n_lvls)
            results[sem_key] = {"means": m, "stds": s, "ses": se, "grouping": let}

        # Main Plot
        proc_eff("Main Plot (A)", self.a_col, "Main Plot (A)", DF_Ea, self.n_a)
        
        # Sub Plot
        proc_eff("Sub Plot (B)", self.b_col, "Sub Plot (B)", DF_Eb, self.n_b)
        proc_eff("Interaction A x B", [self.a_col, self.b_col], "Interaction A x B", DF_Eb, self.n_a*self.n_b)
        
        # Sub-Sub Plot
        proc_eff("Sub-Sub Plot (C)", self.c_col, "Sub-Sub Plot (C)", DF_Ec, self.n_c)
        proc_eff("Interaction A x C", [self.a_col, self.c_col], "Interaction A x C", DF_Ec, self.n_a*self.n_c)
        proc_eff("Interaction B x C", [self.b_col, self.c_col], "Interaction B x C", DF_Ec, self.n_b*self.n_c)
        proc_eff("Interaction A x B x C", [self.a_col, self.b_col, self.c_col], "Interaction A x B x C", DF_Ec, self.n_a*self.n_b*self.n_c)
        
        self.post_hoc_res = results

    def create_report(self):
        doc = Document()
        doc.add_heading('Split-Split Plot Analysis Report', 0)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        doc.add_heading('1. ANOVA Table', 1)
        t = doc.add_table(1, 6)
        t.style = 'Table Grid'
        def set_cell(r, i, txt): r.cells[i].text = str(txt)
        h = t.rows[0]
        for i,c in enumerate(["Source", "DF", "SS", "MS", "F", "Prob"]): set_cell(h, i, c)
        
        keys = ["Replication", "Main Plot (A)", "Error (a)",
                "Sub Plot (B)", "Interaction A x B", "Error (b)",
                "Sub-Sub Plot (C)", "Interaction A x C", "Interaction B x C", "Interaction A x B x C", "Error (c)",
                "Total"]
        for k in keys:
            if k in self.anova_table:
                v = self.anova_table[k]
                r = t.add_row()
                sig = ""
                if v['P'] is not None:
                     if v['P']<0.01: sig="**"
                     elif v['P']<0.05: sig="*"
                     else: sig="ns"
                set_cell(r, 0, k)
                set_cell(r, 1, v['df'])
                set_cell(r, 2, f"{v['SS']:.4f}")
                set_cell(r, 3, f"{v['MS']:.4f}" if v['MS'] else "-")
                set_cell(r, 4, f"{v['F']:.4f}" if v['F'] else "-")
                set_cell(r, 5, f"{v['P']:.4f} {sig}" if v['P'] is not None else "")
        
        doc.add_heading('2. Statistics', 1)
        doc.add_paragraph(f"CV (a) %: {self.stats['CV (a)']:.2f}")
        doc.add_paragraph(f"CV (b) %: {self.stats['CV (b)']:.2f}")
        doc.add_paragraph(f"CV (c) %: {self.stats['CV (c)']:.2f}")
        
        stats_keys = ["Main Plot (A)", "Sub Plot (B)", "Interaction A x B", 
                      "Sub-Sub Plot (C)", "Interaction A x C", "Interaction B x C", "Interaction A x B x C"]
        for k in stats_keys:
             s = self.stats[k]
             doc.add_paragraph(f"{k}: SEm={s['SEm']:.4f}, SEd={s['SEd']:.4f}, CD={s['CD']:.4f}")
             
        doc.add_heading('3. Means and Grouping', 1)
        ph = self.post_hoc_res
        for k in stats_keys:
             ds = ph[k]
             doc.add_heading(f"{k} Means", 2)
             t = doc.add_table(1, 5)
             t.style = 'Table Grid'
             r = t.rows[0]
             for i,c in enumerate(["Level", "Mean", "Std.Dev", "Std.Err", "Group"]): set_cell(r, i, c)
             for lvl, val in ds['means'].items():
                 r = t.add_row()
                 if isinstance(lvl, tuple): l_str = " x ".join(map(str, lvl))
                 else: l_str = str(lvl)
                 set_cell(r, 0, l_str)
                 set_cell(r, 1, f"{val:.4f}")
                 set_cell(r, 2, f"{ds['stds'][lvl]:.4f}")
                 set_cell(r, 3, f"{ds['ses'][lvl]:.4f}")
                 set_cell(r, 4, ds['grouping'].get(lvl, ""))
                 
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
