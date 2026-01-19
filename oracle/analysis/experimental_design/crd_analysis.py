import pandas as pd
import numpy as np
import scipy.stats as stats
import scikit_posthocs as sp
import networkx as nx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

class CRDAnalyzer:
    def __init__(self, df, treat_col, resp_col):
        self.df = df
        self.treat_col = treat_col
        self.resp_col = resp_col
        
        self.anova_table = {}
        self.means = None
        self.grouping = None
        
        # Stats
        self.t = 0
        self.n = 0
        self.dfs = {}
        self.MS_E = 0
        self.r_avergae = 0 # r_bar
        
        self.alpha = 0.05
        
    def validate(self):
        # 1. Type Conversion
        self.df[self.treat_col] = self.df[self.treat_col].astype(str)
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        
        if self.df[self.resp_col].isnull().any():
             raise ValueError("Response variable contains missing or non-numeric values.")
             
        # 2. Dimensions
        self.t = self.df[self.treat_col].nunique()
        self.n = len(self.df)
        
        if self.n <= self.t:
            raise ValueError(f"CRD Error: Total observations ({self.n}) must be greater than number of treatments ({self.t}). Degrees of freedom for error must be > 0.")
            
    def run_anova(self):
        # 4. Grand Total and Mean
        y = self.df[self.resp_col]
        G = y.sum()
        Y_bar = y.mean()
        
        # 5. Correction Factor
        CF = (G**2) / self.n
        
        # 6. Total SS
        SS_Total = (y**2).sum() - CF
        
        # 7. Treatment SS
        # SS_treatment = sum( (T_total[i]^2) / r_i ) - CF
        # Group by treatment, then for each group: sum^2 / count
        treat_groups = self.df.groupby(self.treat_col)[self.resp_col]
        
        # Calculate Term 1: Sum(T_i^2 / r_i)
        term1 = 0
        r_counts = []
        for name, group in treat_groups:
            Ti = group.sum()
            ri = len(group)
            r_counts.append(ri)
            term1 += (Ti**2) / ri
            
        SS_Treat = term1 - CF
        
        # 8. Error SS
        # Direct subtraction method
        SS_Error = SS_Total - SS_Treat
        
        # 9. Degrees of Freedom
        df_treat = self.t - 1
        df_error = self.n - self.t
        df_total = self.n - 1
        
        if df_error <= 0:
             raise ValueError("Degrees of freedom for error is 0 or negative. Cannot perform analysis.")

        # 10. Mean Squares
        MS_Treat = SS_Treat / df_treat
        MS_Error = SS_Error / df_error
        
        # Store for post-hoc
        self.MS_E = MS_Error
        self.dfs = {"treat": df_treat, "error": df_error, "total": df_total}
        self.r_average = np.mean(r_counts) # As per spec 3/13: r_bar for unequal replication approximation?
        # Actually for balanced it's just r. For unbalanced, Harmonic mean is usually preferred for PostHoc, 
        # but spec says "r_bar = average_replications_per_treatment". We will use arithmetic mean if spec implies simple avg,
        # but statistically, harmonic mean is safer for unbalanced Tukey/LSD.
        # Let's use harmonic mean for robustness if unbalanced, or arithmetic if strictly following "average".
        # Given "sciientifically acurate analysis" request, Harmonic Mean is better for unbalanced data (Kramer modification).
        # But let's stick to the prompt's variable `r_bar`.
        # Prompt: "r_bar = average_replications_per_treatment".
        self.r_bar = self.n / self.t
        
        # Harmonic mean for better accuracy in unbalanced SEd
        # h_mean = stats.hmean(r_counts)
        # We will use r_bar from prompt for now, or maybe the specific r per comparison is better?
        # Standard simple formula uses n/t.
        
        # 11. F-Statistic & P-Value
        if MS_Error > 0:
            F_Treat = MS_Treat / MS_Error
            P_Treat = 1 - stats.f.cdf(F_Treat, df_treat, df_error)
        else:
            F_Treat = 0 # Or infinity
            P_Treat = 0 if F_Treat > 0 else 1
            
        self.anova_table = {
            "Treatments": {"df": df_treat, "SS": SS_Treat, "MS": MS_Treat, "F": F_Treat, "P": P_Treat},
            "Error": {"df": df_error, "SS": SS_Error, "MS": MS_Error, "F": None, "P": None},
            "Total": {"df": df_total, "SS": SS_Total, "MS": None, "F": None, "P": None}
        }
        
        # Stats
        self.grand_mean = Y_bar
        
        return self.anova_table

    def run_post_hoc(self, method='lsd', alpha=0.05, order='desc', control_group=None, notation='letters'):
        self.alpha = alpha
        is_ascending = True if order == 'asc' else False
        
        # 13. Precision Stats
        # SEm = sqrt(MS_error / r_bar)
        SEm = np.sqrt(self.MS_E / self.r_bar)
        
        # SEd = sqrt(2 * MS_error / r_bar)
        SEd = np.sqrt(2 * self.MS_E / self.r_bar)
        
        # CV
        CV = (np.sqrt(self.MS_E) / self.grand_mean) * 100
        
        # CD (Default)
        CD = None
        if self.anova_table["Treatments"]["P"] is not None and self.anova_table["Treatments"]["P"] <= alpha:
             if method != 'dunnett':
                 t_crit = stats.t.ppf(1 - alpha/2, self.dfs["error"])
                 CD = t_crit * SEd
        
        results = {
            "SEm": SEm,
            "SEd": SEd,
            "CV": CV,
            "CD": CD
        }
        
        # 1. Means Calculation
        means_calc = self.df.groupby(self.treat_col)[self.resp_col].mean().sort_values(ascending=is_ascending)
        
        # Grouping Logic
        grouping = {}
        if self.anova_table["Treatments"]["P"] is not None and self.anova_table["Treatments"]["P"] <= alpha:
             if method == 'dunnett' and control_group is not None:
                 # Dunnett's Test
                 # Use scikit-posthocs
                 try:
                     p_matrix = sp.posthoc_dunnett(self.df, val_col=self.resp_col, group_col=self.treat_col, control=control_group)
                     # p_matrix columns are treatments, rows are treatments. 
                     # Actually posthoc_dunnett output format might vary. Usually it's a matrix?
                     # Wait, it compares all to control.
                     # Documentation says: "Returns: P values. DataFrame".
                     # Row for each group, Column for each group?
                     # Usually for Dunnett, it returns comparisons against control.
                     
                     # Extract P-values vs Control
                     is_symbol = notation and notation.strip().startswith('symbol')
                     grouping = {control_group: "Control" if is_symbol else "a (Control)"}
                     
                     for trt in means_calc.index:
                         if trt == control_group:
                             continue
                         
                         pval = 1.0
                         if trt in p_matrix.index and control_group in p_matrix.columns:
                             pval = p_matrix.loc[trt, control_group]
                         elif control_group in p_matrix.index and trt in p_matrix.columns:
                             pval = p_matrix.loc[control_group, trt]
                             
                         if is_symbol:
                             if pval < alpha:
                                 stars = "*"
                                 if pval < 0.01: stars = "**"
                                 if pval < 0.001: stars = "***"
                                 grouping[trt] = stars
                             else:
                                 grouping[trt] = "ns"
                         else: # Letters Only
                             if pval < alpha:
                                 grouping[trt] = "b"
                             else:
                                 grouping[trt] = "a"
                                 
                 except Exception as e:
                     print(f"Dunnett Error: {e}")
                     grouping = {t: "err" for t in means_calc.index}
             else:
                 grouping = self._compute_grouping(means_calc, method, alpha, SEm, self.dfs["error"])
        else:
             grouping = {t: "ns" for t in means_calc.index}
             
        # 2. Display Prep (Use calculated sort order)
        means_display = means_calc
        sds = self.df.groupby(self.treat_col)[self.resp_col].std().sort_index()
        counts = self.df.groupby(self.treat_col)[self.resp_col].count().sort_index()
        ses_ind = sds / np.sqrt(counts) # Mean specific SE
        
        # Assemble Rows
        rows = []
        for idx in means_display.index:
            rows.append({
                "level": str(idx),
                "mean": float(means_display[idx]),
                "sd": float(sds[idx]) if not pd.isna(sds[idx]) else 0.0,
                "se": float(ses_ind[idx]) if not pd.isna(ses_ind[idx]) else 0.0,
                "group": grouping.get(idx, "-")
            })
            
        results["means"] = rows
        results["df_means"] = means_display # Keep for reporting
        results["df_sds"] = sds
        results["grouping"] = grouping
        
        self.results = results
        return results

    def _compute_grouping(self, means, method, alpha, SE, df):
        # means is sorted Series
        labels = means.index.tolist()
        vals = means.values
        n = len(vals)
        significance_matrix = set()
        
        # SEd is derived from passed SE (SEm) assuming balanced r_bar?
        SEd = np.sqrt(2) * SE
        
        limit = 0
        
        if method == 'lsd':
            t_crit = stats.t.ppf(1 - alpha/2, df)
            limit = t_crit * SEd
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= limit:
                        significance_matrix.add((i, j))
                        
        elif method == 'tukey':
            # Tukey HSD
            # q_tukey(alpha, t, df) * sqrt(MS_E / r) -> q * SE
            q = stats.studentized_range.ppf(1-alpha, n, df)
            limit = q * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= limit:
                         significance_matrix.add((i, j))
        
        # Clique Cover for Letters
        G_ns = nx.Graph()
        G_ns.add_nodes_from(range(n))
        for i in range(n):
            for j in range(i+1, n):
                if (i, j) not in significance_matrix:
                    G_ns.add_edge(i, j)
                    
        cliques = list(nx.find_cliques(G_ns))
        cliques.sort(key=lambda c: (-max(c), -len(c)))
        
        letters_vocab = "abcdefghijklmnopqrstuvwxyz"
        grouping_letters = {i: "" for i in range(n)}
        
        for idx, clq in enumerate(cliques):
            if idx < len(letters_vocab):
                let = letters_vocab[idx]
                for node in clq:
                    grouping_letters[node] += let
                    
        res = {labels[i]: "".join(sorted(grouping_letters[i])) for i in range(n)}
        return res

    def create_report(self):
        doc = Document()
        doc.add_heading('CRD Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        
        # Model
        doc.add_heading('1. Statistical Model (CRD)', level=1)
        doc.add_paragraph("Y_ij = mu + T_i + e_ij")
        
        # ANOVA
        doc.add_heading('2. ANOVA Summary', level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, t in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'P-value']):
            hdr[i].text = t
            
        for k in ["Treatments", "Error", "Total"]:
            row = table.add_row().cells
            dat = self.anova_table[k]
            row[0].text = k
            row[1].text = str(dat['df'])
            row[2].text = f"{dat['SS']:.2f}"
            row[3].text = f"{dat['MS']:.2f}" if dat['MS'] else ""
            row[4].text = f"{dat['F']:.2f}" if dat['F'] else ""
            if dat['P'] is not None:
                sig = "**" if dat['P']<=0.01 else ("*" if dat['P']<=0.05 else "ns")
                row[5].text = f"{dat['P']:.4f} {sig}"
        
        # Means
        doc.add_heading('3. Mean Comparison', level=1)
        table2 = doc.add_table(rows=1, cols=5)
        table2.style = 'Table Grid'
        h2 = table2.rows[0].cells
        for i, t in enumerate(['Treatment', 'Mean', 'Std Dev', 'Std Err', 'Group']):
            h2[i].text = t
            
        for m in self.results['means']:
             r = table2.add_row().cells
             r[0].text = m['level']
             r[1].text = f"{m['mean']:.2f}"
             r[2].text = f"{m['sd']:.2f}"
             r[3].text = f"{m['se']:.2f}"
             r[4].text = m['group']
             
        # Precision
        doc.add_paragraph(f"\nSE(m): {self.results['SEm']:.2f}")
        doc.add_paragraph(f"SE(d): {self.results['SEd']:.2f}")
        doc.add_paragraph(f"CV%: {self.results['CV']:.2f}%")
        
        # Interpretation
        doc.add_heading('4. Interpretation', level=1)
        p = self.anova_table["Treatments"]["P"]
        if p <= self.alpha:
            doc.add_paragraph(f"The analysis revealed significant differences among treatments (p={p:.4f}).")
            # Find best
            best = self.results['means'][0] # This is sorted by index now! Not by value.
            # Need to look at means dataframe
            means = self.results['df_means']
            best_t = means.idxmax()
            best_val = means.max()
            doc.add_paragraph(f"Treatment {best_t} recorded the highest mean value of {best_val:.2f}.")
        else:
            doc.add_paragraph("No significant differences were observed among treatments.")
            
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
