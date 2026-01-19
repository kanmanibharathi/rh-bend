import pandas as pd
import numpy as np
import scipy.stats as stats
import scikit_posthocs as sp
import networkx as nx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from .duncan_util import get_duncan_q

class StripPlotAnalyzer:
    def __init__(self, df, rep_col, a_col, b_col, resp_col):
        self.df = df
        self.rep_col = rep_col
        self.a_col = a_col
        self.b_col = b_col
        self.resp_col = resp_col
        
        # Dimensions
        self.r = 0
        self.a = 0
        self.b = 0
        
        self.anova_table = {}
        self.means = {}
        self.grouping = {}
        
        # Error Terms
        self.MS_Ea = 0
        self.df_Ea = 0
        self.MS_Eb = 0
        self.df_Eb = 0
        self.MS_Ec = 0
        self.df_Ec = 0
        
        self.alpha = 0.05
        
    def validate(self):
        # Convert factors to string
        self.df[self.rep_col] = self.df[self.rep_col].astype(str)
        self.df[self.a_col] = self.df[self.a_col].astype(str)
        self.df[self.b_col] = self.df[self.b_col].astype(str)
        
        # Numeric response
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        if self.df[self.resp_col].isnull().any():
            raise ValueError("Response variable contains missing or non-numeric values.")
            
        # Check dimensions
        reps = self.df[self.rep_col].unique()
        as_ = self.df[self.a_col].unique()
        bs = self.df[self.b_col].unique()
        
        self.r = len(reps)
        self.a = len(as_)
        self.b = len(bs)
        
        expected_n = self.r * self.a * self.b
        if len(self.df) != expected_n:
            raise ValueError(f"Unbalanced Design: Expected {expected_n} observations (r={self.r}, a={self.a}, b={self.b}), found {len(self.df)}.")
            
        # Check orthogonality (e.g., every Rep has every AxB combination)
        grouped = self.df.groupby([self.rep_col, self.a_col, self.b_col]).size()
        if not all(grouped == 1):
             raise ValueError("Design Validation Failed: Duplicate or missing combinations found. Each Rep*A*B combination must appear exactly once.")

    def run_anova(self):
        # Data vectors
        y = self.df[self.resp_col]
        G = y.sum()
        N = self.r * self.a * self.b
        CF = (G**2) / N
        
        # 6. Total SS
        SS_Total = (y**2).sum() - CF
        
        # 7. Replication SS
        # Sum over fixed Rep i
        R_totals = self.df.groupby(self.rep_col)[self.resp_col].sum()
        SS_Rep = (R_totals**2).sum() / (self.a * self.b) - CF
        
        # 8. Factor A (Horizontal) SS
        A_totals = self.df.groupby(self.a_col)[self.resp_col].sum()
        SS_A = (A_totals**2).sum() / (self.r * self.b) - CF
        
        # 9. Error A (Rep x A Interaction)
        # We need sums for each Rep x A combination
        RA_totals = self.df.groupby([self.rep_col, self.a_col])[self.resp_col].sum()
        SS_RA_cell = (RA_totals**2).sum() / self.b - CF
        SS_Ea = SS_RA_cell - SS_Rep - SS_A
        
        # 10. Factor B (Vertical) SS
        B_totals = self.df.groupby(self.b_col)[self.resp_col].sum()
        SS_B = (B_totals**2).sum() / (self.r * self.a) - CF
        
        # 11. Error B (Rep x B Interaction)
        RB_totals = self.df.groupby([self.rep_col, self.b_col])[self.resp_col].sum()
        SS_RB_cell = (RB_totals**2).sum() / self.a - CF
        SS_Eb = SS_RB_cell - SS_Rep - SS_B
        
        # 12. Interaction AxB SS
        AB_totals = self.df.groupby([self.a_col, self.b_col])[self.resp_col].sum()
        SS_AB_cell = (AB_totals**2).sum() / self.r - CF
        SS_AB = SS_AB_cell - SS_A - SS_B
        
        # 13. Error C (Residual)
        # SS_Ec = SS_Total - SS_Rep - SS_A - SS_Ea - SS_B - SS_Eb - SS_AB
        # Mathematically equivalent to: SS_Total - (SS_RA_cell + SS_RB_cell + SS_AB_cell) + 2*CF + SS_Rep ? 
        # Safer: SS_Total - sum(all others)
        SS_Ec = SS_Total - (SS_Rep + SS_A + SS_Ea + SS_B + SS_Eb + SS_AB)
        
        # 14. Degrees of Freedom
        df_rep = self.r - 1
        df_A = self.a - 1
        df_Ea = (self.r - 1) * (self.a - 1)
        df_B = self.b - 1
        df_Eb = (self.r - 1) * (self.b - 1)
        df_AB = (self.a - 1) * (self.b - 1)
        df_Ec = (self.r - 1) * (self.a - 1) * (self.b - 1)
        df_Total = N - 1
        
        # 15. Mean Squares
        MS_Rep = SS_Rep / df_rep
        MS_A = SS_A / df_A
        MS_Ea = SS_Ea / df_Ea
        MS_B = SS_B / df_B
        MS_Eb = SS_Eb / df_Eb
        MS_AB = SS_AB / df_AB
        MS_Ec = SS_Ec / df_Ec
        
        # Store Error MS for Post-hoc
        self.MS_Ea = MS_Ea
        self.df_Ea = df_Ea
        self.MS_Eb = MS_Eb
        self.df_Eb = df_Eb
        self.MS_Ec = MS_Ec
        self.df_Ec = df_Ec
        
        # 16. F-Statistics
        F_A = MS_A / MS_Ea if MS_Ea > 0 else 0
        F_B = MS_B / MS_Eb if MS_Eb > 0 else 0
        F_AB = MS_AB / MS_Ec if MS_Ec > 0 else 0
        
        # F-Test for Replication (Test against Error A - Horizontal Error)
        F_Rep = MS_Rep / MS_Ea if MS_Ea > 0 else 0
        
        # 17. P-Values
        P_A = 1 - stats.f.cdf(F_A, df_A, df_Ea)
        P_B = 1 - stats.f.cdf(F_B, df_B, df_Eb)
        P_AB = 1 - stats.f.cdf(F_AB, df_AB, df_Ec)
        P_Rep = 1 - stats.f.cdf(F_Rep, df_rep, df_Ea)
        
        self.anova_table = {
            "Replication": {"df": df_rep, "SS": SS_Rep, "MS": MS_Rep, "F": F_Rep, "P": P_Rep},
            "Factor A": {"df": df_A, "SS": SS_A, "MS": MS_A, "F": F_A, "P": P_A},
            "Error A": {"df": df_Ea, "SS": SS_Ea, "MS": MS_Ea, "F": None, "P": None},
            "Factor B": {"df": df_B, "SS": SS_B, "MS": MS_B, "F": F_B, "P": P_B},
            "Error B": {"df": df_Eb, "SS": SS_Eb, "MS": MS_Eb, "F": None, "P": None},
            "Interaction AxB": {"df": df_AB, "SS": SS_AB, "MS": MS_AB, "F": F_AB, "P": P_AB},
            "Error C": {"df": df_Ec, "SS": SS_Ec, "MS": MS_Ec, "F": None, "P": None},
            "Total": {"df": df_Total, "SS": SS_Total, "MS": None, "F": None, "P": None}
        }
        
        return self.anova_table

    def run_post_hoc(self, method='lsd', alpha=0.05, order='desc'):
        self.alpha = alpha
        is_ascending = True if order == 'asc' else False
        grand_mean = self.df[self.resp_col].mean()
        
        results = {}
        
        # --- Factor A ---
        # 1. Sort for Grouping Calculation (Value based)
        means_A_calc = self.df.groupby(self.a_col)[self.resp_col].mean().sort_values(ascending=is_ascending)
        SE_A = np.sqrt(self.MS_Ea / (self.r * self.b))
        SEd_A = np.sqrt(2) * SE_A
        CV_A = (np.sqrt(self.MS_Ea) / grand_mean) * 100
        
        if self.anova_table["Factor A"]["P"] <= alpha:
            eff_rep_A = self.r * self.b
            group_A_calc = self._compute_grouping(means_A_calc, method, alpha, SE_A, self.df_Ea, eff_rep_A)
        else:
             group_A_calc = {k: "ns" for k in means_A_calc.index}

        # 2. Sort for Display (Index/Alphabetical based)
        means_A_display = means_A_calc.sort_index()
        sds_A = self.df.groupby(self.a_col)[self.resp_col].std().sort_index()
        # Individual SEs
        counts_A = self.r * self.b
        ses_A = sds_A / np.sqrt(counts_A)
        
        results["Factor A"] = {
            "means": means_A_display, 
            "sds": sds_A,
            "ses": ses_A,
            "grouping": group_A_calc, 
            "SE": SE_A, 
            "SEd": SEd_A,
            "CV": CV_A,
            "CD": self._get_cd(method, alpha, self.df_Ea, SE_A, self.a)
        }

        # --- Factor B ---
        means_B_calc = self.df.groupby(self.b_col)[self.resp_col].mean().sort_values(ascending=is_ascending)
        SE_B = np.sqrt(self.MS_Eb / (self.r * self.a))
        SEd_B = np.sqrt(2) * SE_B
        CV_B = (np.sqrt(self.MS_Eb) / grand_mean) * 100
        
        if self.anova_table["Factor B"]["P"] <= alpha:
            eff_rep_B = self.r * self.a
            group_B_calc = self._compute_grouping(means_B_calc, method, alpha, SE_B, self.df_Eb, eff_rep_B)
        else:
             group_B_calc = {k: "ns" for k in means_B_calc.index}
             
        means_B_display = means_B_calc.sort_index()
        sds_B = self.df.groupby(self.b_col)[self.resp_col].std().sort_index()
        counts_B = self.r * self.a
        ses_B = sds_B / np.sqrt(counts_B)

        results["Factor B"] = {
            "means": means_B_display, 
            "sds": sds_B,
            "ses": ses_B,
            "grouping": group_B_calc, 
            "SE": SE_B, 
            "SEd": SEd_B,
            "CV": CV_B,
            "CD": self._get_cd(method, alpha, self.df_Eb, SE_B, self.b)
        }

        # --- Interaction AxB ---
        self.df['AxB'] = self.df[self.a_col].astype(str) + " : " + self.df[self.b_col].astype(str)
        means_AB_calc = self.df.groupby('AxB')[self.resp_col].mean().sort_values(ascending=is_ascending)
        
        SE_AB = np.sqrt(self.MS_Ec / self.r)
        SEd_AB = np.sqrt(2) * SE_AB
        CV_AB = (np.sqrt(self.MS_Ec) / grand_mean) * 100

        if self.anova_table["Interaction AxB"]["P"] <= alpha:
            group_AB_calc = self._compute_grouping(means_AB_calc, method, alpha, SE_AB, self.df_Ec, self.r)
        else:
            group_AB_calc = {k: "ns" for k in means_AB_calc.index}
            
        means_AB_display = means_AB_calc.sort_index()
        sds_AB = self.df.groupby('AxB')[self.resp_col].std().sort_index()
        counts_AB = self.r
        ses_AB = sds_AB / np.sqrt(counts_AB)

        results["Interaction AxB"] = {
            "means": means_AB_display, 
            "sds": sds_AB,
            "ses": ses_AB,
            "grouping": group_AB_calc, 
            "SE": SE_AB, 
            "SEd": SEd_AB,
            "CV": CV_AB,
            "CD": self._get_cd(method, alpha, self.df_Ec, SE_AB, self.a * self.b)
        }
            
        self.results = results
        return results

    def _get_cd(self, method, alpha, df, SE, n_means=2):
        if method == 'lsd':
            t_crit = stats.t.ppf(1 - alpha/2, df)
            return t_crit * (np.sqrt(2) * SE)
        elif method == 'tukey':
            # Tukey HSD = q(1-alpha, k, df) * SE
            # stats.studentized_range.ppf uses 1-alpha confidence
            q_crit = stats.studentized_range.ppf(1-alpha, n_means, df)
            return q_crit * SE
        elif method == 'duncan':
            # For display CD, we return the minimum significant range (p=2) or similar
            # D_2 = q(2, df) * SE
            q_crit = get_duncan_q(2, df, alpha)
            return q_crit * SE
        return None 

    def _compute_grouping(self, means, method, alpha, SE, df, eff_reps):
        labels = means.index.tolist()
        vals = means.values
        n = len(vals)
        significance_matrix = set()
        
        # 19. SEd
        SEd = np.sqrt(2) * SE
        
        if method == 'lsd':
            t_crit = stats.t.ppf(1 - alpha/2, df)
            crit_val = t_crit * SEd
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= crit_val:
                        significance_matrix.add((i, j))
                        
        elif method == 'tukey':
            # q_tukey(alpha, t, df)
            # stats.studentized_range.ppf(1-alpha, k, df)
            # HSD = q * SEm (not SEd) ?? 
            # User formula: HSD = q * sqrt(MS_error / eff_reps) -> q * SE
            q_crit = stats.studentized_range.ppf(1-alpha, n, df)
            crit_val = q_crit * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= crit_val:
                        significance_matrix.add((i, j))
        
        elif method == 'duncan':
            # Duncan's logic
            # Range p = number of steps between means
            for i in range(n):
                for j in range(i+1, n):
                    p = j - i + 1 # count of means in range inclusive
                    q_val = get_duncan_q(p, df, alpha)
                    D_p = q_val * SE
                    
                    if abs(vals[i] - vals[j]) >= D_p:
                         significance_matrix.add((i, j))
        
        # Grouping Logic (Clique Cover)
        G_ns = nx.Graph()
        G_ns.add_nodes_from(range(n))
        for i in range(n):
            for j in range(i+1, n):
                if (i, j) not in significance_matrix:
                    G_ns.add_edge(i, j)
                    
        cliques = list(nx.find_cliques(G_ns))
        cliques.sort(key=lambda c: (-max(c), -len(c)))
        
        letters_vocab = "abcdefghijklmnopqrstuvwxyz"
        grouping_letters = {i: "" for i in range(n)}
        
        for idx, clq in enumerate(cliques):
            if idx < len(letters_vocab):
                let = letters_vocab[idx]
                for node in clq:
                    grouping_letters[node] += let
                    
        res = {labels[i]: "".join(sorted(grouping_letters[i])) for i in range(n)}
        return res

    def create_report(self):
        doc = Document()
        doc.add_heading('Strip Plot Design Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 1. ANOVA Table
        doc.add_heading('ANOVA Summary', level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, t in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'P-value']):
            hdr[i].text = t
            
        keys = ["Replication", "Factor A", "Error A", "Factor B", "Error B", "Interaction AxB", "Error C", "Total"]
        for k in keys:
            row = table.add_row().cells
            dat = self.anova_table[k]
            row[0].text = k
            row[1].text = str(dat['df'])
            row[2].text = f"{dat['SS']:.4f}"
            row[3].text = f"{dat['MS']:.4f}" if dat['MS'] else "-"
            row[4].text = f"{dat['F']:.4f}" if dat['F'] else "-"
            
            pv = dat['P']
            if pv is not None:
                sig = "**" if pv <= 0.01 else ("*" if pv <= 0.05 else "ns")
                row[5].text = f"{pv:.4f} {sig}"
            else:
                row[5].text = "-"

        # 2. Main Effects & Interactions
        doc.add_heading('Mean Comparisions', level=1)
        
        def add_mean_table(title, effect_key):
             if effect_key in self.results:
                 doc.add_heading(title, level=2)
                 res = self.results[effect_key]
                 tbl = doc.add_table(rows=1, cols=3)
                 tbl.style = 'Table Grid'
                 h = tbl.rows[0].cells
                 h[0].text = "Level"
                 h[1].text = "Mean"
                 h[2].text = "Group"
                 
                 for level, mean in res['means'].items():
                     r = tbl.add_row().cells
                     r[0].text = str(level)
                     r[1].text = f"{mean:.4f}"
                     r[2].text = res['grouping'].get(level, "-")
                 
                 doc.add_paragraph(f"SE: {res['SE']:.4f} | CD (LSD): {res['CD'] if res['CD'] else 'N/A'}")

        add_mean_table("Factor A (Horizontal)", "Factor A")
        add_mean_table("Factor B (Vertical)", "Factor B")
        add_mean_table("Interaction (AxB)", "Interaction AxB")
        
        # Footer
        footer = doc.sections[0].footer.paragraphs[0]
        footer.text = "Analysis by Research Hub | Strip Plot Module"
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
