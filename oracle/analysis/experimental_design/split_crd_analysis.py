import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
import io
from datetime import datetime

class SplitCRDAnalyzer:
    def __init__(self, df, col_main, col_sub, col_rep, col_resp):
        self.df = df
        self.a_col = col_main
        self.b_col = col_sub
        self.r_col = col_rep
        self.resp_col = col_resp
        
        self.anova_table = {}
        self.stats = {}
        self.post_hoc_res = {}

    def validate(self):
        # Clean column names
        self.df.columns = self.df.columns.str.strip()
        
        # Validate columns exist
        for col in [self.a_col, self.b_col, self.r_col, self.resp_col]:
             if col not in self.df.columns:
                 raise ValueError(f"Column '{col}' not found in CSV.")
        
        # Clean data
        for col in [self.a_col, self.b_col, self.r_col]:
             self.df[col] = self.df[col].astype(str).str.strip()
        
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.resp_col])
        
        self.n_a = len(self.df[self.a_col].unique())
        self.n_b = len(self.df[self.b_col].unique())
        self.n_r = len(self.df[self.r_col].unique())
        
        if any(n < 2 for n in [self.n_a, self.n_b, self.n_r]):
             raise ValueError("All Factors and Replications must have at least 2 levels.")

    def run_anova(self):
        y = self.df[self.resp_col]
        G = y.sum()
        N = len(y)
        CF = (G**2) / N
        SS_Total = (y**2).sum() - CF
        
        def get_ss(cols, divisor):
            grp = self.df.groupby(cols)[self.resp_col].sum()
            return (grp**2).sum() / divisor - CF
            
        nA, nB, nR = self.n_a, self.n_b, self.n_r
        
        # 1. Replication (R)
        SS_R = get_ss(self.r_col, nA*nB)
        
        # 2. Main Plot (A)
        SS_A = get_ss(self.a_col, nB*nR)
        
        # 3. Error (a) - Main Plot Error
        SS_AR_Total = get_ss([self.a_col, self.r_col], nB)
        SS_Err_A = SS_AR_Total - SS_R - SS_A
        if SS_Err_A < 0: SS_Err_A = 0 # Safety
        
        # 4. Sub Plot (B)
        SS_B = get_ss(self.b_col, nA*nR)
        
        # 5. Interaction (A x B)
        SS_AB_Total = get_ss([self.a_col, self.b_col], nR)
        SS_AxB = SS_AB_Total - SS_A - SS_B
        
        # 6. Error (b) - Sub Plot Error
        SS_Err_B = SS_Total - (SS_R + SS_A + SS_Err_A + SS_B + SS_AxB)
        if SS_Err_B < 0: SS_Err_B = 0
        
        # DF
        df_r = int(nR - 1)
        df_a = int(nA - 1)
        df_err_a = int((nR - 1) * (nA - 1))
        
        df_b = int(nB - 1)
        df_axb = int((nA - 1) * (nB - 1))
        df_err_b = int(nA * (nR - 1) * (nB - 1))
        
        df_tot = int(nA * nB * nR - 1)
        
        # MS
        MS_R = SS_R / df_r if df_r > 0 else 0
        MS_A = SS_A / df_a if df_a > 0 else 0
        MS_Err_A = SS_Err_A / df_err_a if df_err_a > 0 else 0
        
        MS_B = SS_B / df_b if df_b > 0 else 0
        MS_AxB = SS_AxB / df_axb if df_axb > 0 else 0
        MS_Err_B = SS_Err_B / df_err_b if df_err_b > 0 else 0
        
        # F Tests
        F_A = MS_A / MS_Err_A if MS_Err_A > 0 else None
        F_B = MS_B / MS_Err_B if MS_Err_B > 0 else None
        F_AxB = MS_AxB / MS_Err_B if MS_Err_B > 0 else None
        
        # P Values
        def get_p(f, dfn, dfd):
            if f is None: return None
            return float(1 - stats.f.cdf(f, dfn, dfd))
        
        P_A = get_p(F_A, df_a, df_err_a)
        P_B = get_p(F_B, df_b, df_err_b)
        P_AxB = get_p(F_AxB, df_axb, df_err_b)
        
        # Helper to force native python types
        def s(df, ss, ms, f, p): 
            return {
                "df": int(df), 
                "SS": float(ss), 
                "MS": float(ms), 
                "F": float(f) if f is not None else None, 
                "P": float(p) if p is not None else None
            }
        
        self.anova_table = {
            "Replication": {"df": df_r, "SS": float(SS_R), "MS": float(MS_R), "F": None, "P": None},
            "Main Plot (A)": s(df_a, SS_A, MS_A, F_A, P_A),
            "Error (a)": {"df": df_err_a, "SS": float(SS_Err_A), "MS": float(MS_Err_A), "F": None, "P": None},
            
            "Sub Plot (B)": s(df_b, SS_B, MS_B, F_B, P_B),
            "Interaction A x B": s(df_axb, SS_AxB, MS_AxB, F_AxB, P_AxB),
            "Error (b)": {"df": df_err_b, "SS": float(SS_Err_B), "MS": float(MS_Err_B), "F": None, "P": None},
            
            "Total": {"df": df_tot, "SS": float(SS_Total), "MS": None, "F": None, "P": None}
        }

    def run_post_hoc(self, method, alpha):
        res = self.anova_table
        MS_Ea = res["Error (a)"]["MS"]
        DF_Ea = int(res["Error (a)"]["df"])
        
        MS_Eb = res["Error (b)"]["MS"]
        DF_Eb = int(res["Error (b)"]["df"])
        
        gm = self.df[self.resp_col].mean()
        
        cv_a = float((np.sqrt(MS_Ea) / gm) * 100) if gm != 0 else 0.0
        cv_b = float((np.sqrt(MS_Eb) / gm) * 100) if gm != 0 else 0.0
        
        def calc_sem(ms, div): return np.sqrt(ms / div) if div > 0 else 0
        
        # Main Plot statistics (Error A)
        sem_a = calc_sem(MS_Ea, self.n_r * self.n_b)
        
        # Sub Plot statistics (Error B)
        sem_b = calc_sem(MS_Eb, self.n_r * self.n_a)
        sem_axb = calc_sem(MS_Eb, self.n_r)
        
        def calc_cd(sem, df, k):
            if df <= 0: return 0.0
            sed = sem * np.sqrt(2)
            if method == 'lsd': return float(stats.t.ppf(1 - alpha/2, df) * sed)
            elif method == 'tukey': return float(stats.studentized_range.ppf(1-alpha, k, df) * sem)
            elif method == 'dmrt': return float(stats.t.ppf(1 - alpha/2, df) * sed)
            return 0.0
        
        stats_map = {
            "CV (a)": cv_a, "CV (b)": cv_b,
            "Main Plot (A)": {"SEm": float(sem_a), "SEd": float(sem_a*np.sqrt(2)), "CD": calc_cd(sem_a, DF_Ea, self.n_a)},
            "Sub Plot (B)": {"SEm": float(sem_b), "SEd": float(sem_b*np.sqrt(2)), "CD": calc_cd(sem_b, DF_Eb, self.n_b)},
            "Interaction A x B": {"SEm": float(sem_axb), "SEd": float(sem_axb*np.sqrt(2)), "CD": calc_cd(sem_axb, DF_Eb, self.n_a*self.n_b)}
        }
        self.stats = stats_map

        # Results
        results = {}
        def get_data(cols):
            grp = self.df.groupby(cols)[self.resp_col]
            means = grp.mean().sort_index()
            stds = grp.std().reindex(means.index)
            ses = (stds / np.sqrt(grp.count())).reindex(means.index)
            return means, stds, ses

        def compute_letters(means, sem, df, k):
             m_desc = means.sort_values(ascending=False)
             vals = m_desc.values
             keys = m_desc.index.tolist()
             n = len(vals)
             sed = sem * np.sqrt(2)
             sig_pairs = set()
             
             if df <= 0: return {k: "ns" for k in keys}

             if method == 'lsd':
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             elif method == 'tukey':
                 crit = stats.studentized_range.ppf(1-alpha, n, df) * sem
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             else: # DMRT
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             
             G = nx.Graph()
             G.add_nodes_from(range(n))
             for i in range(n):
                 for j in range(i+1, n):
                     if (i,j) not in sig_pairs: G.add_edge(i, j)
             
             cliques = list(nx.find_cliques(G))
             cliques.sort(key=lambda x: (min(x), -len(x)))
             vocab = "abcdefghijklmnopqrstuvwxyz"
             res_map = {i: "" for i in range(n)}
             for idx, clq in enumerate(cliques):
                 if idx < len(vocab):
                     c = vocab[idx]
                     for node in clq: res_map[node] += c
             return {keys[i]: "".join(sorted(res_map[i])) for i in range(n)}

        def proc_eff(name, cols, sem_key, df_err, n_lvls):
            m, s, se = get_data(cols)
            p = res[name]["P"]
            let = {k: "ns" for k in m.index}
            if p is not None and p < alpha:
                 let = compute_letters(m, stats_map[sem_key]["SEm"], df_err, n_lvls)
            results[sem_key] = {"means": m, "stds": s, "ses": se, "grouping": let}

        proc_eff("Main Plot (A)", self.a_col, "Main Plot (A)", DF_Ea, self.n_a)
        proc_eff("Sub Plot (B)", self.b_col, "Sub Plot (B)", DF_Eb, self.n_b)
        proc_eff("Interaction A x B", [self.a_col, self.b_col], "Interaction A x B", DF_Eb, self.n_a*self.n_b)
        
        self.post_hoc_res = results

    def create_report(self):
        doc = Document()
        doc.add_heading('Split CRD Analysis Report', 0)
        now_str = datetime.now().strftime('%Y-%m-%d %H:%M')
        doc.add_paragraph(f"Date: {now_str}")
        
        doc.add_heading('1. ANOVA Table', 1)
        t = doc.add_table(1, 6)
        t.style = 'Table Grid'
        def set_cell(r, i, txt): r.cells[i].text = str(txt)
        h = t.rows[0]
        for i,c in enumerate(["Source", "DF", "SS", "MS", "F", "Prob"]): set_cell(h, i, c)
        
        keys = ["Replication", "Main Plot (A)", "Error (a)",
                "Sub Plot (B)", "Interaction A x B", "Error (b)", "Total"]
        for k in keys:
            if k in self.anova_table:
                v = self.anova_table[k]
                r = t.add_row()
                sig = ""
                if v['P'] is not None:
                     if v['P']<0.01: sig="**"
                     elif v['P']<0.05: sig="*"
                     else: sig="ns"
                set_cell(r, 0, k)
                set_cell(r, 1, v['df'])
                set_cell(r, 2, f"{v['SS']:.4f}")
                set_cell(r, 3, f"{v['MS']:.4f}" if v['MS'] else "-")
                set_cell(r, 4, f"{v['F']:.4f}" if v['F'] is not None else "-")
                set_cell(r, 5, f"{v['P']:.4f} {sig}" if v['P'] is not None else "")
        
        doc.add_heading('2. Statistics', 1)
        doc.add_paragraph(f"CV (a) %: {self.stats['CV (a)']:.2f}")
        doc.add_paragraph(f"CV (b) %: {self.stats['CV (b)']:.2f}")
        
        stats_keys = ["Main Plot (A)", "Sub Plot (B)", "Interaction A x B"]
        for k in stats_keys:
             s = self.stats[k]
             doc.add_paragraph(f"{k}: SEm={s['SEm']:.4f}, SEd={s['SEd']:.4f}, CD={s['CD']:.4f}")
             
        doc.add_heading('3. Means and Grouping', 1)
        ph = self.post_hoc_res
        for k in stats_keys:
             ds = ph[k]
             doc.add_heading(f"{k} Means", 2)
             t = doc.add_table(1, 5)
             t.style = 'Table Grid'
             r = t.rows[0]
             for i,c in enumerate(["Level", "Mean", "Std.Dev", "Std.Err", "Group"]): set_cell(r, i, c)
             for lvl, val in ds['means'].items():
                 r = t.add_row()
                 if isinstance(lvl, tuple): l_str = " x ".join(map(str, lvl))
                 else: l_str = str(lvl)
                 set_cell(r, 0, l_str)
                 set_cell(r, 1, f"{val:.4f}")
                 set_cell(r, 2, f"{ds['stds'][lvl]:.4f}")
                 set_cell(r, 3, f"{ds['ses'][lvl]:.4f}")
                 set_cell(r, 4, ds['grouping'].get(lvl, ""))
                 
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
