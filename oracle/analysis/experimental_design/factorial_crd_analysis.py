import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import scikit_posthocs as sp
from .duncan_util import get_duncan_q

class FactorialCRDAnalyzer:
    def __init__(self, df, a_col, b_col, resp_col, rep_col=None):
        self.df = df
        self.a_col = a_col
        self.b_col = b_col
        self.resp_col = resp_col
        self.rep_col = rep_col
        
        self.a = 0
        self.b = 0
        self.n = 0
        
        self.anova_table = {}
        self.results = {}
        
        self.MS_E = 0
        self.df_E = 0
        self.grand_mean = 0
        self.r_bar = 0
        
        self.alpha = 0.05

    def validate(self):
        # Type conversion
        self.df[self.a_col] = self.df[self.a_col].astype(str)
        self.df[self.b_col] = self.df[self.b_col].astype(str)
        if self.rep_col:
            self.df[self.rep_col] = self.df[self.rep_col].astype(str)
            # Check for duplicates: one observation per combination of A, B, Rep
            dups = self.df.duplicated(subset=[self.a_col, self.b_col, self.rep_col])
            if dups.any():
                 raise ValueError(f"Duplicate entries found for the same Factor A, Factor B, and Replication combination. Each combination should have only one value for {self.resp_col}.")

        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        
        if self.df[self.resp_col].isnull().any():
             raise ValueError(f"Response variable '{self.resp_col}' contains missing or non-numeric values.")
             
        self.n = len(self.df)
        self.a = self.df[self.a_col].nunique()
        self.b = self.df[self.b_col].nunique()
        
        if self.n <= self.a * self.b:
            raise ValueError(f"Insufficient degrees of freedom for {self.resp_col}. Total observations ({self.n}) must be greater than number of treatment combinations ({self.a * self.b}).")
            
    def run_anova(self):
        y = self.df[self.resp_col]
        G = y.sum()
        self.grand_mean = y.mean()
        
        CF = (G**2) / self.n
        SS_Total = (y**2).sum() - CF
        
        # Factor A SS
        A_groups = self.df.groupby(self.a_col)[self.resp_col]
        term_A = sum((g.sum()**2 / len(g)) for _, g in A_groups)
        SS_A = term_A - CF
        
        # Factor B SS
        B_groups = self.df.groupby(self.b_col)[self.resp_col]
        term_B = sum((g.sum()**2 / len(g)) for _, g in B_groups)
        SS_B = term_B - CF
        
        # Interaction AB SS
        AB_groups = self.df.groupby([self.a_col, self.b_col])[self.resp_col]
        counts = AB_groups.count()
        
        # Check for balance
        is_balanced = counts.nunique() == 1
        
        term_Cells = sum((g.sum()**2 / len(g)) for _, g in AB_groups)
        SS_Cells = term_Cells - CF
        
        SS_AB = SS_Cells - SS_A - SS_B
        
        # Error SS
        SS_Error = SS_Total - SS_Cells 
        
        # Degrees of Freedom
        df_A = self.a - 1
        df_B = self.b - 1
        df_AB = df_A * df_B
        df_Total = self.n - 1
        df_Error = df_Total - (df_A + df_B + df_AB)
        
        if df_Error <= 0:
            raise ValueError(f"Error degrees of freedom is <= 0 for {self.resp_col}. Insufficient replications.")
            
        # Mean Squares
        MS_A = SS_A / df_A
        MS_B = SS_B / df_B
        MS_AB = SS_AB / df_AB
        MS_Error = SS_Error / df_Error
        
        self.MS_E = MS_Error
        self.df_E = df_Error
        
        # Calculate harmonic mean of replications for post-hoc if unbalanced
        if is_balanced:
            self.r_bar = counts.iloc[0]
        else:
            self.r_bar = stats.hmean(counts)
        
        # F-Statistics
        F_A = MS_A / MS_Error if MS_Error > 0 else 0
        F_B = MS_B / MS_Error if MS_Error > 0 else 0
        F_AB = MS_AB / MS_Error if MS_Error > 0 else 0
        
        # P-Values
        P_A = 1 - stats.f.cdf(F_A, df_A, df_Error)
        P_B = 1 - stats.f.cdf(F_B, df_B, df_Error)
        P_AB = 1 - stats.f.cdf(F_AB, df_AB, df_Error)
        
        self.anova_table = {
            "Factor A": {"df": df_A, "SS": SS_A, "MS": MS_A, "F": F_A, "P": P_A},
            "Factor B": {"df": df_B, "SS": SS_B, "MS": MS_B, "F": F_B, "P": P_B},
            "Interaction AxB": {"df": df_AB, "SS": SS_AB, "MS": MS_AB, "F": F_AB, "P": P_AB},
            "Error": {"df": df_Error, "SS": SS_Error, "MS": MS_Error, "F": None, "P": None},
            "Total": {"df": df_Total, "SS": SS_Total, "MS": None, "F": None, "P": None}
        }
        return self.anova_table

    def run_post_hoc(self, method='lsd', alpha=0.05, is_ascending=False, control_group=None, notation="alphabet"):
        self.alpha = alpha
        results = {}
        
        a_ctrl = None
        b_ctrl = None
        if control_group and " : " in control_group:
            parts = control_group.split(" : ")
            a_ctrl = parts[0]
            b_ctrl = parts[1]

        # --- Factor A ---
        means_A_calc = self.df.groupby(self.a_col)[self.resp_col].mean().sort_values(ascending=is_ascending)
        means_A_Display = means_A_calc.sort_index()
        sds_A = self.df.groupby(self.a_col)[self.resp_col].std().sort_index()
        counts_A = self.df.groupby(self.a_col)[self.resp_col].count().sort_index()
        ses_A = sds_A / np.sqrt(counts_A)
        
        SE_A = np.sqrt(self.MS_E / (self.b * self.r_bar))
        SEd_A = np.sqrt(2) * SE_A
        CV_A = (np.sqrt(self.MS_E) / self.grand_mean) * 100
        
        if self.anova_table["Factor A"]["P"] <= alpha:
             if method == 'dunnett' and a_ctrl is not None:
                  try:
                       p_matrix = sp.posthoc_dunnett(self.df, val_col=self.resp_col, group_col=self.a_col, control=a_ctrl)
                       group_A_calc = {a_ctrl: "a (Control)" if notation == "alphabet" else "Control"}
                       for trt in means_A_calc.index:
                            if trt == a_ctrl: continue
                            pval = 1.0
                            if trt in p_matrix.index and a_ctrl in p_matrix.columns:
                                 pval = p_matrix.loc[trt, a_ctrl]
                            elif a_ctrl in p_matrix.index and trt in p_matrix.columns:
                                 pval = p_matrix.loc[a_ctrl, trt]
                            
                            if pval < alpha:
                                 if notation == "alphabet":
                                      group_A_calc[trt] = "b"
                                 else:
                                      stars = "*"
                                      if pval < 0.01: stars = "**"
                                      if pval < 0.001: stars = "***"
                                      group_A_calc[trt] = stars
                            else:
                                 group_A_calc[trt] = "a" if notation == "alphabet" else "ns"
                  except Exception as e:
                       print(f"Dunnett A Error: {e}")
                       group_A_calc = {k: "err" for k in means_A_calc.index}
             else:
                  group_A_calc = self._compute_grouping(means_A_calc, method, alpha, SE_A, self.df_E, self.a) 
        else:
             group_A_calc = {k: "ns" for k in means_A_calc.index}

        results["Factor A"] = {
            "means": means_A_Display,
            "sds": sds_A,
            "ses": ses_A,
            "grouping": group_A_calc,
            "SE": SE_A,
            "SEd": SEd_A,
            "CV": CV_A,
            "CD": self._get_cd(method, alpha, self.df_E, SE_A, self.a)
        }

        # --- Factor B ---
        means_B_calc = self.df.groupby(self.b_col)[self.resp_col].mean().sort_values(ascending=is_ascending)
        means_B_Display = means_B_calc.sort_index()
        sds_B = self.df.groupby(self.b_col)[self.resp_col].std().sort_index()
        counts_B = self.df.groupby(self.b_col)[self.resp_col].count().sort_index()
        ses_B = sds_B / np.sqrt(counts_B)
        
        SE_B = np.sqrt(self.MS_E / (self.a * self.r_bar))
        SEd_B = np.sqrt(2) * SE_B
        CV_B = (np.sqrt(self.MS_E) / self.grand_mean) * 100
        
        if self.anova_table["Factor B"]["P"] <= alpha:
             if method == 'dunnett' and b_ctrl is not None:
                  try:
                       p_matrix = sp.posthoc_dunnett(self.df, val_col=self.resp_col, group_col=self.b_col, control=b_ctrl)
                       group_B_calc = {b_ctrl: "a (Control)" if notation == "alphabet" else "Control"}
                       for trt in means_B_calc.index:
                            if trt == b_ctrl: continue
                            pval = 1.0
                            if trt in p_matrix.index and b_ctrl in p_matrix.columns:
                                 pval = p_matrix.loc[trt, b_ctrl]
                            elif b_ctrl in p_matrix.index and trt in p_matrix.columns:
                                 pval = p_matrix.loc[b_ctrl, trt]
                            
                            if pval < alpha:
                                 if notation == "alphabet":
                                      group_B_calc[trt] = "b"
                                 else:
                                      stars = "*"
                                      if pval < 0.01: stars = "**"
                                      if pval < 0.001: stars = "***"
                                      group_B_calc[trt] = stars
                            else:
                                 group_B_calc[trt] = "a" if notation == "alphabet" else "ns"
                  except Exception as e:
                       print(f"Dunnett B Error: {e}")
                       group_B_calc = {k: "err" for k in means_B_calc.index}
             else:
                  group_B_calc = self._compute_grouping(means_B_calc, method, alpha, SE_B, self.df_E, self.b)
        else:
             group_B_calc = {k: "ns" for k in means_B_calc.index}

        results["Factor B"] = {
            "means": means_B_Display,
            "sds": sds_B,
            "ses": ses_B,
            "grouping": group_B_calc,
            "SE": SE_B,
            "SEd": SEd_B,
            "CV": CV_B,
            "CD": self._get_cd(method, alpha, self.df_E, SE_B, self.b)
        }

        # --- Interaction AB ---
        self.df['AxB'] = self.df[self.a_col].astype(str) + " : " + self.df[self.b_col].astype(str)
        means_AB_calc = self.df.groupby('AxB')[self.resp_col].mean().sort_values(ascending=is_ascending)
        means_AB_Display = means_AB_calc.sort_index()
        sds_AB = self.df.groupby('AxB')[self.resp_col].std().sort_index()
        counts_AB = self.df.groupby('AxB')[self.resp_col].count().sort_index()
        ses_AB = sds_AB / np.sqrt(counts_AB)
        
        SE_AB = np.sqrt(self.MS_E / self.r_bar)
        SEd_AB = np.sqrt(2) * SE_AB
        CV_AB = (np.sqrt(self.MS_E) / self.grand_mean) * 100
        
        if self.anova_table["Interaction AxB"]["P"] <= alpha:
             if method == 'dunnett' and control_group is not None:
                  try:
                       p_matrix = sp.posthoc_dunnett(self.df, val_col=self.resp_col, group_col='AxB', control=control_group)
                       group_AB_calc = {control_group: "a (Control)" if notation == "alphabet" else "Control"}
                       for trt in means_AB_calc.index:
                            if trt == control_group: continue
                            pval = 1.0
                            if trt in p_matrix.index and control_group in p_matrix.columns:
                                 pval = p_matrix.loc[trt, control_group]
                            elif control_group in p_matrix.index and trt in p_matrix.columns:
                                 pval = p_matrix.loc[control_group, trt]
                            
                            if pval < alpha:
                                 if notation == "alphabet":
                                      group_AB_calc[trt] = "b"
                                 else:
                                      stars = "*"
                                      if pval < 0.01: stars = "**"
                                      if pval < 0.001: stars = "***"
                                      group_AB_calc[trt] = stars
                            else:
                                 group_AB_calc[trt] = "a" if notation == "alphabet" else "ns"
                  except Exception as e:
                       print(f"Dunnett Error: {e}")
                       group_AB_calc = {k: "err" for k in means_AB_calc.index}
             else:
                  group_AB_calc = self._compute_grouping(means_AB_calc, method, alpha, SE_AB, self.df_E, self.a * self.b)
        else:
             group_AB_calc = {k: "ns" for k in means_AB_calc.index}

        results["Interaction AxB"] = {
            "means": means_AB_Display,
            "sds": sds_AB,
            "ses": ses_AB,
            "grouping": group_AB_calc,
            "SE": SE_AB,
            "SEd": SEd_AB,
            "CV": CV_AB,
            "CD": self._get_cd(method, alpha, self.df_E, SE_AB, self.a * self.b)
        }
        
        self.results = results
        return results

    def _get_cd(self, method, alpha, df, SE, n_means=2):
        if method == 'lsd':
            t_crit = stats.t.ppf(1 - alpha/2, df)
            return t_crit * (np.sqrt(2) * SE)
        elif method == 'tukey':
            q_crit = stats.studentized_range.ppf(1-alpha, n_means, df)
            return q_crit * SE
        elif method == 'duncan':
            q_crit = get_duncan_q(2, df, alpha)
            return q_crit * SE
        elif method == 'dunnett':
            # Simplified Dunnett critical value approximation:
            # Using t-distribution with alpha/k adjustment for 1-sided vs control
            # or approx t * 1.1. Let's use a slightly more conservative t-value.
            k = n_means - 1
            if k <= 0: return 0
            t_crit = stats.t.ppf(1 - alpha/k, df) # Bonferroni-style adjustment
            return t_crit * (np.sqrt(2) * SE)
        return None

    def _compute_grouping(self, means, method, alpha, SE, df, n_means_total=None):
        labels = means.index.tolist()
        vals = means.values
        n = len(vals)
        if n == 0: return {}
        significance_matrix = set()
        
        SEd = np.sqrt(2) * SE
        
        if method == 'lsd':
            t_crit = stats.t.ppf(1 - alpha/2, df)
            crit_val = t_crit * SEd
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= crit_val:
                        significance_matrix.add((i, j))
                        
        elif method == 'tukey':
            q_crit = stats.studentized_range.ppf(1-alpha, n, df)
            crit_val = q_crit * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= crit_val:
                         significance_matrix.add((i, j))
                         
        elif method == 'duncan':
            for i in range(n):
                for j in range(i+1, n):
                    p = j - i + 1
                    q_val = get_duncan_q(p, df, alpha)
                    D_p = q_val * SE
                    if abs(vals[i] - vals[j]) >= D_p:
                        significance_matrix.add((i, j))
                        
        G_ns = nx.Graph()
        G_ns.add_nodes_from(range(n))
        for i in range(n):
            for j in range(i+1, n):
                if (i, j) not in significance_matrix:
                    G_ns.add_edge(i, j)
        
        cliques = list(nx.find_cliques(G_ns))
        cliques.sort(key=lambda c: (-max(c), -len(c)))
        
        letters_vocab = "abcdefghijklmnopqrstuvwxyz"
        grouping_letters = {i: "" for i in range(n)}
        
        for idx, clq in enumerate(cliques):
            if idx < len(letters_vocab):
                let = letters_vocab[idx]
                for node in clq:
                    grouping_letters[node] += let
                    
        res = {labels[i]: "".join(sorted(grouping_letters[i])) for i in range(n)}
        return res

    def append_to_report(self, doc):
        doc.add_heading(f"Analysis for Response: {self.resp_col}", level=1)
        
        # ANOVA
        doc.add_heading('ANOVA Summary', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, t in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'P-value']):
            hdr[i].text = t
            
        keys = ["Factor A", "Factor B", "Interaction AxB", "Error", "Total"]
        for k in keys:
            row = table.add_row().cells
            dat = self.anova_table[k]
            row[0].text = k
            row[1].text = str(dat['df'])
            row[2].text = f"{dat['SS']:.2f}"
            row[3].text = f"{dat['MS']:.2f}" if dat['MS'] else ""
            row[4].text = f"{dat['F']:.2f}" if dat['F'] else ""
            if dat['P'] is not None:
                pv = dat['P']
                sig = "**" if pv <= 0.01 else ("*" if pv <= 0.05 else "ns")
                row[5].text = f"{pv:.4f} {sig}"
        
        def add_mean_table(title, effect_key):
             if effect_key in self.results:
                 doc.add_heading(title, level=2)
                 res = self.results[effect_key]
                 tbl = doc.add_table(rows=1, cols=4)
                 tbl.style = 'Table Grid'
                 h = tbl.rows[0].cells
                 h[0].text = "Level"
                 h[1].text = "Mean"
                 h[2].text = "Std Err"
                 h[3].text = "Group"
                 
                 for level, mean in res['means'].items():
                     r = tbl.add_row().cells
                     se = res['ses'].get(level, 0.0)
                     r[0].text = str(level)
                     r[1].text = f"{mean:.2f}"
                     r[2].text = f"{se:.2f}"
                     r[3].text = res['grouping'].get(level, '-')
                 
                 doc.add_paragraph(f"SE(m): {res['SE']:.2f} | SE(d): {res['SEd']:.2f}")
                 doc.add_paragraph(f"CD ({self.alpha}): {res['CD']:.2f} | CV%: {res['CV']:.2f}%")

        add_mean_table("Factor A Means", "Factor A")
        add_mean_table("Factor B Means", "Factor B")
        add_mean_table("Interaction (AxB) Means", "Interaction AxB")
        doc.add_page_break()

    def create_report(self):
        doc = Document()
        doc.add_heading('Factorial CRD Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        self.append_to_report(doc)
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

