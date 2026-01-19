import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from .duncan_util import get_duncan_q

class TwoFactorPooledCRDAnalyzer:
    def __init__(self, df, factor_a_col, factor_b_col, year_col, resp_col):
        self.df = df
        self.a_col = factor_a_col
        self.b_col = factor_b_col
        self.y_col = year_col
        self.resp_col = resp_col
        
        self.bartlett_res = {}
        self.anova_table = {}
        self.post_hoc_res = {}
        
        # Metadata
        self.a_levels = []
        self.b_levels = []
        self.y_levels = []
        self.n_reps_harm = 1

    def validate(self):
        # Cleanup
        for col in [self.a_col, self.b_col, self.y_col]:
            self.df[col] = self.df[col].astype(str).str.strip()
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.resp_col])

        self.a_levels = sorted(self.df[self.a_col].unique())
        self.b_levels = sorted(self.df[self.b_col].unique())
        self.y_levels = sorted(self.df[self.y_col].unique())
        
        if len(self.a_levels) < 2: raise ValueError("Factor A must have at least 2 levels.")
        if len(self.b_levels) < 2: raise ValueError("Factor B must have at least 2 levels.")
        if len(self.y_levels) < 2: raise ValueError("Year/Location must have at least 2 levels.")
        
        # Calculate harmonic mean of replicates per cell (A x B x Y)
        # Model assumes 'r' reps per cell
        grp = self.df.groupby([self.a_col, self.b_col, self.y_col])[self.resp_col].count()
        if len(grp) > 0:
            recip = sum(1.0/c for c in grp)
            self.n_reps_harm = len(grp) / recip
        else:
            self.n_reps_harm = 1

    def run_bartlett_test(self):
        # Calculate MSE for each Year individually
        # For each Year, we have a Factorial CRD (A x B)
        # But simpler: For homogeneity of error variance, we just look at the Residuals of the model within each year
        # Or even simpler: Treat each Year as a dataset. 
        # In each Year, Model is Y = mu + A + B + AB + E
        # We need the MSE from that local ANOVA.
        
        mse_list = []
        df_list = []
        
        for yr, group in self.df.groupby(self.y_col):
            # Run local 2-way factorial CRD to get MSE
            # SS_Total_Local
            y = group[self.resp_col]
            N_i = len(y)
            CF_i = (y.sum()**2) / N_i
            SS_Tot_i = (y**2).sum() - CF_i
            
            # SS_Cells (A x B combinations in this year)
            grp_cells = group.groupby([self.a_col, self.b_col])[self.resp_col]
            SS_Cells_i = sum((g.sum()**2 / len(g)) for _, g in grp_cells) - CF_i
            
            # SS_Error = Total - Cells
            SS_Err_i = SS_Tot_i - SS_Cells_i
            
            # DF Error = N - (a * b)  (assuming full cells)
            n_cells = len(grp_cells)
            DF_Err_i = N_i - n_cells
            
            if DF_Err_i > 0 and SS_Err_i > 1e-9:
                mse_list.append(SS_Err_i / DF_Err_i)
                df_list.append(DF_Err_i)
                
        k = len(mse_list)
        if k < 2:
            self.bartlett_res = {"valid": True, "stat": 0.0, "p": 1.0, "df": 0, "msg": "Insufficient years for test."}
            return

        sum_df = sum(df_list)
        numerator = sum(d * s for d, s in zip(df_list, mse_list))
        sp_sq = numerator / sum_df
        
        term1 = sum_df * np.log(sp_sq)
        term2 = sum(d * np.log(s) for d, s in zip(df_list, mse_list))
        chi_sq = term1 - term2
        
        inv_sum = sum(1/d for d in df_list)
        C = 1 + (1 / (3 * (k - 1))) * (inv_sum - (1 / sum_df))
        
        corr_chi_sq = chi_sq / C
        p_val = 1 - stats.chi2.cdf(corr_chi_sq, k - 1)
        
        is_homo = bool(p_val >= 0.05)
        msg = "Variances are homogeneous." if is_homo else "Variances are heterogeneous."
        
        self.bartlett_res = {
            "valid": is_homo,
            "stat": float(corr_chi_sq),
            "p": float(p_val),
            "df": int(k - 1),
            "msg": msg
        }

    def run_anova(self):
        # Y_ijkl = mu + Yi + Aj + (AY)ij + Bk + (BY)ik + (AB)jk + (ABY)ijk + e
        
        y = self.df[self.resp_col]
        N = len(y)
        G = y.sum()
        CF = (G**2) / N
        SS_Total = (y**2).sum() - CF
        
        # Dimensions (Observed)
        a = len(self.a_levels)
        b = len(self.b_levels)
        y_n = len(self.y_levels)
        
        # 1. Main Effects
        # Year
        grp_y = self.df.groupby(self.y_col)[self.resp_col]
        SS_Y = sum((g.sum()**2 / len(g)) for _, g in grp_y) - CF
        
        # A
        grp_a = self.df.groupby(self.a_col)[self.resp_col]
        SS_A = sum((g.sum()**2 / len(g)) for _, g in grp_a) - CF
        
        # B
        grp_b = self.df.groupby(self.b_col)[self.resp_col]
        SS_B = sum((g.sum()**2 / len(g)) for _, g in grp_b) - CF
        
        # 2. Two-Way Interactions
        # A x Y (Sum(YiAj)^2 / br - CF - SSY - SSA)
        grp_ay = self.df.groupby([self.a_col, self.y_col])[self.resp_col]
        SS_AY_sub = sum((g.sum()**2 / len(g)) for _, g in grp_ay) - CF
        SS_AY = SS_AY_sub - SS_Y - SS_A
        
        # B x Y
        grp_by = self.df.groupby([self.b_col, self.y_col])[self.resp_col]
        SS_BY_sub = sum((g.sum()**2 / len(g)) for _, g in grp_by) - CF
        SS_BY = SS_BY_sub - SS_Y - SS_B
        
        # A x B
        grp_ab = self.df.groupby([self.a_col, self.b_col])[self.resp_col]
        SS_AB_sub = sum((g.sum()**2 / len(g)) for _, g in grp_ab) - CF
        SS_AB = SS_AB_sub - SS_A - SS_B
        
        # 3. Three-Way Interaction
        # A x B x Y (Cells)
        grp_aby = self.df.groupby([self.a_col, self.b_col, self.y_col])[self.resp_col]
        SS_Cells = sum((g.sum()**2 / len(g)) for _, g in grp_aby) - CF
        
        SS_ABY = SS_Cells - (SS_Y + SS_A + SS_B + SS_AY + SS_BY + SS_AB)
        
        # 4. Error
        SS_Error = SS_Total - SS_Cells
        
        # Degrees of Freedom
        # Assuming Data is somewhat balanced for DF calculation, or using generalized logic
        # For simple reporting we use standard factorial DFs
        DF_Y = y_n - 1
        DF_A = a - 1
        DF_B = b - 1
        
        DF_AY = DF_A * DF_Y
        DF_BY = DF_B * DF_Y
        DF_AB = DF_A * DF_B
        DF_ABY = DF_A * DF_B * DF_Y
        
        # Error DF: N - number of cells
        n_cells = len(grp_aby)
        DF_Error = N - n_cells
        DF_Total = N - 1
        
        # Mean Squares & F
        # All tested against Error
        MS_E = SS_Error / DF_Error
        
        sources = [
            ("Year", DF_Y, SS_Y),
            ("Factor A", DF_A, SS_A),
            ("Factor B", DF_B, SS_B),
            ("A x Year", DF_AY, SS_AY),
            ("B x Year", DF_BY, SS_BY),
            ("A x B", DF_AB, SS_AB),
            ("A x B x Year", DF_ABY, SS_ABY)
        ]
        
        res = {}
        for name, df, ss in sources:
            ms = ss / df if df > 0 else 0
            f = ms / MS_E if MS_E > 0 else 0
            p = 1 - stats.f.cdf(f, df, DF_Error)
            res[name] = {"df": int(df), "SS": float(ss), "MS": float(ms), "F": float(f), "P": float(p)}
            
        res["Error"] = {"df": int(DF_Error), "SS": float(SS_Error), "MS": float(MS_E), "F": None, "P": None}
        res["Total"] = {"df": int(DF_Total), "SS": float(SS_Total), "MS": None, "F": None, "P": None}
        
        self.anova_table = res

    def run_post_hoc(self, method, alpha):
        # Decision Logic:
        # Check ABY -> if Sig, report ABY means, stop.
        # Check AB -> if Sig, report AB means.
        # Check Main Effects -> report if sig.
        
        res = self.anova_table
        MS_E = res["Error"]["MS"]
        DF_E = res["Error"]["df"]
        
        p_aby = res["A x B x Year"]["P"]
        p_ab = res["A x B"]["P"]
        
        sig_aby = p_aby < alpha
        sig_ab = p_ab < alpha
        
        ph_results = {}
        
        # Helper for stats
        def calc_stats(means, n_reps_for_mean, k_means):
            sem = np.sqrt(MS_E / n_reps_for_mean)
            sed = sem * np.sqrt(2)
            cd = 0
            if method == 'lsd':
                cd = stats.t.ppf(1 - alpha/2, DF_E) * sed
            elif method == 'tukey':
                cd = stats.studentized_range.ppf(1-alpha, k_means, DF_E) * sem
            return sem, sed, cd

        # 1. Check 3-way
        if sig_aby:
            # Report A x B x Y Interaction Means
            grp = self.df.groupby([self.a_col, self.b_col, self.y_col])[self.resp_col]
            means = grp.mean()
            # r = n_reps_harm
            sem, sed, cd = calc_stats(means, self.n_reps_harm, len(means))
            
            # Since 3-way table is huge, we will structure it by Year
            # List of {year: ..., data: table}
            tables = []
            for yr in self.y_levels:
                 sub = means.xs(yr, level=2) if len(means.index.names) == 3 else means # Handle potential index issues
                 # Reconstruct proper sub-table if needed or just dump lines
                 # Actually xs on multiindex: Level 0=A, 1=B, 2=Y
                 # means index is (A, B, Y)
                 # Let's pivot for display: Row=A, Col=B
                 try:
                    # Select data for this year
                    # We can't easily use xs if index types mismatch, safer to filter df
                    sub_df = self.df[self.df[self.y_col] == yr]
                    sub_piv = sub_df.groupby([self.a_col, self.b_col])[self.resp_col].mean().unstack()
                    tables.append({"year": yr, "pivot": sub_piv})
                 except: continue

            ph_results["Decision"] = "Significant A x B x Year Interaction detected. Lower order comparisons suppressed."
            ph_results["ABY"] = {
                "tables": tables,
                "sem": float(sem), "sed": float(sed), "cd": float(cd), "sig": True
            }
            # STOP
        
        elif sig_ab:
            # Report A x B Means
            grp = self.df.groupby([self.a_col, self.b_col])[self.resp_col]
            means = grp.mean() # Index (A, B)
            
            # n_reps for AB mean = r * y
            n_r = self.n_reps_harm * len(self.y_levels)
            sem, sed, cd = calc_stats(means, n_r, len(means))
            
            # Pivot for neat table
            piv = means.unstack()
            
            ph_results["Decision"] = "Significant A x B Interaction detected. Main effects comparison suppressed."
            ph_results["AB"] = {
                "pivot": piv,
                "sem": float(sem), "sed": float(sed), "cd": float(cd), "sig": True,
                "grouping": self._compute_interaction_grouping(means, method, alpha, DF_E, MS_E, n_r)
            }
            
        else:
            # Main Effects
            ph_results["Decision"] = "No significant High-Order Interactions. Analyzing Main Effects."
            
            # Factor A
            if res["Factor A"]["P"] < alpha:
                # Calculate mean, std, sem per level
                stats_a = self.df.groupby(self.a_col)[self.resp_col].agg(['mean', 'std', 'sem']).sort_index()
                means_a = stats_a['mean']
                
                # n_reps = r * b * y
                n_r_a = self.n_reps_harm * len(self.b_levels) * len(self.y_levels)
                sem, sed, cd = calc_stats(means_a, n_r_a, len(means_a))
                grp = self._compute_grouping(means_a, method, alpha, DF_E, MS_E, n_r_a)
                
                ph_results["A"] = {
                    "means": means_a,
                    "stds": stats_a['std'],
                    "ses": stats_a['sem'],
                    "grouping": grp, 
                    "sem": float(sem), "sed": float(sed), "cd": float(cd), "sig": True
                }
            
            # Factor B
            if res["Factor B"]["P"] < alpha:
                stats_b = self.df.groupby(self.b_col)[self.resp_col].agg(['mean', 'std', 'sem']).sort_index()
                means_b = stats_b['mean']
                
                n_r_b = self.n_reps_harm * len(self.a_levels) * len(self.y_levels)
                sem, sed, cd = calc_stats(means_b, n_r_b, len(means_b))
                grp = self._compute_grouping(means_b, method, alpha, DF_E, MS_E, n_r_b)
                
                ph_results["B"] = {
                    "means": means_b, 
                    "stds": stats_b['std'],
                    "ses": stats_b['sem'],
                    "grouping": grp, 
                    "sem": float(sem), "sed": float(sed), "cd": float(cd), "sig": True
                }
                
            # Year
            if res["Year"]["P"] < alpha:
                stats_y = self.df.groupby(self.y_col)[self.resp_col].agg(['mean', 'std', 'sem']).sort_index()
                means_y = stats_y['mean']
                
                n_r_y = self.n_reps_harm * len(self.a_levels) * len(self.b_levels)
                sem, sed, cd = calc_stats(means_y, n_r_y, len(means_y))
                grp = self._compute_grouping(means_y, method, alpha, DF_E, MS_E, n_r_y)
                
                ph_results["Year"] = {
                    "means": means_y,
                    "stds": stats_y['std'],
                    "ses": stats_y['sem'],
                    "grouping": grp, 
                    "sem": float(sem), "sed": float(sed), "cd": float(cd), "sig": True
                }

        # Global CV
        gm = self.df[self.resp_col].mean()
        cv = (np.sqrt(MS_E) / gm) * 100
        ph_results["CV"] = float(cv)
        ph_results["Method"] = method.upper()
        
        self.post_hoc_res = ph_results

    def _compute_grouping(self, means, method, alpha, df, MSE, n_reps):
        # Standard logic as before
        vals = means.values
        labels = means.index.tolist()
        n = len(vals)
        sig = set()
        
        SE = np.sqrt(MSE / n_reps)
        
        if method == 'lsd':
            crit = stats.t.ppf(1 - alpha/2, df)
            LSD = crit * np.sqrt(2) * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= LSD: sig.add((i, j))
        elif method == 'tukey':
            crit = stats.studentized_range.ppf(1-alpha, n, df)
            HSD = crit * SE
            for i in range(n):
                for j in range(i+1, n):
                    if abs(vals[i] - vals[j]) >= HSD: sig.add((i, j))
        elif method == 'dmrt':
            for i in range(n):
                for j in range(i+1, n):
                    p = j - i + 1
                    q_val = get_duncan_q(p, df, alpha)
                    Rp = q_val * SE
                    if abs(vals[i] - vals[j]) >= Rp: sig.add((i, j))
                    
        G = nx.Graph()
        G.add_nodes_from(range(n))
        for i in range(n):
            for j in range(i+1, n):
                if (i, j) not in sig: G.add_edge(i, j)
        
        cliques = list(nx.find_cliques(G))
        cliques.sort(key=lambda c: (min(c), -len(c)))
        vocab = "abcdefghijklmnopqrstuvwxyz"
        res = {i: "" for i in range(n)}
        
        for idx, clq in enumerate(cliques):
            if idx < len(vocab):
                char = vocab[idx]
                for node in clq: res[node] += char
                
        return {labels[i]: "".join(sorted(res[i])) for i in range(n)}

    def _compute_interaction_grouping(self, means_series, method, alpha, df, MSE, n_reps):
         # Flattened interaction grouping
         # Means series index is MultiIndex (A, B)
         # We treat each (Ai, Bk) as a treatment
         return self._compute_grouping(means_series, method, alpha, df, MSE, n_reps)

    def create_report(self):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        doc.add_heading('Two-Factor Pooled CRD Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # Bartlett
        doc.add_heading('1. Homogeneity of Variance', level=1)
        if self.bartlett_res:
             doc.add_paragraph(f"Statistic: {self.bartlett_res['stat']:.4f}, P: {self.bartlett_res['p']:.4f}")
             doc.add_paragraph(f"Decision: {self.bartlett_res['msg']}")
             
        # ANOVA
        doc.add_heading('2. ANOVA Table', level=1)
        if self.anova_table:
            t = doc.add_table(1, 6)
            t.style = 'Table Grid'
            h = t.rows[0].cells
            for i, c in enumerate(["Source", "DF", "SS", "MS", "F", "Prob"]): h[i].text = c
            order = ["Year", "Factor A", "Factor B", "A x Year", "B x Year", "A x B", "A x B x Year", "Error", "Total"]
            for k in order:
                if k in self.anova_table:
                    v = self.anova_table[k]
                    r = t.add_row().cells
                    r[0].text = k
                    r[1].text = str(v['df'])
                    r[2].text = f"{v['SS']:.4f}"
                    r[3].text = f"{v['MS']:.4f}" if v['MS'] else "-"
                    r[4].text = f"{v['F']:.4f}" if v['F'] else "-"
                    if v['P'] is not None:
                        sig = "**" if v['P'] < 0.01 else "*" if v['P'] < 0.05 else "ns"
                        r[5].text = f"{v['P']:.4f} {sig}"
        
        # PostHoc
        doc.add_heading('3. Means & Interactions', level=1)
        ph = self.post_hoc_res
        if ph:
            doc.add_paragraph(f"Decision Rule Applied: {ph.get('Decision', '')}")
            doc.add_paragraph(f"CV%: {ph.get('CV', 0):.2f}")
            
            # Depending on structure
            if "ABY" in ph:
                dat = ph["ABY"]
                doc.add_heading('A x B x Year Interaction', 2)
                doc.add_paragraph(f"SEm: {dat['sem']:.4f} | SEd: {dat['sed']:.4f} | CD: {dat['cd']:.4f}")
                
                for item in dat["tables"]:
                    doc.add_paragraph(f"Year: {item['year']}")
                    # Pivot table dump
                    piv = item['pivot'] # DataFrame
                    # Add table
                    t = doc.add_table(piv.shape[0]+1, piv.shape[1]+1)
                    t.style = 'Table Grid'
                    # Header
                    t.cell(0,0).text = "A/B"
                    for j, col in enumerate(piv.columns):
                        t.cell(0, j+1).text = str(col)
                    # Rows
                    for i, (idx, row) in enumerate(piv.iterrows()):
                        t.cell(i+1, 0).text = str(idx)
                        for j, val in enumerate(row):
                            t.cell(i+1, j+1).text = f"{val:.4f}"
                    doc.add_paragraph("")
                    
            elif "AB" in ph:
                 dat = ph["AB"]
                 doc.add_heading('A x B Interaction', 2)
                 doc.add_paragraph(f"SEm: {dat['sem']:.4f} | SEd: {dat['sed']:.4f} | CD: {dat['cd']:.4f}")
                 # Pivot Table
                 piv = dat['pivot']
                 t = doc.add_table(piv.shape[0]+1, piv.shape[1]+1)
                 t.style = 'Table Grid'
                 t.cell(0,0).text = "A/B"
                 for j, col in enumerate(piv.columns):
                     t.cell(0, j+1).text = str(col)
                 for i, (idx, row) in enumerate(piv.iterrows()):
                        t.cell(i+1, 0).text = str(idx)
                        for j, val in enumerate(row):
                            t.cell(i+1, j+1).text = f"{val:.4f}"
            else:
                # Main Effects
                for eff in ["A", "B", "Year"]:
                    if eff in ph:
                        d = ph[eff]
                        doc.add_heading(f"{eff} Means", 2)
                        doc.add_paragraph(f"SEm: {d['sem']:.4f} | SEd: {d['sed']:.4f} | CD: {d['cd']:.4f}")
                        t = doc.add_table(1, 5) # Added 2 columns
                        t.style = 'Table Grid'
                        h = t.rows[0].cells
                        h[0].text = "Level"
                        h[1].text = "Mean"
                        h[2].text = "St.Dev"
                        h[3].text = "St.Err"
                        h[4].text = "Group"
                        for lvl, val in d['means'].items():
                            r = t.add_row().cells
                            r[0].text = str(lvl)
                            r[1].text = f"{val:.4f}"
                            r[2].text = f"{d['stds'][lvl]:.4f}" # Add Std
                            r[3].text = f"{d['ses'][lvl]:.4f}" # Add SE
                            r[4].text = d['grouping'].get(lvl, '')

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
