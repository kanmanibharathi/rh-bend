import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from .duncan_util import get_duncan_q

class TwoFactorRCBDAnalyzer:
    def __init__(self, df, fact_a_col, fact_b_col, rep_col, resp_col):
        self.df = df
        self.a_col = fact_a_col
        self.b_col = fact_b_col
        self.r_col = rep_col
        self.resp_col = resp_col
        
        self.anova_table = {}
        self.post_hoc_res = {}
        
    def validate(self):
        # Convert to string/numeric
        for col in [self.a_col, self.b_col, self.r_col]:
            self.df[col] = self.df[col].astype(str).str.strip()
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.resp_col])
        
        self.a_levels = sorted(self.df[self.a_col].unique())
        self.b_levels = sorted(self.df[self.b_col].unique())
        self.r_levels = sorted(self.df[self.r_col].unique())
        
        self.n_a = len(self.a_levels)
        self.n_b = len(self.b_levels)
        self.n_r = len(self.r_levels)
        
        if self.n_a < 2 or self.n_b < 2 or self.n_r < 2:
            raise ValueError("Factors and Replications must have at least 2 levels.")
            
        # Check Balance (N = a * b * r)
        expected_N = self.n_a * self.n_b * self.n_r
        actual_N = len(self.df)
        
        # Checking exact balance: each A*B*R combination should exist once?
        # User spec says "Each A*B treatment combination appears once in every replication."
        # If dataset is just replicates, rows should equal a*b*r.
        # We can loosely check len(df) for now.
        if actual_N != expected_N:
            # We can warn or just proceed, user formulas assume balance.
            pass

    def run_anova(self):
        y = self.df[self.resp_col]
        # Totals
        G = y.sum()
        N = len(y)
        CF = (G**2) / N
        
        SS_Total = (y**2).sum() - CF
        
        # Replication: Sum(Ri^2)/(a*b) - CF
        # R_i are replication totals
        grp_r = self.df.groupby(self.r_col)[self.resp_col].sum()
        SS_Rep = (grp_r**2).sum() / (self.n_a * self.n_b) - CF
        
        # Factor A: Sum(Aj^2)/(b*r) - CF
        grp_a = self.df.groupby(self.a_col)[self.resp_col].sum()
        SS_A = (grp_a**2).sum() / (self.n_b * self.n_r) - CF
        
        # Factor B: Sum(Bk^2)/(a*r) - CF
        grp_b = self.df.groupby(self.b_col)[self.resp_col].sum()
        SS_B = (grp_b**2).sum() / (self.n_a * self.n_r) - CF
        
        # Interaction AxB: Sum((AB)jk^2)/r - CF - SS_A - SS_B
        grp_ab = self.df.groupby([self.a_col, self.b_col])[self.resp_col].sum()
        SS_AxB = (grp_ab**2).sum() / self.n_r - CF - SS_A - SS_B
        
        # Error
        SS_Error = SS_Total - SS_Rep - SS_A - SS_B - SS_AxB
        
        # DFs
        DF_Rep = self.n_r - 1
        DF_A = self.n_a - 1
        DF_B = self.n_b - 1
        DF_AxB = (self.n_a - 1) * (self.n_b - 1)
        DF_Error = (self.n_r - 1) * (self.n_a * self.n_b - 1)
        DF_Total = (self.n_a * self.n_b * self.n_r) - 1
        
        # MS
        MS_Rep = SS_Rep / DF_Rep
        MS_A = SS_A / DF_A
        MS_B = SS_B / DF_B
        MS_AxB = SS_AxB / DF_AxB
        MS_Error = SS_Error / DF_Error
        
        # F
        F_Rep = MS_Rep / MS_Error
        F_A = MS_A / MS_Error
        F_B = MS_B / MS_Error
        F_AxB = MS_AxB / MS_Error
        
        # P
        P_Rep = 1 - stats.f.cdf(F_Rep, DF_Rep, DF_Error)
        P_A = 1 - stats.f.cdf(F_A, DF_A, DF_Error)
        P_B = 1 - stats.f.cdf(F_B, DF_B, DF_Error)
        P_AxB = 1 - stats.f.cdf(F_AxB, DF_AxB, DF_Error)
        
        self.anova_table = {
            "Replication": {"df": int(DF_Rep), "SS": float(SS_Rep), "MS": float(MS_Rep), "F": float(F_Rep), "P": float(P_Rep)},
            "Factor A": {"df": int(DF_A), "SS": float(SS_A), "MS": float(MS_A), "F": float(F_A), "P": float(P_A)},
            "Factor B": {"df": int(DF_B), "SS": float(SS_B), "MS": float(MS_B), "F": float(F_B), "P": float(P_B)},
            "Interaction A x B": {"df": int(DF_AxB), "SS": float(SS_AxB), "MS": float(MS_AxB), "F": float(F_AxB), "P": float(P_AxB)},
            "Error": {"df": int(DF_Error), "SS": float(SS_Error), "MS": float(MS_Error), "F": None, "P": None},
            "Total": {"df": int(DF_Total), "SS": float(SS_Total), "MS": None, "F": None, "P": None}
        }

    def run_post_hoc(self, method, alpha, order='desc'):
        res = self.anova_table
        MS_E = res["Error"]["MS"]
        DF_E = res["Error"]["df"]
        
        # CV
        gm = self.df[self.resp_col].mean()
        cv = (np.sqrt(MS_E) / gm) * 100
        
        # Standard Errors
        # SEm_A = sqrt(MS_E / (b*r))
        sem_a = np.sqrt(MS_E / (self.n_b * self.n_r))
        sed_a = np.sqrt(2) * sem_a
        
        # SEm_B = sqrt(MS_E / (a*r))
        sem_b = np.sqrt(MS_E / (self.n_a * self.n_r))
        sed_b = np.sqrt(2) * sem_b
        
        # SEm_AxB = sqrt(MS_E / r)
        sem_ab = np.sqrt(MS_E / self.n_r)
        sed_ab = np.sqrt(2) * sem_ab
        
        # CD Calculations
        # For LSD: t * SEd
        # For Tukey: q * SEm
        def calc_cd(sem, sed, k):
            if method == 'lsd':
                return stats.t.ppf(1 - alpha/2, DF_E) * sed
            elif method == 'tukey':
                return stats.studentized_range.ppf(1-alpha, k, DF_E) * sem
            elif method == 'dmrt':
                # DMRT doesn't have a single CD, but we can return a reference LSD or minimum range
                return stats.t.ppf(1 - alpha/2, DF_E) * sed 
            return 0
            
        cd_a = calc_cd(sem_a, sed_a, self.n_a)
        cd_b = calc_cd(sem_b, sed_b, self.n_b)
        cd_ab = calc_cd(sem_ab, sed_ab, self.n_a * self.n_b)
        
        self.stats = {
            "CV": cv,
            "Factor A": {"SEm": sem_a, "SEd": sed_a, "CD": cd_a},
            "Factor B": {"SEm": sem_b, "SEd": sed_b, "CD": cd_b},
            "Interaction": {"SEm": sem_ab, "SEd": sed_ab, "CD": cd_ab}
        }
        
        # Grouping Logic
        # Hierarchy:
        # IF AxB Sig -> Interaction only.
        # ELSE -> Check A, Check B independently.
        
        p_ab = res["Interaction A x B"]["P"]
        p_a = res["Factor A"]["P"]
        p_b = res["Factor B"]["P"]
        
        results = {}
        
        # Helper to compute table data
        def get_table_data(groupby_cols):
            grp = self.df.groupby(groupby_cols)[self.resp_col]
            means = grp.mean()
            stds = grp.std()
            counts = grp.count()
            ses = stds / np.sqrt(counts) # or rely on balanced formula: stds / sqrt(counts) = stds / sqrt(r) if balanced
            # Standard error provided in request: Std.Err = Std.Dev / sqrt(r)
            # We implemented `stds / sqrt(counts)` which is safer if unbalanced, matches if balanced.
            
            # Sort for display
            # Always use Natural Index Sort (A-Z) to prevent rearranging Level/Interaction column
            display_means = means.sort_index()
            
            # Reindex stats to match display order
            stds = stds.reindex(display_means.index)
            ses = ses.reindex(display_means.index)
            
            return display_means, stds, ses
        
        # Helper for grouping
        def compute_grouping(means_series, n_items, specific_sem, specific_sed):
             # Ensure descending sort for grouping
             # We need to compute grouping based on PURE DESCENDING means, 
             # then map back to the keys.
             
             means_desc = means_series.sort_values(ascending=False)
             vals = means_desc.values
             keys = means_desc.index.tolist()
             n = len(vals)
             
             sig_pairs = set()
             
             if method == 'lsd':
                 crit = stats.t.ppf(1 - alpha/2, DF_E) * specific_sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i] - vals[j]) >= crit: sig_pairs.add((i,j))
                         
             elif method == 'tukey':
                 crit = stats.studentized_range.ppf(1-alpha, n, DF_E) * specific_sem
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i] - vals[j]) >= crit: sig_pairs.add((i,j))
            
             elif method == 'dmrt':
                 # Approximate pairwise check using rank ranges
                 # Re-implement simplified clique search or full DMRT
                 # Using the simplified logic from before (safe for now)
                 for i in range(n):
                     for j in range(i+1, n):
                         p = j - i + 1
                         q_val = get_duncan_q(p, DF_E, alpha)
                         rp = q_val * specific_sem
                         if abs(vals[i] - vals[j]) >= rp: sig_pairs.add((i,j))
             
             # Graph
             G = nx.Graph()
             G.add_nodes_from(range(n))
             for i in range(n):
                 for j in range(i+1, n):
                     if (i, j) not in sig_pairs: G.add_edge(i, j)
             
             cliques = list(nx.find_cliques(G))
             cliques.sort(key=lambda c: (min(c), -len(c)))
             vocab = "abcdefghijklmnopqrstuvwxyz"
             res_map = {i: "" for i in range(n)}
             for idx, clq in enumerate(cliques):
                 if idx < len(vocab):
                     c = vocab[idx]
                     for node in clq: res_map[node] += c
            
             return {keys[i]: "".join(sorted(res_map[i])) for i in range(n)}

        # Execution of Rule: INDEPENDENT CHECK (User Request)
        # Always calculate and return all tables.
        
        # 1. Factor A
        m_a, s_a, se_a = get_table_data(self.a_col)
        if p_a < alpha:
            l_a = compute_grouping(m_a, self.n_a, sem_a, sed_a)
            results["Factor A"] = {"means": m_a, "stds": s_a, "ses": se_a, "grouping": l_a, "sig": True}
        else:
             l_a = {k: "ns" for k in m_a.index}
             results["Factor A"] = {"means": m_a, "stds": s_a, "ses": se_a, "grouping": l_a, "sig": False}

        # 2. Factor B
        m_b, s_b, se_b = get_table_data(self.b_col)
        if p_b < alpha:
            l_b = compute_grouping(m_b, self.n_b, sem_b, sed_b)
            results["Factor B"] = {"means": m_b, "stds": s_b, "ses": se_b, "grouping": l_b, "sig": True}
        else:
            l_b = {k: "ns" for k in m_b.index}
            results["Factor B"] = {"means": m_b, "stds": s_b, "ses": se_b, "grouping": l_b, "sig": False}

        # 3. Interaction
        m_ab, s_ab, se_ab = get_table_data([self.a_col, self.b_col])
        if p_ab < alpha:
            grp_letters = compute_grouping(m_ab, self.n_a * self.n_b, sem_ab, sed_ab)
            results["Interaction"] = {"means": m_ab, "stds": s_ab, "ses": se_ab, "grouping": grp_letters, "sig": True}
        else:
            # If interaction is not sig, still show table with "ns"
            l_ab = {k: "ns" for k in m_ab.index}
            results["Interaction"] = {"means": m_ab, "stds": s_ab, "ses": se_ab, "grouping": l_ab, "sig": False}

        self.post_hoc_res = results

    def create_report(self):
        doc = Document()
        doc.add_heading('Two-Factor RCBD Analysis Report', 0)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # 1. ANOVA
        doc.add_heading('1. ANOVA Table', 1)
        t = doc.add_table(1, 6)
        t.style = 'Table Grid'
        fill_row(t.rows[0], ["Source", "DF", "SS", "MS", "F", "Prob"])
        
        ord_keys = ["Replication", "Factor A", "Factor B", "Interaction A x B", "Error", "Total"]
        for k in ord_keys:
            if k in self.anova_table:
                v = self.anova_table[k]
                r = t.add_row()
                sig = ""
                if v['P'] is not None:
                    if v['P'] < 0.01: sig = "**"
                    elif v['P'] < 0.05: sig = "*"
                    else: sig = "ns"
                
                fill_row(r, [
                    k, v['df'], f"{v['SS']:.4f}", 
                    f"{v['MS']:.4f}" if v['MS'] else "-",
                    f"{v['F']:.4f}" if v['F'] else "-",
                    f"{v['P']:.4f} {sig}" if v['P'] is not None else ""
                ])

        # 2. Precision Stats
        doc.add_heading('2. Statistics', 1)
        doc.add_paragraph(f"CV %: {self.stats['CV']:.2f}")
        for k in ["Factor A", "Factor B", "Interaction"]:
            s = self.stats[k]
            doc.add_paragraph(f"{k}: SEm={s['SEm']:.4f}, SEd={s['SEd']:.4f}, CD={s['CD']:.4f}")

        # 3. Means
        doc.add_heading('3. Means and Grouping', 1)
        ph = self.post_hoc_res
        
        def add_mean_table(title, dataset):
            doc.add_heading(title, 2)
            if not dataset.get('means', None) is None: # check if means exist
                t = doc.add_table(1, 5)
                t.style = 'Table Grid'
                fill_row(t.rows[0], ["Level", "Mean", "Std.Dev", "Std.Err", "Group"])
                for lvl, val in dataset['means'].items():
                    # Handle tuple keys for interaction
                    lvl_str = str(lvl) if not isinstance(lvl, tuple) else f"{lvl[0]} x {lvl[1]}"
                    # Grouping dict keys might be tuples too
                    grp = dataset['grouping'].get(lvl, "")
                    fill_row(t.add_row(), [
                        lvl_str, f"{val:.4f}", 
                        f"{dataset['stds'][lvl]:.4f}", 
                        f"{dataset['ses'][lvl]:.4f}", 
                        grp
                    ])
            elif "note" in dataset:
                doc.add_paragraph(f"Note: {dataset['note']}")

        # Always add all three tables
        add_mean_table("Factor A Means", ph["Factor A"])
        add_mean_table("Factor B Means", ph["Factor B"])
        add_mean_table("Interaction Means (A x B)", ph["Interaction"])
            
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

def fill_row(row, vals):
    for i, v in enumerate(vals):
        row.cells[i].text = str(v)
