import pandas as pd
import numpy as np
import scipy.stats as stats
import networkx as nx
from docx import Document
from datetime import datetime
import io

class PooledRCBDAnalyzer:
    def __init__(self, df, col_year, col_treat, col_rep, col_resp):
        self.df = df
        self.y_col = col_year
        self.t_col = col_treat
        self.r_col = col_rep
        self.resp_col = col_resp
        
        self.anova_table = {}
        self.stats = {}
        self.post_hoc_res = {}
        self.bartlett_res = {}
        
    def validate(self):
        # Clean header to avoid whitespace issues
        self.df.columns = self.df.columns.str.strip()
        
        # Check columns
        for col in [self.y_col, self.t_col, self.r_col, self.resp_col]:
            if col not in self.df.columns:
                raise ValueError(f"Column '{col}' not found in CSV.")
                
        # String cleanup
        for col in [self.y_col, self.t_col, self.r_col]:
             self.df[col] = self.df[col].astype(str).str.strip()
             
        self.df[self.resp_col] = pd.to_numeric(self.df[self.resp_col], errors='coerce')
        self.df = self.df.dropna(subset=[self.resp_col])
        
        # Counts
        self.years = self.df[self.y_col].unique()
        self.treats = self.df[self.t_col].unique()
        self.reps = self.df[self.r_col].unique()
        
        self.n_y = len(self.years)
        self.n_t = len(self.treats)
        self.n_r = len(self.reps)
        
        if self.n_y < 2: raise ValueError("Pooled analysis requires at least 2 years/locations.")
        if self.n_t < 2: raise ValueError("Treatment factor must have at least 2 levels.")
        if self.n_r < 2: raise ValueError("Replication factor must have at least 2 levels.")
        
    def run_bartlett(self):
        # Calculate MSE for each year separately (using RCBD model per year)
        ms_errors = []
        dfs = []
        
        for yr in self.years:
            sub = self.df[self.df[self.y_col] == yr]
            # RCBD ANOVA for this year
            # SSTot, SSRep, SSTrt, SSErr
            y_data = sub[self.resp_col]
            G = y_data.sum()
            N = len(y_data)
            CF = (G**2)/N
            SST = (y_data**2).sum() - CF
            
            # Rep SS
            r_grp = sub.groupby(self.r_col)[self.resp_col].sum()
            SSR = (r_grp**2).sum() / self.n_t - CF
            
            # Trt SS
            t_grp = sub.groupby(self.t_col)[self.resp_col].sum()
            SSTrt = (t_grp**2).sum() / self.n_r - CF
            
            SSE = SST - SSR - SSTrt
            if SSE < 0: SSE = 0
            
            df_e = (self.n_r - 1) * (self.n_t - 1)
            
            if df_e > 0:
                ms_errors.append(SSE / df_e)
                dfs.append(df_e)
        
        if len(ms_errors) < 2:
             self.bartlett_res = {"result": "Insufficient groups"}
             return

        # Regular Bartlett's Test using SciPy if raw data? No, we have summary stats (variances).
        # Use Bartlett's formula or F-max if simpler. 
        # User prompt mentions: "11. Homogeneity of Variance Test: Bartlett's Test"
        # Since I have variances (MS_E is variance estimator), I can use scipy.stats.bartlett if I had raw residuals?
        # But I only computed MS. Let's use Bartlett's formula for variances.
        
        k = len(ms_errors)
        Si2 = ms_errors
        ni_1 = dfs # degrees of freedom (nu)
        
        N_tot = sum(ni_1)
        Sp2 = sum([v*df for v,df in zip(Si2, ni_1)]) / N_tot # Pooled variance
        
        # M = N_tot * ln(Sp2) - Sum(nu_i * ln(Si^2))
        try:
            ln_Sp2 = np.log(Sp2)
            sum_nulnSi2 = sum([df * np.log(s) for df, s in zip(ni_1, Si2)])
            M = N_tot * ln_Sp2 - sum_nulnSi2
            
            # Correction C
            inv_sum = sum([1/df for df in ni_1])
            C = 1 + (1/(3*(k-1))) * (inv_sum - 1/N_tot)
            
            chi_sq = M / C
            p_val = 1 - stats.chi2.cdf(chi_sq, k-1)
            
            self.bartlett_res = {
                "stat": float(chi_sq),
                "p": float(p_val),
                "df": int(k - 1),
                "result": "Homogeneous" if p_val > 0.05 else "Heterogeneous"
            }
        except:
             self.bartlett_res = {"result": "Calculation Error"}

    def run_anova(self):
        y = self.df[self.resp_col]
        G = y.sum()
        N = len(y)
        CF = (G**2) / N
        SS_Total = (y**2).sum() - CF
        
        def get_ss(cols, divisor):
            grp = self.df.groupby(cols)[self.resp_col].sum()
            return (grp**2).sum() / divisor - CF
            
        nY, nR, nT = self.n_y, self.n_r, self.n_t
        
        # 1. Year (Y)
        SS_Y = get_ss(self.y_col, nR*nT)
        
        # 2. Replication within Year R(Y)
        # Note: Reps might just be labeled "1, 2, 3" in all years. 
        # So we must group by [Year, Rep] to get unique Rep-Year combos.
        SS_RY_Total = get_ss([self.y_col, self.r_col], nT)
        SS_RY = SS_RY_Total - SS_Y
        
        # 3. Treatment (T)
        SS_T = get_ss(self.t_col, nY*nR)
        
        # 4. Year x Treatment (YT)
        SS_YT_Total = get_ss([self.y_col, self.t_col], nR)
        SS_YT = SS_YT_Total - SS_Y - SS_T
        
        # 5. Error
        # SS_E = SS_Total - (SS_RY + SS_Y + SS_T + SS_YT)
        SS_E = SS_Total - (SS_RY + SS_Y + SS_T + SS_YT)
        if SS_E < 0: SS_E = 0
        
        # DF
        df_y = int(nY - 1)
        df_ry = int(nY * (nR - 1))
        df_t = int(nT - 1)
        df_yt = int((nY - 1) * (nT - 1))
        df_e = int(nY * (nR - 1) * (nT - 1))
        
        df_tot = int(nY * nR * nT - 1)
        
        # MS
        ms = lambda ss, df: float(ss / df) if df > 0 else 0.0
        
        MS_Y = ms(SS_Y, df_y)
        MS_RY = ms(SS_RY, df_ry)
        MS_T = ms(SS_T, df_t)
        MS_YT = ms(SS_YT, df_yt)
        MS_E = ms(SS_E, df_e)
        
        # F Tests
        # F_Y = MS_Y / MS_R(Y)
        F_Y = MS_Y / MS_RY if MS_RY > 0 else None
        
        # F_T = MS_T / MS_E
        F_T = MS_T / MS_E if MS_E > 0 else None
        
        # F_YT = MS_YT / MS_E
        F_YT = MS_YT / MS_E if MS_E > 0 else None
        
        # P Values
        def get_p(f, dfn, dfd):
            if f is None or dfd <= 0: return None
            return float(1 - stats.f.cdf(f, dfn, dfd))
            
        P_Y = get_p(F_Y, df_y, df_ry)
        P_T = get_p(F_T, df_t, df_e)
        P_YT = get_p(F_YT, df_yt, df_e)
        
        def s(df, ss, ms, f, p): 
            return {
                "df": df, "SS": float(ss), "MS": float(ms), 
                "F": float(f) if f is not None else None, 
                "P": float(p) if p is not None else None
            }
        
        self.anova_table = {
            "Year (Y)": s(df_y, SS_Y, MS_Y, F_Y, P_Y),
            "Rep (within Y)": s(df_ry, SS_RY, MS_RY, None, None), # Rep is error for Year? No, Rep(Y) is error term for Year.
            "Treatment (T)": s(df_t, SS_T, MS_T, F_T, P_T),
            "Year x Treatment (Y x T)": s(df_yt, SS_YT, MS_YT, F_YT, P_YT),
            "Error": {"df": df_e, "SS": float(SS_E), "MS": float(MS_E), "F": None, "P": None},
            "Total": {"df": df_tot, "SS": float(SS_Total), "MS": None, "F": None, "P": None}
        }

    def run_post_hoc(self, method, alpha):
        res = self.anova_table
        MS_RY = res["Rep (within Y)"]["MS"]
        MS_E = res["Error"]["MS"]
        DF_E = int(res["Error"]["df"])
        
        gm = self.df[self.resp_col].mean()
        cv = float((np.sqrt(MS_E) / gm) * 100) if gm != 0 and MS_E >=0 else 0.0
        
        nY, nR, nT = self.n_y, self.n_r, self.n_t
        
        # SEm
        def calc_sem(ms, div): return float(np.sqrt(ms / div)) if div > 0 and ms>=0 else 0.0
        
        sem_y = calc_sem(MS_RY, nR * nT) 
        sem_t = calc_sem(MS_E, nY * nR)
        sem_yt = calc_sem(MS_E, nR)
        
        def calc_cd(sem, df, k):
            if df <= 0: return 0.0
            sed = sem * np.sqrt(2)
            if method == 'lsd': return float(stats.t.ppf(1 - alpha/2, df) * sed)
            elif method == 'tukey': return float(stats.studentized_range.ppf(1-alpha, k, df) * sem)
            elif method == 'dmrt': return float(stats.t.ppf(1 - alpha/2, df) * sed)
            return 0.0
            
        stats_map = {
            "CV": cv,
            "Year (Y)": {"SEm": sem_y, "SEd": sem_y*np.sqrt(2), "CD": calc_cd(sem_y, int(res["Rep (within Y)"]["df"]), nY)},
            "Treatment (T)": {"SEm": sem_t, "SEd": sem_t*np.sqrt(2), "CD": calc_cd(sem_t, DF_E, nT)},
            "Year x Treatment": {"SEm": sem_yt, "SEd": sem_yt*np.sqrt(2), "CD": calc_cd(sem_yt, DF_E, nY*nT)}
        }
        self.stats = stats_map
        
        results = {}
        def get_data(cols):
            grp = self.df.groupby(cols)[self.resp_col]
            means = grp.mean().sort_index()
            stds = grp.std().reindex(means.index)
            ses = (stds / np.sqrt(grp.count())).reindex(means.index)
            return means, stds, ses
            
        def compute_letters(means, sem, df, k):
             m_desc = means.sort_values(ascending=False)
             vals = m_desc.values
             keys = m_desc.index.tolist()
             n = len(vals)
             sed = sem * np.sqrt(2)
             sig_pairs = set()
             
             if df <= 0: return {k: "ns" for k in keys}

             if method == 'lsd':
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             elif method == 'tukey':
                 crit = stats.studentized_range.ppf(1-alpha, n, df) * sem
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             else: 
                 crit = stats.t.ppf(1-alpha/2, df) * sed
                 for i in range(n):
                     for j in range(i+1, n):
                         if abs(vals[i]-vals[j]) >= crit: sig_pairs.add((i,j))
             
             G = nx.Graph()
             G.add_nodes_from(range(n))
             for i in range(n):
                 for j in range(i+1, n):
                     if (i,j) not in sig_pairs: G.add_edge(i, j)
             
             cliques = list(nx.find_cliques(G))
             cliques.sort(key=lambda x: (min(x), -len(x)))
             vocab = "abcdefghijklmnopqrstuvwxyz"
             res_map = {i: "" for i in range(n)}
             for idx, clq in enumerate(cliques):
                 if idx < len(vocab):
                     c = vocab[idx]
                     for node in clq: res_map[node] += c
             return {keys[i]: "".join(sorted(res_map[i])) for i in range(n)}

        def proc_eff(name, cols, sem_key, df_err, n_lvls):
            m, s, se = get_data(cols)
            p = res[name]["P"]
            let = {k: "ns" for k in m.index}
            if p is not None and p < alpha:
                # Special Year case
                if name == "Year (Y)":
                    let = compute_letters(m, stats_map[sem_key]["SEm"], int(res["Rep (within Y)"]["df"]), n_lvls)
                else:
                    let = compute_letters(m, stats_map[sem_key]["SEm"], df_err, n_lvls)
            
            results[sem_key] = {"means": m, "stds": s, "ses": se, "grouping": let}

        proc_eff("Year (Y)", self.y_col, "Year (Y)", 0, nY)
        proc_eff("Treatment (T)", self.t_col, "Treatment (T)", DF_E, nT)
        proc_eff("Year x Treatment (Y x T)", [self.y_col, self.t_col], "Year x Treatment", DF_E, nY*nT)
        
        self.post_hoc_res = results

    def create_report(self):
        doc = Document()
        doc.add_heading('Pooled ONE FACTOR RCBD Analysis Report', 0)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        doc.add_heading('1. Bartlett\'s Test', 1)
        if self.bartlett_res:
             doc.add_paragraph(f"Chi-Sq: {self.bartlett_res.get('stat', 0):.4f}, P-value: {self.bartlett_res.get('p', 0):.4f}")
             doc.add_paragraph(f"Result: {self.bartlett_res.get('result', '-')}")
             
        doc.add_heading('2. ANOVA Table', 1)
        t = doc.add_table(1, 6)
        t.style = 'Table Grid'
        def set_cell(r, i, txt): r.cells[i].text = str(txt)
        h = t.rows[0]
        for i,c in enumerate(["Source", "DF", "SS", "MS", "F", "Prob"]): set_cell(h, i, c)
        
        keys = ["Year (Y)", "Rep (within Y)", "Treatment (T)", "Year x Treatment (Y x T)", "Error", "Total"]
        for k in keys:
            if k in self.anova_table:
                v = self.anova_table[k]
                r = t.add_row()
                sig = ""
                if v['P'] is not None:
                     if v['P']<0.01: sig="**"
                     elif v['P']<0.05: sig="*"
                     else: sig="ns"
                set_cell(r, 0, k)
                set_cell(r, 1, v['df'])
                set_cell(r, 2, f"{v['SS']:.4f}")
                set_cell(r, 3, f"{v['MS']:.4f}" if v['MS'] else "-")
                set_cell(r, 4, f"{v['F']:.4f}" if v['F'] is not None else "-")
                set_cell(r, 5, f"{v['P']:.4f} {sig}" if v['P'] is not None else "")
                
        doc.add_heading('3. Statistics', 1)
        doc.add_paragraph(f"CV %: {self.stats['CV']:.2f}")
        for k in ["Year (Y)", "Treatment (T)", "Year x Treatment"]:
             s = self.stats[k]
             doc.add_paragraph(f"{k}: SEm={s['SEm']:.4f}, SEd={s['SEd']:.4f}, CD={s['CD']:.4f}")
             
        doc.add_heading('4. Means', 1)
        for k in ["Year (Y)", "Treatment (T)", "Year x Treatment"]:
             ds = self.post_hoc_res[k]
             doc.add_heading(f"{k} Means", 2)
             t = doc.add_table(1, 5)
             t.style = 'Table Grid'
             r = t.rows[0]
             for i,c in enumerate(["Level", "Mean", "Std.Dev", "Std.Err", "Group"]): set_cell(r, i, c)
             for lvl, val in ds['means'].items():
                 r = t.add_row()
                 if isinstance(lvl, tuple): l_str = " x ".join(map(str, lvl))
                 else: l_str = str(lvl)
                 set_cell(r, 0, l_str)
                 set_cell(r, 1, f"{val:.4f}")
                 set_cell(r, 2, f"{ds['stds'][lvl]:.4f}")
                 set_cell(r, 3, f"{ds['ses'][lvl]:.4f}")
                 set_cell(r, 4, ds['grouping'].get(lvl, ""))
                 
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
