import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import matplotlib.pyplot as plt
import seaborn as sns

class OneSampleTTestAnalyzer:
    def __init__(self, df, value_col, mu_0=0.0, alpha=0.05):
        self.df = df
        self.value_col = value_col
        self.mu_0 = float(mu_0)
        self.alpha = float(alpha)
        
        self.data = None
        self.descriptive_stats = {}
        self.normality_results = {}
        self.t_test_results = {}
        
    def validate(self):
        # Numeric checks
        self.df[self.value_col] = pd.to_numeric(self.df[self.value_col], errors='coerce')
        
        # Missing values removal (listwise)
        self.data = self.df[self.value_col].dropna().values
        
        # Check alpha
        if not (0 < self.alpha <= 0.1):
            raise ValueError("Alpha must be between 0 and 0.1.")
            
        # Sample size check
        if len(self.data) < 3:
            raise ValueError(f"Sample has too few observations ({len(self.data)}). Minimum 3 required.")

    def run_analysis(self):
        # 1. Descriptive Statistics
        n = len(self.data)
        mean = np.mean(self.data)
        var = np.var(self.data, ddof=1)
        std = np.sqrt(var)
        se = std / np.sqrt(n)
        cv = (std / mean * 100) if mean != 0 else 0
        
        self.descriptive_stats = {
            "n": n,
            "Mean": mean,
            "Median": np.median(self.data),
            "Variance": var,
            "StdDev": std,
            "StdError": se,
            "CV": cv,
            "Skewness": stats.skew(self.data),
            "Kurtosis": stats.kurtosis(self.data)
        }
        
        # 2. Normality Test (Shapiro-Wilk)
        w_stat, p_val_norm = stats.shapiro(self.data)
        self.normality_results = {
            "W": w_stat,
            "p_value": p_val_norm,
            "Interpretation": "Normal" if p_val_norm > self.alpha else "Not Normal"
        }
        
        # 3. One-Sample t-Test
        t_stat, p_val = stats.ttest_1samp(self.data, self.mu_0)
        df = n - 1
        
        # Confidence Interval
        t_crit = stats.t.ppf(1 - self.alpha/2, df)
        lower_ci = mean - t_crit * se
        upper_ci = mean + t_crit * se
        
        self.t_test_results = {
            "HypotheticalMean": self.mu_0,
            "SampleMean": mean,
            "StdError": se,
            "t_value": t_stat,
            "df": df,
            "p_value": p_val,
            "Lower_CI": lower_ci,
            "Upper_CI": upper_ci,
            "Conclusion": "Fail to reject Ho" if p_val > self.alpha else "Reject Ho"
        }

    def get_interpretation(self):
        p_val = self.t_test_results["p_value"]
        if p_val > self.alpha:
            return (
                f"The p-value ({p_val:.4f}) is greater than the significance level (α = {self.alpha}). "
                "The result is non-significant and the null hypothesis (Ho: μ = μ₀) is failed to be rejected. "
                f"The population mean is statistically equal to the hypothetical mean ({self.mu_0})."
            )
        else:
            return (
                f"The p-value ({p_val:.4f}) is less than the significance level (α = {self.alpha}). "
                "The result is statistically significant and the null hypothesis (Ho: μ = μ₀) is rejected. "
                f"The population mean differs significantly from the hypothetical mean ({self.mu_0})."
            )

    def create_report(self):
        doc = Document()
        
        # Style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Title
        title = doc.add_heading("One-Sample t-Test Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 1. Objective
        doc.add_heading("1. Objective", level=1)
        doc.add_paragraph(
            f"The objective of this analysis is to test whether the population mean of '{self.value_col}' "
            f"differs significantly from a hypothetical mean (μ₀ = {self.mu_0}) using a One-Sample t-test "
            f"at a {self.alpha*100}% significance level."
        )
        
        # 2. Descriptive Statistics
        doc.add_heading("2. Descriptive Statistics", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        s = self.descriptive_stats
        desc_data = [
            ("Sample Size (n)", str(s['n'])),
            ("Mean", f"{s['Mean']:.4f}"),
            ("Median", f"{s['Median']:.4f}"),
            ("Variance", f"{s['Variance']:.4f}"),
            ("Std. Deviation", f"{s['StdDev']:.4f}"),
            ("Standard Error", f"{s['StdError']:.4f}"),
            ("CV (%)", f"{s['CV']:.2f}"),
            ("Skewness", f"{s['Skewness']:.4f}"),
            ("Kurtosis", f"{s['Kurtosis']:.4f}")
        ]
        for label, val in desc_data:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val
            
        # 3. Normality Test
        doc.add_heading("3. Normality Test (Shapiro-Wilk)", level=1)
        doc.add_paragraph("Assumption: Data should follow a normal distribution for a valid t-test.")
        table_norm = doc.add_table(rows=1, cols=4)
        table_norm.style = 'Table Grid'
        headers = ["Test", "Statistic (W)", "P-value", "Interpretation"]
        for i, h in enumerate(headers): table_norm.cell(0, i).text = h
        
        n_res = self.normality_results
        row = table_norm.add_row().cells
        row[0].text = "Shapiro-Wilk"
        row[1].text = f"{n_res['W']:.4f}"
        row[2].text = f"{n_res['p_value']:.4f}"
        row[3].text = n_res['Interpretation']
            
        # 4. One-Sample t-Test Results
        doc.add_heading("4. One-Sample t-Test Results", level=1)
        table_t = doc.add_table(rows=1, cols=2)
        table_t.style = 'Table Grid'
        r = self.t_test_results
        t_data = [
            ("Hypothetical Mean (μ₀)", str(r['HypotheticalMean'])),
            ("Sample Mean", f"{r['SampleMean']:.4f}"),
            ("Standard Error", f"{r['StdError']:.4f}"),
            ("t-value", f"{r['t_value']:.4f}"),
            ("Degrees of Freedom (df)", str(r['df'])),
            ("P-value (Two-tailed)", f"{r['p_value']:.4f}"),
            ("Lower 95% CI", f"{r['Lower_CI']:.4f}"),
            ("Upper 95% CI", f"{r['Upper_CI']:.4f}")
        ]
        for label, val in t_data:
            row = table_t.add_row().cells
            row[0].text = label
            row[1].text = val
            
        # 5. Visualizations
        doc.add_heading("5. Visualizations", level=1)
        
        # Plot 1: Histogram with Density Curve
        plt.figure(figsize=(8, 5))
        sns.histplot(self.data, kde=True, color='skyblue', edgecolor='black')
        plt.title(f"Histogram and Density Curve of {self.value_col}")
        plt.xlabel(self.value_col)
        buf1 = io.BytesIO()
        plt.savefig(buf1, format='png', dpi=300)
        buf1.seek(0)
        doc.add_picture(buf1, width=Inches(5.5))
        plt.close()
        
        # Plot 2: Box Plot
        plt.figure(figsize=(6, 5))
        sns.boxplot(y=self.data, color='lightgreen')
        plt.title(f"Box Plot of {self.value_col}")
        plt.ylabel(self.value_col)
        buf2 = io.BytesIO()
        plt.savefig(buf2, format='png', dpi=300)
        buf2.seek(0)
        doc.add_picture(buf2, width=Inches(4))
        plt.close()
        
        # 6. Conclusion
        doc.add_heading("6. Conclusion & Interpretation", level=1)
        doc.add_paragraph(self.get_interpretation())
        doc.add_paragraph(f"Decision: {r['Conclusion']}")
        
        doc.add_paragraph("\nREFERENCES:")
        references = [
            "Student. (1908). The probable error of a mean. Biometrika.",
            "Shapiro, S. S. & Wilk, M. B. (1965). An analysis of variance test for normality. Biometrika.",
            "Snedecor, G. W. & Cochran, W. G. (1989). Statistical Methods. Iowa State University Press."
        ]
        for ref in references:
            doc.add_paragraph(ref, style='List Bullet')

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            pd.DataFrame([self.descriptive_stats]).to_excel(writer, sheet_name='Descriptive Statistics', index=False)
            pd.DataFrame([self.normality_results]).to_excel(writer, sheet_name='Normality Test', index=False)
            pd.DataFrame([self.t_test_results]).to_excel(writer, sheet_name='t-Test Results', index=False)
            pd.DataFrame(self.data, columns=[self.value_col]).to_excel(writer, sheet_name='Raw Data', index=False)
        output.seek(0)
        return output
