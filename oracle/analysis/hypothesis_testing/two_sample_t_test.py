import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import matplotlib.pyplot as plt
import seaborn as sns

class TwoSampleTTestAnalyzer:
    def __init__(self, df, category_col, value_col, alpha=0.05, variance_option="bartlett"):
        self.df = df
        self.category_col = category_col
        self.value_col = value_col
        self.alpha = float(alpha)
        self.variance_option = variance_option  # "equal", "unequal", or "bartlett"
        
        self.groups = []
        self.group_data = {}
        self.descriptive_stats = {}
        self.normality_results = {}
        self.bartlett_results = {}
        self.t_test_results = {}
        
    def validate(self):
        # Numeric checks
        self.df[self.value_col] = pd.to_numeric(self.df[self.value_col], errors='coerce')
        
        # Missing values removal (listwise)
        self.df = self.df.dropna(subset=[self.category_col, self.value_col])
        
        # Check alpha
        if not (0 < self.alpha <= 0.1):
            raise ValueError("Alpha must be between 0 and 0.1.")
            
        # Exactly two groups check
        unique_groups = self.df[self.category_col].unique()
        if len(unique_groups) != 2:
            raise ValueError(f"Two-Sample t-Test requires exactly two groups. Found: {len(unique_groups)}")
            
        self.groups = sorted([str(g) for g in unique_groups])
        for group in self.groups:
            data = self.df[self.df[self.category_col].astype(str) == group][self.value_col].values
            if len(data) < 3:
                raise ValueError(f"Group '{group}' has too few observations (minimum 3 required).")
            self.group_data[group] = data

    def run_analysis(self):
        g1, g2 = self.groups
        data1 = self.group_data[g1]
        data2 = self.group_data[g2]
        
        # 1. Descriptive Statistics
        for group, data in zip(self.groups, [data1, data2]):
            mean = np.mean(data)
            var = np.var(data, ddof=1)
            std = np.sqrt(var)
            n = len(data)
            se = std / np.sqrt(n)
            cv = (std / mean * 100) if mean != 0 else 0
            
            self.descriptive_stats[group] = {
                "n": n,
                "Mean": mean,
                "Variance": var,
                "StdDev": std,
                "StdError": se,
                "CV": cv,
                "Median": np.median(data),
                "Skewness": stats.skew(data),
                "Kurtosis": stats.kurtosis(data)
            }
            
            # 2. Normality Test (Shapiro-Wilk)
            w_stat, p_val_norm = stats.shapiro(data)
            self.normality_results[group] = {
                "W": w_stat,
                "p_value": p_val_norm,
                "Interpretation": "Normal" if p_val_norm > self.alpha else "Not Normal"
            }
            
        # 3. Bartlett's Test for Homogeneity of Variances
        b_stat, p_bartlett = stats.bartlett(data1, data2)
        self.bartlett_results = {
            "Statistic": b_stat,
            "p_value": p_bartlett,
            "Interpretation": "Equal variances" if p_bartlett > self.alpha else "Unequal variances"
        }
        
        # 4. Two-Sample t-Test Selection
        if self.variance_option == "equal":
            equal_var = True
            test_type = "Student's t-test (Equal Variances)"
        elif self.variance_option == "unequal":
            equal_var = False
            test_type = "Welch's t-test (Unequal Variances)"
        else:  # bartlett
            equal_var = (p_bartlett > self.alpha)
            test_type = "Student's t-test (Equal Variances)" if equal_var else "Welch's t-test (Unequal Variances)"
            
        # Perform t-test
        t_stat, p_val = stats.ttest_ind(data1, data2, equal_var=equal_var)
        
        # Calculate DF and SE for manual CI
        s1_sq = self.descriptive_stats[g1]["Variance"]
        s2_sq = self.descriptive_stats[g2]["Variance"]
        n1 = self.descriptive_stats[g1]["n"]
        n2 = self.descriptive_stats[g2]["n"]
        mean1 = self.descriptive_stats[g1]["Mean"]
        mean2 = self.descriptive_stats[g2]["Mean"]
        
        if equal_var:
            df = n1 + n2 - 2
            sp2 = ((n1 - 1) * s1_sq + (n2 - 1) * s2_sq) / df
            se_diff = np.sqrt(sp2 * (1/n1 + 1/n2))
        else:
            # Welch-Satterthwaite DF
            se1_sq = s1_sq / n1
            se2_sq = s2_sq / n2
            num = (se1_sq + se2_sq)**2
            den = (se1_sq**2 / (n1 - 1)) + (se2_sq**2 / (n2 - 1))
            df = num / den
            se_diff = np.sqrt(se1_sq + se2_sq)
            
        # Confidence Interval
        t_crit = stats.t.ppf(1 - self.alpha/2, df)
        diff = mean1 - mean2
        lower_ci = diff - t_crit * se_diff
        upper_ci = diff + t_crit * se_diff
        
        self.t_test_results = {
            "TestType": test_type,
            "Mean1": mean1,
            "Mean2": mean2,
            "MeanDiff": diff,
            "t_value": t_stat,
            "df": df,
            "p_value": p_val,
            "Lower_CI": lower_ci,
            "Upper_CI": upper_ci,
            "Conclusion": "Fail to reject Ho" if p_val > self.alpha else "Reject Ho"
        }

    def get_interpretation(self):
        p_val = self.t_test_results["p_value"]
        if p_val > self.alpha:
            return (
                "The p-value is greater than the level of significance. "
                "The result is non-significant and the null hypothesis is failed to be rejected. "
                "There is no significant difference between the means of the two populations."
            )
        else:
            return (
                "The p-value is less than the level of significance. "
                "The result is significant and the null hypothesis is rejected. "
                "There is a significant difference between the means of the two populations."
            )

    def create_report(self):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        title = doc.add_heading("Two Independent Sample t-Test Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 1. Objective
        doc.add_heading("1. Objective", level=1)
        doc.add_paragraph(
            f"The objective of this analysis is to compare the means of two independent groups ('{self.groups[0]}' and '{self.groups[1]}') "
            f"for the variable '{self.value_col}' using a Two-Sample t-test at a {self.alpha*100}% significance level."
        )
        
        # 2. Descriptive Statistics
        doc.add_heading("2. Descriptive Statistics", level=1)
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        headers = ["Group", "Mean", "SD", "SE", "CV (%)", "Skewness", "Kurtosis"]
        for i, h in enumerate(headers): table.cell(0, i).text = h
        for g in self.groups:
            s = self.descriptive_stats[g]
            row = table.add_row().cells
            row[0].text = str(g)
            row[1].text = f"{s['Mean']:.4f}"
            row[2].text = f"{s['StdDev']:.4f}"
            row[3].text = f"{s['StdError']:.4f}"
            row[4].text = f"{s['CV']:.2f}"
            row[5].text = f"{s['Skewness']:.4f}"
            row[6].text = f"{s['Kurtosis']:.4f}"
            
        # 3. Normality Test
        doc.add_heading("3. Normality Test (Shapiro-Wilk)", level=1)
        table_norm = doc.add_table(rows=1, cols=4)
        table_norm.style = 'Table Grid'
        headers = ["Group", "W Statistic", "P-value", "Interpretation"]
        for i, h in enumerate(headers): table_norm.cell(0, i).text = h
        for g in self.groups:
            n = self.normality_results[g]
            row = table_norm.add_row().cells
            row[0].text = str(g)
            row[1].text = f"{n['W']:.4f}"
            row[2].text = f"{n['p_value']:.4f}"
            row[3].text = n['Interpretation']
            
        # 4. Bartlett's Test
        doc.add_heading("4. Homogeneity of Variances (Bartlett's Test)", level=1)
        table_b = doc.add_table(rows=1, cols=4)
        table_b.style = 'Table Grid'
        headers = ["Test", "Statistic", "P-value", "Interpretation"]
        for i, h in enumerate(headers): table_b.cell(0, i).text = h
        b = self.bartlett_results
        row = table_b.add_row().cells
        row[0].text = "Bartlett's"
        row[1].text = f"{b['Statistic']:.4f}"
        row[2].text = f"{b['p_value']:.4f}"
        row[3].text = b['Interpretation']
        
        # 5. t-Test Results
        doc.add_heading("5. Two-Sample t-Test Analysis", level=1)
        table_t = doc.add_table(rows=1, cols=2)
        table_t.style = 'Table Grid'
        r = self.t_test_results
        t_data = [
            ("Test Type", r['TestType']),
            (f"Mean ({self.groups[0]})", f"{r['Mean1']:.4f}"),
            (f"Mean ({self.groups[1]})", f"{r['Mean2']:.4f}"),
            ("Mean Difference", f"{r['MeanDiff']:.4f}"),
            ("t-value", f"{r['t_value']:.4f}"),
            ("Degrees of Freedom (df)", f"{r['df']:.4f}"),
            ("P-value (Two-tailed)", f"{r['p_value']:.4f}"),
            ("Lower 95% CI", f"{r['Lower_CI']:.4f}"),
            ("Upper 95% CI", f"{r['Upper_CI']:.4f}")
        ]
        for label, val in t_data:
            row = table_t.add_row().cells
            row[0].text = label
            row[1].text = val
            
        # 6. Visualizations
        doc.add_heading("6. Visualizations", level=1)
        
        # Plot 1: Boxplot
        plt.figure(figsize=(8, 6))
        sns.boxplot(x=self.category_col, y=self.value_col, data=self.df, palette="Set2")
        plt.title(f"Box Plot of {self.value_col} by {self.category_col}")
        buf1 = io.BytesIO()
        plt.savefig(buf1, format='png', dpi=300)
        buf1.seek(0)
        doc.add_picture(buf1, width=Inches(5))
        plt.close()
        
        # Plot 2: Histogram
        plt.figure(figsize=(10, 6))
        for g in self.groups:
            sns.histplot(self.group_data[g], kde=True, label=str(g), alpha=0.5)
        plt.title(f"Distribution Comparison: {self.value_col}")
        plt.legend()
        buf2 = io.BytesIO()
        plt.savefig(buf2, format='png', dpi=300)
        buf2.seek(0)
        doc.add_picture(buf2, width=Inches(5))
        plt.close()
        
        # 7. Conclusion
        doc.add_heading("7. Conclusion & Interpretation", level=1)
        doc.add_paragraph(self.get_interpretation())
        doc.add_paragraph(f"Decision: {r['Conclusion']}")
        
        references = [
            "Student (1908). The probable error of a mean. Biometrika.",
            "Welch (1947). The generalization of 'Student's' problem when several different population variances are involved. Biometrika.",
            "Bartlett (1937). Properties of sufficiency and statistical tests. Proceedings of the Royal Society of London.",
            "Shapiro & Wilk (1965). An analysis of variance test for normality. Biometrika."
        ]
        doc.add_paragraph("\nREFERENCES:")
        for ref in references: doc.add_paragraph(ref, style='List Bullet')

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            pd.DataFrame(self.descriptive_stats).T.to_excel(writer, sheet_name='Descriptive Statistics')
            pd.DataFrame(self.normality_results).T.to_excel(writer, sheet_name='Normality Test')
            pd.DataFrame([self.bartlett_results]).to_excel(writer, sheet_name='Bartlett Test', index=False)
            pd.DataFrame([self.t_test_results]).to_excel(writer, sheet_name='t-Test Results', index=False)
        output.seek(0)
        return output
