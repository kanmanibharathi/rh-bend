import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import matplotlib.pyplot as plt
import seaborn as sns

class PairedTTestAnalyzer:
    def __init__(self, df, col1, col2, alpha=0.05, d0=0.0):
        self.df = df
        self.col1 = col1
        self.col2 = col2
        self.alpha = float(alpha)
        self.d0 = float(d0)
        
        self.data1 = None
        self.data2 = None
        self.differences = None
        self.descriptive_stats = {}
        self.diff_stats = {}
        self.normality_results = {}
        self.t_test_results = {}
        
    def validate(self):
        # Numeric checks
        self.df[self.col1] = pd.to_numeric(self.df[self.col1], errors='coerce')
        self.df[self.col2] = pd.to_numeric(self.df[self.col2], errors='coerce')
        
        # Row-wise deletion of missing values
        cleaned_df = self.df.dropna(subset=[self.col1, self.col2])
        
        self.data1 = cleaned_df[self.col1].values
        self.data2 = cleaned_df[self.col2].values
        
        if len(self.data1) < 3:
            raise ValueError(f"Too few valid paired observations ({len(self.data1)}). Minimum 3 required.")
            
        if not (0 < self.alpha <= 0.1):
            raise ValueError("Alpha must be between 0 and 0.1.")
            
        self.differences = self.data1 - self.data2

    def run_analysis(self):
        # 1. Descriptive Statistics for each sample
        for col_name, data in [(self.col1, self.data1), (self.col2, self.data2)]:
            mean = np.mean(data)
            var = np.var(data, ddof=1)
            std = np.sqrt(var)
            cv = (std / mean * 100) if mean != 0 else 0
            
            self.descriptive_stats[col_name] = {
                "n": len(data),
                "Mean": mean,
                "Variance": var,
                "StdDev": std,
                "CV": cv,
                "Median": np.median(data),
                "Skewness": stats.skew(data),
                "Kurtosis": stats.kurtosis(data)
            }
            
        # 2. Descriptive Stats for differences
        n = len(self.differences)
        mean_d = np.mean(self.differences)
        var_d = np.var(self.differences, ddof=1)
        sd_d = np.sqrt(var_d)
        se_d = sd_d / np.sqrt(n)
        
        self.diff_stats = {
            "n": n,
            "Mean_d": mean_d,
            "Var_d": var_d,
            "SD_d": sd_d,
            "SE_d": se_d
        }
        
        # 3. Normality Test on differences
        w_stat, p_norm = stats.shapiro(self.differences)
        self.normality_results = {
            "W": w_stat,
            "p_value": p_norm,
            "Interpretation": "Normal" if p_norm > self.alpha else "Not Normal"
        }
        
        # 4. Paired t-test
        t_stat, p_val = stats.ttest_rel(self.data1, self.data2)
        # Note: stats.ttest_rel tests against d0=0. 
        # For custom d0: t = (mean_d - d0) / se_d
        if self.d0 != 0:
            t_stat = (mean_d - self.d0) / se_d
            df = n - 1
            p_val = 2 * (1 - stats.t.cdf(abs(t_stat), df))
        else:
            df = n - 1
            
        # Confidence Interval
        t_crit = stats.t.ppf(1 - self.alpha/2, df)
        lower_ci = mean_d - t_crit * se_d
        upper_ci = mean_d + t_crit * se_d
        
        self.t_test_results = {
            "d0": self.d0,
            "MeanDiff": mean_d,
            "StdError": se_d,
            "t_value": t_stat,
            "df": df,
            "p_value": p_val,
            "Lower_CI": lower_ci,
            "Upper_CI": upper_ci,
            "Conclusion": "Fail to reject Ho" if p_val > self.alpha else "Reject Ho"
        }

    def get_interpretation(self):
        p_val = self.t_test_results["p_value"]
        if p_val > self.alpha:
            return (
                f"The p-value ({p_val:.4f}) is greater than the level of significance (α = {self.alpha}). "
                "The result is non-significant and the null hypothesis (Ho: μ_d = d₀) is accepted. "
                f"The mean difference between the paired samples is statistically equal to {self.d0}."
            )
        else:
            return (
                f"The p-value ({p_val:.4f}) is less than the level of significance (α = {self.alpha}). "
                "The result is statistically significant and the null hypothesis (Ho: μ_d = d₀) is rejected. "
                f"The mean difference between the paired samples differs significantly from {self.d0}."
            )

    def create_report(self):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        title = doc.add_heading("Paired Sample t-Test Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 1. Objective
        doc.add_heading("1. Objective", level=1)
        doc.add_paragraph(
            f"The objective of this analysis is to test whether the mean difference between two related samples "
            f"('{self.col1}' and '{self.col2}') differs significantly from a hypothesized value (d₀ = {self.d0}) "
            f"at a {self.alpha*100}% significance level."
        )
        
        # 2. Descriptive Statistics
        doc.add_heading("2. Descriptive Statistics", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        headers = ["Sample", "Mean", "SD", "CV (%)", "Skewness", "Kurtosis"]
        for i, h in enumerate(headers): table.cell(0, i).text = h
        
        for col in [self.col1, self.col2]:
            s = self.descriptive_stats[col]
            row = table.add_row().cells
            row[0].text = str(col)
            row[1].text = f"{s['Mean']:.4f}"
            row[2].text = f"{s['StdDev']:.4f}"
            row[3].text = f"{s['CV']:.2f}"
            row[4].text = f"{s['Skewness']:.4f}"
            row[5].text = f"{s['Kurtosis']:.4f}"
            
        # 3. Paired Differences
        doc.add_heading("3. Paired Differences Analysis", level=1)
        table_diff = doc.add_table(rows=1, cols=2)
        table_diff.style = 'Table Grid'
        d_data = [
            ("Number of Pairs (n)", str(self.diff_stats['n'])),
            ("Mean Difference", f"{self.diff_stats['Mean_d']:.4f}"),
            ("SD of Differences", f"{self.diff_stats['SD_d']:.4f}"),
            ("Standard Error of Mean Difference", f"{self.diff_stats['SE_d']:.4f}")
        ]
        for label, val in d_data:
            row = table_diff.add_row().cells
            row[0].text = label
            row[1].text = val
            
        # 4. Normality Test on Differences
        doc.add_heading("4. Normality Test (Shapiro-Wilk) on Differences", level=1)
        table_norm = doc.add_table(rows=1, cols=4)
        table_norm.style = 'Table Grid'
        h_norm = ["Test", "Statistic (W)", "P-value", "Interpretation"]
        for i, h in enumerate(h_norm): table_norm.cell(0, i).text = h
        n_res = self.normality_results
        row = table_norm.add_row().cells
        row[0].text = "Shapiro-Wilk"
        row[1].text = f"{n_res['W']:.4f}"
        row[2].text = f"{n_res['p_value']:.4f}"
        row[3].text = n_res['Interpretation']
        
        # 5. Paired t-Test Results
        doc.add_heading("5. Paired t-Test Results", level=1)
        table_t = doc.add_table(rows=1, cols=2)
        table_t.style = 'Table Grid'
        r = self.t_test_results
        t_data = [
            ("Hypothesized Difference (d₀)", str(r['d0'])),
            ("Observed Mean Difference", f"{r['MeanDiff']:.4f}"),
            ("Standard Error", f"{r['StdError']:.4f}"),
            ("t-value", f"{r['t_value']:.4f}"),
            ("Degrees of Freedom (df)", str(r['df'])),
            ("P-value (Two-tailed)", f"{r['p_value']:.4f}"),
            ("Lower 95% CI", f"{r['Lower_CI']:.4f}"),
            ("Upper 95% CI", f"{r['Upper_CI']:.4f}")
        ]
        for label, val in t_data:
            row = table_t.add_row().cells
            row[0].text = label
            row[1].text = val
            
        # 6. Visualizations
        doc.add_heading("6. Visualizations", level=1)
        
        # Plot 1: Histogram of Differences
        plt.figure(figsize=(8, 5))
        sns.histplot(self.differences, kde=True, color='salmon', edgecolor='black')
        plt.title(f"Distribution of Paired Differences ({self.col1} - {self.col2})")
        plt.xlabel("Difference")
        buf1 = io.BytesIO()
        plt.savefig(buf1, format='png', dpi=300, bbox_inches='tight')
        buf1.seek(0)
        doc.add_picture(buf1, width=Inches(5.5))
        plt.close()
        
        # Plot 2: Side-by-side Boxplot
        plt.figure(figsize=(8, 6))
        plot_df = pd.DataFrame({
            'Value': np.concatenate([self.data1, self.data2]),
            'Sample': [self.col1]*len(self.data1) + [self.col2]*len(self.data2)
        })
        sns.boxplot(x='Sample', y='Value', data=plot_df, palette="husl")
        plt.title("Sample Distributions Comparison")
        buf2 = io.BytesIO()
        plt.savefig(buf2, format='png', dpi=300, bbox_inches='tight')
        buf2.seek(0)
        doc.add_picture(buf2, width=Inches(5))
        plt.close()
        
        # Plot 3: Paired Line Plot
        plt.figure(figsize=(8, 6))
        for i in range(len(self.data1)):
            plt.plot([0, 1], [self.data1[i], self.data2[i]], color='gray', alpha=0.3, marker='o')
        plt.xticks([0, 1], [self.col1, self.col2])
        plt.title("Visualizing Paired Changes")
        plt.ylabel("Value")
        buf3 = io.BytesIO()
        plt.savefig(buf3, format='png', dpi=300, bbox_inches='tight')
        buf3.seek(0)
        doc.add_picture(buf3, width=Inches(5))
        plt.close()
        
        # 7. Conclusion
        doc.add_heading("7. Conclusion & Interpretation", level=1)
        doc.add_paragraph(self.get_interpretation())
        doc.add_paragraph(f"Decision: {r['Conclusion']}")
        
        doc.add_paragraph("\nREFERENCES:")
        refs = [
            "Student (1908). The probable error of a mean. Biometrika.",
            "Snedecor, G. W. & Cochran, W. G. (1989). Statistical Methods. Iowa State University Press.",
            "Shapiro, S. S. & Wilk, M. B. (1965). An analysis of variance test for normality. Biometrika."
        ]
        for ref in refs: doc.add_paragraph(ref, style='List Bullet')

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            pd.DataFrame(self.descriptive_stats).T.to_excel(writer, sheet_name='Descriptive Statistics')
            pd.DataFrame([self.diff_stats]).to_excel(writer, sheet_name='Difference Stats', index=False)
            pd.DataFrame([self.normality_results]).to_excel(writer, sheet_name='Normality Test', index=False)
            pd.DataFrame([self.t_test_results]).to_excel(writer, sheet_name='t-Test Results', index=False)
        output.seek(0)
        return output
