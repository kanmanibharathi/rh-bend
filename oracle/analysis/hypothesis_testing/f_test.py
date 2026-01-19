import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import matplotlib.pyplot as plt
import seaborn as sns

class FTestAnalyzer:
    def __init__(self, df, category_col, value_col, alpha=0.05, mode='long'):
        self.df = df
        self.category_col = category_col  # In 'wide' mode, this is Group 1 Column
        self.value_col = value_col        # In 'wide' mode, this is Group 2 Column
        self.alpha = float(alpha)
        self.mode = mode # 'long' (grouped) or 'wide' (separate columns)
        
        self.groups = []
        self.group_data = {}
        self.descriptive_stats = {}
        self.normality_results = {}
        self.f_test_results = {}
        
    def validate(self):
        if self.mode == 'long':
            # Numeric checks
            self.df[self.value_col] = pd.to_numeric(self.df[self.value_col], errors='coerce')
            
            # Missing values removal (listwise for long format)
            self.df = self.df.dropna(subset=[self.category_col, self.value_col])
            
            # Check alpha
            if not (0 < self.alpha <= 0.1):
                raise ValueError("Alpha must be between 0 and 0.1.")
                
            # Exactly two groups check
            unique_groups = self.df[self.category_col].unique()
            if len(unique_groups) != 2:
                raise ValueError(f"F-Test requires exactly two groups in the category column. Found: {len(unique_groups)}")
                
            self.groups = sorted([str(g) for g in unique_groups])
            for group in self.groups:
                data = self.df[self.df[self.category_col].astype(str) == group][self.value_col].values
                if len(data) < 3:
                    raise ValueError(f"Group '{group}' has too few observations (minimum 3 required).")
                self.group_data[group] = data
        else:
            # Wide mode: category_col is G1, value_col is G2
            g1_name = self.category_col
            g2_name = self.value_col
            
            self.groups = [g1_name, g2_name]
            
            g1_data = pd.to_numeric(self.df[g1_name], errors='coerce').dropna().values
            g2_data = pd.to_numeric(self.df[g2_name], errors='coerce').dropna().values
            
            if len(g1_data) < 3 or len(g2_data) < 3:
                raise ValueError("At least 3 numeric observations are required per group.")
                
            self.group_data[g1_name] = g1_data
            self.group_data[g2_name] = g2_data
            
            # Create a combined cleaned dataframe for plotting (balanced if needed, or just long format)
            # For plotting, long format is easier
            plot_df = []
            for g in self.groups:
                for val in self.group_data[g]:
                    plot_df.append({self.category_col: g, self.value_col: val})
            self.plot_df = pd.DataFrame(plot_df)

    def run_analysis(self):
        # 1. Descriptive Statistics
        for group in self.groups:
            data = self.group_data[group]
            mean = np.mean(data)
            var = np.var(data, ddof=1)
            std = np.sqrt(var)
            cv = (std / mean * 100) if mean != 0 else 0
            
            self.descriptive_stats[group] = {
                "Mean": mean,
                "Median": np.median(data),
                "Variance": var,
                "StdDev": std,
                "CV": cv,
                "Skewness": stats.skew(data),
                "Kurtosis": stats.kurtosis(data),
                "n": len(data)
            }
            
            # 2. Normality Test (Shapiro-Wilk)
            w_stat, p_val = stats.shapiro(data)
            self.normality_results[group] = {
                "W": w_stat,
                "p_value": p_val,
                "Interpretation": "Normal" if p_val > self.alpha else "Not Normal"
            }
            
        # 3. F-Test (Variance Ratio Test)
        g1, g2 = self.groups
        s1_sq = self.descriptive_stats[g1]["Variance"]
        s2_sq = self.descriptive_stats[g2]["Variance"]
        n1 = self.descriptive_stats[g1]["n"]
        n2 = self.descriptive_stats[g2]["n"]
        
        # F = max(s1^2, s2^2) / min(s1^2, s2^2)
        if s1_sq >= s2_sq:
            f_calc = s1_sq / s2_sq
            df1 = n1 - 1
            df2 = n2 - 1
        else:
            f_calc = s2_sq / s1_sq
            df1 = n2 - 1
            df2 = n1 - 1
            
        # Two-tailed p-value
        p_val = 2 * (1 - stats.f.cdf(f_calc, df1, df2))
        
        # Confidence Interval for Variance Ratio (s1^2 / s2^2)
        ratio = s1_sq / s2_sq
        f_crit_upper = stats.f.ppf(1 - self.alpha/2, n1-1, n2-1)
        f_crit_lower = stats.f.ppf(self.alpha/2, n1-1, n2-1)
        
        lower_ci = ratio / f_crit_upper
        upper_ci = ratio / f_crit_lower
        
        self.f_test_results = {
            "Variable": "Comparison",
            "Group1": g1,
            "Group2": g2,
            "Variance1": s1_sq,
            "Variance2": s2_sq,
            "F_value": f_calc,
            "p_value": p_val,
            "df1": df1,
            "df2": df2,
            "Lower_CI": lower_ci,
            "Upper_CI": upper_ci,
            "Conclusion": "Fail to reject Ho (variances are equal)" if p_val > self.alpha else "Reject Ho (variances are not equal)"
        }

    def get_interpretation(self):
        p_val = self.f_test_results["p_value"]
        if p_val > self.alpha:
            return (
                f"The p-value ({p_val:.4f}) is greater than the significance level (α = {self.alpha}). "
                "The result is non-significant, and the null hypothesis (Ho: σ₁² = σ₂²) is failed to be rejected. "
                "There is no statistically significant difference between the variances of the two groups."
            )
        else:
            return (
                f"The p-value ({p_val:.4f}) is less than or equal to the significance level (α = {self.alpha}). "
                "The result is statistically significant, and the null hypothesis (Ho: σ₁² = σ₂²) is rejected. "
                "The population variances of the two samples are significantly different."
            )

    def create_report(self):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        title = doc.add_heading("F-Test for Equality of Variances (Variance Ratio Test)", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading("1. Objective", level=1)
        doc.add_paragraph(
            f"The objective of this analysis is to statistically compare the variances of two independent samples "
            f"('{self.groups[0]}' and '{self.groups[1]}') using the F-test at a {self.alpha*100}% significance level."
        )
        
        # 2. Descriptive Statistics
        doc.add_heading("2. Descriptive Statistics", level=1)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        headers = ["Group", "Mean", "Median", "Variance", "Std. Dev", "CV (%)", "Skewness", "Kurtosis"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            
        for group in self.groups:
            s = self.descriptive_stats[group]
            row = table.add_row().cells
            row[0].text = str(group)
            row[1].text = f"{s['Mean']:.4f}"
            row[2].text = f"{s['Median']:.4f}"
            row[3].text = f"{s['Variance']:.4f}"
            row[4].text = f"{s['StdDev']:.4f}"
            row[5].text = f"{s['CV']:.2f}"
            row[6].text = f"{s['Skewness']:.4f}"
            row[7].text = f"{s['Kurtosis']:.4f}"
            
        # 3. Normality Test
        doc.add_heading("3. Normality Test (Shapiro-Wilk)", level=1)
        table_norm = doc.add_table(rows=1, cols=4)
        table_norm.style = 'Table Grid'
        hdr_n = ["Sample", "W Statistic", "P-value", "Interpretation"]
        for i, h in enumerate(hdr_n): table_norm.cell(0, i).text = h
        for group in self.groups:
            n = self.normality_results[group]
            row = table_norm.add_row().cells
            row[0].text = str(group)
            row[1].text = f"{n['W']:.4f}"
            row[2].text = f"{n['p_value']:.4f}"
            row[3].text = n['Interpretation']
            
        # 4. F-Test Results
        doc.add_heading("4. F-Test Analysis", level=1)
        table_f = doc.add_table(rows=1, cols=2)
        table_f.style = 'Table Grid'
        r = self.f_test_results
        f_data = [
            ("Sample 1 Group", r['Group1']),
            ("Sample 2 Group", r['Group2']),
            ("Variance 1 (s₁²)", f"{r['Variance1']:.4f}"),
            ("Variance 2 (s₂²)", f"{r['Variance2']:.4f}"),
            ("F Statistic (Calculated)", f"{r['F_value']:.4f}"),
            ("P-value (Two-tailed)", f"{r['p_value']:.4f}"),
            ("Numerator DF (df1)", str(r['df1'])),
            ("Denominator DF (df2)", str(r['df2'])),
            ("Lower CI (Ratio)", f"{r['Lower_CI']:.4f}"),
            ("Upper CI (Ratio)", f"{r['Upper_CI']:.4f}")
        ]
        for label, val in f_data:
            row = table_f.add_row().cells
            row[0].text = label
            row[1].text = str(val)
            
        doc.add_heading("5. Visualizations", level=1)
        
        # Prepare plotting data
        if self.mode == 'long':
            p_df = self.df
            p_cat = self.category_col
            p_val = self.value_col
        else:
            p_df = self.plot_df
            p_cat = self.category_col
            p_val = self.value_col

        # Plot 1: Boxplot
        plt.figure(figsize=(8, 6))
        sns.boxplot(x=p_cat, y=p_val, data=p_df, palette="Set2")
        plt.title(f"Box Plot Comparison")
        buf1 = io.BytesIO()
        plt.savefig(buf1, format='png', dpi=300)
        buf1.seek(0)
        doc.add_picture(buf1, width=Inches(5))
        plt.close()
        
        # Plot 2: Histogram
        plt.figure(figsize=(10, 6))
        for group in self.groups:
            sns.histplot(self.group_data[group], kde=True, label=str(group), alpha=0.5)
        plt.title(f"Distribution Comparison")
        plt.legend()
        buf2 = io.BytesIO()
        plt.savefig(buf2, format='png', dpi=300)
        buf2.seek(0)
        doc.add_picture(buf2, width=Inches(5))
        plt.close()
        
        doc.add_heading("6. Conclusion & Interpretation", level=1)
        doc.add_paragraph(self.get_interpretation())
        doc.add_paragraph(f"Decision: {r['Conclusion']}")
        
        doc.add_paragraph("\nREFERENCES:")
        doc.add_paragraph("Snedecor, G. W. & Cochran, W. G. (1989). Statistical Methods. Iowa State University Press.", style='List Bullet')
        doc.add_paragraph("Shapiro, S. S. & Wilk, M. B. (1965). An analysis of variance test for normality. Biometrika.", style='List Bullet')

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            pd.DataFrame(self.descriptive_stats).T.to_excel(writer, sheet_name='Descriptive Statistics')
            pd.DataFrame(self.normality_results).T.to_excel(writer, sheet_name='Normality Test')
            pd.DataFrame([self.f_test_results]).to_excel(writer, sheet_name='F-Test Results', index=False)
        output.seek(0)
        return output
