"""
Three-Factor CRD API Endpoints
"""

from fastapi import UploadFile, File, Form, Request, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
from security import limiter, validate_file
import pandas as pd
import numpy as np
import io
import traceback
import base64
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from doe.three_factor_crd_analyzer import ThreeFactorCRDAnalyzer


@limiter.limit("10/minute")
async def analyze_three_factor_crd(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    c_col: str = Form(..., max_length=100),
    rep_col: str = Form(None, max_length=100),
    resp_cols: list[str] = Form(...),
    alpha: float = Form(0.05),
    post_hoc: str = Form("lsd", max_length=10),
    mean_order: str = Form("desc", max_length=10),
    control_col: str = Form(None, max_length=100),
    notation: str = Form("letters", max_length=20)
):
    """Analyze three-factor CRD experiment."""
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        for col in resp_cols:
            if len(col) > 100:
                raise HTTPException(status_code=400, detail=f"Response column name too long: {col[:50]}...")

        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        results = {}
        
        for resp_col in resp_cols:
            if resp_col not in df.columns:
                continue
            
            analyzer = ThreeFactorCRDAnalyzer(df, a_col, b_col, c_col, resp_col, rep_col)
            analyzer.validate()
            anova_table = analyzer.run_anova()
            
            effects = ["Factor A", "Factor B", "Factor C", "Interaction AxB", 
                      "Interaction AxC", "Interaction BxC", "Interaction AxBxC"]
            
            means_data = {}
            for effect in effects:
                means_result = analyzer.calculate_means_and_comparisons(
                    effect=effect, alpha=alpha, method=post_hoc,
                    control=control_col, notation=notation
                )
                means_data[effect] = means_result
            
            results[resp_col] = {"anova": anova_table, "means_data": means_data}
        
        return {"status": "success", "results": results}
    
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


@limiter.limit("20/minute")
async def generate_three_factor_crd_plot(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    c_col: str = Form(..., max_length=100),
    rep_col: str = Form(None, max_length=100),
    resp_col: str = Form(..., max_length=100),
    effect: str = Form("Factor A", max_length=50),
    plot_type: str = Form("bar", max_length=20),
    error_bar: str = Form("se", max_length=20),
    alpha: float = Form(0.05),
    post_hoc: str = Form("lsd", max_length=10),
    mean_order: str = Form("desc", max_length=10),
    control_col: str = Form(None, max_length=100)
):
    """Generate plot for three-factor CRD effect."""
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = ThreeFactorCRDAnalyzer(df, a_col, b_col, c_col, resp_col, rep_col)
        analyzer.validate()
        analyzer.run_anova()
        
        means_result = analyzer.calculate_means_and_comparisons(
            effect=effect, alpha=alpha, method=post_hoc,
            control=control_col, notation='letters'
        )
        
        means_list = means_result["means"]
        labels = [m["level"] for m in means_list]
        values = [m["mean"] for m in means_list]
        groups = [m["group"] for m in means_list]
        
        if mean_order == "desc":
            sorted_data = sorted(zip(labels, values, groups), key=lambda x: x[1], reverse=True)
        else:
            sorted_data = sorted(zip(labels, values, groups), key=lambda x: x[1])
        
        labels, values, groups = zip(*sorted_data) if sorted_data else ([], [], [])
        
        plt.figure(figsize=(12, 6), facecolor='none')
        palette = sns.color_palette("husl", len(labels))
        
        if plot_type == 'bar':
            if error_bar == 'se':
                errors = [means_result["se_pooled"]] * len(values)
            elif error_bar == 'sd':
                errors = [means_result["se_pooled"] * np.sqrt(len(df) / len(labels))] * len(values)
            else:
                t_crit = stats.t.ppf(0.975, analyzer.df_error)
                errors = [means_result["se_pooled"] * t_crit] * len(values)
            
            bars = plt.bar(range(len(labels)), values, color=palette, edgecolor='white', linewidth=1.5, alpha=0.85)
            plt.errorbar(range(len(labels)), values, yerr=errors, fmt='none', ecolor='#333333', capsize=5, capthick=2, alpha=0.7)
            
            ylim = plt.ylim()
            y_range = ylim[1] - ylim[0]
            offset = y_range * 0.05
            
            for i, (bar, group) in enumerate(zip(bars, groups)):
                height = bar.get_height()
                err = errors[i]
                plt.text(bar.get_x() + bar.get_width()/2., height + err + offset, group, 
                        ha='center', va='bottom', color='#333333', fontweight='600', fontsize=11)
            
            plt.xticks(range(len(labels)), labels, rotation=45, ha='right')
            plt.ylabel(resp_col, labelpad=10, color='#555555', fontweight='bold')
            plt.title(f"{effect} - {resp_col}", pad=20, color='#333333', fontweight='bold', fontsize=14)
        
        sns.despine(left=True, bottom=False)
        plt.grid(axis='y', linestyle='--', alpha=0.5)
        ax = plt.gca()
        ax.patch.set_alpha(0)
        plt.tight_layout()
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=600, bbox_inches='tight', transparent=True, facecolor='none')
        plt.close()
        buf.seek(0)
        img_str = base64.b64encode(buf.read()).decode('utf-8')
        
        return {"status": "success", "image": img_str}
    
    except Exception as e:
        traceback.print_exc()
        plt.close()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


@limiter.limit("5/minute")
async def report_three_factor_crd(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    c_col: str = Form(..., max_length=100),
    rep_col: str = Form(None, max_length=100),
    resp_cols: list[str] = Form(...),
    alpha: float = Form(0.05),
    post_hoc: str = Form("lsd", max_length=10),
    mean_order: str = Form("desc", max_length=10),
    control_col: str = Form(None, max_length=100),
    notation: str = Form("letters", max_length=20)
):
    """Generate DOCX report for three-factor CRD."""
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        for col in resp_cols:
            if len(col) > 100:
                raise HTTPException(status_code=400, detail=f"Response column name too long: {col[:50]}...")

        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        doc = Document()
        doc.add_heading('Three-Factor CRD Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Factor A: {a_col}, Factor B: {b_col}, Factor C: {c_col}")
        
        for resp_col in resp_cols:
            if resp_col not in df.columns:
                continue
            
            doc.add_page_break()
            doc.add_heading(f"Analysis for: {resp_col}", level=1)
            
            analyzer = ThreeFactorCRDAnalyzer(df, a_col, b_col, c_col, resp_col, rep_col)
            analyzer.validate()
            anova_table = analyzer.run_anova()
            
            doc.add_heading('ANOVA Table', level=2)
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, h in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'P-value', 'Sig']):
                hdr[i].text = h
            
            sources = ["Factor A", "Factor B", "Factor C", "Interaction AxB", 
                      "Interaction AxC", "Interaction BxC", "Interaction AxBxC", "Error", "Total"]
            
            for source in sources:
                if source not in anova_table:
                    continue
                row = table.add_row().cells
                data = anova_table[source]
                row[0].text = source
                row[1].text = str(data['df'])
                row[2].text = f"{data['SS']:.4f}"
                row[3].text = f"{data['MS']:.4f}" if data['MS'] is not None else ""
                row[4].text = f"{data['F']:.4f}" if data['F'] is not None else ""
                row[5].text = f"{data['P']:.4f}" if data['P'] is not None else ""
                if data['P'] is not None:
                    sig = "**" if data['P'] <= 0.01 else ("*" if data['P'] <= 0.05 else "ns")
                    row[6].text = sig
            
            # Add means tables for each effect
            effects = ["Factor A", "Factor B", "Factor C", "Interaction AxB", 
                      "Interaction AxC", "Interaction BxC", "Interaction AxBxC"]
            
            for effect in effects:
                doc.add_heading(f'{effect} Means', level=2)
                means_result = analyzer.calculate_means_and_comparisons(
                    effect=effect, alpha=alpha, method=post_hoc,
                    control=control_col, notation=notation
                )
                
                table2 = doc.add_table(rows=1, cols=4)
                table2.style = 'Table Grid'
                h2 = table2.rows[0].cells
                for i, t in enumerate(['Level', 'Mean', 'Std Err', 'Group']):
                    h2[i].text = t
                
                for m in means_result["means"]:
                    r = table2.add_row().cells
                    r[0].text = m['level']
                    r[1].text = f"{m['mean']:.4f}"
                    r[2].text = f"{m['se']:.4f}"
                    r[3].text = m['group']
                
                doc.add_paragraph(f"SE(m): {means_result['se_pooled']:.4f}")
                doc.add_paragraph(f"SE(d): {means_result['sed']:.4f}")
                doc.add_paragraph(f"CV%: {means_result['cv']:.2f}%")
                if means_result['cd']:
                    doc.add_paragraph(f"CD ({alpha}): {means_result['cd']:.4f}")
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return StreamingResponse(
            f,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=ThreeFactorCRD_Report.docx"}
        )
    
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})
