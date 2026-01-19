import pandas as pd
import numpy as np
import statsmodels.api as sm
from sklearn.preprocessing import PolynomialFeatures
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import matplotlib.pyplot as plt
import seaborn as sns

class RegressionAnalyzer:
    def __init__(self, df, y_col, x_cols, model_type='linear', degree=2, alpha=0.05):
        self.df = df
        self.y_col = y_col
        self.x_cols = x_cols if isinstance(x_cols, list) else [x_cols]
        self.model_type = model_type  # 'linear', 'multiple', 'polynomial'
        self.degree = int(degree)
        self.alpha = float(alpha)
        
        self.model = None
        self.results = None
        self.summary_stats = {}
        self.coefficient_table = []
        self.anova_table = []
        
    def validate(self):
        # Numeric check
        cols_to_check = [self.y_col] + self.x_cols
        for col in cols_to_check:
            self.df[col] = pd.to_numeric(self.df[col], errors='coerce')
        
        # Drop missing values
        self.df = self.df.dropna(subset=cols_to_check)
        
        if len(self.df) < 2:
            raise ValueError("At least 2 observations required for regression.")
            
        if self.model_type == 'polynomial':
            if len(self.x_cols) > 1:
                raise ValueError("Polynomial regression implemented for a single independent variable in this version.")
            if self.degree < 1:
                raise ValueError("Polynomial degree must be at least 1.")
            if (self.degree + 1) >= len(self.df):
                raise ValueError(f"Insufficient data for polynomial degree {self.degree}. Need more than {self.degree + 1} observations.")
        else:
            if (len(self.x_cols) + 1) >= len(self.df):
                raise ValueError(f"Number of predictors ({len(self.x_cols)}) must be less than number of observations ({len(self.df)}) minus 1.")

    def run_analysis(self):
        Y = self.df[self.y_col]
        
        if self.model_type == 'polynomial':
            X_orig = self.df[self.x_cols[0]].values.reshape(-1, 1)
            poly = PolynomialFeatures(degree=self.degree, include_bias=True)
            X_transformed = poly.fit_transform(X_orig)
            # Create feature names for the summary
            feature_names = ['Intercept'] + [f"{self.x_cols[0]}^{i}" for i in range(1, self.degree + 1)]
            X = pd.DataFrame(X_transformed, columns=feature_names)
        else:
            X = sm.add_constant(self.df[self.x_cols])
            
        self.model = sm.OLS(Y, X)
        self.results = self.model.fit()
        
        # 1. Model Summary Stats
        self.summary_stats = {
            "R-Squared": self.results.rsquared,
            "Adj. R-Squared": self.results.rsquared_adj,
            "F-Statistic": self.results.fvalue,
            "Prob (F-statistic)": self.results.f_pvalue,
            "MSE (Residual)": self.results.mse_resid,
            "RMSE": np.sqrt(self.results.mse_resid),
            "N": int(self.results.nobs),
            "RegressionEquation": self._generate_equation()
        }
        
        # 2. Coefficient Table
        for idx, row in self.results.summary2().tables[1].iterrows():
            self.coefficient_table.append({
                "Variable": idx,
                "Coefficient": row['Coef.'],
                "StdError": row['Std.Err.'],
                "t_value": row['t'],
                "p_value": row['P>|t|'],
                "ConfLower": row['[0.025'],
                "ConfUpper": row['0.975]']
            })
            
        # 3. ANOVA Table
        anova_results = sm.stats.anova_lm(self.results, typ=1)
        for idx, row in anova_results.iterrows():
            self.anova_table.append({
                "Source": idx,
                "df": int(row['df']),
                "SS": row['sum_sq'],
                "MS": row['mean_sq'],
                "F": row['F'] if not pd.isna(row['F']) else None,
                "p": row['PR(>F)'] if not pd.isna(row['PR(>F)']) else None
            })

    def _generate_equation(self):
        params = self.results.params
        if self.model_type == 'polynomial':
            terms = [f"({params[0]:.4f})"]
            for i in range(1, len(params)):
                terms.append(f"({params[i]:.4f} * {self.x_cols[0]}^{i})")
            return f"{self.y_col} = " + " + ".join(terms)
        else:
            terms = [f"({params['const']:.4f})"]
            for col in self.x_cols:
                terms.append(f"({params[col]:.4f} * {col})")
            return f"{self.y_col} = " + " + ".join(terms)

    def get_interpretation(self):
        sig_vars = [c['Variable'] for c in self.coefficient_table if c['Variable'] != 'const' and c['p_value'] < self.alpha]
        r2 = self.summary_stats['R-Squared']
        
        interpretation = f"The regression model explains {r2*100:.2f}% of the variance in {self.y_col}. "
        
        if self.summary_stats['Prob (F-statistic)'] < self.alpha:
            interpretation += "The overall model is statistically significant. "
        else:
            interpretation += "The overall model is not statistically significant. "
            
        if sig_vars:
            interpretation += f"The following variables have a significant effect: {', '.join(sig_vars)}."
        else:
            interpretation += "No individual predictors proved to be significant at the selected alpha level."
            
        return interpretation

    def create_report(self):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        title = doc.add_heading("Regression Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading("1. Model Objective & Setup", level=1)
        doc.add_paragraph(f"Dependent Variable (Y): {self.y_col}")
        doc.add_paragraph(f"Independent Variables (X): {', '.join(self.x_cols)}")
        doc.add_paragraph(f"Model Type: {self.model_type.title()}")
        if self.model_type == 'polynomial': doc.add_paragraph(f"Degree: {self.degree}")
        
        doc.add_heading("2. Model Summary Statistics", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        stats_data = [
            ("R-Squared", f"{self.summary_stats['R-Squared']:.4f}"),
            ("Adjusted R-Squared", f"{self.summary_stats['Adj. R-Squared']:.4f}"),
            ("F-Statistic", f"{self.summary_stats['F-Statistic']:.4f}"),
            ("Prob (F-statistic)", f"{self.summary_stats['Prob (F-statistic)']:.4g}"),
            ("Residual Std. Error", f"{self.summary_stats['RMSE']:.4f}"),
            ("Observations (n)", str(self.summary_stats['N']))
        ]
        for label, val in stats_data:
            row = table.add_row().cells
            row[0].text, row[1].text = label, val
            
        doc.add_heading("3. Regression Coefficients", level=1)
        table_cf = doc.add_table(rows=1, cols=7)
        table_cf.style = 'Table Grid'
        heads = ["Variable", "Coef", "Std.Error", "t-value", "P-value", "Lower CI", "Upper CI"]
        for i, h in enumerate(heads): table_cf.cell(0, i).text = h
        for c in self.coefficient_table:
            row = table_cf.add_row().cells
            row[0].text = c['Variable']
            row[1].text = f"{c['Coefficient']:.4f}"
            row[2].text = f"{c['StdError']:.4f}"
            row[3].text = f"{c['t_value']:.4f}"
            row[4].text = f"{c['p_value']:.4g}"
            row[5].text = f"{c['ConfLower']:.4f}"
            row[6].text = f"{c['ConfUpper']:.4f}"
            
        doc.add_heading("4. Regression Equation", level=1)
        p = doc.add_paragraph()
        run = p.add_run(self.summary_stats['RegressionEquation'])
        run.bold = True
        run.italic = True
        
        doc.add_heading("5. Visualizations", level=1)
        
        # Plot 1: Actual vs Predicted
        plt.figure(figsize=(10, 6))
        sns.regplot(x=self.results.fittedvalues, y=self.df[self.y_col], scatter_kws={'alpha':0.5})
        plt.title("Actual vs Predicted")
        plt.xlabel("Predicted Values")
        plt.ylabel("Actual Values")
        buf1 = io.BytesIO()
        plt.savefig(buf1, format='png', dpi=300)
        buf1.seek(0)
        doc.add_picture(buf1, width=Inches(5.5))
        plt.close()
        
        # Plot 2: Residual Plot
        plt.figure(figsize=(10, 6))
        sns.residplot(x=self.results.fittedvalues, y=self.results.resid, lowess=True, color="g")
        plt.title("Residuals vs Fitted")
        plt.xlabel("Fitted Values")
        plt.ylabel("Residuals")
        buf2 = io.BytesIO()
        plt.savefig(buf2, format='png', dpi=300)
        buf2.seek(0)
        doc.add_picture(buf2, width=Inches(5.5))
        plt.close()
        
        doc.add_heading("6. Conclusion & Interpretation", level=1)
        doc.add_paragraph(self.get_interpretation())
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_excel(self):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            pd.DataFrame([self.summary_stats]).to_excel(writer, sheet_name='Model Summary', index=False)
            pd.DataFrame(self.coefficient_table).to_excel(writer, sheet_name='Coefficients', index=False)
            pd.DataFrame(self.anova_table).to_excel(writer, sheet_name='ANOVA', index=False)
        output.seek(0)
        return output
