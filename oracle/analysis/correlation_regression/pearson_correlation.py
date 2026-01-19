
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import io
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from scipy import stats

class PearsonCorrelationAnalyzer:
    def __init__(self, df, vars):
        self.df = df
        self.vars = vars
        self.results = {}
        self.heatmap_buf = None
        
    def validate(self):
        # 1. Check columns
        missing = [v for v in self.vars if v not in self.df.columns]
        if missing:
            raise ValueError(f"Missing variables: {missing}")
            
        # 2. Check numeric and missing
        self.data = self.df[self.vars].apply(pd.to_numeric, errors='coerce')
        if self.data.isnull().any().any():
            raise ValueError("Selected variables contain missing values. Please handle missing data first.")
            
        # 3. Check sample size
        n = self.data.shape[0]
        if n < 3:
            raise ValueError("At least 3 observations are required for correlation analysis.")
            
        # 4. Check variance
        if (self.data.var() == 0).any():
             raise ValueError("One or more variables have zero variance (constant values).")

        self.n = n
        
    def run_analysis(self):
        # Compute Correlation Matrix (r)
        corr_matrix = self.data.corr(method='pearson')
        
        # Compute p-values
        # scipy.stats.pearsonr returns (r, p) for two arrays
        # We need a matrix of p-values
        p_matrix = pd.DataFrame(index=self.vars, columns=self.vars)
        sig_matrix = pd.DataFrame(index=self.vars, columns=self.vars) # For stars
        
        for r in self.vars:
            for c in self.vars:
                if r == c:
                    p_matrix.loc[r, c] = 1.0 # Diagonal p-value irrelevant, usually 1 or 0
                    sig_matrix.loc[r, c] = ""
                else:
                    stat, p = stats.pearsonr(self.data[r], self.data[c])
                    p_matrix.loc[r, c] = p
                    
                    # Sig code
                    if p <= 0.01: sig = "**"
                    elif p <= 0.05: sig = "*"
                    else: sig = ""
                    sig_matrix.loc[r, c] = sig
                    
        self.results = {
            "corr_matrix": corr_matrix,
            "p_matrix": p_matrix,
            "sig_matrix": sig_matrix
        }
        
    def generate_heatmap(self):
        plt.figure(figsize=(10, 8))
        corr = self.results['corr_matrix']
        
        # Heatmap with annotations
        sns.heatmap(corr, annot=True, cmap='coolwarm', vmin=-1, vmax=1, center=0, fmt=".2f",
                    linewidths=0.5, linecolor='gray', square=True)
                    
        plt.title('Pearson Correlation Heatmap')
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        buf.seek(0)
        self.heatmap_buf = buf
        plt.close()
        return buf
        
    def create_report_doc(self):
        doc = Document()
        doc.add_heading('Pearson Correlation Report', 0)
        
        doc.add_heading('1. Method', level=1)
        doc.add_paragraph(f"Pearson correlation coefficients were computed for {len(self.vars)} variables using {self.n} observations. Significance was tested at p <= 0.05 (*) and p <= 0.01 (**).")
        
        # 2. Correlation Matrix Table
        doc.add_heading('2. Correlation Matrix', level=1)
        table = doc.add_table(rows=1, cols=len(self.vars) + 1)
        hdr = table.rows[0].cells
        hdr[0].text = "Variable"
        for i, v in enumerate(self.vars):
            hdr[i+1].text = v
            
        corr = self.results['corr_matrix']
        sig = self.results['sig_matrix']
        
        for i, r_var in enumerate(self.vars):
            row = table.add_row().cells
            row[0].text = r_var
            for j, c_var in enumerate(self.vars):
                val = corr.loc[r_var, c_var]
                s = sig.loc[r_var, c_var]
                # Lower triangle only? Usually full matrix is fine for report, or full symmetric
                if i == j:
                    txt = "1.00"
                else:
                    txt = f"{val:.2f} {s}".strip()
                row[j+1].text = txt
                
        # 3. Heatmap
        if self.heatmap_buf:
            doc.add_heading('3. Heatmap', level=1)
            self.heatmap_buf.seek(0)
            doc.add_picture(self.heatmap_buf, width=Inches(6))
            
        # 4. Interpretation
        doc.add_heading('4. Interpretation', level=1)
        
        # Generate pairwise interpretation for strong correlations
        # To avoid permutations, we iterate upper triangle
        pairs_reported = []
        
        for i in range(len(self.vars)):
            for j in range(i+1, len(self.vars)):
                v1 = self.vars[i]
                v2 = self.vars[j]
                r = corr.loc[v1, v2]
                p = self.results['p_matrix'].loc[v1, v2]
                
                # Check criteria
                desc = ""
                relation = "positive" if r > 0 else "negative"
                strength = ""
                significance = ""
                
                if abs(r) >= 0.7: strength = "strong"
                elif abs(r) >= 0.4: strength = "moderate"
                else: strength = "weak"
                
                if p <= 0.01: significance = "highly significant"
                elif p <= 0.05: significance = "significant"
                else: significance = "non-significant"
                
                # Compose text for interesting ones (sig or mod/strong)
                if p <= 0.05 or abs(r) >= 0.4:
                    sig_txt = "**" if p <= 0.01 else "*" if p <= 0.05 else ""
                    doc.add_paragraph(f"{v1} showed a {strength} and {significance} {relation} association with {v2} (r = {r:.2f}{sig_txt}).")
                    
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_output_excel(self):
        wb = Workbook()
        
        # 1. Correlation Matrix
        ws1 = wb.active
        ws1.title = "Correlation_Matrix_r"
        ws1.append(["Variable"] + self.vars)
        
        for r_var in self.vars:
            row = [r_var] + [self.results['corr_matrix'].loc[r_var, c] for c in self.vars]
            ws1.append(row)
            
        # 2. P-values
        ws2 = wb.create_sheet("P_Values")
        ws2.append(["Variable"] + self.vars)
        for r_var in self.vars:
            row = [r_var] + [self.results['p_matrix'].loc[r_var, c] for c in self.vars]
            ws2.append(row)
            
        # 3. Significance Codes
        ws3 = wb.create_sheet("Significance_Codes")
        ws3.append(["Code", "Meaning", "P-value Range"])
        ws3.append(["**", "Highly Significant", "<= 0.01"])
        ws3.append(["*", "Significant", "<= 0.05"])
        
        f = io.BytesIO()
        wb.save(f)
        f.seek(0)
        return f
