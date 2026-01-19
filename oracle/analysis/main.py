from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from slowapi import _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from pydantic import BaseModel, Field
from dotenv import load_dotenv
import shutil
import os

# Load environment variables
load_dotenv()
import pandas as pd
import numpy as np
import io
import traceback
import base64
import matplotlib
matplotlib.use('Agg') # Set backend before importing pyplot
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH



from experimental_design.lsd_analysis import LSDAnalyzer
from experimental_design.strip_plot_analysis import StripPlotAnalyzer
from experimental_design.crd_analysis import CRDAnalyzer
from experimental_design.factorial_crd_analysis import FactorialCRDAnalyzer
from experimental_design.three_factor_crd_analysis import ThreeFactorCRDAnalyzer
from experimental_design.pooled_crd_analysis import PooledCRDAnalyzer
from experimental_design.two_factor_pooled_crd import TwoFactorPooledCRDAnalyzer
from breeding.griffing_method1 import GriffingMethod1Analyzer
from breeding.griffing_method1_check import GriffingMethod1CheckAnalyzer
from breeding.griffing_method2 import GriffingMethod2Analyzer
from breeding.griffing_method2_check import GriffingMethod2CheckAnalyzer
from breeding.line_tester_analyzer import LineTesterAnalyzer
from breeding.genotypic_correlation import GenotypicCorrelationAnalyzer
from breeding.phenotypic_correlation import PhenotypicCorrelationAnalyzer
from breeding.genotypic_path_analysis import GenotypicPathAnalyzer
from breeding.phenotypic_path_analysis import PhenotypicPathAnalyzer
from breeding.mahalanobis_d2 import MahalanobisD2Analyzer
from breeding.genetic_parameters import GeneticParameterAnalyzer
from breeding.eberhart_russell import EberhartRussellAnalyzer
from hypothesis_testing.f_test import FTestAnalyzer
from hypothesis_testing.one_sample_t_test import OneSampleTTestAnalyzer
from hypothesis_testing.two_sample_t_test import TwoSampleTTestAnalyzer
from hypothesis_testing.paired_t_test import PairedTTestAnalyzer
from correlation_regression.regression_analyzer import RegressionAnalyzer
from api.models import *


from security import limiter
# Initialize Rate Limiter (previously here, now in security.py)
app = FastAPI()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Security Headers Middleware
class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        response.headers["Content-Security-Policy"] = "default-src 'self'; script-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com https://cdn.jsdelivr.net https://fonts.googleapis.com; style-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com https://fonts.googleapis.com; font-src 'self' https://fonts.gstatic.com; img-src 'self' data:; connect-src 'self' http://localhost:8000;"
        return response

app.add_middleware(SecurityHeadersMiddleware)

# CORS configuration
allow_origins_env = os.getenv("ALLOW_ORIGINS", "*")
allow_origins = [origin.strip() for origin in allow_origins_env.split(",")]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Helper function for significance notation
def get_sig(p_value):
    """Convert p-value to significance notation"""
    if p_value is None:
        return ""
    if p_value <= 0.01:
        return "**"
    elif p_value <= 0.05:
        return "*"
    else:
        return "ns"

# Helper for CRD
# Helper for CRD
def perform_crd_analysis(df, treat_col, resp_col, post_hoc, alpha, mean_order, control_group=None, notation='letters'):
    analyzer = CRDAnalyzer(df, treat_col, resp_col)
    analyzer.validate()
    anova = analyzer.run_anova()
    results = analyzer.run_post_hoc(method=post_hoc, alpha=alpha, order=mean_order, control_group=control_group, notation=notation)
    return analyzer, anova, results

from security import limiter, validate_file

@app.post("/analyze_crd")
@limiter.limit("10/minute")
async def analyze_crd(
    request: Request,
    file: UploadFile = File(...),
    treat_col: str = Form(..., max_length=100),
    rep_col: str = Form("", max_length=100),
    resp_col: list[str] = Form(...),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10),
    control_group: str = Form(None, max_length=100),
    notation: str = Form(None, max_length=50),
    comparison_mode: str = Form(None, max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
            raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        # Validate resp_col list contents
        for col in resp_col:
            if len(col) > 100:
                raise HTTPException(status_code=400, detail="Response column name too long.")

        final_notation = notation if notation is not None else (comparison_mode if comparison_mode is not None else "letters")
        
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        final_results = {}
        
        for col in resp_col:
             # Sanity check: ensure col exists
             if col not in df.columns: continue
             
             analyzer, anova, results = perform_crd_analysis(
                df, treat_col, col, post_hoc, alpha, mean_order, control_group, final_notation
             )
             
             final_results[col] = {
                 "anova": {k: {**v, "sig": get_sig(v['P'])} for k, v in anova.items()},
                 "stats": {
                    "means": results["means"],
                    "se": results["SEm"],
                    "sed": results["SEd"],
                    "cv": results["CV"],
                    "cd": results["CD"]
                 }
             }
        
        return {
            "status": "success",
            "results": final_results # Keyed by variable name
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_crd")
@limiter.limit("5/minute")
async def report_crd(
    request: Request,
    file: UploadFile = File(...),
    treat_col: str = Form(..., max_length=100),
    rep_col: str = Form("", max_length=100),
    resp_col: list[str] = Form(...),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10),
    control_group: str = Form(None, max_length=100),
    notation: str = Form(None, max_length=50),
    comparison_mode: str = Form(None, max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        for col in resp_col:
            if len(col) > 100:
                raise HTTPException(status_code=400, detail="Response column name too long.")

        final_notation = notation if notation is not None else (comparison_mode if comparison_mode is not None else "letters")
        
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        # Create a single document
        doc = Document()
        doc.add_heading('CRD Analysis Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Treatment Factor: {treat_col}")
        if rep_col: doc.add_paragraph(f"Replication Factor: {rep_col}")
        
        for col in resp_col:
            if col not in df.columns: continue
            
            analyzer, _, _ = perform_crd_analysis(
                df, treat_col, col, post_hoc, alpha, mean_order, control_group, final_notation
            )
            
            # Use a helper to append to existing doc instead of creating new one
            # We need to expose a method in analyzer or manually replicate
            # For simplicity, let's manually add a Section Break or Heading
            doc.add_page_break()
            doc.add_heading(f"Analysis for Response: {col}", level=1)
            
            # Re-implement report parts here or modify Analyzer to accept doc
            # Let's modify logic to call a helper that takes `doc`
            
            # Model & ANOVA
            doc.add_heading('1. ANOVA Summary', level=2)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, t in enumerate(['Source', 'DF', 'SS', 'MS', 'F-value', 'P-value']):
                hdr[i].text = t
            
            for k in ["Treatments", "Error", "Total"]:
                row = table.add_row().cells
                dat = analyzer.anova_table[k]
                row[0].text = k
                row[1].text = str(dat['df'])
                row[2].text = f"{dat['SS']:.2f}"
                row[3].text = f"{dat['MS']:.2f}" if dat['MS'] else ""
                row[4].text = f"{dat['F']:.2f}" if dat['F'] else ""
                if dat['P'] is not None:
                    sig = "**" if dat['P']<=0.01 else ("*" if dat['P']<=0.05 else "ns")
                    row[5].text = f"{dat['P']:.4f} {sig}"
            
            # Means
            doc.add_heading('2. Mean Comparison', level=2)
            table2 = doc.add_table(rows=1, cols=5)
            table2.style = 'Table Grid'
            h2 = table2.rows[0].cells
            for i, t in enumerate(['Treatment', 'Mean', 'Std Dev', 'Std Err', 'Group']):
                h2[i].text = t
            
            for m in analyzer.results['means']:
                 r = table2.add_row().cells
                 r[0].text = m['level']
                 r[1].text = f"{m['mean']:.2f}"
                 r[2].text = f"{m['sd']:.2f}"
                 r[3].text = f"{m['se']:.2f}"
                 r[4].text = m['group']
                 
            doc.add_paragraph(f"\nSE(m): {analyzer.results['SEm']:.2f}")
            doc.add_paragraph(f"SE(d): {analyzer.results['SEd']:.2f}")
            doc.add_paragraph(f"CV%: {analyzer.results['CV']:.2f}%")
            if analyzer.results['CD']:
                doc.add_paragraph(f"CD ({alpha}): {analyzer.results['CD']:.2f}")
            
            # Interpret
            doc.add_heading('3. Interpretation', level=2)
            p = analyzer.anova_table["Treatments"]["P"]
            if p <= analyzer.alpha:
                doc.add_paragraph(f"The analysis revealed significant differences among treatments (p={p:.4f}).")
                # Find best
                means_series = analyzer.results['df_means']
                best_t = means_series.idxmax()
                best_val = means_series.max()
                doc.add_paragraph(f"Treatment {best_t} recorded the highest mean value of {best_val:.2f}.")
            else:
                doc.add_paragraph("No significant differences were observed among treatments.")

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return StreamingResponse(
             f,
             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
             headers={"Content-Disposition": "attachment; filename=CRD_Multi_Variable_Report.docx"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/generate_crd_plot")
@limiter.limit("20/minute")
async def generate_crd_plot(
    request: Request,
    file: UploadFile = File(...),
    treat_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    plot_type: str = Form("bar", max_length=20),
    error_bar: str = Form("se", max_length=10),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10),
    control_group: str = Form(None, max_length=100),
    notation: str = Form(None, max_length=50),
    comparison_mode: str = Form(None, max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        # Backward compatibility
        final_notation = notation if notation is not None else (comparison_mode if comparison_mode is not None else "letters")
        
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        # We need to run analysis to get Means & Letters (for Bar plot)
        analyzer, anova, results = perform_crd_analysis(
            df, treat_col, resp_col, post_hoc, alpha, mean_order, control_group, final_notation
        )
        
        plt.figure(figsize=(10, 6), facecolor='none')
        sns.set_style("whitegrid", {'grid.linestyle': '--', 'axes.edgecolor': '#cccccc'})
        sns.set_context("talk", font_scale=0.9)
        
        # Consistent Color Palette (Blue/Orange/Green theme similar to image)
        # Using a custom palette or a standard one that looks "modern"
        palette = sns.color_palette("viridis", n_colors=len(results['means']))
        
        # Prepare Data Labels (Sorted by mean_order)
        means_data = analyzer.results['means'] 
        labels = [m['level'] for m in means_data]
            
        if plot_type == 'bar':
             # Data preparation
            values = [m['mean'] for m in means_data]
            groups = [m['group'] for m in means_data]
            
            # Error Bars
            errors = []
            for m in means_data:
                if error_bar == 'sd':
                    errors.append(m['sd'])
                elif error_bar == 'ci':
                     t_val = 1.96 
                     errors.append(t_val * m['se'])
                else: # SE
                    errors.append(m['se'])
                    
            x_pos = np.arange(len(labels))
            
            # Bar Plot with grid and softer look
            bars = plt.bar(x_pos, values, yerr=errors, align='center', alpha=0.85, capsize=5, color=palette, edgecolor='white', linewidth=1.5)
            plt.xticks(x_pos, labels, rotation=0) # Keep rotation 0 if possible, or 45 if long
            
            # Add Letters
            ylim = plt.ylim()
            y_range = ylim[1] - ylim[0]
            offset = y_range * 0.05
            
            for i, bar in enumerate(bars):
                height = bar.get_height()
                err = errors[i]
                label = groups[i]
                plt.text(bar.get_x() + bar.get_width()/2., height + err + offset, label, ha='center', va='bottom', color='#333333', fontweight='600', fontsize=12)
                
            plt.xlabel(treat_col, labelpad=10, color='#555555', fontweight='bold')
            plt.ylabel(resp_col, labelpad=10, color='#555555', fontweight='bold')
            plt.title(f"{resp_col} by {treat_col}", pad=20, color='#333333', fontweight='bold')

        elif plot_type == 'box':
            # Boxplot with "Raincloud" style dots
            ax = sns.boxplot(x=treat_col, y=resp_col, data=df, palette=palette, order=labels, width=0.5, fliersize=0) # fliersize=0 hide outliers from box, shown in strip
            
            # Make box transparent/lighter to see style
            for patch in ax.patches:
                r, g, b, a = patch.get_facecolor()
                patch.set_facecolor((r, g, b, 0.7)) # 70% opacity
                patch.set_edgecolor((r, g, b, 1))

            # Add Strip Plot (The dots)
            sns.stripplot(x=treat_col, y=resp_col, data=df, order=labels, size=5, color=".3", linewidth=0, alpha=0.6, jitter=True)
            
            plt.xlabel(treat_col, labelpad=10, color='#555555', fontweight='bold')
            plt.ylabel(resp_col, labelpad=10, color='#555555', fontweight='bold')
            plt.title(f"Distribution of {resp_col}", pad=20, color='#333333', fontweight='bold')
            
        elif plot_type == 'violin':
            sns.violinplot(x=treat_col, y=resp_col, data=df, palette=palette, order=labels, inner="quartile", alpha=0.7)
            sns.stripplot(x=treat_col, y=resp_col, data=df, order=labels, size=4, color="white", linewidth=0, alpha=0.5, jitter=True)
            
            plt.xlabel(treat_col, labelpad=10, color='#555555', fontweight='bold')
            plt.ylabel(resp_col, labelpad=10, color='#555555', fontweight='bold')
            plt.title(f"Distribution of {resp_col}", pad=20, color='#333333', fontweight='bold')

        # Clean Spines
        sns.despine(left=True, bottom=False)
        plt.grid(axis='y', linestyle='--', alpha=0.5) # Explicit vertical grid
        
        # Set transparent background for axes
        ax = plt.gca()
        ax.patch.set_alpha(0)
        
        plt.tight_layout()
        
        # Save to buffer with transparent background and high quality
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=600, bbox_inches='tight', transparent=True, facecolor='none')
        plt.close()
        buf.seek(0)
        img_str = base64.b64encode(buf.read()).decode('utf-8')
        
        return {"status": "success", "image": img_str}

    except Exception as e:
        traceback.print_exc()
        plt.close()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Three-Factor CRD Endpoints
from api.three_factor_crd_endpoints import (
    analyze_three_factor_crd,
    generate_three_factor_crd_plot,
    report_three_factor_crd
)

app.post("/analyze_three_factor_crd")(analyze_three_factor_crd)
app.post("/generate_three_factor_crd_plot")(generate_three_factor_crd_plot)
app.post("/report_three_factor_crd")(report_three_factor_crd)

# Helper for Strip Plot
def perform_strip_analysis(df, rep_col, a_col, b_col, resp_col, post_hoc, alpha, mean_order):
    analyzer = StripPlotAnalyzer(df, rep_col, a_col, b_col, resp_col)
    analyzer.validate()
    anova = analyzer.run_anova()
    results = analyzer.run_post_hoc(method=post_hoc, alpha=alpha, order=mean_order)
    return analyzer, anova, results

@app.post("/analyze_strip_plot")
@limiter.limit("10/minute")
async def analyze_strip_plot(
    request: Request,
    file: UploadFile = File(...),
    rep_col: str = Form(..., max_length=100),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer, anova, results = perform_strip_analysis(
            df, rep_col, a_col, b_col, resp_col, post_hoc, alpha, mean_order
        )
        
        # Serialize Results
        # Results is a dictionary { "Factor A": {means, grouping, SE, CD}, ... }
        # We need to convert Pandas Series to lists/dicts
        
        serialized_results = {}
        for key, val in results.items():
            means_list = []
            sds_series = val['sds'].to_dict() # helper
            ses_series = val['ses'].to_dict() # helper
            
            for idx, mean_val in val['means'].items():
                # idx is likely string, but let's be safe
                sd_val = sds_series.get(idx, 0.0)
                se_val = ses_series.get(idx, 0.0)
                
                means_list.append({
                    "level": str(idx),
                    "mean": float(mean_val),
                    "sd": float(sd_val) if not pd.isna(sd_val) else 0.0,
                    "se": float(se_val) if not pd.isna(se_val) else 0.0,
                    "group": val['grouping'].get(idx, "-")
                })
            
            serialized_results[key] = {
                "means": means_list,
                "se_pooled": float(val['SE']) if val['SE'] else 0, # Renamed to avoid confusion, but frontend uses result.se currently
                "se": float(val['SE']) if val['SE'] else 0, # Keep for backward compat with JS if needed
                "sed": float(val['SEd']) if val['SEd'] else 0,
                "cv": float(val['CV']) if val['CV'] else 0,
                "cd": float(val['CD']) if val['CD'] else 0
            }

        return {
            "status": "success",
            "anova": {k: {**v, "sig": get_sig(v['P'])} for k, v in anova.items()},
            "post_hoc": serialized_results
        }

    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_strip_plot")
@limiter.limit("5/minute")
async def report_strip_plot(
    request: Request,
    file: UploadFile = File(...),
    rep_col: str = Form(..., max_length=100),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer, _, _ = perform_strip_analysis(
            df, rep_col, a_col, b_col, resp_col, post_hoc, alpha, mean_order
        )
        
        report_buffer = analyzer.create_report()
        filename = "StripPlot_Analysis_Report.docx"
        
        return StreamingResponse(
            report_buffer, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# ... (Existing LSD code and mounts)


# Serve Frontend
# Assuming we run this from 'backend' folder, frontend is at '../frontend'
# But better to use absolute or relative to this file
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# ROOT of the project, two levels up from oracle/analysis/
ROOT_DIR = os.path.dirname(os.path.dirname(BASE_DIR))
FRONTEND_DIR = os.path.join(ROOT_DIR, "cloudflare", "data-anal")

# Helper to avoid code duplication
def perform_analysis(df, row_col, col_col, treat_col, resp_col, post_hoc, alpha, mean_order):
    analyzer = LSDAnalyzer(df, row_col, col_col, treat_col, resp_col)
    analyzer.validate()
    anova = analyzer.run_anova()
    grouping = analyzer.run_post_hoc(method=post_hoc, alpha=alpha, order=mean_order)
    interpretation = analyzer.interpret()
    return analyzer, anova, grouping, interpretation

# API Routes first
@app.post("/analyze")
@limiter.limit("10/minute")
async def analyze_data(
    request: Request,
    file: UploadFile = File(...),
    row_col: str = Form(..., max_length=100),
    col_col: str = Form(..., max_length=100),
    treat_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        # Read file
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer, anova, grouping, interpretation = perform_analysis(
            df, row_col, col_col, treat_col, resp_col, post_hoc, alpha, mean_order
        )
        
        # Format results for JSON
        means_data = []
        sds = analyzer.df.groupby(treat_col)[resp_col].std()
        
        # Helper dictionary for fast lookup of sorted means
        means_dict = analyzer.means.to_dict()
        
        # Iterate over treatments alphabetically/numerically to preserve stable order in table
        all_treats = sorted([str(t) for t in analyzer.means.index])
        
        for t in all_treats:
            # We need to find the original key type if it wasn't string, but we cast to str in validate
            # However, analyzer.means index is what grouping uses.
            # Let's ensure we access correctly.
            mean_val = means_dict.get(t, means_dict.get(int(t) if t.isdigit() else t))
            
            means_data.append({
                "treatment": t,
                "mean": float(mean_val) if mean_val is not None else 0.0,
                "sd": float(sds.get(t, sds.get(int(t) if t.isdigit() else t, 0))),
                "se": float(analyzer.SE_m), # SE is constant for balanced LSD
                "group": grouping.get(t, grouping.get(int(t) if t.isdigit() else t, "-"))
            })
            
        return {
            "status": "success",
            "anova": {k: {**v, "sig": get_sig(v['P'])} for k, v in anova.items()},
            "means": means_data,
            "precision": {
                "sem": analyzer.SE_m,
                "sed": analyzer.SE_d,
                "cv": analyzer.CV,
                "cd": getattr(analyzer, 'CD', 0)
            },
            "interpretation": interpretation
        }

    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report")
@limiter.limit("5/minute")
async def download_report(
    request: Request,
    file: UploadFile = File(...),
    row_col: str = Form(..., max_length=100),
    col_col: str = Form(..., max_length=100),
    treat_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer, _, _, _ = perform_analysis(
            df, row_col, col_col, treat_col, resp_col, post_hoc, alpha, mean_order
        )
        
        report_buffer = analyzer.create_report()
        
        filename = f"LSD_Analysis_Report.docx"
        return StreamingResponse(
            report_buffer, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Helper for Factorial CRD
def perform_factorial_crd_analysis(df, a_col, b_col, resp_col, post_hoc="lsd", alpha=0.05, mean_order="desc", rep_col=None, control_group=None, notation="alphabet"):
    analyzer = FactorialCRDAnalyzer(df, a_col, b_col, resp_col, rep_col=rep_col)
    analyzer.validate()
    anova = analyzer.run_anova()
    results = analyzer.run_post_hoc(method=post_hoc, alpha=alpha, is_ascending=(mean_order == "asc"), control_group=control_group, notation=notation)
    return analyzer, anova, results

@app.post("/analyze_factorial_crd")
@limiter.limit("10/minute")
async def analyze_factorial_crd(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    rep_col: str = Form(None, max_length=100),
    resp_col: list[str] = Form(...),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10),
    control_group: str = Form(None, max_length=100),
    notation: str = Form("alphabet", max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        for col in resp_col:
            if len(col) > 100:
                raise HTTPException(status_code=400, detail="Response column name too long.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        final_results = {}
        for col in resp_col:
            if col not in df.columns: continue
            
            analyzer, anova, results = perform_factorial_crd_analysis(
                df, a_col, b_col, col, post_hoc, alpha, mean_order, rep_col=rep_col, control_group=control_group, notation=notation
            )
            
            # Serialize results
            serialized_results = {}
            for key, val in results.items():
                means_list = []
                sds_series = val['sds'].to_dict()
                ses_series = val['ses'].to_dict()
                
                for idx, mean_val in val['means'].items():
                    sd_val = sds_series.get(idx, 0.0)
                    se_val = ses_series.get(idx, 0.0)
                    
                    means_list.append({
                        "level": str(idx),
                        "mean": float(mean_val),
                        "sd": float(sd_val) if not pd.isna(sd_val) else 0.0,
                        "se": float(se_val) if not pd.isna(se_val) else 0.0,
                        "group": val['grouping'].get(idx, "-")
                    })
                
                serialized_results[key] = {
                    "means": means_list,
                    "se_pooled": float(val['SE']) if val['SE'] else 0,
                    "se": float(val['SE']) if val['SE'] else 0, # Legacy
                    "sed": float(val['SEd']) if val['SEd'] else 0,
                    "cv": float(val['CV']) if val['CV'] else 0,
                    "cd": float(val['CD']) if val['CD'] else 0
                }
            
            final_results[col] = {
                "anova": {k: {**v, "sig": get_sig(v['P'])} for k, v in anova.items()},
                "results": serialized_results
            }

        return {
            "status": "success",
            "results": final_results
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/generate_factorial_crd_plot")
@limiter.limit("20/minute")
async def generate_factorial_crd_plot(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    rep_col: str = Form(None, max_length=100),
    resp_col: str = Form(..., max_length=100),
    plot_type: str = Form("bar", max_length=20), 
    error_bar: str = Form("se", max_length=10), 
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10),
    control_group: str = Form(None, max_length=100),
    notation: str = Form("alphabet", max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        # Focus on Interaction (AxB)
        analyzer, anova, results = perform_factorial_crd_analysis(
            df, a_col, b_col, resp_col, post_hoc, alpha, mean_order, rep_col=rep_col, control_group=control_group, notation=notation
        )
        
        plt.figure(figsize=(10, 6), facecolor='none')
        sns.set_style("whitegrid", {'grid.linestyle': '--', 'axes.edgecolor': '#cccccc'})
        sns.set_context("talk", font_scale=0.8)
        
        interaction_res = results["Interaction AxB"]
        means_data = interaction_res['means'] 
        labels = [str(l) for l in means_data.index]
        df['Interaction'] = df[a_col].astype(str) + " : " + df[b_col].astype(str)
        palette = sns.color_palette("viridis", n_colors=len(labels))
            
        if plot_type == 'bar':
            values = [float(v) for v in means_data.values]
            groups = [interaction_res['grouping'].get(l, "-") for l in labels]
            
            errors = []
            for l in labels:
                if error_bar == 'sd':
                    errors.append(float(interaction_res['sds'].get(l, 0)))
                elif error_bar == 'ci':
                     from scipy import stats as scipy_stats
                     t_val = scipy_stats.t.ppf(1 - alpha/2, analyzer.df_E)
                     errors.append(t_val * float(interaction_res['ses'].get(l, 0)))
                else: # SE
                    errors.append(float(interaction_res['ses'].get(l, 0)))
            
            x_pos = np.arange(len(labels))
            bars = plt.bar(x_pos, values, yerr=errors, align='center', alpha=0.85, capsize=6, color=palette, edgecolor='white', linewidth=1.5)
            plt.xticks(x_pos, labels, rotation=45, ha='right')
            
            ylim = plt.ylim()
            y_range = ylim[1] - ylim[0]
            offset = y_range * 0.05
            for i, bar in enumerate(bars):
                height = bar.get_height()
                err = errors[i]
                label = groups[i]
                plt.text(bar.get_x() + bar.get_width()/2., height + err + offset, label, ha='center', va='bottom', color='#333333', fontweight='600', fontsize=11)
                
            plt.xlabel(f"{a_col} x {b_col}", labelpad=10, color='#555555', fontweight='bold')
            plt.ylabel(resp_col, labelpad=10, color='#555555', fontweight='bold')
            plt.title(f"Mean {resp_col} (±{error_bar.upper()})", pad=20, color='#333333', fontweight='bold')

        elif plot_type == 'box':
            ax = sns.boxplot(x='Interaction', y=resp_col, data=df, palette=palette, order=labels, width=0.6, fliersize=4)
            for patch in ax.patches:
                r, g, b, a = patch.get_facecolor()
                patch.set_facecolor((r, g, b, 0.6))
            sns.stripplot(x='Interaction', y=resp_col, data=df, order=labels, size=4, color=".3", alpha=0.5, jitter=True)
            plt.xticks(rotation=45, ha='right')
            plt.title(f"Box Plot: {resp_col} Distribution", pad=20, color='#333333', fontweight='bold')

        elif plot_type == 'violin':
            sns.violinplot(x='Interaction', y=resp_col, data=df, palette=palette, order=labels, inner="quartile", alpha=0.6)
            sns.stripplot(x='Interaction', y=resp_col, data=df, order=labels, size=3, color="white", alpha=0.4, jitter=True)
            plt.xticks(rotation=45, ha='right')
            plt.title(f"Violin Plot: {resp_col} Density", pad=20, color='#333333', fontweight='bold')

        sns.despine(left=True, bottom=False)
        plt.grid(axis='y', linestyle='--', alpha=0.4)
        ax = plt.gca()
        ax.patch.set_alpha(0)
        plt.tight_layout()
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight', transparent=True)
        plt.close()
        buf.seek(0)
        img_str = base64.b64encode(buf.read()).decode('utf-8')
        return {"status": "success", "image": img_str}

    except Exception as e:
        traceback.print_exc()
        if 'plt' in locals(): plt.close()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


@app.post("/report_factorial_crd")
@limiter.limit("5/minute")
async def report_factorial_crd(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    rep_col: str = Form(None, max_length=100),
    resp_col: list[str] = Form(...),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10),
    control_group: str = Form(None, max_length=100),
    notation: str = Form("alphabet", max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        for col in resp_col:
            if len(col) > 100:
                raise HTTPException(status_code=400, detail="Response column name too long.")
        
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        doc = Document()
        doc.add_heading('Factorial CRD Analysis Multi-Variable Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Factor A: {a_col} | Factor B: {b_col}")
        if rep_col: doc.add_paragraph(f"Replication Column: {rep_col}")

        for col in resp_col:
            if col not in df.columns: continue
            
            analyzer, _, _ = perform_factorial_crd_analysis(
                 df, a_col, b_col, col, post_hoc, alpha, mean_order, rep_col=rep_col, control_group=control_group, notation=notation
            )
            analyzer.append_to_report(doc)
            
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return StreamingResponse(
             f,
             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
             headers={"Content-Disposition": "attachment; filename=Factorial_CRD_Multi_Report.docx"}
        )
    except Exception as e:
         traceback.print_exc()
         return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Helper for Three Factor CRD
# Helper for Three Factor CRD
def perform_three_factor_crd_analysis(df, a_col, b_col, c_col, resp_col, post_hoc, alpha, mean_order, rep_col=None, control_col=None, notation=None):
    analyzer = ThreeFactorCRDAnalyzer(df, a_col, b_col, c_col, resp_col, rep_col=rep_col)
    analyzer.control_col = control_col
    if notation: analyzer.notation = notation
    analyzer.validate()
    anova = analyzer.run_anova()
    results = analyzer.run_post_hoc(method=post_hoc, alpha=alpha, order=mean_order)
    return analyzer, anova, results

@app.post("/analyze_three_factor_crd")
@limiter.limit("10/minute")
async def analyze_three_factor_crd(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    c_col: str = Form(..., max_length=100),
    rep_col: str = Form(None, max_length=100),
    resp_cols: list[str] = Form(...),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10),
    control_col: str = Form(None, max_length=100),
    notation: str = Form(None, max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        for col in resp_cols:
            if len(col) > 100:
                raise HTTPException(status_code=400, detail="Response column name too long.")
        data = await file.read()
        df = pd.read_csv(io.BytesIO(data))
        # Ensure column names are stripped of whitespace to match frontend
        df.columns = df.columns.str.strip()
        
        full_results = {}
        
        for col in resp_cols:
            if col not in df.columns: continue
            
            # Pass control_col to the helper
            analyzer, anova, results = perform_three_factor_crd_analysis(
                df, a_col, b_col, c_col, col, post_hoc, alpha, mean_order, rep_col=rep_col, control_col=control_col, notation=notation
            )
            
            # Serialize
            serial_res = {}
            for key, val in results.items():
                 means_list = []
                 sds = val['sds'].to_dict()
                 ses = val['ses'].to_dict()
                 grp = val['grouping']
                 
                 for idx, m_val in val['means'].items():
                     means_list.append({
                         "level": str(idx),
                         "mean": float(m_val),
                         "sd": float(sds.get(idx, 0)),
                         "se": float(ses.get(idx, 0)),
                         "group": grp.get(idx, "-")
                     })
                     
                 serial_res[key] = {
                     "means": means_list,
                     "se_pooled": float(val['SE']) if val['SE'] else 0,
                     "sed": float(val['SEd']) if val['SEd'] else 0,
                     "cv": float(val['CV']) if val['CV'] else 0,
                     "cd": float(val['CD']) if val['CD'] else 0
                 }

            # Serialize ANOVA
            serial_anova = {k: {**v, "sig": get_sig(v['P'])} for k, v in anova.items()}
            
            full_results[col] = {
                "anova": serial_anova,
                "means_data": serial_res 
            }
            
        return {"status": "success", "results": full_results}
        
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


@app.post("/generate_three_factor_crd_plot")
@limiter.limit("20/minute")
async def generate_three_factor_crd_plot(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    c_col: str = Form(..., max_length=100),
    rep_col: str = Form(None, max_length=100),
    resp_col: str = Form(..., max_length=100),
    plot_type: str = Form(..., max_length=20),
    error_bar: str = Form("se", max_length=10),
    effect: str = Form(..., max_length=50),
    alpha: float = Form(0.05),
    post_hoc: str = Form("lsd", max_length=50),
    mean_order: str = Form("desc", max_length=10),
    control_col: str = Form(None, max_length=100),
    notation: str = Form(None, max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        content = await file.read()
        df = pd.read_csv(io.BytesIO(content))
        
        # Run analysis to get means and letters
        analyzer, anova, results = perform_three_factor_crd_analysis(
            df, a_col, b_col, c_col, resp_col, post_hoc, alpha, mean_order, rep_col=rep_col, control_col=control_col, notation=notation
        )
        
        effect_key = effect # e.g. "Factor A", "Interaction AxB", etc.
        if effect_key not in results:
             return JSONResponse(status_code=400, content={"status": "error", "message": f"Effect '{effect}' not found in results."})
        
        eff_res = results[effect_key]
        means = eff_res['means']
        sds = eff_res['sds']
        ses = eff_res['ses']
        grouping = eff_res['grouping']
        
        plt.figure(figsize=(10, 6))
        
        # Decide data for plotting
        x_levels = [str(x) for x in means.index]
        y_values = means.values
        
        # Color palette
        colors = sns.color_palette("viridis", len(x_levels))
        
        if plot_type == 'bar':
             # Error bars
             yerr = None
             if error_bar == 'se':
                 yerr = ses.values
             elif error_bar == 'sd':
                 yerr = sds.values
             elif error_bar == 'ci':
                 # calculate 95% CI roughly using t or z approx? Or just 1.96 * SE
                 # Better to use t-dist if possible, but df_E is available in analyzer
                 # t_val = stats.t.ppf(0.975, analyzer.df_E)
                 # ci = t_val * ses
                 # For simplicity in this rough plot endpoint:
                 yerr = 1.96 * ses.values
            
             bars = plt.bar(x_levels, y_values, yerr=yerr, capsize=5, color=colors, alpha=0.8, edgecolor='black')
             
             # Add letters
             y_max = max(y_values + (yerr if yerr is not None else 0))
             offset = y_max * 0.05
             for i, rect in enumerate(bars):
                 height = rect.get_height()
                 label = grouping.get(means.index[i], "")
                 plt.text(rect.get_x() + rect.get_width()/2., height + (yerr[i] if yerr is not None else 0) + offset/2,
                          label, ha='center', va='bottom', fontsize=11, fontweight='bold')
             
             plt.ylabel(f"Mean {resp_col}")
             
        elif plot_type == 'box':
             # Need raw data subset for this effect
             # We need to construct the grouping column again
             # The result dictionary doesn't hold raw data, only summaries.
             # We can reconstruct group column in df
             
             group_cols = []
             if effect == "Factor A": group_cols = [a_col]
             elif effect == "Factor B": group_cols = [b_col]
             elif effect == "Factor C": group_cols = [c_col]
             
             # Interactions
             # Mapping names back to cols is tricky if we don't have a map.
             # But we can infer from effect string
             # "Interaction AxB" -> [a_col, b_col]
             # "Interaction AxC" -> [a_col, c_col]
             # "Interaction BxC" -> [b_col, c_col]
             # "Interaction AxBxC" -> [a_col, b_col, c_col]
             
             if effect == "Factor A": group_cols = [a_col]
             elif effect == "Factor B": group_cols = [b_col]
             elif effect == "Factor C": group_cols = [c_col]
             elif effect == "Interaction AxB": group_cols = [a_col, b_col]
             elif effect == "Interaction AxC": group_cols = [a_col, c_col]
             elif effect == "Interaction BxC": group_cols = [b_col, c_col]
             elif effect == "Interaction AxBxC": group_cols = [a_col, b_col, c_col]
             
             plot_data = df.copy()
             if len(group_cols) > 1:
                 plot_data['ActiveGroup'] = plot_data.apply(lambda x: " : ".join([str(x[c]) for c in group_cols]), axis=1)
             else:
                 plot_data['ActiveGroup'] = plot_data[group_cols[0]]
                 
             sns.boxplot(x='ActiveGroup', y=resp_col, data=plot_data, palette="viridis")
             plt.ylabel(resp_col)

        elif plot_type == 'violin':
             # Similar logic to box
             group_cols = []
             if effect == "Factor A": group_cols = [a_col]
             elif effect == "Factor B": group_cols = [b_col]
             elif effect == "Factor C": group_cols = [c_col]
             elif effect == "Interaction AxB": group_cols = [a_col, b_col]
             elif effect == "Interaction AxC": group_cols = [a_col, c_col]
             elif effect == "Interaction BxC": group_cols = [b_col, c_col]
             elif effect == "Interaction AxBxC": group_cols = [a_col, b_col, c_col]
             
             plot_data = df.copy()
             if len(group_cols) > 1:
                 plot_data['ActiveGroup'] = plot_data.apply(lambda x: " : ".join([str(x[c]) for c in group_cols]), axis=1)
             else:
                 plot_data['ActiveGroup'] = plot_data[group_cols[0]]
             
             sns.violinplot(x='ActiveGroup', y=resp_col, data=plot_data, palette="viridis", inner="quartile")
             plt.ylabel(resp_col)

        plt.xlabel(effect)
        plt.title(f"{resp_col} by {effect}", pad=20)
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        
        # Save to buffer
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, transparent=True)
        buf.seek(0)
        img_str = base64.b64encode(buf.read()).decode('utf-8')
        plt.close()
        
        return {"status": "success", "image": img_str}
        
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


@app.post("/report_three_factor_crd")
@limiter.limit("5/minute")
async def report_three_factor_crd(
    request: Request,
    file: UploadFile = File(...),
    a_col: str = Form(..., max_length=100),
    b_col: str = Form(..., max_length=100),
    c_col: str = Form(..., max_length=100),
    resp_cols: list[str] = Form(...),
    rep_col: str = Form(None, max_length=100),
    control_col: str = Form(None, max_length=100),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10),
    notation: str = Form(None, max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        for col in resp_cols:
            if len(col) > 100:
                raise HTTPException(status_code=400, detail="Response column name too long.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        doc = Document()
        doc.add_heading('Three-Factor CRD Multi-Variable Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Factors: {a_col} x {b_col} x {c_col}")
        if rep_col: doc.add_paragraph(f"Replication: {rep_col}")
        
        for col in resp_cols:
            if col not in df.columns: continue
            
            analyzer, _, _ = perform_three_factor_crd_analysis(
                df, a_col, b_col, c_col, col, post_hoc, alpha, mean_order, rep_col=rep_col, control_col=control_col, notation=notation
            )
            analyzer.append_to_report(doc)
            
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return StreamingResponse(
             f,
             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
             headers={"Content-Disposition": "attachment; filename=ThreeFactor_CRD_Report.docx"}
        )
    except Exception as e:
         traceback.print_exc()
         return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# Helper for Pooled CRD
def perform_pooled_crd_analysis(df, treat_col, year_col, resp_col, post_hoc, alpha, mean_order):
    analyzer = PooledCRDAnalyzer(df, treat_col, year_col, resp_col)
    analyzer.validate()
    analyzer.run_bartlett_test()
    # Proceed even if heterogeneous, but warn
    analyzer.run_pooled_anova()
    analyzer.run_post_hoc(method=post_hoc, alpha=alpha, order=mean_order)
    return analyzer

@app.post("/analyze_pooled_crd")
@limiter.limit("10/minute")
async def analyze_pooled_crd(
    request: Request,
    file: UploadFile = File(...),
    treat_col: str = Form(..., max_length=100),
    year_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        data = await file.read()
        df = pd.read_csv(io.BytesIO(data))
        
        analyzer = perform_pooled_crd_analysis(
            df, treat_col, year_col, resp_col, post_hoc, alpha, mean_order
        )
        
        # Serialize Response
        # Bartlett
        b = analyzer.bartlett_res
        
        # ANOVA
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"],
                "SS": v["SS"],
                "MS": v["MS"],
                "F": v["F"],
                "P": v["P"],
                "sig": get_sig(v["P"])
            }
            
        # Post Hoc
        ph = analyzer.post_hoc_res
        ph_serial = None
        if ph:
             def serialize_factor(factor_dict):
                 means_list = []
                 for tr, val in factor_dict['means'].items():
                     means_list.append({
                         "level": str(tr),
                         "mean": float(val),
                         "sd": float(factor_dict['sds'].get(tr, 0)),
                         "se": float(factor_dict['ses'].get(tr, 0)),
                         "group": factor_dict['grouping'].get(tr, "")
                     })
                 return {
                     "means": means_list,
                     "sem_pooled": float(factor_dict['sem_pooled']),
                     "sed": float(factor_dict.get('sed', 0)),
                     "cd": float(factor_dict.get('cd', 0)),
                     "test_performed": factor_dict['test_performed'],
                     "reason": factor_dict['reason']
                 }

             ph_serial = {
                 "Treatment": serialize_factor(ph["Treatment"]),
                 "Year": serialize_factor(ph["Year"]),
                 "cv": float(ph['CV']),
                 "method": ph['method']
             }

        return {
            "status": "success",
            "bartlett": b,
            "anova": a,
            "post_hoc": ph_serial
        }
        
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_pooled_crd")
@limiter.limit("5/minute")
async def report_pooled_crd(
    request: Request,
    file: UploadFile = File(...),
    treat_col: str = Form(..., max_length=100),
    year_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    post_hoc: str = Form("lsd", max_length=50),
    alpha: float = Form(0.05),
    mean_order: str = Form("desc", max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        data = await file.read()
        df = pd.read_csv(io.BytesIO(data))
        analyzer = perform_pooled_crd_analysis(
            df, treat_col, year_col, resp_col, post_hoc, alpha, mean_order
        )
        buf = analyzer.create_report()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=PooledCRD_Report.docx"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# Two Factor Pooled CRD
from experimental_design.two_factor_pooled_crd import TwoFactorPooledCRDAnalyzer

# Multivariate
from multivariate_analysis.pca_analysis import PCAAnalyzer
from multivariate_analysis.path_analysis import PathAnalyzer

# Correlation
from correlation_regression.pearson_correlation import PearsonCorrelationAnalyzer
from correlation_regression.spearman_correlation import SpearmanCorrelationAnalyzer

@app.post("/analyze_two_factor_pooled_crd")
@limiter.limit("10/minute")
async def analyze_two_factor_pooled_crd(
    request: Request,
    file: UploadFile = File(...),
    treat_a_col: str = Form(..., max_length=100),
    treat_b_col: str = Form(..., max_length=100),
    year_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = TwoFactorPooledCRDAnalyzer(df, treat_a_col, treat_b_col, year_col, resp_col)
        analyzer.validate()
        analyzer.run_bartlett_test()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        # Serialize with handling for Pivot Tables / DataFrames
        ph = analyzer.post_hoc_res
        ph_serial = {}
        
        # Copy basic scalar fields
        for k, v in ph.items():
            if k not in ["ABY", "AB", "A", "B", "Year"]:
                ph_serial[k] = v
        
        # Helper to serialize pivots
        def ser_pivot(piv):
            cols = list(piv.columns)
            rows = []
            for idx, row in piv.iterrows():
                rows.append({
                    "label": str(idx),
                    "values": [float(x) for x in row.values]
                })
            return {"cols": [str(c) for c in cols], "rows": rows}
            
        if "ABY" in ph:
            dat = ph["ABY"]
            tables = []
            for t in dat["tables"]:
                tables.append({
                    "year": str(t["year"]),
                    "pivot": ser_pivot(t["pivot"])
                })
            ph_serial["ABY"] = {
                "tables": tables,
                "sem": float(dat["sem"]), "sed": float(dat["sed"]), "cd": float(dat["cd"]),
                "sig": dat["sig"]
            }
            
        if "AB" in ph:
            dat = ph["AB"]
            ph_serial["AB"] = {
                "pivot": ser_pivot(dat["pivot"]),
                "sem": float(dat["sem"]), "sed": float(dat["sed"]), "cd": float(dat["cd"]),
                "sig": dat["sig"],
                "grouping": dat["grouping"]
            }
            
        for eff in ["A", "B", "Year"]:
            if eff in ph:
                dat = ph[eff]
                ph_serial[eff] = {
                    "means": [
                        {
                            "level": str(k), 
                            "mean": float(v),
                            "std": float(dat["stds"][k]) if "stds" in dat else 0,
                            "se": float(dat["ses"][k]) if "ses" in dat else 0
                        } 
                        for k,v in dat["means"].items()
                    ],
                    "grouping": dat["grouping"],
                    "sem": float(dat["sem"]), "sed": float(dat["sed"]), "cd": float(dat["cd"]), "sig": True
                }

        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }

        return {
            "status": "success",
            "bartlett": analyzer.bartlett_res,
            "anova": a,
            "post_hoc": ph_serial
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_two_factor_pooled_crd")
@limiter.limit("5/minute")
async def report_two_factor_pooled_crd(
    request: Request,
    file: UploadFile = File(...),
    treat_a_col: str = Form(..., max_length=100),
    treat_b_col: str = Form(..., max_length=100),
    year_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = TwoFactorPooledCRDAnalyzer(df, treat_a_col, treat_b_col, year_col, resp_col)
        analyzer.validate()
        analyzer.run_bartlett_test()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=TwoFactorPooledCRD_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}

def get_sig(p_val):
    if p_val is None or np.isnan(p_val): return ""
    if p_val <= 0.01: return "**"
    if p_val <= 0.05: return "*"
    return ""

def safe_float(val):
    if val is None or np.isnan(val) or np.isinf(val):
        return 0.0
    return float(val)


from experimental_design.rcbd_analysis import RCBDAnalyzer
from experimental_design.two_factor_rcbd import TwoFactorRCBDAnalyzer
from experimental_design.three_factor_rcbd import ThreeFactorRCBDAnalyzer
from experimental_design.split_plot_analysis import SplitPlotAnalyzer
from experimental_design.split_plot_21 import SplitPlot21Analyzer
from experimental_design.split_plot_12 import SplitPlot12Analyzer
from experimental_design.split_split_plot import SplitSplitPlotAnalyzer
from experimental_design.split_crd_analysis import SplitCRDAnalyzer
from experimental_design.split_plot_pooled import SplitPlotPooledAnalyzer
from experimental_design.pooled_rcbd_analysis import PooledRCBDAnalyzer
from experimental_design.pooled_two_factor_rcbd import PooledTwoFactorRCBDAnalyzer

@app.post("/analyze_one_factor_rcbd")
@limiter.limit("10/minute")
async def analyze_one_factor_rcbd(
    request: Request,
    file: UploadFile = File(...),
    treat_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50),
    mean_order: str = Form('desc', max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = RCBDAnalyzer(df, treat_col, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha, mean_order)
        
        # Serialize
        ph = analyzer.post_hoc_res
        ph_serial = {}
        
        # Basic scalar fields
        for k, v in ph.items():
            if k != "Treatment":
                ph_serial[k] = v
                
        # Treatment Data
        if "Treatment" in ph:
            dat = ph["Treatment"]
            ph_serial["Treatment"] = {
                "means": [
                    {
                        "level": str(k), 
                        "mean": float(v),
                        "std": float(dat["stds"][k]),
                        "se": float(dat["ses"][k])
                    } for k, v in dat["means"].items()
                ],
                "grouping": dat["grouping"],
                "sig": dat["sig"]
            }

        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }

        return {
            "status": "success",
            "anova": a,
            "post_hoc": ph_serial
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_one_factor_rcbd")
@limiter.limit("5/minute")
async def report_one_factor_rcbd(
    request: Request,
    file: UploadFile = File(...),
    treat_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50),
    mean_order: str = Form('desc', max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = RCBDAnalyzer(df, treat_col, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha, mean_order)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=RCBD_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}



@app.post("/analyze_two_factor_rcbd")
@limiter.limit("10/minute")
async def analyze_two_factor_rcbd(
    request: Request,
    file: UploadFile = File(...),
    fact_a: str = Form(..., max_length=100),
    fact_b: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50),
    mean_order: str = Form('desc', max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = TwoFactorRCBDAnalyzer(df, fact_a, fact_b, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha, mean_order)
        
        # Serialize Post Hoc
        # Helper to serialize a dataset (Factor A/B/Interaction)
        def ser_ds(ds):
            if "means" not in ds: return ds # e.g. info notes
            return {
                "sig": ds["sig"],
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else f"{k[0]} x {k[1]}",
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else f"{k[0]} x {k[1]}"): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
            
        ph = analyzer.post_hoc_res
        ph_serial = {
            "stats": analyzer.stats,
            "Interaction": ser_ds(ph["Interaction"]),
            "Factor A": ser_ds(ph["Factor A"]),
            "Factor B": ser_ds(ph["Factor B"])
        }
        
        # Serialize ANOVA
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "anova": a, "post_hoc": ph_serial }
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_two_factor_rcbd")
@limiter.limit("5/minute")
async def report_two_factor_rcbd(
    request: Request,
    file: UploadFile = File(...),
    fact_a: str = Form(..., max_length=100),
    fact_b: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50),
    mean_order: str = Form('desc', max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = TwoFactorRCBDAnalyzer(df, fact_a, fact_b, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha, mean_order)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=TwoFactor_RCBD_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}



@app.post("/analyze_three_factor_rcbd")
@limiter.limit("10/minute")
async def analyze_three_factor_rcbd(
    request: Request,
    file: UploadFile = File(...),
    fact_a: str = Form(..., max_length=100),
    fact_b: str = Form(..., max_length=100),
    fact_c: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50),
    mean_order: str = Form('desc', max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = ThreeFactorRCBDAnalyzer(df, fact_a, fact_b, fact_c, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha, mean_order)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "sig": ds["sig"],
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_three_factor_rcbd")
@limiter.limit("5/minute")
async def report_three_factor_rcbd(
    request: Request,
    file: UploadFile = File(...),
    fact_a: str = Form(..., max_length=100),
    fact_b: str = Form(..., max_length=100),
    fact_c: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50),
    mean_order: str = Form('desc', max_length=10)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = ThreeFactorRCBDAnalyzer(df, fact_a, fact_b, fact_c, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha, mean_order)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=ThreeFactor_RCBD_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.post("/analyze_split_plot")
@limiter.limit("10/minute")
async def analyze_split_plot(
    request: Request,
    file: UploadFile = File(...),
    main_col: str = Form(..., max_length=100),
    sub_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = SplitPlotAnalyzer(df, main_col, sub_col, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_split_plot")
@limiter.limit("5/minute")
async def report_split_plot(
    request: Request,
    file: UploadFile = File(...),
    main_col: str = Form(..., max_length=100),
    sub_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = SplitPlotAnalyzer(df, main_col, sub_col, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Split_Plot_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.post("/analyze_split_plot_21")
@limiter.limit("10/minute")
async def analyze_split_plot_21(
    request: Request,
    file: UploadFile = File(...),
    main_a: str = Form(..., max_length=100),
    main_b: str = Form(..., max_length=100),
    sub_c: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = SplitPlot21Analyzer(df, main_a, main_b, sub_c, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_split_plot_21")
@limiter.limit("5/minute")
async def report_split_plot_21(
    request: Request,
    file: UploadFile = File(...),
    main_a: str = Form(..., max_length=100),
    main_b: str = Form(..., max_length=100),
    sub_c: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = SplitPlot21Analyzer(df, main_a, main_b, sub_c, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Split_Plot_21_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.post("/analyze_split_plot_12")
@limiter.limit("10/minute")
async def analyze_split_plot_12(
    request: Request,
    file: UploadFile = File(...),
    main_a: str = Form(..., max_length=100),
    sub_b: str = Form(..., max_length=100),
    sub_c: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = SplitPlot12Analyzer(df, main_a, sub_b, sub_c, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_split_plot_12")
@limiter.limit("5/minute")
async def report_split_plot_12(
    request: Request,
    file: UploadFile = File(...),
    main_a: str = Form(..., max_length=100),
    sub_b: str = Form(..., max_length=100),
    sub_c: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = SplitPlot12Analyzer(df, main_a, sub_b, sub_c, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Split_Plot_12_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.post("/analyze_split_split_plot")
@limiter.limit("10/minute")
async def analyze_split_split_plot(
    request: Request,
    file: UploadFile = File(...),
    main_a: str = Form(..., max_length=100),
    sub_b: str = Form(..., max_length=100),
    sub_c: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = SplitSplitPlotAnalyzer(df, main_a, sub_b, sub_c, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_split_split_plot")
@limiter.limit("5/minute")
async def report_split_split_plot(
    request: Request,
    file: UploadFile = File(...),
    main_a: str = Form(..., max_length=100),
    sub_b: str = Form(..., max_length=100),
    sub_c: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = SplitSplitPlotAnalyzer(df, main_a, sub_b, sub_c, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Split_Split_Plot_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}
        

@app.post("/analyze_split_crd")
@limiter.limit("10/minute")
async def analyze_split_crd(
    request: Request,
    file: UploadFile = File(...),
    main_a: str = Form(..., max_length=100),
    sub_b: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = SplitCRDAnalyzer(df, main_a, sub_b, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_split_crd")
@limiter.limit("5/minute")
async def report_split_crd(
    request: Request,
    file: UploadFile = File(...),
    main_a: str = Form(..., max_length=100),
    sub_b: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = SplitCRDAnalyzer(df, main_a, sub_b, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Split_CRD_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/analyze_split_pooled")
@limiter.limit("10/minute")
async def analyze_split_pooled(
    request: Request,
    file: UploadFile = File(...),
    year_col: str = Form(..., max_length=100),
    main_a: str = Form(..., max_length=100),
    sub_b: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = SplitPlotPooledAnalyzer(df, year_col, main_a, sub_b, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_bartlett()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "bartlett": analyzer.bartlett_res, "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_split_pooled")
@limiter.limit("5/minute")
async def report_split_pooled(
    request: Request,
    file: UploadFile = File(...),
    year_col: str = Form(..., max_length=100),
    main_a: str = Form(..., max_length=100),
    sub_b: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = SplitPlotPooledAnalyzer(df, year_col, main_a, sub_b, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_bartlett()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Pooled_Split_Plot_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/analyze_pooled_rcbd")
@limiter.limit("10/minute")
async def analyze_pooled_rcbd(
    request: Request,
    file: UploadFile = File(...),
    year_col: str = Form(..., max_length=100),
    treat_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = PooledRCBDAnalyzer(df, year_col, treat_col, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_bartlett()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "bartlett": analyzer.bartlett_res, "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_pooled_rcbd")
@limiter.limit("5/minute")
async def report_pooled_rcbd(
    request: Request,
    file: UploadFile = File(...),
    year_col: str = Form(..., max_length=100),
    treat_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    resp_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = PooledRCBDAnalyzer(df, year_col, treat_col, rep_col, resp_col)
        analyzer.validate()
        analyzer.run_bartlett()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Pooled_RCBD_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/analyze_pooled_two_factor_rcbd")
@limiter.limit("10/minute")
async def analyze_pooled_two_factor_rcbd(
    request: Request,
    file: UploadFile = File(...),
    col_year: str = Form(..., max_length=100),
    col_a: str = Form(..., max_length=100),
    col_b: str = Form(..., max_length=100),
    col_rep: str = Form(..., max_length=100),
    col_resp: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = PooledTwoFactorRCBDAnalyzer(df, col_year, col_a, col_b, col_rep, col_resp)
        analyzer.validate()
        analyzer.run_homogeneity_test()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        def ser_ds(ds):
            if "means" not in ds: return ds
            return {
                "means": [
                    {
                        "level": str(k) if not isinstance(k, tuple) else " x ".join(map(str, k)),
                        "mean": float(v),
                        "std": float(ds["stds"][k]),
                        "se": float(ds["ses"][k])
                    } for k, v in ds["means"].items()
                ],
                "grouping": {
                    (str(k) if not isinstance(k, tuple) else " x ".join(map(str, k))): v 
                    for k, v in ds.get("grouping", {}).items()
                }
            }
        
        ph = analyzer.post_hoc_res
        ph_serial = { "stats": analyzer.stats }
        for k in ph.keys(): ph_serial[k] = ser_ds(ph[k])
        
        a = {}
        for k, v in analyzer.anova_table.items():
            a[k] = {
                "df": v["df"], "SS": v["SS"], "MS": v["MS"], "F": v["F"], "P": v["P"],
                "sig": get_sig(v["P"]) if v["P"] is not None else ""
            }
            
        return { "status": "success", "bartlett": analyzer.bartlett_res, "anova": a, "post_hoc": ph_serial }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/report_pooled_two_factor_rcbd")
@limiter.limit("5/minute")
async def report_pooled_two_factor_rcbd(
    request: Request,
    file: UploadFile = File(...),
    col_year: str = Form(..., max_length=100),
    col_a: str = Form(..., max_length=100),
    col_b: str = Form(..., max_length=100),
    col_rep: str = Form(..., max_length=100),
    col_resp: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    post_hoc: str = Form('lsd', max_length=50)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = PooledTwoFactorRCBDAnalyzer(df, col_year, col_a, col_b, col_rep, col_resp)
        analyzer.validate()
        analyzer.run_homogeneity_test()
        analyzer.run_anova()
        analyzer.run_post_hoc(post_hoc, alpha)
        
        docx = analyzer.create_report()
        return StreamingResponse(
            docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Pooled_Two_Factor_Report.docx"}
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.get("/")
@limiter.limit("20/minute")
async def read_root(request: Request):
    return FileResponse(os.path.join(FRONTEND_DIR, "data-analyzer.html"))


# ==========================================
# MULTIVARIATE ANALYSIS (PCA)
# ==========================================

# Helper
def perform_pca(df, obs_col, var_cols):
    var_list = var_cols.split(',') if isinstance(var_cols, str) else var_cols
    analyzer = PCAAnalyzer(df, obs_col, var_list)
    analyzer.validate()
    analyzer.run_pca()
    return analyzer

@app.post("/analyze_pca")
@limiter.limit("10/minute")
async def analyze_pca(
    request: Request,
    file: UploadFile = File(...),
    obs_col: str = Form(..., max_length=100),
    var_cols: str = Form(..., max_length=500) # Comma separated list, allow more
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_pca(df, obs_col, var_cols)
        analyzer.generate_plots() # Generate plot buffers
        
        # Return basic stats
        res = analyzer.pca_res
        
        return {
            "status": "success",
            "eigenvalues": res['eigenvalues'].tolist(),
            "variance_pct": res['variance_pct'].tolist(),
            "cum_variance_pct": res['cum_variance_pct'].tolist()
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_pca_plots")
@limiter.limit("5/minute")
async def report_pca_plots(
    request: Request,
    file: UploadFile = File(...),
    obs_col: str = Form(..., max_length=100),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_pca(df, obs_col, var_cols)
        analyzer.generate_plots()
        
        buf = analyzer.create_report_plots()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=PCA_Plots.docx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_pca_text")
@limiter.limit("5/minute")
async def report_pca_text(
    request: Request,
    file: UploadFile = File(...),
    obs_col: str = Form(..., max_length=100),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_pca(df, obs_col, var_cols)
        
        buf = analyzer.create_report_interpretation()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=PCA_Interpretation.docx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_pca_excel")
@limiter.limit("5/minute")
async def report_pca_excel(
    request: Request,
    file: UploadFile = File(...),
    obs_col: str = Form(..., max_length=100),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_pca(df, obs_col, var_cols)
        
        buf = analyzer.create_output_excel()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=PCA_Output.xlsx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# ==========================================
# MULTIVARIATE ANALYSIS (PATH ANALYSIS)
# ==========================================

def perform_path_analysis(df, dep_var, indep_vars):
    indep_list = indep_vars.split(',') if isinstance(indep_vars, str) else indep_vars
    # Clean list
    indep_list = [x.strip() for x in indep_list if x.strip()]
    analyzer = PathAnalyzer(df, dep_var, indep_list)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_path")
@limiter.limit("10/minute")
async def analyze_path(
    request: Request,
    file: UploadFile = File(...),
    dep_var: str = Form(..., max_length=100),
    indep_vars: str = Form(..., max_length=500) # Comma separated
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_path_analysis(df, dep_var, indep_vars)
        analyzer.generate_diagram() # Gen diagram to ensure no errors
        
        res = analyzer.results
        
        return {
            "status": "success",
            "R2": res['R2'],
            "residual": res['residual'],
            "direct_effects": res['direct_effects']
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_path_doc")
@limiter.limit("5/minute")
async def report_path_doc(
    request: Request,
    file: UploadFile = File(...),
    dep_var: str = Form(..., max_length=100),
    indep_vars: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_path_analysis(df, dep_var, indep_vars)
        analyzer.generate_diagram()
        
        buf = analyzer.create_report_doc()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Path_Analysis_Report.docx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_path_excel")
@limiter.limit("5/minute")
async def report_path_excel(
    request: Request,
    file: UploadFile = File(...),
    dep_var: str = Form(..., max_length=100),
    indep_vars: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_path_analysis(df, dep_var, indep_vars)
        
        buf = analyzer.create_output_excel()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=Path_Analysis_Output.xlsx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# ==========================================
# CORRELATION ANALYSIS (PEARSON)
# ==========================================

def perform_pearson(df, var_cols):
    var_list = var_cols.split(',') if isinstance(var_cols, str) else var_cols
    var_list = [x.strip() for x in var_list if x.strip()]
    
    analyzer = PearsonCorrelationAnalyzer(df, var_list)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_pearson")
@limiter.limit("10/minute")
async def analyze_pearson(
    request: Request,
    file: UploadFile = File(...),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_pearson(df, var_cols)
        analyzer.generate_heatmap()
        
        # Serialize matrices for frontend table (just r and sig)
        res = analyzer.results
        r_dict = res['corr_matrix'].to_dict()
        sig_dict = res['sig_matrix'].to_dict()
        
        return {
            "status": "success",
            "vars": analyzer.vars,
            "corr_matrix": r_dict,
            "sig_matrix": sig_dict
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_pearson_doc")
@limiter.limit("5/minute")
async def report_pearson_doc(
    request: Request,
    file: UploadFile = File(...),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_pearson(df, var_cols)
        analyzer.generate_heatmap()
        
        buf = analyzer.create_report_doc()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Pearson_Correlation_Report.docx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_pearson_excel")
@limiter.limit("5/minute")
async def report_pearson_excel(
    request: Request,
    file: UploadFile = File(...),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_pearson(df, var_cols)
        analyzer.run_analysis() # Run explicitly if skipped
        
        buf = analyzer.create_output_excel()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=Pearson_Correlation_Output.xlsx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# ==========================================
# SPEARMAN RANK CORRELATION
# ==========================================

def perform_spearman(df, var_cols):
    var_list = var_cols.split(',') if isinstance(var_cols, str) else var_cols
    var_list = [x.strip() for x in var_list if x.strip()]
    
    analyzer = SpearmanCorrelationAnalyzer(df, var_list)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_spearman")
@limiter.limit("10/minute")
async def analyze_spearman(
    request: Request,
    file: UploadFile = File(...),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_spearman(df, var_cols)
        analyzer.generate_heatmap()
        
        # Serialize matrices for frontend table (just r and sig)
        res = analyzer.results
        r_dict = res['corr_matrix'].to_dict()
        sig_dict = res['sig_matrix'].to_dict()
        
        return {
            "status": "success",
            "vars": analyzer.vars,
            "corr_matrix": r_dict,
            "sig_matrix": sig_dict
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_spearman_doc")
@limiter.limit("5/minute")
async def report_spearman_doc(
    request: Request,
    file: UploadFile = File(...),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_spearman(df, var_cols)
        analyzer.generate_heatmap()
        
        buf = analyzer.create_report_doc()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Spearman_Correlation_Report.docx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_spearman_excel")
@limiter.limit("5/minute")
async def report_spearman_excel(
    request: Request,
    file: UploadFile = File(...),
    var_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_spearman(df, var_cols)
        analyzer.run_analysis()
        
        buf = analyzer.create_output_excel()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=Spearman_Correlation_Output.xlsx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Griffing Method 1
def perform_griffing_method1(df, female_col, male_col, rep_col, trait_cols):
    # trait_cols might be a comma separated string if coming from Form
    if isinstance(trait_cols, str):
        trait_cols = [t.strip() for t in trait_cols.split(",") if t.strip()]
    analyzer = GriffingMethod1Analyzer(df, female_col, male_col, rep_col, trait_cols)
    analyzer.validate()
    analyzer.run_all()
    return analyzer

@app.post("/analyze_griffing1")
@limiter.limit("10/minute")
async def analyze_griffing1(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500) # Comma separated
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_griffing_method1(df, female_col, male_col, rep_col, trait_cols)
        
        # Serialize results
        serialized = {}
        for trait in analyzer.trait_cols:
            res = analyzer.results[trait]
            # Convert matrices to dictionaries or lists for JSON
            s_mat = res['sca_effects']
            r_mat = res['rca_effects']
            
            # GCA is already a list of dicts
            
            serialized[trait] = {
                "anova_geno": {k: {**v, "sig": get_sig(v['P'])} for k, v in res['anova_geno'].items()},
                "anova_comb": {k: {**v, "sig": get_sig(v['P'])} for k, v in res['anova_comb'].items()},
                "gca_effects": res['gca_effects'],
                "sca_matrix": s_mat.tolist(),
                "rca_matrix": r_mat.tolist(),
                "se_sca": float(res['se_sca']),
                "se_rca": float(res['se_rca']),
                "variances": res['variances'],
                "heterosis": res['heterosis'],
                "parents": analyzer.parents
            }
            
        return {"status": "success", "results": serialized}
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_griffing1_doc")
@limiter.limit("5/minute")
async def report_griffing1_doc(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method1(df, female_col, male_col, rep_col, trait_cols)
        buf = analyzer.create_report()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Griffing_Method1_Report.docx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_griffing1_excel")
@limiter.limit("5/minute")
async def report_griffing1_excel(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method1(df, female_col, male_col, rep_col, trait_cols)
        buf = analyzer.create_excel()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=Griffing_Method1_Output.xlsx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Heatmap endpoint
@app.get("/griffing1_heatmap/{trait}/{type}")
async def griffing1_heatmap(trait: str, type: str):
    # This would require caching the analyzer result or re-running.
    # For now, we can skip or implement if needed. 
    # Usually heatmaps are better generated on frontend with D3/Chart.js if possible,
    # but the request asked for heatmaps. I can generate them on the fly if I have the data.
    # But since it's a POST with file upload, we usually return it in the main analyze response as base64 or separate.
    return {"status": "not_implemented_separately"}


# Griffing Method 1 WITH CHECK
def perform_griffing_method1_check(df, female_col, male_col, rep_col, check_col, trait_cols):
    if isinstance(trait_cols, str):
        trait_cols = [t.strip() for t in trait_cols.split(",") if t.strip()]
    analyzer = GriffingMethod1CheckAnalyzer(df, female_col, male_col, rep_col, check_col, trait_cols)
    analyzer.validate()
    for trait in analyzer.trait_cols:
        analyzer.analyze_trait(trait)
    return analyzer

@app.post("/analyze_griffing1_check")
@limiter.limit("10/minute")
async def analyze_griffing1_check(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    check_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method1_check(df, female_col, male_col, rep_col, check_col, trait_cols)
        
        serialized = {}
        for trait in analyzer.trait_cols:
            res = analyzer.results[trait]
            serialized[trait] = {
                "anova_geno": {k: {**v, "sig": get_sig(v['P'])} for k, v in res['anova_geno'].items()},
                "anova_comb": {k: {**v, "sig": get_sig(v['P'])} for k, v in res['anova_comb'].items()},
                "gca_effects": res['gca_effects'],
                "sca_matrix": res['sca_matrix'],
                "rca_matrix": res['rca_matrix'],
                "se_sca": float(res['se_sca']),
                "se_rca": float(res['se_rca']),
                "variances": res['variances'],
                "std_heterosis": res['std_heterosis'],
                "check_means": res['check_means'],
                "parents": res['parents']
            }
        return {"status": "success", "results": serialized}
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_griffing1_check_doc")
@limiter.limit("5/minute")
async def report_griffing1_check_doc(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    check_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method1_check(df, female_col, male_col, rep_col, check_col, trait_cols)
        buf = analyzer.create_report()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Griffing_Method1_WithCheck_Report.docx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_griffing1_check_excel")
@limiter.limit("5/minute")
async def report_griffing1_check_excel(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    check_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method1_check(df, female_col, male_col, rep_col, check_col, trait_cols)
        buf = analyzer.create_excel()
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=Griffing_Method1_WithCheck_Output.xlsx"}
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# Griffing Method 2
def perform_griffing_method2(df, female_col, male_col, rep_col, trait_cols):
    if isinstance(trait_cols, str):
        trait_cols = [t.strip() for t in trait_cols.split(",") if t.strip()]
    analyzer = GriffingMethod2Analyzer(df, female_col, male_col, rep_col, trait_cols)
    analyzer.validate()
    for trait in analyzer.trait_cols:
        analyzer.analyze_trait(trait)
    return analyzer

@app.post("/analyze_griffing2")
@limiter.limit("10/minute")
async def analyze_griffing2(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method2(df, female_col, male_col, rep_col, trait_cols)
        serialized = {}
        for trait in analyzer.trait_cols:
            res = analyzer.results[trait]
            serialized[trait] = {
                "anova_geno": {k: {**v, "sig": get_sig(v['P'])} for k, v in res['anova_geno'].items()},
                "anova_comb": {k: {**v, "sig": get_sig(v['P'])} for k, v in res['anova_comb'].items()},
                "gca_effects": res['gca_effects'],
                "sca_matrix": res['sca_matrix'],
                "variances": res['variances'],
                "heterosis": res['heterosis'],
                "parents": res['parents']
            }
        return {"status": "success", "results": serialized}
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_griffing2_doc")
@limiter.limit("5/minute")
async def report_griffing2_doc(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method2(df, female_col, male_col, rep_col, trait_cols)
        buf = analyzer.create_report()
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=Griffing_Method2_Report.docx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_griffing2_excel")
@limiter.limit("5/minute")
async def report_griffing2_excel(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method2(df, female_col, male_col, rep_col, trait_cols)
        buf = analyzer.create_excel()
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=Griffing_Method2_Output.xlsx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# Griffing Method 2 WITH CHECK
def perform_griffing_method2_check(df, female_col, male_col, rep_col, check_col, trait_cols):
    if isinstance(trait_cols, str):
        trait_cols = [t.strip() for t in trait_cols.split(",") if t.strip()]
    analyzer = GriffingMethod2CheckAnalyzer(df, female_col, male_col, rep_col, check_col, trait_cols)
    analyzer.validate()
    for trait in analyzer.trait_cols:
        analyzer.analyze_trait(trait)
    return analyzer

@app.post("/analyze_griffing2_check")
@limiter.limit("10/minute")
async def analyze_griffing2_check(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    check_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method2_check(df, female_col, male_col, rep_col, check_col, trait_cols)
        serialized = {}
        for trait in analyzer.trait_cols:
            res = analyzer.results[trait]
            serialized[trait] = {
                "anova_comb": {k: {**v, "sig": get_sig(v.get('P'))} for k, v in res['anova_comb'].items()},
                "gca_effects": res['gca_effects'],
                "sca_matrix": res['sca_matrix'],
                "variances": res['variances'],
                "heterosis": res['heterosis'],
                "parents": res['parents']
            }
        return {"status": "success", "results": serialized}
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_griffing2_check_doc")
@limiter.limit("5/minute")
async def report_griffing2_check_doc(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    check_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method2_check(df, female_col, male_col, rep_col, check_col, trait_cols)
        buf = analyzer.create_report()
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=Griffing_Method2_WithCheck_Report.docx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_griffing2_check_excel")
@limiter.limit("5/minute")
async def report_griffing2_check_excel(
    request: Request,
    file: UploadFile = File(...),
    female_col: str = Form(..., max_length=100),
    male_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    check_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_griffing_method2_check(df, female_col, male_col, rep_col, check_col, trait_cols)
        buf = analyzer.create_excel()
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=Griffing_Method2_WithCheck_Output.xlsx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# Genotypic Correlation
def perform_genotypic_correlation(df, genotype_col, rep_col, trait_cols):
    if isinstance(trait_cols, str):
        trait_cols = [t.strip() for t in trait_cols.split(",") if t.strip()]
    analyzer = GenotypicCorrelationAnalyzer(df, genotype_col, rep_col, trait_cols)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_genotypic_correlation")
@limiter.limit("10/minute")
async def analyze_genotypic_correlation(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_genotypic_correlation(df, genotype_col, rep_col, trait_cols)
        
        return {
            "status": "success",
            "traits": analyzer.trait_cols,
            "variances": analyzer.variances,
            "corr_matrix": analyzer.correlation_matrix.to_dict(),
            "sig_matrix": analyzer.p_values.to_dict(),
            "interpretations": analyzer.get_interpretation()
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_genotypic_correlation_doc")
@limiter.limit("5/minute")
async def report_genotypic_correlation_doc(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_genotypic_correlation(df, genotype_col, rep_col, trait_cols)
        buf = analyzer.create_report()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={"Content-Disposition": "attachment; filename=Genotypic_Correlation_Report.docx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_genotypic_correlation_excel")
@limiter.limit("5/minute")
async def report_genotypic_correlation_excel(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_genotypic_correlation(df, genotype_col, rep_col, trait_cols)
        buf = analyzer.create_excel()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            headers={"Content-Disposition": "attachment; filename=Genotypic_Correlation_Output.xlsx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# Phenotypic Correlation
def perform_phenotypic_correlation(df, genotype_col, rep_col, trait_cols):
    if isinstance(trait_cols, str):
        trait_cols = [t.strip() for t in trait_cols.split(",") if t.strip()]
    analyzer = PhenotypicCorrelationAnalyzer(df, genotype_col, rep_col, trait_cols)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_phenotypic_correlation")
@limiter.limit("10/minute")
async def analyze_phenotypic_correlation(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_phenotypic_correlation(df, genotype_col, rep_col, trait_cols)
        
        return {
            "status": "success",
            "traits": analyzer.trait_cols,
            "variances": {k: v for k, v in analyzer.variances.items()},
            "corr_matrix": analyzer.correlation_matrix.to_dict(),
            "sig_matrix": analyzer.p_values.to_dict(),
            "interpretations": analyzer.get_interpretation()
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_phenotypic_correlation_doc")
@limiter.limit("5/minute")
async def report_phenotypic_correlation_doc(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_phenotypic_correlation(df, genotype_col, rep_col, trait_cols)
        buf = analyzer.create_report()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={"Content-Disposition": "attachment; filename=Phenotypic_Correlation_Report.docx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_phenotypic_correlation_excel")
@limiter.limit("5/minute")
async def report_phenotypic_correlation_excel(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_phenotypic_correlation(df, genotype_col, rep_col, trait_cols)
        buf = analyzer.create_excel()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            headers={"Content-Disposition": "attachment; filename=Phenotypic_Correlation_Output.xlsx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Mount Frontend
# --- Genotypic Path Analysis ---
def perform_genotypic_path(df, genotype_col, rep_col, dependent_var, independent_vars):
    analyzer = GenotypicPathAnalyzer(df, genotype_col, rep_col, dependent_var, independent_vars)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_genotypic_path")
@limiter.limit("10/minute")
async def analyze_genotypic_path(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    dependent_var: str = Form(..., max_length=100),
    independent_vars: str = Form(..., max_length=500) # Comma separated
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        indep_list = [x.strip() for x in independent_vars.split(',')]
        
        analyzer = perform_genotypic_path(df, genotype_col, rep_col, dependent_var, indep_list)
        path_table = analyzer.get_path_table()
        
        # Generate Diagram in base64
        import base64
        diag_buf = analyzer.generate_path_diagram()
        diag_base64 = base64.b64encode(diag_buf.read()).decode('utf-8')
        
        return {
            "status": "success",
            "traits": analyzer.independent_vars,
            "dependent": analyzer.dependent_var,
            "path_table": path_table.to_dict(orient='records'),
            "residual": float(analyzer.residual_effect),
            "explained": float(analyzer.explained_variation),
            "unexplained": float(analyzer.unexplained_variation),
            "corr_matrix": analyzer.correlation_matrix.to_dict(),
            "diagram": diag_base64
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_genotypic_path_doc")
@limiter.limit("5/minute")
async def report_genotypic_path_doc(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    dependent_var: str = Form(..., max_length=100),
    independent_vars: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        indep_list = [x.strip() for x in independent_vars.split(',')]
        analyzer = perform_genotypic_path(df, genotype_col, rep_col, dependent_var, indep_list)
        buf = analyzer.create_report()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Genotypic_Path_Report.docx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_genotypic_path_excel")
@limiter.limit("5/minute")
async def report_genotypic_path_excel(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    dependent_var: str = Form(..., max_length=100),
    independent_vars: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        indep_list = [x.strip() for x in independent_vars.split(',')]
        analyzer = perform_genotypic_path(df, genotype_col, rep_col, dependent_var, indep_list)
        buf = analyzer.create_excel()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            headers={"Content-Disposition": "attachment; filename=Genotypic_Path_Output.xlsx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# --- Phenotypic Path Analysis ---
def perform_phenotypic_path(df, genotype_col, rep_col, dependent_var, independent_vars):
    analyzer = PhenotypicPathAnalyzer(df, genotype_col, rep_col, dependent_var, independent_vars)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_phenotypic_path")
@limiter.limit("10/minute")
async def analyze_phenotypic_path(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    dependent_var: str = Form(..., max_length=100),
    independent_vars: str = Form(..., max_length=500) # Comma separated
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        indep_list = [x.strip() for x in independent_vars.split(',')]
        
        analyzer = perform_phenotypic_path(df, genotype_col, rep_col, dependent_var, indep_list)
        path_table = analyzer.get_path_table()
        
        import base64
        diag_buf = analyzer.generate_path_diagram()
        diag_base64 = base64.b64encode(diag_buf.read()).decode('utf-8')
        
        return {
            "status": "success",
            "traits": analyzer.independent_vars,
            "dependent": analyzer.dependent_var,
            "path_table": path_table.to_dict(orient='records'),
            "residual": float(analyzer.residual_effect),
            "explained": float(analyzer.explained_variation),
            "unexplained": float(analyzer.unexplained_variation),
            "corr_matrix": analyzer.correlation_matrix.to_dict(),
            "diagram": diag_base64
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_phenotypic_path_doc")
@limiter.limit("5/minute")
async def report_phenotypic_path_doc(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    dependent_var: str = Form(..., max_length=100),
    independent_vars: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        indep_list = [x.strip() for x in independent_vars.split(',')]
        analyzer = perform_phenotypic_path(df, genotype_col, rep_col, dependent_var, indep_list)
        buf = analyzer.create_report()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Phenotypic_Path_Report.docx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_phenotypic_path_excel")
@limiter.limit("5/minute")
async def report_phenotypic_path_excel(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    dependent_var: str = Form(..., max_length=100),
    independent_vars: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        indep_list = [x.strip() for x in independent_vars.split(',')]
        analyzer = perform_phenotypic_path(df, genotype_col, rep_col, dependent_var, indep_list)
        buf = analyzer.create_excel()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            headers={"Content-Disposition": "attachment; filename=Phenotypic_Path_Output.xlsx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# --- Mahalanobis D2 Statistics ---
def perform_mahalanobis_d2(df, genotype_col, rep_col, trait_cols):
    analyzer = MahalanobisD2Analyzer(df, genotype_col, rep_col, trait_cols)
    # Note: rep_col is used for averaging data in the analyzer if needed
    analyzer.validate()
    # We pass rep_col check here if we want to ensure averaging correctly
    # For now, the analyzer groups by genotype and averages.
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_mahalanobis_d2")
@limiter.limit("10/minute")
async def analyze_mahalanobis_d2(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500) # Comma separated
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        traits_list = [x.strip() for x in trait_cols.split(',')]
        
        analyzer = perform_mahalanobis_d2(df, genotype_col, rep_col, traits_list)
        
        import base64
        dendro_buf = analyzer.generate_dendrogram()
        dendro_b64 = base64.b64encode(dendro_buf.read()).decode('utf-8')
        
        plot_buf = analyzer.generate_cluster_plot()
        plot_b64 = base64.b64encode(plot_buf.read()).decode('utf-8')
        
        # Data Cleaning for JSON serialization
        def clean_data(obj):
            if isinstance(obj, dict):
                return {k: clean_data(v) for k, v in obj.items()}
            elif isinstance(obj, list):
                return [clean_data(x) for x in obj]
            elif hasattr(obj, "item"): # Handle NumPy scalars
                return obj.item()
            elif isinstance(obj, float):
                if np.isnan(obj) or np.isinf(obj): return 0.0
                return obj
            return obj

        response_data = {
            "status": "success",
            "clusters": analyzer.clusters,
            "intra_distances": analyzer.intra_distances,
            "inter_distances": analyzer.inter_distances.to_dict(),
            "cluster_means": analyzer.cluster_means.to_dict(orient='index'),
            "trait_contributions": analyzer.trait_contributions.to_dict(orient='records'),
            "d2_matrix": analyzer.d2_matrix.to_dict(),
            "dendrogram": dendro_b64,
            "cluster_plot": plot_b64
        }
        
        return clean_data(response_data)
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_mahalanobis_d2_doc")
@limiter.limit("5/minute")
async def report_mahalanobis_d2_doc(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        traits_list = [x.strip() for x in trait_cols.split(',')]
        analyzer = perform_mahalanobis_d2(df, genotype_col, rep_col, traits_list)
        buf = analyzer.create_report()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Mahalanobis_D2_Report.docx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_mahalanobis_d2_excel")
@limiter.limit("5/minute")
async def report_mahalanobis_d2_excel(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_cols: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        traits_list = [x.strip() for x in trait_cols.split(',')]
        analyzer = perform_mahalanobis_d2(df, genotype_col, rep_col, traits_list)
        buf = analyzer.create_excel()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            headers={"Content-Disposition": "attachment; filename=Mahalanobis_D2_Output.xlsx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# --- Genetic Parameter Estimation ---
def perform_genetic_parameters(df, genotype_col, rep_col, traits):
    analyzer = GeneticParameterAnalyzer(df, genotype_col, rep_col, traits)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_genetic_parameters")
@limiter.limit("10/minute")
async def analyze_genetic_parameters(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    traits: str = Form(..., max_length=500) # Comma separated
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        traits_list = [x.strip() for x in traits.split(',')]
        
        analyzer = perform_genetic_parameters(df, genotype_col, rep_col, traits_list)
        
        # Data Cleaning for JSON serialization
        def clean_data(obj):
            if isinstance(obj, dict):
                return {k: clean_data(v) for k, v in obj.items()}
            elif isinstance(obj, list):
                return [clean_data(x) for x in obj]
            elif hasattr(obj, "item"): # Handle NumPy scalars
                return obj.item()
            elif isinstance(obj, float):
                if np.isnan(obj) or np.isinf(obj): return 0.0
                return obj
            return obj

        response_data = {
            "status": "success",
            "results": analyzer.results,
            "summary": analyzer.get_summary_table().to_dict(orient='records')
        }
        
        return clean_data(response_data)
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_genetic_parameters_doc")
@limiter.limit("5/minute")
async def report_genetic_parameters_doc(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    traits: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        traits_list = [x.strip() for x in traits.split(',')]
        analyzer = perform_genetic_parameters(df, genotype_col, rep_col, traits_list)
        buf = analyzer.create_report()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Genetic_Parameters_Report.docx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_genetic_parameters_excel")
@limiter.limit("5/minute")
async def report_genetic_parameters_excel(
    request: Request,
    file: UploadFile = File(...),
    genotype_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    traits: str = Form(..., max_length=500)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        traits_list = [x.strip() for x in traits.split(',')]
        analyzer = perform_genetic_parameters(df, genotype_col, rep_col, traits_list)
        buf = analyzer.create_excel()
        return StreamingResponse(buf, 
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            headers={"Content-Disposition": "attachment; filename=Genetic_Parameters_Output.xlsx"})
    except Exception as e: return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})


# Eberhart-Russell Stability Analysis
def perform_eberhart_russell_analysis(df, geno_col, env_col, rep_col, trait_col, model_type):
    analyzer = EberhartRussellAnalyzer(df, geno_col, env_col, rep_col, trait_col, model_type)
    analyzer.validate()
    results = analyzer.run_analysis()
    return analyzer, results

@app.post("/analyze_eberhart_russell")
@limiter.limit("10/minute")
async def analyze_eberhart_russell(
    request: Request,
    file: UploadFile = File(...),
    geno_col: str = Form(..., max_length=100),
    env_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_col: str = Form(..., max_length=100),
    model_type: str = Form("fixed", max_length=20)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer, results = perform_eberhart_russell_analysis(
            df, geno_col, env_col, rep_col, trait_col, model_type
        )
        
        # Serialize Response
        # Pooled ANOVA
        pa = {}
        for k, v in results['pooled_anova'].items():
            pa[k] = {
                "df": int(v["df"]),
                "SS": safe_float(v["SS"]),
                "MS": safe_float(v["MS"]),
                "F": safe_float(v["F"]),
                "P": safe_float(v["P"]),
                "sig": get_sig(v["P"])
            }
            
        # Stability ANOVA
        sa = {}
        for k, v in results['stability_anova'].items():
            sa[k] = {
                "df": int(v["df"]),
                "SS": safe_float(v["SS"]),
                "MS": safe_float(v["MS"]),
                "F": safe_float(v.get("F", 0)),
                "P": safe_float(v.get("P", 1)),
                "sig": get_sig(v.get("P", 1))
            }
            
        # Environmental Indices
        ei = []
        env_means = df.groupby(env_col)[trait_col].mean()
        for env, idx in results['env_indices'].items():
            ei.append({
                "env": str(env),
                "mean": safe_float(env_means[env]),
                "index": safe_float(idx)
            })
            
        # Stability Parameters
        sp = []
        for p in results['stability_parameters']:
            sp.append({
                "genotype": str(p["Genotype"]),
                "mean": safe_float(p["Mean"]),
                "bi": safe_float(p["bi"]),
                "se_bi": safe_float(p["SE_bi"]),
                "t_b0": safe_float(p["t_b0"]),
                "p_b0": safe_float(p["p_b0"]),
                "t_b1": safe_float(p["t_b1"]),
                "p_b1": safe_float(p["p_b1"]),
                "ms_di": safe_float(p["MS_di"]),
                "s2di": safe_float(p["S2di"]),
                "f_s2di": safe_float(p["F_S2di"]),
                "p_s2di": safe_float(p["p_S2di"]),
                "inference": p["Inference"]
            })
            
        return {
            "status": "success",
            "bartlett": {
                "stat": safe_float(results['bartlett']['stat']),
                "p": safe_float(results['bartlett']['p'])
            },
            "pooled_anova": pa,
            "stability_anova": sa,
            "env_indices": ei,
            "stability_parameters": sp,
            "grand_mean": safe_float(results['grand_mean'])
        }
        
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_eberhart_russell")
@limiter.limit("5/minute")
async def report_eberhart_russell(
    request: Request,
    file: UploadFile = File(...),
    geno_col: str = Form(..., max_length=100),
    env_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_col: str = Form(..., max_length=100),
    model_type: str = Form("fixed", max_length=20)
):
    try:
        validate_file(file)
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer, _ = perform_eberhart_russell_analysis(
            df, geno_col, env_col, rep_col, trait_col, model_type
        )
        
        report_buffer = analyzer.create_report()
        filename = "Eberhart_Russell_Stability_Report.docx"
        
        return StreamingResponse(
            report_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Helper for One-Sample t-Test
def perform_one_sample_t_test(df, value_col, mu_0, alpha):
    analyzer = OneSampleTTestAnalyzer(df, value_col, mu_0, alpha)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_one_sample_t_test")
@limiter.limit("10/minute")
async def analyze_one_sample_t_test(
    request: Request,
    file: UploadFile = File(...),
    value_col: str = Form(..., max_length=100),
    mu_0: float = Form(0.0),
    alpha: float = Form(0.05)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        
        return {
            "status": "success",
            "descriptive": {k: safe_float(v) for k, v in analyzer.descriptive_stats.items()},
            "normality": {
                "W": safe_float(analyzer.normality_results["W"]),
                "p": safe_float(analyzer.normality_results["p_value"]),
                "interpretation": analyzer.normality_results["Interpretation"]
            },
            "t_test": {
                "mu_0": safe_float(analyzer.t_test_results["HypotheticalMean"]),
                "mean": safe_float(analyzer.t_test_results["SampleMean"]),
                "se": safe_float(analyzer.t_test_results["StdError"]),
                "t_value": safe_float(analyzer.t_test_results["t_value"]),
                "df": int(analyzer.t_test_results["df"]),
                "p_value": safe_float(analyzer.t_test_results["p_value"]),
                "lower_ci": safe_float(analyzer.t_test_results["Lower_CI"]),
                "upper_ci": safe_float(analyzer.t_test_results["Upper_CI"]),
                "conclusion": analyzer.t_test_results["Conclusion"]
            },
            "interpretation": analyzer.get_interpretation()
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_one_sample_t_test")
@limiter.limit("5/minute")
async def report_one_sample_t_test(
    request: Request,
    file: UploadFile = File(...),
    value_col: str = Form(..., max_length=100),
    mu_0: float = Form(0.0),
    alpha: float = Form(0.05)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_one_sample_t_test(df, value_col, mu_0, alpha)
        
        report_buffer = analyzer.create_report()
        filename = f"One_Sample_t_Test_Report_{value_col}.docx"
        
        return StreamingResponse(
            report_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Helper for Two-Sample t-Test
def perform_two_sample_t_test(df, category_col, value_col, alpha, variance_option):
    analyzer = TwoSampleTTestAnalyzer(df, category_col, value_col, alpha, variance_option)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_two_sample_t_test")
@limiter.limit("10/minute")
async def analyze_two_sample_t_test(
    request: Request,
    file: UploadFile = File(...),
    category_col: str = Form(..., max_length=100),
    value_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    variance_option: str = Form("bartlett", max_length=20)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_two_sample_t_test(df, category_col, value_col, alpha, variance_option)
        
        return {
            "status": "success",
            "descriptive": {g: {k: safe_float(v) for k, v in s.items()} for g, s in analyzer.descriptive_stats.items()},
            "normality": {g: {
                "W": safe_float(n["W"]),
                "p": safe_float(n["p_value"]),
                "interpretation": n["Interpretation"]
            } for g, n in analyzer.normality_results.items()},
            "bartlett": {
                "statistic": safe_float(analyzer.bartlett_results["Statistic"]),
                "p_value": safe_float(analyzer.bartlett_results["p_value"]),
                "interpretation": analyzer.bartlett_results["Interpretation"]
            },
            "t_test": {
                "test_type": analyzer.t_test_results["TestType"],
                "mean1": safe_float(analyzer.t_test_results["Mean1"]),
                "mean2": safe_float(analyzer.t_test_results["Mean2"]),
                "diff": safe_float(analyzer.t_test_results["MeanDiff"]),
                "t_value": safe_float(analyzer.t_test_results["t_value"]),
                "df": safe_float(analyzer.t_test_results["df"]),
                "p_value": safe_float(analyzer.t_test_results["p_value"]),
                "lower_ci": safe_float(analyzer.t_test_results["Lower_CI"]),
                "upper_ci": safe_float(analyzer.t_test_results["Upper_CI"]),
                "conclusion": analyzer.t_test_results["Conclusion"]
            },
            "interpretation": analyzer.get_interpretation()
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_two_sample_t_test")
@limiter.limit("5/minute")
async def report_two_sample_t_test(
    request: Request,
    file: UploadFile = File(...),
    category_col: str = Form(..., max_length=100),
    value_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    variance_option: str = Form("bartlett", max_length=20)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_two_sample_t_test(df, category_col, value_col, alpha, variance_option)
        
        report_buffer = analyzer.create_report()
        filename = f"Two_Sample_t_Test_Report_{value_col}.docx"
        
        return StreamingResponse(
            report_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Helper for Paired t-Test
def perform_paired_t_test(df, col1, col2, alpha, d0):
    analyzer = PairedTTestAnalyzer(df, col1, col2, alpha, d0)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_paired_t_test")
@limiter.limit("10/minute")
async def analyze_paired_t_test(
    request: Request,
    file: UploadFile = File(...),
    col1: str = Form(..., max_length=100),
    col2: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    d0: float = Form(0.0)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_paired_t_test(df, col1, col2, alpha, d0)
        
        return {
            "status": "success",
            "descriptive": {k: {m: safe_float(v) for m, v in s.items()} for k, s in analyzer.descriptive_stats.items()},
            "diff_stats": {k: safe_float(v) for k, v in analyzer.diff_stats.items()},
            "normality": {
                "W": safe_float(analyzer.normality_results["W"]),
                "p": safe_float(analyzer.normality_results["p_value"]),
                "interpretation": analyzer.normality_results["Interpretation"]
            },
            "t_test": {
                "d0": safe_float(analyzer.t_test_results["d0"]),
                "mean_diff": safe_float(analyzer.t_test_results["MeanDiff"]),
                "se": safe_float(analyzer.t_test_results["StdError"]),
                "t_value": safe_float(analyzer.t_test_results["t_value"]),
                "df": int(analyzer.t_test_results["df"]),
                "p_value": safe_float(analyzer.t_test_results["p_value"]),
                "lower_ci": safe_float(analyzer.t_test_results["Lower_CI"]),
                "upper_ci": safe_float(analyzer.t_test_results["Upper_CI"]),
                "conclusion": analyzer.t_test_results["Conclusion"]
            },
            "interpretation": analyzer.get_interpretation()
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_paired_t_test")
@limiter.limit("5/minute")
async def report_paired_t_test(
    request: Request,
    file: UploadFile = File(...),
    col1: str = Form(..., max_length=100),
    col2: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    d0: float = Form(0.0)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_paired_t_test(df, col1, col2, alpha, d0)
        
        report_buffer = analyzer.create_report()
        filename = f"Paired_t_Test_Report_{col1}_vs_{col2}.docx"
        
        return StreamingResponse(
            report_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Helper for Regression Analysis
def perform_regression_analysis(df, y_col, x_cols, model_type, degree, alpha):
    analyzer = RegressionAnalyzer(df, y_col, x_cols, model_type, degree, alpha)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_regression")
@limiter.limit("10/minute")
async def analyze_regression(
    request: Request,
    file: UploadFile = File(...),
    y_col: str = Form(..., max_length=100),
    x_cols: str = Form(..., max_length=500),  # Comma separated
    model_type: str = Form("linear", max_length=20),
    degree: int = Form(2),
    alpha: float = Form(0.05)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        x_list = [x.strip() for x in x_cols.split(",") if x.strip()]
        analyzer = perform_regression_analysis(df, y_col, x_list, model_type, degree, alpha)
        
        return {
            "status": "success",
            "summary": {k: (safe_float(v) if isinstance(v, (float, int, np.float64, np.int64)) else v) for k, v in analyzer.summary_stats.items()},
            "coefficients": [{k: (safe_float(v) if k != 'Variable' else v) for k, v in c.items()} for c in analyzer.coefficient_table],
            "anova": [{k: (safe_float(v) if k != 'Source' else v) for k, v in a.items()} for a in analyzer.anova_table],
            "interpretation": analyzer.get_interpretation()
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_regression")
@limiter.limit("5/minute")
async def report_regression(
    request: Request,
    file: UploadFile = File(...),
    y_col: str = Form(..., max_length=100),
    x_cols: str = Form(..., max_length=500),
    model_type: str = Form("linear", max_length=20),
    degree: int = Form(2),
    alpha: float = Form(0.05)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        x_list = [x.strip() for x in x_cols.split(",") if x.strip()]
        analyzer = perform_regression_analysis(df, y_col, x_list, model_type, degree, alpha)
        
        report_buffer = analyzer.create_report()
        filename = f"Regression_Report_{y_col}.docx"
        
        return StreamingResponse(
            report_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Helper for Line x Tester Analysis
def perform_line_tester_analysis(df, line_col, tester_col, rep_col, trait_col, alpha):
    analyzer = LineTesterAnalyzer(df, line_col, tester_col, rep_col, trait_col, alpha)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_line_tester")
@limiter.limit("10/minute")
async def analyze_line_tester(
    request: Request,
    file: UploadFile = File(...),
    line_col: str = Form(..., max_length=100),
    tester_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_line_tester_analysis(df, line_col, tester_col, rep_col, trait_col, alpha)
        
        return {
            "status": "success",
            "anova": [{k: (safe_float(v) if k not in ['Source', 'DF'] else v) for k, v in row.items()} for row in analyzer.anova_table],
            "gca_lines": [{k: (safe_float(v) if k not in ['Line', 'Sig'] else v) for k, v in row.items()} for row in analyzer.gca_lines],
            "gca_testers": [{k: (safe_float(v) if k not in ['Tester', 'Sig'] else v) for k, v in row.items()} for row in analyzer.gca_testers],
            "sca": [{k: (safe_float(v) if k not in ['Hybrid', 'Sig'] else v) for k, v in row.items()} for row in analyzer.sca_effects],
            "variances": {k: (safe_float(v) if k != 'GeneAction' else v) for k, v in analyzer.genetic_variances.items()},
            "summary": {k: safe_float(v) for k, v in analyzer.summary_stats.items()},
            "interpretation": analyzer.get_interpretation()
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_line_tester")
@limiter.limit("5/minute")
async def report_line_tester(
    request: Request,
    file: UploadFile = File(...),
    line_col: str = Form(..., max_length=100),
    tester_col: str = Form(..., max_length=100),
    rep_col: str = Form(..., max_length=100),
    trait_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        analyzer = perform_line_tester_analysis(df, line_col, tester_col, rep_col, trait_col, alpha)
        
        report_buffer = analyzer.create_report()
        filename = f"Line_Tester_Report_{trait_col}.docx"
        
        return StreamingResponse(
            report_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

# Helper for F-Test
def perform_f_test_analysis(df, category_col, value_col, alpha, mode):
    analyzer = FTestAnalyzer(df, category_col, value_col, alpha, mode)
    analyzer.validate()
    analyzer.run_analysis()
    return analyzer

@app.post("/analyze_f_test")
@limiter.limit("10/minute")
async def analyze_f_test(
    request: Request,
    file: UploadFile = File(...),
    category_col: str = Form(..., max_length=100),
    value_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    mode: str = Form("long", max_length=20)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_f_test_analysis(df, category_col, value_col, alpha, mode)
        
        # Serialize Response
        res = {
            "status": "success",
            "descriptive": {g: {k: safe_float(v) for k, v in s.items()} for g, s in analyzer.descriptive_stats.items()},
            "normality": {g: {
                "W": safe_float(n["W"]),
                "p": safe_float(n["p_value"]),
                "interpretation": n["Interpretation"]
            } for g, n in analyzer.normality_results.items()},
            "f_test": {
                "var1": safe_float(analyzer.f_test_results["Variance1"]),
                "var2": safe_float(analyzer.f_test_results["Variance2"]),
                "f_value": safe_float(analyzer.f_test_results["F_value"]),
                "p_value": safe_float(analyzer.f_test_results["p_value"]),
                "df1": int(analyzer.f_test_results["df1"]),
                "df2": int(analyzer.f_test_results["df2"]),
                "lower_ci": safe_float(analyzer.f_test_results["Lower_CI"]),
                "upper_ci": safe_float(analyzer.f_test_results["Upper_CI"]),
                "conclusion": analyzer.f_test_results["Conclusion"]
            },
            "interpretation": analyzer.get_interpretation()
        }
        return res
        
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

@app.post("/report_f_test")
@limiter.limit("5/minute")
async def report_f_test(
    request: Request,
    file: UploadFile = File(...),
    category_col: str = Form(..., max_length=100),
    value_col: str = Form(..., max_length=100),
    alpha: float = Form(0.05),
    mode: str = Form("long", max_length=20)
):
    try:
        validate_file(file)
        if not (0 < alpha < 1):
             raise HTTPException(status_code=400, detail="Alpha must be between 0 and 1.")
        contents = await file.read()
        df = pd.read_csv(io.BytesIO(contents))
        
        analyzer = perform_f_test_analysis(df, category_col, value_col, alpha, mode)
        
        report_buffer = analyzer.create_report()
        filename = f"F_Test_Report_{value_col}.docx" if mode == 'long' else f"F_Test_Report_{category_col}_vs_{value_col}.docx"
        
        return StreamingResponse(
            report_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(status_code=400, content={"status": "error", "message": str(e)})

app.mount("/vendor", StaticFiles(directory=os.path.join(ROOT_DIR, "cloudflare", "vendor")), name="vendor")
app.mount("/assets", StaticFiles(directory=os.path.join(ROOT_DIR, "cloudflare", "assets")), name="assets")

app.mount("/", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
