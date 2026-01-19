
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import networkx as nx
import io
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from scipy import stats

class PathAnalyzer:
    def __init__(self, df, dep_var, indep_vars):
        self.df = df
        self.dep_var = dep_var
        self.indep_vars = indep_vars
        
        # Results
        self.results = {}
        self.diagram_buf = None
        
    def validate(self):
        # 1. Check columns
        if self.dep_var not in self.df.columns:
            raise ValueError(f"Dependent variable '{self.dep_var}' not found.")
        
        missing = [v for v in self.indep_vars if v not in self.df.columns]
        if missing:
            raise ValueError(f"Missing independent variables: {missing}")
            
        # 2. Check numeric
        all_vars = [self.dep_var] + self.indep_vars
        self.data = self.df[all_vars].apply(pd.to_numeric, errors='coerce')
        
        if self.data.isnull().any().any():
            raise ValueError("Data contains missing values or non-numeric characters.")
            
        # 3. Sample size check (n >= 10 + k)
        n, k = self.data.shape[0], len(self.indep_vars)
        if n < 10 + k:
            # Minimal scientific warning, though we might allow it with a warning
            pass 
            
        self.n = n
        self.k = k
        
    def run_analysis(self):
        # 1. Standardization (Z-score)
        # However, correlation based path coefficients can be computed directly from correlation matrix
        # as standardization is implicit in correlation.
        
        # 2. Correlation Matrix
        corr_matrix = self.data.corr(method='pearson')
        self.corr_matrix = corr_matrix
        
        # 3. Partition Matrices
        # R_xx : Correlation among Independent Variables (k x k)
        # r_yx : Correlation between Dependent and Independent (k x 1)
        
        R_xx = corr_matrix.loc[self.indep_vars, self.indep_vars].values
        r_yx = corr_matrix.loc[self.indep_vars, self.dep_var].values
        
        # 4. Compute Direct Effects (Path Coefficients)
        # P = inv(R_xx) * r_yx
        try:
            R_xx_inv = np.linalg.inv(R_xx)
        except np.linalg.LinAlgError:
            raise ValueError("Matrix R_xx is singular. High multicollinearity detected among independent variables.")
            
        P = R_xx_inv @ r_yx  # Shape (k,)
        
        # 5. Compute Indirect Effects
        # IE_Xi->Y (via Xj) = r_XiXj * P_YXj
        # We need a table of indirect effects for each Xi
        
        indirect_effects = {} # Key: Xi, Value: List of (via Xj, value)
        total_indirect = np.zeros(self.k)
        
        for i, var_i in enumerate(self.indep_vars):
            indirect_effects[var_i] = []
            sum_ind = 0.0
            for j, var_j in enumerate(self.indep_vars):
                if i == j: continue
                
                r_ij = R_xx[i, j]
                p_j = P[j]
                ie_ij = r_ij * p_j
                
                indirect_effects[var_i].append({
                    "via": var_j,
                    "val": ie_ij
                })
                sum_ind += ie_ij
            
            total_indirect[i] = sum_ind
            
        # 6. Total Effects & Validation
        # Total = Direct + Indirect
        # Should approx match correlation r_yxi
        total_effects = P + total_indirect
        
        # 7. Residual Effect
        # R^2 = Sum(r_yxi * P_yxi)
        # Residual = sqrt(1 - R^2)
        
        R2 = np.sum(r_yx * P)
        residual = np.sqrt(1 - R2) if R2 < 1 else 0.0
        
        # Store results
        self.results = {
            "direct_effects": dict(zip(self.indep_vars, P)),
            "indirect_effects": indirect_effects,
            "total_indirect": dict(zip(self.indep_vars, total_indirect)),
            "total_effects": dict(zip(self.indep_vars, total_effects)),
            "correlations_y": dict(zip(self.indep_vars, r_yx)),
            "R2": R2,
            "residual": residual
        }
        
    def generate_diagram(self):
        plt.figure(figsize=(10, 8))
        G = nx.DiGraph()
        
        # Nodes
        G.add_node(self.dep_var, layer=1)
        for var in self.indep_vars:
            G.add_node(var, layer=0)
            
        # Positions
        pos = {}
        # Central Dependent Variable
        pos[self.dep_var] = np.array([1.0, 0.5])
        
        # Independent variables in a semi-circle or vertical line on left
        y_range = np.linspace(0.1, 0.9, self.k)
        for i, var in enumerate(self.indep_vars):
            pos[var] = np.array([0.0, y_range[i]])
            
        # Draw Nodes
        nx.draw_networkx_nodes(G, pos, nodelist=[self.dep_var], node_color='lightblue', node_size=3000, edgecolors='black')
        nx.draw_networkx_nodes(G, pos, nodelist=self.indep_vars, node_color='lightgreen', node_size=2500, edgecolors='black')
        nx.draw_networkx_labels(G, pos, font_size=10, font_weight='bold')
        
        # Draw Direct Effects (Solid Arrows)
        directs = self.results['direct_effects']
        for var, val in directs.items():
            # Edge color matches strength? Or fixed blue as req
            col = 'blue'
            # Width prop to magnitude
            width = 1.0 + abs(val) * 3
            nx.draw_networkx_edges(G, pos, edgelist=[(var, self.dep_var)], edge_color=col, width=width, arrowsize=20)
            
            # Label: Path Coeff
            mid_x = (pos[var][0] + pos[self.dep_var][0]) / 2
            mid_y = (pos[var][1] + pos[self.dep_var][1]) / 2
            plt.text(mid_x, mid_y, f"{val:.3f}", color='blue', fontsize=9, fontweight='bold', bbox=dict(facecolor='white', alpha=0.7, edgecolor='none'))

        # Draw Residual
        res_val = self.results['residual']
        # Phantom node for residual
        res_pos = np.array([1.0, 0.2]) # Below Y
        plt.arrow(res_pos[0], res_pos[1], 0, 0.25, color='red', linestyle='--', head_width=0.03)
        plt.text(res_pos[0], res_pos[1]-0.05, f"Residual\n{res_val:.3f}", color='red', ha='center', fontsize=9)
        
        plt.axis('off')
        plt.title("Path Diagram", fontsize=14)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        buf.seek(0)
        self.diagram_buf = buf
        plt.close()
        return buf
        
    def create_report_doc(self):
        doc = Document()
        doc.add_heading('Path Analysis Report', 0)
        
        res = self.results
        
        # 1. Summary
        doc.add_heading('1. Model Summary', level=1)
        doc.add_paragraph(f"Dependent Variable: {self.dep_var}")
        doc.add_paragraph(f"Independent Variables: {', '.join(self.indep_vars)}")
        doc.add_paragraph(f"Sample Size (n): {self.n}")
        doc.add_paragraph(f"Coefficient of Determination (R²): {res['R2']:.4f}")
        doc.add_paragraph(f"Residual Effect: {res['residual']:.4f} ({(res['residual']**2)*100:.2f}% of variation unexplained)")
        
        # 2. Direct & Indirect Effects Table
        doc.add_heading('2. Path Coefficients (Direct & Indirect)', level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Variable"
        hdr[1].text = "Total Correlation (r)"
        hdr[2].text = "Direct Effect"
        hdr[3].text = "Total Indirect"
        hdr[4].text = "Total Effect (Calc)"
        hdr[5].text = "Interpretation"
        
        for var in self.indep_vars:
            row = table.add_row().cells
            direct = res['direct_effects'][var]
            indirect = res['total_indirect'][var]
            total_corr = res['correlations_y'][var]
            calc_total = res['total_effects'][var]
            
            row[0].text = var
            row[1].text = f"{total_corr:.3f}"
            row[2].text = f"{direct:.3f}"
            row[3].text = f"{indirect:.3f}"
            row[4].text = f"{calc_total:.3f}"
            
            # Interpretation Rule
            interp = []
            if abs(direct) >= 0.3:
                interp.append("Strong Direct")
            if abs(total_corr) > 0.3 and abs(direct) < 0.1:
                interp.append("Indirect Influence")
            if not interp:
                interp.append("Weak/Moderate")
            
            row[5].text = ", ".join(interp)
            
        # 3. Path Diagram Image
        if self.diagram_buf:
            doc.add_heading('3. Path Diagram', level=1)
            self.diagram_buf.seek(0)
            doc.add_picture(self.diagram_buf, width=Inches(6))
            
        # 4. Detailed Narrative
        doc.add_heading('4. Interpretation', level=1)
        for var in self.indep_vars:
            direct = res['direct_effects'][var]
            if abs(direct) >= 0.3:
                doc.add_paragraph(f"{var} has a strong direct influence on {self.dep_var} (Path Coeff = {direct:.3f}). This trait can be used as a primary selection criterion.")
            elif res['correlations_y'][var] > 0.3 and abs(direct) < 0.2:
                doc.add_paragraph(f"{var} has a high correlation with {self.dep_var} (r={res['correlations_y'][var]:.3f}) but a low direct effect ({direct:.3f}). Its influence is primarily indirect (via other traits).")
                
        # Residual
        if res['residual'] > 0.3: # >30% roughly? 
            doc.add_paragraph(f"WARNING: The residual effect is {res['residual']:.3f}, indicating that important traits influencing {self.dep_var} may have been omitted from the study.")
            
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_output_excel(self):
        wb = Workbook()
        
        # 1. Summary
        ws1 = wb.active
        ws1.title = "Summary"
        res = self.results
        ws1.append(["Metric", "Value"])
        ws1.append(["Dependent Variable", self.dep_var])
        ws1.append(["R Square", res['R2']])
        ws1.append(["Residual", res['residual']])
        
        # 2. Effects Table
        ws2 = wb.create_sheet("Effects_Summary")
        ws2.append(["Variable", "Direct Effect", "Total Indirect", "Total Correlation", "Check (Diff)"])
        for var in self.indep_vars:
            diff = res['correlations_y'][var] - res['total_effects'][var]
            ws2.append([
                var,
                res['direct_effects'][var],
                res['total_indirect'][var],
                res['correlations_y'][var],
                diff
            ])

        # 3. Indirect Breakdown
        ws3 = wb.create_sheet("Indirect_Breakdown")
        ws3.append(["Variable", "Via Mediator", "Effect Contribution"])
        for var in self.indep_vars:
            for item in res['indirect_effects'][var]:
                ws3.append([var, item['via'], item['val']])
                
        # 4. Correlation Matrix
        ws4 = wb.create_sheet("Correlation_Matrix")
        df_corr = self.corr_matrix
        ws4.append(["Variable"] + list(df_corr.columns))
        for idx, row in df_corr.iterrows():
            ws4.append([idx] + list(row))
            
        f = io.BytesIO()
        wb.save(f)
        f.seek(0)
        return f
