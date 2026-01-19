import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
import io

class PCAAnalyzer:
    def __init__(self, df, obs_col, var_cols):
        self.df = df
        self.obs_col = obs_col
        self.var_cols = var_cols
        
        # Results container
        self.pca_res = {}
        self.plots = {}
        
    def validate(self):
        # 1. Check columns exist
        if self.obs_col not in self.df.columns:
            raise ValueError(f"Observation column '{self.obs_col}' not found.")
        
        missing_vars = [v for v in self.var_cols if v not in self.df.columns]
        if missing_vars:
            raise ValueError(f"Missing variables: {missing_vars}")
            
        # 2. Check for missing values
        data = self.df[self.var_cols]
        if data.isnull().any().any():
            raise ValueError("Dataset contains missing values. Please handle missing data before running PCA.")
            
        # 3. Check numeric
        try:
            data = data.apply(pd.to_numeric)
        except:
            raise ValueError("Selected variables must be numeric.")
            
        # 4. Handle duplicates - Aggregate by Mean
        # Instead of error, we group by obs_col and take mean
        if self.df[self.obs_col].duplicated().any():
            # Create a combined DF properly
            temp_df = self.df[[self.obs_col] + self.var_cols].copy()
            # Group by Observation ID and take Mean of variables
            # Convert Obs ID to string to ensure clean grouping if mixed types
            temp_df[self.obs_col] = temp_df[self.obs_col].astype(str)
            aggregated = temp_df.groupby(self.obs_col)[self.var_cols].mean().reset_index()
            
            # Update data matrix and obs labels
            self.data_matrix = aggregated[self.var_cols]
            self.obs_labels = aggregated[self.obs_col].values
        else:
            self.data_matrix = data
            self.obs_labels = self.df[self.obs_col].astype(str).values

        # 5. Check dimensions
        n, p = self.data_matrix.shape
        if p < 2:
            raise ValueError("At least 2 variables are required for PCA.")
        if n < 2:
            raise ValueError("At least 2 unique observations are required.")
            
        self.n, self.p = n, p

    def run_pca(self):
        # Data Matrix is already set in validate (possibly aggregated)
        X = self.data_matrix.values
        
        # 1. Standardization (Z-score)
        # Z_ij = (X_ij - Mean_j) / Std_j
        means = np.mean(X, axis=0)
        stds = np.std(X, axis=0, ddof=1) # Sample std deviation
        
        # Check for zero variance
        if np.any(stds == 0):
             raise ValueError("One or more variables have zero variance (constant values). Remove them.")
             
        Z = (X - means) / stds
        self.Z = Z
        
        # 2. SVD
        # Z = U * Sigma * Vt
        # full_matrices=False -> U is (n, k), Vt is (k, p) where k = min(n, p)
        U, S, Vt = np.linalg.svd(Z, full_matrices=False)
        
        # 3. Eigenvalues & Variance
        # Eigenvalue_k = S_k^2 / (n - 1)
        eig_vals = (S ** 2) / (self.n - 1)
        total_var = np.sum(eig_vals)
        var_pct = (eig_vals / total_var) * 100
        cum_var_pct = np.cumsum(var_pct)
        
        # 4. Scores & Loadings (Raw)
        # Scores = Z * V  OR  U * S
        # Loadings (Eigenvectors) = V^T (columns of V matrix, rows of Vt)
        # Note: In standard PCA text, Loadings often refers to Correlation(X, PC) = Eigenvector * sqrt(Eigenvalue)
        # Let's clarify:
        # V (from SVD) are Principal Axes / Eigenvectors.
        # "Loadings" in many contexts (SPSS/SAS) are Eigenvectors scaled by sqrt(Eigenvalue).
        # Prompt says: "Loading_jk = Correlation(X_j, PC_k)"
        # This confirms we need Scaled Loadings for the table/interpretation.
        
        V = Vt.T
        scores_raw = U @ np.diag(S) # or Z @ V
        loadings_corr = V @ np.diag(np.sqrt(eig_vals)) # Correlation Loadings
        
        # 5. Adjusted Scores & Loadings (Biplot alpha=0.5)
        # Adj Scores = U * S^0.5
        # Adj Loadings = V * S^0.5
        scores_adj = U @ np.diag(S**0.5)
        loadings_adj = V @ np.diag(S**0.5)
        
        # Save results
        self.pca_res = {
             "eigenvalues": eig_vals,
             "variance_pct": var_pct,
             "cum_variance_pct": cum_var_pct,
             "scores": scores_raw,
             "loadings_corr": loadings_corr, # Use this for variable correlations
             "eigenvectors": V, # Raw V
             "scores_adj": scores_adj,
             "loadings_adj": loadings_adj
        }
        
    def generate_plots(self):
        plt.style.use('ggplot')
        plots = {}
        
        res = self.pca_res
        pc1_expl = res['variance_pct'][0]
        pc2_expl = res['variance_pct'][1] if len(res['variance_pct']) > 1 else 0
        
        xlabel = f"PC1 ({pc1_expl:.2f}%)"
        ylabel = f"PC2 ({pc2_expl:.2f}%)"
        
        def save_plot(name):
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
            buf.seek(0)
            plots[name] = buf
            plt.close()

        # Helper: Biplot with Auto-Scaling
        def draw_biplot(scores, loadings, title, draw_obs=True, draw_vars=True, labels=None, circle=False):
            plt.figure(figsize=(10, 8))
            
            # Determine Scaling Factor if Mixing (Raw Biplot)
            scale_factor = 1.0
            if draw_obs and draw_vars:
                # If absolute max score is much larger than absolute max loading, scale loadings up
                max_score = np.max(np.abs(scores[:, :2])) if scores.size > 0 else 1.0
                max_loading = np.max(np.abs(loadings[:, :2])) if loadings.size > 0 else 1.0
                
                # If ratios differ significantly, scale loadings to 80% of score range
                if max_loading > 0 and max_score > 0:
                     scale_factor = (max_score / max_loading) * 0.8
            
            # Draw Observations
            if draw_obs:
                plt.scatter(scores[:, 0], scores[:, 1], alpha=0.7, c='steelblue')
                if labels is not None:
                    # Limit labels if too many?
                    for i, txt in enumerate(labels):
                        plt.annotate(txt, (scores[i, 0], scores[i, 1]), size=9, alpha=0.8)
                        
            # Draw Variables
            if draw_vars:
                loadings_scaled = loadings * scale_factor
                
                for i in range(loadings_scaled.shape[0]):
                    x_v = loadings_scaled[i, 0]
                    y_v = loadings_scaled[i, 1]
                    
                    plt.arrow(0, 0, x_v, y_v, color='darkred', alpha=0.9, width=0.005, head_width=0.05*scale_factor)
                    
                    # Adjust label position
                    plt.text(x_v * 1.15, y_v * 1.15, self.var_cols[i], color='darkred', ha='center', va='center', weight='semibold')
                    
            # Draw Unit Circle for Correlation Circle Plot
            if circle and draw_vars and not draw_obs:
                circ = plt.Circle((0, 0), 1, color='gray', fill=False, linestyle='--')
                plt.gca().add_patch(circ)
                plt.xlim(-1.1, 1.1)
                plt.ylim(-1.1, 1.1)
                plt.gca().set_aspect('equal')
                
            plt.xlabel(xlabel)
            plt.ylabel(ylabel)
            plt.title(title)
            plt.grid(True, linestyle='--', alpha=0.6)
            plt.axhline(0, color='black', linewidth=0.8)
            plt.axvline(0, color='black', linewidth=0.8)
            
            if scale_factor != 1.0 and draw_vars:
                plt.figtext(0.99, 0.01, f'Variables Scaled x{scale_factor:.2f} for visibility', ha='right', size=8, style='italic')

        # 1. Biplot PC1 vs PC2 (Raw)
        draw_biplot(res['scores'], res['loadings_corr'], "Biplot (PC1 vs PC2)", labels=self.obs_labels)
        save_plot('biplot_raw')
        
        # 2. Biplot Adjusted (Symmetric)
        # Usually symmetric biplot doesn't need extra scaling as they are in similar space
        draw_biplot(res['scores_adj'], res['loadings_adj'], "Biplot - Adjusted (Symmetric)", labels=self.obs_labels)
        save_plot('biplot_adj')
        
        # 3. Biplot Variables (Correlation Circle)
        # Standard "Variables Only" plot is essentially the Correlation Circle
        draw_biplot(res['scores'], res['loadings_corr'], "Correlation Circle (Variables)", draw_obs=False, circle=True)
        save_plot('biplot_vars')
        
        # 4. Biplot Variables Adj
        draw_biplot(res['scores_adj'], res['loadings_adj'], "Biplot - Variables Adjusted", draw_obs=False, labels=None)
        save_plot('biplot_vars_adj')
        
        # 5. Biplot Obs
        draw_biplot(res['scores'], res['loadings_corr'], "Biplot - Observations Only", draw_vars=False, labels=self.obs_labels)
        save_plot('biplot_obs')
        
        # 6. Biplot Obs Adj
        draw_biplot(res['scores_adj'], res['loadings_adj'], "Biplot - Observations Adjusted", draw_vars=False, labels=self.obs_labels)
        save_plot('biplot_obs_adj')
        
        # SCREE PLOTS
        n_comps = len(res['eigenvalues'])
        x_ticks = np.arange(1, n_comps + 1)
        
        # 7. Eigenvalues
        plt.figure(figsize=(8, 5))
        plt.plot(x_ticks, res['eigenvalues'], 'o-', color='purple', linewidth=2)
        plt.title('Scree Plot - Eigenvalues')
        plt.xlabel('Principal Component')
        plt.ylabel('Eigenvalue')
        plt.xticks(x_ticks)
        plt.axhline(1, color='r', linestyle='--', label='Kaiser Criterion (1.0)')
        plt.legend()
        plt.grid(True, alpha=0.3)
        save_plot('scree_eig')
        
        # 8. Cumulative Eigenvalues
        plt.figure(figsize=(8, 5))
        plt.plot(x_ticks, np.cumsum(res['eigenvalues']), 'o-', color='purple', linewidth=2)
        plt.title('Scree Plot - Cumulative Eigenvalues')
        plt.xlabel('Principal Component')
        plt.ylabel('Cumulative Eigenvalue')
        plt.grid(True, alpha=0.3)
        save_plot('scree_eig_cum')
        
        # 9. Variance %
        plt.figure(figsize=(8, 5))
        plt.bar(x_ticks, res['variance_pct'], color='teal', alpha=0.8)
        plt.title('Scree Plot - % Variance Explained')
        plt.xlabel('Principal Component')
        plt.ylabel('Variance (%)')
        plt.grid(True, alpha=0.3)
        save_plot('scree_var')
        
        # 10. Cumulative Variance %
        plt.figure(figsize=(8, 5))
        plt.plot(x_ticks, res['cum_variance_pct'], 'o-', color='teal', linewidth=2)
        plt.title('Scree Plot - Cumulative Variance (%)')
        plt.xlabel('Principal Component')
        plt.ylabel('Cumulative Variance (%)')
        plt.ylim(0, 105)
        plt.grid(True, alpha=0.3)
        save_plot('scree_var_cum')
        
        # 11. Correlation Plot (Heatmap of Loadings)
        plt.figure(figsize=(8, max(5, self.p * 0.5))) # Auto-height
        k = min(5, self.p)
        sns.heatmap(res['loadings_corr'][:, :k], annot=True, cmap='RdBu_r', center=0, fmt=".3f",
                    yticklabels=self.var_cols, xticklabels=[f"PC{i+1}" for i in range(k)])
        plt.title('Correlation Plot (Loadings Heatmap)')
        save_plot('corr_plot')
        

        
        self.plots = plots
        return plots

    def create_report_plots(self):
        doc = Document()
        doc.add_heading('Principal Component Analysis - Plots', 0)
        
        # Add all 11 plots
        plot_order = [
            ('biplot_raw', '1. Biplot (PC1 vs PC2)'),
            ('biplot_adj', '2. Biplot - Adjusted (Symmetric)'),
            ('biplot_vars', '3. Biplot - Variables'),
            ('biplot_vars_adj', '4. Biplot - Variables Adjusted'),
            ('biplot_obs', '5. Biplot - Observations'),
            ('biplot_obs_adj', '6. Biplot - Observations Adjusted'),
            ('scree_eig', '7. Scree Plot - Eigenvalues'),
            ('scree_eig_cum', '8. Scree Plot - Cumulative Eigenvalues'),
            ('scree_var', '9. Scree Plot - Variance Explained (%)'),
            ('scree_var_cum', '10. Scree Plot - Cumulative Variance (%)'),
            ('corr_plot', '11. Correlation Plot (Loadings)')
        ]
        
        for key, title in plot_order:
            if key in self.plots:
                doc.add_heading(title, level=2)
                doc.add_picture(self.plots[key], width=Inches(6))
                doc.add_paragraph("")
                
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_report_interpretation(self):
        doc = Document()
        doc.add_heading('PCA Interpretation Report', 0)
        
        res = self.pca_res
        
        # Intro
        doc.add_paragraph(f"Principal Component Analysis (PCA) was performed on {self.n} observations and {self.p} variables using Singular Value Decomposition (SVD) on standardized data.")
        
        # Eigenvalues & Variance
        doc.add_heading('1. Eigenvalues and Variance', level=1)
        kaiser_pcs = np.sum(res['eigenvalues'] >= 1)
        doc.add_paragraph(f"Based on the Kaiser criterion (Eigenvalue >= 1), {kaiser_pcs} principal components were retained. Together, they explain {res['cum_variance_pct'][kaiser_pcs-1]:.2f}% of the total variation.")
        
        # Detailed Comp Interpretation
        doc.add_heading('2. Component Interpretation', level=1)
        
        for i in range(min(3, self.p)):
            pc_num = i + 1
            var_exp = res['variance_pct'][i]
            doc.add_heading(f"Principal Component {pc_num} (PC{pc_num})", level=2)
            doc.add_paragraph(f"PC{pc_num} explains {var_exp:.2f}% of the total variance.")
            
            # Loadings analysis
            loadings = res['loadings_corr'][:, i]
            # Threshold 0.4
            strong_pos = []
            strong_neg = []
            for j, val in enumerate(loadings):
                if val >= 0.4: strong_pos.append(self.var_cols[j])
                elif val <= -0.4: strong_neg.append(self.var_cols[j])
            
            if strong_pos:
                doc.add_paragraph(f"Variables positively associated with PC{pc_num}: {', '.join(strong_pos)}.")
            if strong_neg:
                doc.add_paragraph(f"Variables negatively associated with PC{pc_num}: {', '.join(strong_neg)}.")
                
            if not strong_pos and not strong_neg:
                doc.add_paragraph(f"No variables showed strong correlations (|r| >= 0.4) with PC{pc_num}.")

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    def create_output_excel(self):
        wb = Workbook()
        res = self.pca_res
        
        # 1. Eigenvalues & Variance
        ws1 = wb.active
        ws1.title = "Eigenvalues_Variance"
        ws1.append(["PC", "Eigenvalue", "Variance %", "Cumulative %"])
        for i in range(len(res['eigenvalues'])):
            ws1.append([i+1, res['eigenvalues'][i], res['variance_pct'][i], res['cum_variance_pct'][i]])
            
        # 2. Loadings (Correlation)
        ws2 = wb.create_sheet("Loadings_Correlation")
        ws2.append(["Variable"] + [f"PC{i+1}" for i in range(self.p)])
        for i, var in enumerate(self.var_cols):
            row = [var] + list(res['loadings_corr'][i, :])
            ws2.append(row)
            
        # 3. Scores
        ws3 = wb.create_sheet("Scores")
        ws3.append(["Observation"] + [f"PC{i+1}" for i in range(self.p)])
        for i, obs in enumerate(self.obs_labels):
            row = [obs] + list(res['scores'][i, :])
            ws3.append(row)
           
        # 4. Correlation Matrix (Input Data)
        ws4 = wb.create_sheet("Input_Correlation_Matrix")
        corr_matrix = self.df[self.var_cols].corr()
        ws4.append(["Variable"] + list(corr_matrix.columns))
        for idx, row in corr_matrix.iterrows():
            ws4.append([idx] + list(row))
            
        f = io.BytesIO()
        wb.save(f)
        f.seek(0)
        return f
